import os
import sys
from docx import Document
from unittest.mock import patch
import coverpage_generator
from coverpage_generator import generate_cover_page, get_all_placeholders

# Setup paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLES_DIR = os.path.join(os.path.dirname(BASE_DIR), 'Cover Pages', 'poorly formatted samples')
OUTPUT_DIR = os.path.join(os.path.dirname(BASE_DIR), 'outputs', 'repaired_samples')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# The specific file the user mentioned
SAMPLE_FILE = 'CoverPage_Arith Methrics of Algorithms_20260102_115613.docx'
SAMPLE_PATH = os.path.join(SAMPLES_DIR, SAMPLE_FILE)

def test_repair_sample():
    print(f"Testing repair on actual document: {SAMPLE_FILE}")
    
    if not os.path.exists(SAMPLE_PATH):
        print(f"❌ Sample file not found: {SAMPLE_PATH}")
        return

    # Data to fill the specific fields found in this document
    # Based on inspect_samples.py: 
    # {{SCHOOL/FACULTY}}, {{Field Supervisor’s name}}, {{PROJECT TOPIC}}, 
    # {{Supervisor’s Name}}, {{DEPARMENT}}, {{ACADEMIC YEAR}}
    data = {
        "documentType": "Project Report", # This triggers the right logic
        "institution": "The University of Bamenda",
        "faculty": "National Higher Polytechnic Institute",
        "department": "Computer Engineering",
        "title": "Arith Methrics of Algorithms", # Matching the filename hint
        "studentName": "Test Student",
        "studentId": "UBa21E999",
        "academicSupervisor": "Dr. Academic Sup", # Should map to Supervisor's Name
        "fieldSupervisor": "Engr. Field Sup",     # Should map to Field Supervisor's name
        "date": "2026-01-02",
        "level": "400 Level"
    }

    # Mock get_template_path to return our sample file
    # We need to patch it where it is used, which is inside coverpage_generator
    with patch('coverpage_generator.get_template_path') as mock_get_path:
        mock_get_path.return_value = SAMPLE_PATH
        
        # Inspect the document structure first
        print("Inspecting document structure...")
        doc_orig = Document(SAMPLE_PATH)
        for i, p in enumerate(doc_orig.paragraphs):
            if "{{" in p.text or "}}" in p.text:
                print(f"Para {i}: {repr(p.text)}")
        
        # Check tables too
        for t in doc_orig.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        if "{{" in p.text or "}}" in p.text:
                            print(f"Table Para: {repr(p.text)}")
                            
        # Check text boxes (via XML)
        from docx.oxml.ns import qn
        if doc_orig.element.body is not None:
            for txbx in doc_orig.element.body.iter(qn('w:txbxContent')):
                for p in txbx.iter(qn('w:p')):
                    # Reconstruct text from runs to see what python-docx sees
                    text = ""
                    for r in p.iter(qn('w:r')):
                        for t in r.iter(qn('w:t')):
                            if t.text: text += t.text
                    if "{{" in text or "}}" in text:
                        print(f"TextBox Para: {repr(text)}")

        print("Generating cover page using sample as template...")
        output_path, error = generate_cover_page(data)
        
        if error:
            print(f"❌ Error: {error}")
            return

        print(f"✅ Generated: {output_path}")
        
        # Verify the output
        doc = Document(output_path)
        placeholders = get_all_placeholders(doc)
        
        if not placeholders:
            print("✅ SUCCESS: All placeholders resolved!")
        else:
            print(f"❌ FAILED: Remaining placeholders: {placeholders}")
            
        # Check specific content
        full_text = []
        for p in doc.paragraphs: full_text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: full_text.append(p.text)
                    
        # Add text box content
        from docx.oxml.ns import qn
        if doc.element.body is not None:
            for txbx in doc.element.body.iter(qn('w:txbxContent')):
                for p in txbx.iter(qn('w:p')):
                    text = ""
                    for r in p.iter(qn('w:r')):
                        for t in r.iter(qn('w:t')):
                            if t.text: text += t.text
                    full_text.append(text)
        
        text = "\n".join(full_text)
        
        checks = {
            "Dr. Academic Sup": "Academic Supervisor",
            "Engr. Field Sup": "Field Supervisor",
            "National Higher Polytechnic Institute": "Faculty",
            "COMPUTER ENGINEERING": "Department (Upper)",
            "Arith Methrics of Algorithms": "Title"
        }
        
        for val, label in checks.items():
            if val in text:
                print(f"   ✅ Found {label}: '{val}'")
            else:
                print(f"   ❌ Missing {label}: '{val}'")

if __name__ == "__main__":
    test_repair_sample()