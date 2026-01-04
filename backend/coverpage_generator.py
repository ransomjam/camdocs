import os
import json
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
TEMPLATES_DIR = os.path.join(os.path.dirname(BASE_DIR), 'Cover Pages')
OUTPUT_DIR = os.path.join(os.path.dirname(BASE_DIR), 'outputs', 'Cover Pages')

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

def load_json(filename):
    filepath = os.path.join(DATA_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def get_template_path(document_type):
    """
    Map document type to template file.
    """
    mapping = {
        'Assignment': 'Assignments Cover Page Template.docx',
        'Thesis': 'Dissertation Cover Page Template.docx',
        'Dissertation': 'Dissertation Cover Page Template.docx', # Added explicit mapping
        'Research Proposal': 'Dissertation Cover Page Template.docx', # Same as Dissertation
        'Internship Report': 'Internship Cover Page Template.docx', # Renamed from Project Report
        'Project Report': 'Internship Cover Page Template.docx', # Keep for backward compatibility
        'Research Paper': 'Assignments Cover Page Template.docx', # Fallback
        'Lab Report': 'Assignments Cover Page Template.docx', # Fallback
        'Term Paper': 'Assignments Cover Page Template.docx', # Fallback
    }
    
    filename = mapping.get(document_type, 'Assignments Cover Page Template.docx')
    return os.path.join(TEMPLATES_DIR, filename)

def get_all_placeholders(doc):
    """
    Scan document for all {{...}} placeholders, including those with newlines.
    Returns a list of unique placeholder strings found in the doc.
    """
    text_content = []
    
    # Paragraphs
    for p in doc.paragraphs:
        text_content.append(p.text)
        
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content.append(p.text)
                    
    # Text Boxes
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                para_text = ""
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text:
                            para_text += t.text
                if para_text:
                    text_content.append(para_text)
                            
    full_text = "\n".join(text_content)
    # Find all {{...}} patterns, allowing for newlines and whitespace inside
    # The regex matches {{ followed by any char (including newline) until }}
    matches = re.findall(r'\{\{[^}]+\}\}', full_text, re.DOTALL)
    return list(set(matches))

def replace_text_in_paragraph(paragraph, replacements, formatting=None):
    """
    Replace text in a paragraph preserving formatting as much as possible.
    formatting: dict of {key: {'font': 'Name', 'size': Pt(x)}}
    """
    # First, try exact match replacement in runs (best for formatting)
    for key, value in replacements.items():
        if key in paragraph.text:
            replaced = False
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))
                    replaced = True
                    # Apply formatting if needed
                    if formatting and key in formatting:
                        fmt = formatting[key]
                        if 'font' in fmt: run.font.name = fmt['font']
                        if 'size' in fmt: run.font.size = fmt['size']
                        if 'bold' in fmt: run.font.bold = fmt['bold']
            
            # If not replaced in runs, replace in paragraph text
            if not replaced and key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
                # Note: Formatting applied to paragraph.text directly is harder to control for specific runs
                # But we can try to apply to all runs if the paragraph was just the placeholder
                if formatting and key in formatting and len(paragraph.runs) > 0:
                    fmt = formatting[key]
                    for run in paragraph.runs:
                        if 'font' in fmt: run.font.name = fmt['font']
                        if 'size' in fmt: run.font.size = fmt['size']
                        if 'bold' in fmt: run.font.bold = fmt['bold']

def replace_in_textboxes(doc, replacements, formatting=None):
    """
    Replace text in text boxes (shapes) by iterating over XML.
    Handles placeholders split across multiple runs.
    """
    if doc.element.body is None:
        return

    for txbx in doc.element.body.iter(qn('w:txbxContent')):
        for p in txbx.iter(qn('w:p')):
            # 1. Reconstruct full paragraph text
            full_text = ""
            runs = []
            # We need to map runs to their text content to apply formatting later
            # But since we are reconstructing, we only keep the first run active.
            # So we will apply formatting to the first run.
            
            # Collect runs
            xml_runs = []
            for r in p.iter(qn('w:r')):
                xml_runs.append(r)
                for t in r.iter(qn('w:t')):
                    if t.text:
                        full_text += t.text
                        runs.append(t) # Store the text element
            
            # 2. Check if any replacement key is in the full text
            original_text = full_text
            matched_key = None
            replacement_value = None
            
            for key, value in replacements.items():
                if key in full_text:
                    matched_key = key
                    replacement_value = str(value)
                    break
            
            if matched_key:
                print(f"DEBUG: Replaced '{matched_key}' in textbox")
                
                # Check if we need complex split logic (formatting + surrounding text)
                # If formatting is requested, we should split to apply formatting ONLY to the replacement
                # REVERTED: Complex split logic caused issues. Reverting to simple replacement.
                if False and formatting and matched_key in formatting:
                    # SPLIT LOGIC
                    if not runs: continue
                    
                    first_run = runs[0].getparent()
                    p_element = first_run.getparent()
                    
                    # Get base properties (rPr) from the first run
                    base_rPr = first_run.find(qn('w:rPr'))
                    
                    # Remove all existing runs from this paragraph
                    # We iterate over a copy of children to avoid modification issues during iteration
                    for child in list(p_element):
                        if child.tag == qn('w:r'):
                            p_element.remove(child)
                            
                    # Split text by key
                    parts = original_text.split(matched_key)
                    
                    for i, part in enumerate(parts):
                        # Add part (normal)
                        if part:
                            new_r = OxmlElement('w:r')
                            if base_rPr is not None:
                                new_r.append(deepcopy(base_rPr))
                            
                            # Handle newlines in part
                            if '\n' in part:
                                subparts = part.split('\n')
                                for j, subpart in enumerate(subparts):
                                    if j > 0: new_r.append(OxmlElement('w:br'))
                                    t = OxmlElement('w:t')
                                    t.text = subpart
                                    if subpart.strip() == '': t.set(qn('xml:space'), 'preserve')
                                    new_r.append(t)
                            else:
                                t = OxmlElement('w:t')
                                t.text = part
                                if part.startswith(' ') or part.endswith(' '):
                                    t.set(qn('xml:space'), 'preserve')
                                new_r.append(t)
                            
                            p_element.append(new_r)
                        
                        # Add replacement (formatted)
                        if i < len(parts) - 1:
                            new_r = OxmlElement('w:r')
                            # Copy base props first
                            if base_rPr is not None:
                                new_r.append(deepcopy(base_rPr))
                            else:
                                rPr = OxmlElement('w:rPr')
                                new_r.append(rPr)
                            
                            # Apply formatting
                            rPr = new_r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                new_r.insert(0, rPr)
                            
                            fmt = formatting[matched_key]
                            
                            if 'font' in fmt:
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None:
                                    rFonts = OxmlElement('w:rFonts')
                                    rPr.append(rFonts)
                                rFonts.set(qn('w:ascii'), fmt['font'])
                                rFonts.set(qn('w:hAnsi'), fmt['font'])
                                rFonts.set(qn('w:cs'), fmt['font'])
                            
                            if 'size' in fmt:
                                sz = rPr.find(qn('w:sz'))
                                if sz is None:
                                    sz = OxmlElement('w:sz')
                                    rPr.append(sz)
                                szCs = rPr.find(qn('w:szCs'))
                                if szCs is None:
                                    szCs = OxmlElement('w:szCs')
                                    rPr.append(szCs)
                                val = str(int(fmt['size'].pt * 2))
                                sz.set(qn('w:val'), val)
                                szCs.set(qn('w:val'), val)
                                
                            if 'bold' in fmt and fmt['bold']:
                                b = rPr.find(qn('w:b'))
                                if b is None:
                                    b = OxmlElement('w:b')
                                    rPr.append(b)
                                b.set(qn('w:val'), '1')
                            
                            # Handle newlines in replacement value
                            if '\n' in replacement_value:
                                subparts = replacement_value.split('\n')
                                for j, subpart in enumerate(subparts):
                                    if j > 0: new_r.append(OxmlElement('w:br'))
                                    t = OxmlElement('w:t')
                                    t.text = subpart
                                    new_r.append(t)
                            else:
                                t = OxmlElement('w:t')
                                t.text = replacement_value
                                new_r.append(t)
                                
                            p_element.append(new_r)

                else:
                    # OLD LOGIC (Replace in place)
                    full_text = full_text.replace(matched_key, replacement_value)
                    
                    if runs:
                        # Check if we need newline handling
                        if '\n' in full_text:
                            # We need to handle newlines properly by modifying the parent run
                            first_t = runs[0]
                            parent_r = first_t.getparent()
                            
                            # Clear other runs
                            for t in runs[1:]:
                                t.text = ""
                            
                            # Update the first run with new text, handling newlines
                            # Clear existing content of the run (except rPr)
                            for child in parent_r.getchildren():
                                if child.tag != qn('w:rPr'):
                                    parent_r.remove(child)
                            
                            # Append new content
                            lines = full_text.split('\n')
                            for i, line in enumerate(lines):
                                if i > 0:
                                    br = OxmlElement('w:br')
                                    parent_r.append(br)
                                t = OxmlElement('w:t')
                                t.text = line
                                if line.startswith(' ') or line.endswith(' '):
                                    t.set(qn('xml:space'), 'preserve')
                                parent_r.append(t)
                        else:
                            # Simple logic for single line text
                            runs[0].text = full_text
                            for t in runs[1:]:
                                t.text = ""

def replace_placeholders(doc, replacements, formatting=None):
    """
    Replace placeholders in the document (paragraphs, tables, and text boxes).
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements, formatting)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements, formatting)
                    
    # Replace in text boxes
    replace_in_textboxes(doc, replacements, formatting)

def generate_cover_page(data):
    """
    Generate a cover page based on the provided data.
    """
    try:
        document_type = data.get('documentType', 'Assignment')
        template_path = get_template_path(document_type)
        
        if not os.path.exists(template_path):
            logger.error(f"Template not found: {template_path}")
            return None, f"Template for {document_type} not found."

        doc = Document(template_path)
        
        # Force document defaults to Times New Roman to ensure consistency after merge
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        except Exception as e:
            logger.warning(f"Could not set default font: {e}")

        # Standardize date
        date_obj = datetime.now()
        if data.get('date'):
            try:
                date_obj = datetime.strptime(data.get('date'), '%Y-%m-%d')
            except:
                pass
        
        date_str = date_obj.strftime('%B %d, %Y')
        month_year = date_obj.strftime('%B %Y')
        year_str = date_obj.strftime('%Y')
        academic_year = f"{date_obj.year}/{date_obj.year + 1}" if date_obj.month >= 9 else f"{date_obj.year - 1}/{date_obj.year}"
        
        # Helper to get value or empty string
        def get_val(k): return str(data.get(k, '') or '')

        # Infer degree based on faculty
        faculty = get_val('faculty')
        degree = "Bachelor of Science" # Default
        if "Arts" in faculty:
            degree = "Bachelor of Arts"
        elif "Technology" in faculty:
            degree = "Bachelor of Technology"
        elif "Engineering" in faculty or "Polytechnic" in faculty:
            degree = "Bachelor of Engineering"
        elif "Education" in faculty:
            degree = "Bachelor of Education"
        elif "Commerce" in faculty or "Management" in faculty:
            degree = "Bachelor of Science"
            
        # Override for Thesis/Dissertation if needed (usually Masters)
        if document_type in ['Thesis', 'Dissertation']:
            degree = "Master of Science" # Simplified assumption
            if "Arts" in faculty: degree = "Master of Arts"

        # Base values map
        values_map = {
            'studentName': get_val('studentName'),
            'studentId': get_val('studentId'),
            'courseCode': get_val('courseCode'),
            'courseTitle': get_val('courseTitle'),
            'department': get_val('department'),
            'faculty': get_val('faculty'),
            'institution': get_val('institution'),
            'title': get_val('title'),
            'date': date_str,
            'instructor': get_val('instructor'),
            'supervisor': get_val('supervisor') or get_val('academicSupervisor'), # Fallback for Project Reports
            'coSupervisor': get_val('coSupervisor'),
            'academicSupervisor': get_val('academicSupervisor') or get_val('supervisor'), # Fallback
            'fieldSupervisor': get_val('fieldSupervisor'),
            'level': get_val('level'),
            'assignmentNumber': get_val('assignmentNumber'),
            'degree': degree,
            'monthYear': month_year,
            'year': year_str,
            'academicYear': academic_year
        }

        # Simple French Faculty Translation (can be expanded)
        french_faculties = {
            "Faculty of Science": "Faculté des Sciences",
            "Faculty of Arts": "Faculté des Arts",
            "Faculty of Education": "Faculté d'Éducation",
            "Faculty of Health Sciences": "Faculté des Sciences de la Santé",
            "Faculty of Laws and Political Science": "Faculté des Lois et Sciences Politiques",
            "College of Technology": "Collège de Technologie",
            "Higher Technical Teacher Training College": "École Normale Supérieure de l'Enseignement Technique",
            "National Higher Polytechnic Institute": "École Nationale Supérieure Polytechnique",
            "Faculty of Economics and Management Sciences": "Faculté des Sciences Économiques et de Gestion"
        }
        faculty_fr = french_faculties.get(values_map['faculty'], values_map['faculty']) # Default to English if not found

        # Scan document for actual placeholders
        found_placeholders = get_all_placeholders(doc)
        print(f"DEBUG: Found placeholders: {found_placeholders}")
        
        replacements = {}
        formatting_rules = {} # Map placeholder -> {font: 'Name', size: Pt(x)}
        
        for ph in found_placeholders:
            # Clean the placeholder key (remove {{, }}, newlines, whitespace)
            clean_key = re.sub(r'[{}]', '', ph).strip()
            # Normalize key for matching (lowercase, remove extra spaces)
            norm_key = re.sub(r'\s+', ' ', clean_key).lower()
            
            val = ""
            
            # Map normalized keys to values
            if 'student' in norm_key and 'name' in norm_key and 'lecturer' not in norm_key and 'supervisor' not in norm_key:
                val = values_map['studentName']
            elif 'matricule' in norm_key or 'id' in norm_key:
                val = values_map['studentId']
                # Apply Times New Roman 12 for Matricule
                formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12)}
            elif 'course' in norm_key:
                if 'code' in norm_key: val = values_map['courseCode']
                else: val = values_map['courseTitle']
            elif 'department' in norm_key or 'deparment' in norm_key: # Handle typo
                val = values_map['department'].upper() if 'DEPARMENT' in clean_key or 'DEPARTMENT' in clean_key else values_map['department']
            elif 'faculty' in norm_key or 'schoo' in norm_key:
                # Check for French translation request
                if 'french' in norm_key or 'translation' in norm_key:
                    val = faculty_fr
                else:
                    val = values_map['faculty'].upper() if 'FACULTY' in clean_key else values_map['faculty']
            elif 'institution' in norm_key:
                val = values_map['institution'].upper()
            elif 'title' in norm_key or 'topic' in norm_key:
                val = values_map['title']
            elif 'lecturer' in norm_key or 'instructor' in norm_key:
                # Only bold the name, not the title "Instructor:" if it's part of the placeholder (unlikely for {{INSTRUCTOR}})
                # But if the placeholder is {{Instructor: Name}}, we might have issues.
                # Assuming placeholder is just {{INSTRUCTOR}} or {{Lecturer Name}}
                val = " " + values_map['instructor'] if values_map['instructor'] else ""
                formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12), 'bold': True}
            elif 'field' in norm_key and 'supervisor' in norm_key:
                raw_val = values_map['fieldSupervisor'] or values_map.get('coSupervisor', '')
                # Force new line for all supervisors
                val = "\n" + raw_val if raw_val else ""
            elif 'academic' in norm_key and 'supervisor' in norm_key:
                val = "\n" + values_map['academicSupervisor'] if values_map['academicSupervisor'] else ""
            elif 'supervisor' in norm_key: # Generic supervisor (fallback)
                val = "\n" + values_map['supervisor'] if values_map['supervisor'] else ""
            elif 'level' in norm_key:
                val = values_map['level']
            elif 'degree' in norm_key:
                val = values_map['degree']
            elif 'month' in norm_key:
                val = values_map['monthYear']
            elif 'academic' in norm_key and 'year' in norm_key:
                val = values_map['academicYear']
                formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12)}
            elif 'year' in norm_key and 'month' not in norm_key:
                # If it's just "Year", prefer Academic Year for assignments, or just Year?
                # User said "Academic Year for assignment is not on the final cover page".
                # I'll assume they want Academic Year.
                val = values_map['academicYear']
                formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12)}
            elif 'session' in norm_key: # Handle 'Session' for Academic Year
                val = values_map['academicYear']
                formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12)}
            elif 'date' in norm_key:
                val = values_map['date']
            
            # Add to replacements
            # We replace even if empty to remove the placeholder tag
            replacements[ph] = val
            if val:
                print(f"DEBUG: Mapped '{ph}' -> '{val}'")
            else:
                print(f"DEBUG: Mapped '{ph}' -> '' (Empty)")
                
            # Apply formatting to Student Name as well
            if 'student' in norm_key and 'name' in norm_key:
                 formatting_rules[ph] = {'font': 'Times New Roman', 'size': Pt(12)}

        # Static Text Replacement (e.g. "Field Supervisor" -> "Co-Supervisor" for Dissertations)
        if document_type in ['Thesis', 'Dissertation', 'Research Proposal']:
            # We need to replace the static label "Field Supervisor" with "Co-Supervisor"
            # This is tricky because it's not a placeholder. We'll do a global text replacement.
            # We'll add it to replacements but with a special key that matches the text.
            # Note: replace_placeholders iterates paragraphs/runs, so we can add plain text to replacements.
            replacements['Field Supervisor'] = 'Co-Supervisor'
            replacements['FIELD SUPERVISOR'] = 'CO-SUPERVISOR'
            replacements['Field supervisor'] = 'Co-Supervisor'
        
        # Also add standard clean keys just in case
        replacements.update({
            '{{degree_selected}}': degree,
            '{{Deparment}}': values_map['department'],
            '{{School/Faculty}}': values_map['faculty'],
            '{{Month and Year}}': month_year,
            '{{Academic Year}}': academic_year,
            '{{ACADEMIC YEAR}}': academic_year,
            '{{academic year}}': academic_year,
            '{{Year}}': academic_year,
            '{{Session}}': academic_year,
            '{{Date}}': date_str,
        })
        
        replace_placeholders(doc, replacements, formatting_rules)
        
        # Generate output filename

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_title = "".join([c for c in data.get('title', 'cover_page') if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        filename = f"CoverPage_{safe_title}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        
        doc.save(output_path)
        
        return output_path, None
        
    except Exception as e:
        logger.error(f"Error generating cover page: {str(e)}")
        return None, str(e)
