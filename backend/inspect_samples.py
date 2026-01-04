import os
import re
from docx import Document
from docx.oxml.ns import qn

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLES_DIR = os.path.join(os.path.dirname(BASE_DIR), 'Cover Pages', 'poorly formatted samples')

def get_placeholders(doc_path):
    if not os.path.exists(doc_path):
        return f"File not found: {doc_path}"
        
    doc = Document(doc_path)
    text_content = []
    
    # Paragraphs
    for p in doc.paragraphs:
        text_content.append(p.text)
        
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content.append(p.text)
                    
    # Text Boxes
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text:
                            text_content.append(t.text)
                            
    full_text = "\n".join(text_content)
    # Find all {{...}} patterns, allowing for newlines
    matches = re.findall(r'\{\{[^}]+\}\}', full_text, re.DOTALL)
    return list(set(matches))

print("Inspecting Poorly Formatted Samples for Unresolved Placeholders...")
print("=" * 60)

files = [
    'CoverPage_Arith Methrics of Algorithms_20260102_115613.docx',
    'CoverPage_Arith Methrics of Algorithms_20260102_115702.docx',
    'CoverPage_Arith Methrics of Algorithms_20260102_115738.docx'
]

for t in files:
    path = os.path.join(SAMPLES_DIR, t)
    print(f"\nFile: {t}")
    placeholders = get_placeholders(path)
    if isinstance(placeholders, list):
        if placeholders:
            for p in placeholders:
                print(f"  - {repr(p)}") # Use repr to see newlines
        else:
            print("  (No placeholders found)")
    else:
        print(f"  Error: {placeholders}")

print("\n" + "=" * 60)