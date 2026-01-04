# pattern_formatter_backend.py
# Ultra-Precise Pattern-Based Academic Document Formatter
# NO AI - 100% Rule-Based - Lightning Fast

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import re
import os
import json
import uuid
from datetime import datetime
from io import BytesIO
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Serve frontend files directly from the backend for simple deployment
app = Flask(__name__, static_folder='../frontend', static_url_path='')
CORS(app)

@app.route('/')
def serve_frontend():
    return app.send_static_file('index.html')

# Configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def update_toc_with_word(doc_path):
    """
    Update Table of Contents, List of Figures, and List of Tables in a Word document 
    using Microsoft Word COM automation.
    This opens the document in Word, updates all fields (including TOC, LOF, LOT), and saves it.
    Also ensures LOF and LOT entries are plain text (not bold/italic).
    
    Args:
        doc_path: Absolute path to the Word document
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM
        pythoncom.CoInitialize()
        
        # Create Word application instance
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False  # Run in background
        word.DisplayAlerts = False  # Suppress dialogs
        
        try:
            # Open the document
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Update all fields in the document (including TOC, LOF, LOT)
            # wdStory = 6 (entire document)
            word.Selection.WholeStory()
            word.Selection.Fields.Update()
            
            # Also specifically update TOC if present
            for toc in doc.TablesOfContents:
                toc.Update()
            
            # Update List of Figures and List of Tables if present
            # Both use TablesOfFigures collection (they're all caption-based tables)
            for caption_table in doc.TablesOfFigures:
                caption_table.Update()
            
            # Format LOF and LOT entries to be plain text (not bold, not italic)
            # The "Table of Figures" style controls these entries
            try:
                tof_style = doc.Styles("Table of Figures")
                tof_style.Font.Bold = False
                tof_style.Font.Italic = False
            except:
                pass  # Style may not exist
            
            # Save and close
            doc.Save()
            doc.Close()
            
            logger.info(f"TOC, LOF, and LOT updated successfully in {doc_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error updating TOC/LOF/LOT: {str(e)}")
            return False
            
        finally:
            # Quit Word application
            word.Quit()
            pythoncom.CoUninitialize()
            
    except ImportError:
        logger.warning("win32com not available - TOC/LOF/LOT will need manual update")
        logger.warning("Install with: pip install pywin32")
        return False
    except Exception as e:
        logger.error(f"Failed to update TOC/LOF/LOT with Word: {str(e)}")
        return False


# ============================================================
# IMAGE EXTRACTION AND REINSERTION SYSTEM
# ============================================================

class ImageExtractor:
    """
    Extract images from Word documents with full metadata.
    Preserves position, dimensions, format, and caption associations.
    """
    
    # Supported image formats
    SUPPORTED_FORMATS = {
        'image/png': 'png',
        'image/jpeg': 'jpeg',
        'image/jpg': 'jpg',
        'image/gif': 'gif',
        'image/bmp': 'bmp',
        'image/tiff': 'tiff',
        'image/x-emf': 'emf',
        'image/x-wmf': 'wmf',
    }
    
    # Caption patterns (reuse existing patterns)
    CAPTION_PATTERNS = [
        re.compile(r'^Figure\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Fig\.\s*\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Image\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Diagram\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Chart\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Graph\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Illustration\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Photo\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Plate\s+\d+[\.\:\s]', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.images = []
        self.image_count = 0
        self.extracted_rIds = set()  # Track extracted image rIds to prevent duplicates
        
    def extract_all_images(self, doc_path):
        """
        Extract all images from a Word document.
        
        Args:
            doc_path: Path to the .docx file
            
        Returns:
            list: List of image metadata dictionaries
        """
        self.images = []
        self.image_count = 0
        self.extracted_rIds = set()  # Reset for new document
        
        try:
            doc = Document(doc_path)
            
            # Track paragraph index for position mapping
            paragraph_index = 0
            element_index = 0
            
            # Process document body elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # This is a paragraph - check for inline images
                    for para in doc.paragraphs:
                        if para._element is element:
                            images_in_para = self._extract_images_from_paragraph(
                                para, paragraph_index, element_index
                            )
                            self.images.extend(images_in_para)
                            paragraph_index += 1
                            break
                
                elif element.tag.endswith('tbl'):
                    # This is a table - check cells for images
                    table_index = 0
                    for table in doc.tables:
                        if table._element is element:
                            images_in_table = self._extract_images_from_table(
                                table, table_index, element_index
                            )
                            self.images.extend(images_in_table)
                            break
                        table_index += 1
                
                element_index += 1
            
            # Also check for floating images (anchored drawings)
            floating_images = self._extract_floating_images(doc)
            self.images.extend(floating_images)
            
            logger.info(f"Extracted {len(self.images)} images from document")
            return self.images
            
        except Exception as e:
            logger.error(f"Error extracting images: {str(e)}")
            return []
    
    def _extract_images_from_paragraph(self, para, para_index, element_index):
        """Extract inline images from a paragraph."""
        images = []
        
        try:
            # Get paragraph text for caption detection
            para_text = para.text.strip()
            
            # Look for inline shapes (images) in the paragraph
            for run in para.runs:
                # Check if run contains inline shapes
                drawing_elements = run._element.findall('.//' + qn('a:blip'))
                
                for drawing in drawing_elements:
                    # Get the relationship ID (rId) for the image
                    embed_attr = drawing.get(qn('r:embed'))
                    if embed_attr and embed_attr not in self.extracted_rIds:
                        image_data = self._get_image_from_rId(para.part, embed_attr)
                        if image_data:
                            self.extracted_rIds.add(embed_attr)  # Mark as extracted
                            # Get dimensions from inline shape
                            width, height = self._get_inline_dimensions(run._element)
                            
                            # Detect caption
                            caption = self._detect_caption(para, para_text)
                            
                            image_meta = {
                                'image_id': f'img_{self.image_count:04d}',
                                'position_type': 'paragraph',
                                'paragraph_index': para_index,
                                'element_index': element_index,
                                'table_location': None,
                                'data': image_data['bytes'],
                                'format': image_data['format'],
                                'width': width,
                                'height': height,
                                'width_emu': image_data.get('width_emu'),
                                'height_emu': image_data.get('height_emu'),
                                'caption': caption,
                                'caption_position': 'below' if caption else None,
                                'is_inline': True,
                                'anchor_type': 'inline',
                            }
                            images.append(image_meta)
                            self.image_count += 1
                            logger.info(f"Extracted image {image_meta['image_id']} at paragraph {para_index}")
                
                # Also check for drawing elements with pictures (inside the run loop)
                inline_shapes = run._element.findall('.//' + qn('wp:inline'))
                for inline in inline_shapes:
                    blips = inline.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        embed = blip.get(qn('r:embed'))
                        if embed and embed not in self.extracted_rIds:
                            image_data = self._get_image_from_rId(para.part, embed)
                            if image_data:
                                self.extracted_rIds.add(embed)  # Mark as extracted
                                # Get dimensions
                                extent = inline.find(qn('wp:extent'))
                                width = self._emu_to_inches(int(extent.get('cx', 0))) if extent is not None else 3.0
                                height = self._emu_to_inches(int(extent.get('cy', 0))) if extent is not None else 2.0
                                
                                caption = self._detect_caption(para, para_text)
                                
                                image_meta = {
                                    'image_id': f'img_{self.image_count:04d}',
                                    'position_type': 'paragraph',
                                    'paragraph_index': para_index,
                                    'element_index': element_index,
                                    'table_location': None,
                                    'data': image_data['bytes'],
                                    'format': image_data['format'],
                                    'width': width,
                                    'height': height,
                                    'caption': caption,
                                    'caption_position': 'below' if caption else None,
                                    'is_inline': True,
                                    'anchor_type': 'inline',
                                }
                                images.append(image_meta)
                                self.image_count += 1
                            
        except Exception as e:
            logger.warning(f"Error extracting images from paragraph {para_index}: {str(e)}")
        
        return images
    
    def _extract_images_from_table(self, table, table_index, element_index):
        """Extract images from table cells."""
        images = []
        
        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        for run in para.runs:
                            # Look for embedded images
                            blips = run._element.findall('.//' + qn('a:blip'))
                            for blip in blips:
                                embed = blip.get(qn('r:embed'))
                                if embed:
                                    image_data = self._get_image_from_rId(table.part if hasattr(table, 'part') else cell.part, embed)
                                    if image_data:
                                        # Get dimensions
                                        width, height = self._get_inline_dimensions(run._element)
                                        
                                        image_meta = {
                                            'image_id': f'img_{self.image_count:04d}',
                                            'position_type': 'table',
                                            'paragraph_index': None,
                                            'element_index': element_index,
                                            'table_location': {
                                                'table_index': table_index,
                                                'row_index': row_idx,
                                                'cell_index': cell_idx,
                                                'para_index': para_idx,
                                            },
                                            'data': image_data['bytes'],
                                            'format': image_data['format'],
                                            'width': min(width, 2.0),  # Limit size for table cells
                                            'height': min(height, 2.0),
                                            'caption': None,
                                            'is_inline': True,
                                            'anchor_type': 'inline',
                                        }
                                        images.append(image_meta)
                                        self.image_count += 1
                                        logger.info(f"Extracted table image {image_meta['image_id']} at table {table_index}, row {row_idx}, cell {cell_idx}")
        
        except Exception as e:
            logger.warning(f"Error extracting images from table {table_index}: {str(e)}")
        
        return images
    
    def _extract_floating_images(self, doc):
        """Extract floating/anchored images not inline with text."""
        images = []
        
        try:
            # Look for anchored drawings in the document
            for i, para in enumerate(doc.paragraphs):
                for run in para.runs:
                    anchors = run._element.findall('.//' + qn('wp:anchor'))
                    for anchor in anchors:
                        blips = anchor.findall('.//' + qn('a:blip'))
                        for blip in blips:
                            embed = blip.get(qn('r:embed'))
                            # Skip if already extracted as inline image
                            if embed and embed not in self.extracted_rIds:
                                image_data = self._get_image_from_rId(para.part, embed)
                                if image_data:
                                    self.extracted_rIds.add(embed)  # Mark as extracted
                                    # Get dimensions
                                    extent = anchor.find(qn('wp:extent'))
                                    width = self._emu_to_inches(int(extent.get('cx', 0))) if extent is not None else 3.0
                                    height = self._emu_to_inches(int(extent.get('cy', 0))) if extent is not None else 2.0
                                    
                                    image_meta = {
                                        'image_id': f'img_{self.image_count:04d}',
                                        'position_type': 'floating',
                                        'paragraph_index': i,
                                        'element_index': i,
                                        'table_location': None,
                                        'data': image_data['bytes'],
                                        'format': image_data['format'],
                                        'width': width,
                                        'height': height,
                                        'caption': None,
                                        'is_inline': False,
                                        'anchor_type': 'floating',
                                    }
                                    images.append(image_meta)
                                    self.image_count += 1
                                    logger.info(f"Extracted floating image {image_meta['image_id']}")
        
        except Exception as e:
            logger.warning(f"Error extracting floating images: {str(e)}")
        
        return images
    
    def _get_image_from_rId(self, part, rId):
        """Get image binary data from relationship ID."""
        try:
            rel = part.rels.get(rId)
            if rel and rel.target_part:
                # Get the image part
                image_part = rel.target_part
                content_type = image_part.content_type
                
                # Determine format
                img_format = self.SUPPORTED_FORMATS.get(content_type, 'png')
                
                # Get binary data
                image_bytes = image_part.blob
                
                return {
                    'bytes': image_bytes,
                    'format': img_format,
                    'content_type': content_type,
                }
        except Exception as e:
            logger.warning(f"Could not get image from rId {rId}: {str(e)}")
        
        return None
    
    def _get_inline_dimensions(self, run_element):
        """Get dimensions from inline shape element."""
        try:
            # Try to find extent element
            extent = run_element.find('.//' + qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                cy = int(extent.get('cy', 0))
                return self._emu_to_inches(cx), self._emu_to_inches(cy)
            
            # Try to find a:ext element
            ext = run_element.find('.//' + qn('a:ext'))
            if ext is not None:
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                return self._emu_to_inches(cx), self._emu_to_inches(cy)
                
        except Exception as e:
            logger.warning(f"Could not get dimensions: {str(e)}")
        
        # Default dimensions
        return 4.0, 3.0
    
    def _emu_to_inches(self, emu):
        """Convert EMU (English Metric Units) to inches."""
        # 914400 EMU = 1 inch
        if emu <= 0:
            return 3.0  # Default
        return emu / 914400
    
    def _detect_caption(self, para, para_text):
        """Detect if paragraph text is a figure caption."""
        if not para_text:
            return None
        
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(para_text):
                return para_text
        
        return None
    
    def get_images_by_position(self):
        """Get images organized by their position in document."""
        position_map = {}
        
        for img in self.images:
            if img['position_type'] == 'paragraph':
                key = ('paragraph', img['paragraph_index'])
            elif img['position_type'] == 'table':
                loc = img['table_location']
                key = ('table', loc['table_index'], loc['row_index'], loc['cell_index'])
            else:
                key = ('floating', img['element_index'])
            
            if key not in position_map:
                position_map[key] = []
            position_map[key].append(img)
        
        return position_map


class ImageInserter:
    """
    Insert images into Word documents at correct positions.
    """
    
    def __init__(self, doc, images):
        """
        Initialize with document and extracted images.
        
        Args:
            doc: python-docx Document object
            images: List of image metadata dicts from ImageExtractor
        """
        self.doc = doc
        self.images = images
        self.image_lookup = {img['image_id']: img for img in images}
        
    def insert_image(self, image_id, after_paragraph=None):
        """
        Insert an image into the document.
        
        Args:
            image_id: The ID of the image to insert
            after_paragraph: Paragraph object to insert after (optional)
            
        Returns:
            The paragraph containing the image
        """
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found in lookup")
            return None
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Create paragraph for image
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image from bytes
            run = para.add_run()
            
            # Create BytesIO stream from image data
            image_stream = BytesIO(img_data['data'])
            
            # Determine width and height
            width = img_data.get('width', 4.0)
            height = img_data.get('height', 3.0)
            
            # Limit maximum dimensions
            max_width = 6.0  # Max 6 inches wide
            max_height = 8.0  # Max 8 inches tall
            
            if width > max_width:
                ratio = max_width / width
                width = max_width
                height = height * ratio
            
            if height > max_height:
                ratio = max_height / height
                height = max_height
                width = width * ratio
            
            # Add picture
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            # Add caption if exists
            if img_data.get('caption'):
                caption_para = self.doc.add_paragraph()
                caption_run = caption_para.add_run(img_data['caption'])
                caption_run.italic = True
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = Pt(10)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_after = Pt(12)
            
            logger.info(f"Inserted image {image_id} ({width:.2f}x{height:.2f} inches)")
            return para
            
        except Exception as e:
            logger.error(f"Error inserting image {image_id}: {str(e)}")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Could not insert]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return para
    
    def insert_image_in_table_cell(self, image_id, cell):
        """
        Insert an image into a table cell.
        
        Args:
            image_id: The ID of the image to insert
            cell: The table cell to insert into
        """
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found for table insertion")
            return
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Get or create paragraph in cell
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run()
            image_stream = BytesIO(img_data['data'])
            
            # Use smaller dimensions for table cells
            width = min(img_data.get('width', 2.0), 2.0)
            height = min(img_data.get('height', 1.5), 1.5)
            
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            logger.info(f"Inserted image {image_id} in table cell")
            
        except Exception as e:
            logger.error(f"Error inserting table image {image_id}: {str(e)}")
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.add_run(f"[IMAGE: {image_id}]")


import re

# --- Hierarchy Correction Utility ---
hierarchical_pairs = {
    'WEEK': ['DAY', 'SESSION', 'CLASS', 'LECTURE'],
    'MONTH': ['WEEK', 'PHASE'],
    'YEAR': ['QUARTER', 'SEMESTER', 'TERM'],
    'CHAPTER': ['SECTION', 'TOPIC', 'SUBTOPIC'],
    'UNIT': ['LESSON', 'MODULE', 'EXERCISE'],
    'PART': ['CHAPTER', 'SECTION'],
    'MODULE': ['TOPIC', 'SUBTOPIC', 'ACTIVITY'],
    'THEORY': ['PRINCIPLE', 'CONCEPT', 'MODEL'],
    'METHOD': ['STEP', 'PROCEDURE', 'TECHNIQUE'],
    'ANALYSIS': ['RESULT', 'FINDING', 'INTERPRETATION'],
    'FRAMEWORK': ['COMPONENT', 'ELEMENT', 'DIMENSION'],
    'SYSTEM': ['SUBSYSTEM', 'MODULE', 'COMPONENT'],
    'PROCESS': ['STAGE', 'PHASE', 'STEP'],
    'MODEL': ['VARIABLE', 'COMPONENT', 'ELEMENT'],
    'STRATEGY': ['TACTIC', 'APPROACH', 'METHOD'],
    'CATEGORY': ['TYPE', 'CLASS', 'FORM'],
    'PRINCIPLE': ['RULE', 'GUIDELINE', 'STANDARD'],
}

class HierarchyCorrector:
    """Detect and correct hierarchical numbering issues in heading lines."""
    
    HIERARCHICAL_PAIRS = {
        'WEEK': ['DAY', 'SESSION', 'CLASS', 'LECTURE'],
        'MONTH': ['WEEK', 'PHASE'],
        'YEAR': ['QUARTER', 'SEMESTER', 'TERM'],
        'CHAPTER': ['SECTION', 'TOPIC', 'SUBTOPIC'],
        'UNIT': ['LESSON', 'MODULE', 'EXERCISE'],
        'PART': ['CHAPTER', 'SECTION'],
        'MODULE': ['TOPIC', 'SUBTOPIC', 'ACTIVITY'],
        'THEORY': ['PRINCIPLE', 'CONCEPT', 'MODEL'],
        'METHOD': ['STEP', 'PROCEDURE', 'TECHNIQUE'],
        'ANALYSIS': ['RESULT', 'FINDING', 'INTERPRETATION'],
        'FRAMEWORK': ['COMPONENT', 'ELEMENT', 'DIMENSION'],
        'SYSTEM': ['SUBSYSTEM', 'MODULE', 'COMPONENT'],
        'PROCESS': ['STAGE', 'PHASE', 'STEP'],
        'MODEL': ['VARIABLE', 'COMPONENT', 'ELEMENT'],
        'STRATEGY': ['TACTIC', 'APPROACH', 'METHOD'],
        'CATEGORY': ['TYPE', 'CLASS', 'FORM'],
        'PRINCIPLE': ['RULE', 'GUIDELINE', 'STANDARD'],
    }

    def __init__(self):
        self.patterns = [
            # Pattern A: Sequential Major/Minor Topics (Placeholder for specific logic)
            # (re.compile(r'(?i)^\s*(\d+(?:\.\d+)*)\s+([A-Z\s]{2,})(?:\s+|$)\n^\s*(\d+(?:\.\d+)*)\s+([A-Z\s]{2,})(?:\s+|$)', re.MULTILINE),
            #  lambda m: m.group(0)), 

            # Pattern B: Temporal/Categorical Relationships
            (re.compile(r'(?i)(\d+(?:\.\d+)*)\s+(WEEK|MONTH|YEAR|QUARTER|TERM|SEMESTER)\s+([A-Z\d]+)(?:\s+|$)\n(\d+(?:\.\d+)*)\s+(DAY|SESSION|CLASS|PERIOD|LECTURE)\s+([A-Z\d]+)(?:\s+|$)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)} {m.group(3)}\n{m.group(1)}.1 {m.group(5)} {m.group(6)}"),

            # Pattern D: Week/Day Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+WEEK\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+DAY\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} WEEK {m.group(2)}\n{m.group(1)}.1 DAY {m.group(4)}"),

            # Pattern E: Unit/Lesson Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+UNIT\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+LESSON\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} UNIT {m.group(2)}\n{m.group(1)}.1 LESSON {m.group(4)}"),
             
            # Pattern F: Chapter/Section Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+CHAPTER\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+SECTION\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} CHAPTER {m.group(2)}\n{m.group(1)}.1 SECTION {m.group(4)}"),

            # Pattern G: Module/Topic Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+MODULE\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+TOPIC\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} MODULE {m.group(2)}\n{m.group(1)}.1 TOPIC {m.group(4)}"),

            # Pattern H: Lettered Hierarchies
            (re.compile(r'(?i)^(\d+\.\d+)\s+((?:PART\s+)?[A-Z])\b.*?\n^(\d+\.\d+)\s+(\d+[\.\)]?)\s+(.*)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)} {m.group(5)}"),

            # Pattern J: Short Title Followed by Specific Title
            (re.compile(r'^(\d+\.\d+)\s+([A-Z]{2,15})\s*$\n^(\d+\.\d+)\s+([A-Z].{10,})', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)}"),

            # Pattern K: Category/Subcategory Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+(TYPES|CATEGORIES|CLASSIFICATIONS|FORMS|MODELS).*?\n^(\d+\.\d+)\s+((?:.*?MODEL|.*?TYPE|.*?FORM).*)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)}"), # Simplified replacement

            # Pattern Parent/Child with same starting words (Pattern L/M/N combined logic)
            (re.compile(r'^(\d+\.\d+)\s+(.*?\b\w+\b).*?\n^(\d+\.\d+)\s+\2.*?', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(2)}"),
             
             # Pattern R: Convert Flat to Hierarchical (Generic)
            (re.compile(r'(?i)^(\d+\.\d+)\s+(WEEK|UNIT|MODULE|CHAPTER|PART)\s+([A-Z\d]+)(?:\s+|$)\n^(\d+\.\d+)\s+(DAY|LESSON|TOPIC|SECTION|SESSION)\s+([A-Z\d]+)(?:\s+|$)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)} {m.group(3)}\n{m.group(1)}.1 {m.group(5)} {m.group(6)}"),
        ]

    def correct(self, text):
        corrected = text
        for pattern, repl in self.patterns:
            corrected = pattern.sub(repl, corrected)
        return corrected

    def is_hierarchical_pair(self, parent, child):
        parent_upper = parent.upper()
        child_upper = child.upper()
        
        # Check dictionary
        for p, children in self.HIERARCHICAL_PAIRS.items():
            if p in parent_upper:
                for c in children:
                    if c in child_upper:
                        return True
        
        # Check general/specific (short parent, long child)
        if len(parent.split()) < 4 and len(child.split()) > 3:
            common_words = set(parent_upper.split()) & set(child_upper.split())
            if len(common_words) > 0:
                return True
                
        return False

    def correct_lines(self, lines):
        """
        Correct hierarchical numbering issues in a list of heading lines.
        Implements the logic from 'correct_hierarchical_numbering' and 'smart_hierarchy_correction'.
        """
        corrected_lines = []
        i = 0
        while i < len(lines):
            current_line = lines[i]
            
            # Skip if not a numbered heading (simple check)
            m1 = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)$', current_line)
            if not m1:
                corrected_lines.append(current_line)
                i += 1
                continue
                
            if i == len(lines) - 1:
                corrected_lines.append(current_line)
                break
                
            next_line = lines[i+1]
            m2 = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)$', next_line)
            
            if m2:
                current_num, current_title = m1.group(1), m1.group(2)
                next_num, next_title = m2.group(1), m2.group(2)
                
                if self.is_hierarchical_pair(current_title, next_title):
                    # Convert to hierarchical
                    # If parent is 3.6, child becomes 3.6.1
                    child_num = f"{current_num}.1"
                    corrected_lines.append(current_line)
                    corrected_lines.append(f"{child_num} {next_title}")
                    i += 2  # Skip next line
                    continue
            
            corrected_lines.append(current_line)
            i += 1
            
        return corrected_lines

# --- End Hierarchy Correction Utility ---

class HeadingNumberer:
    """
    Auto-number headings based on chapter context with semantic hierarchy detection.
    Tracks chapter numbers and assigns hierarchical numbering to subheadings.
    
    Features:
    - Detects parent-child relationships based on heading text patterns
    - Handles "Main X" / "Specific X" as children of "X" sections
    - Properly nests definition terms under "Operational Definition of Terms"
    - Fixes flat numbering issues (1.3, 1.4, 1.5 → 1.3, 1.3.1, 1.3.2)
    
    Example transformations:
    - "CHAPTER ONE" → stays as is (sets chapter = 1)
    - "Research Objectives" → "1.4 Research Objectives"
    - "Main Research Objective" → "1.4.1 Main Research Objective"
    - "Specific Research Objectives" → "1.4.2 Specific Research Objectives"
    """
    
    # Roman numeral to integer mapping
    ROMAN_TO_INT = {
        'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
        'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10,
        'XI': 11, 'XII': 12, 'XIII': 13, 'XIV': 14, 'XV': 15
    }
    
    # Word to integer mapping
    WORD_TO_INT = {
        'ONE': 1, 'TWO': 2, 'THREE': 3, 'FOUR': 4, 'FIVE': 5,
        'SIX': 6, 'SEVEN': 7, 'EIGHT': 8, 'NINE': 9, 'TEN': 10,
        'ELEVEN': 11, 'TWELVE': 12, 'THIRTEEN': 13, 'FOURTEEN': 14, 'FIFTEEN': 15
    }
    
    # Front matter sections that should NOT be numbered
    UNNUMBERED_SECTIONS = {
        'DECLARATION', 'CERTIFICATION', 'DEDICATION', 'ACKNOWLEDGEMENTS',
        'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT', 'ABSTRACT', 'RESUME', 'RÉSUMÉ',
        'TABLE OF CONTENTS', 'CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES',
        'LIST OF ABBREVIATIONS', 'ABBREVIATIONS', 'GLOSSARY', 'REFERENCES',
        'BIBLIOGRAPHY', 'APPENDIX', 'APPENDICES', 'INDEX', 'PREFACE', 'FOREWORD'
    }
    
    # Section titles that are typically ALL CAPS and follow chapter headings (not numbered)
    CHAPTER_TITLE_SECTIONS = {
        'GENERAL INTRODUCTION', 'INTRODUCTION', 'REVIEW OF RELATED LITERATURE',
        'LITERATURE REVIEW', 'RESEARCH METHODOLOGY', 'METHODOLOGY',
        'DATA ANALYSIS AND INTERPRETATION', 'DATA ANALYSIS', 'FINDINGS AND DISCUSSION',
        'DISCUSSION', 'SUMMARY CONCLUSION AND RECOMMENDATIONS', 'SUMMARY AND CONCLUSION',
        'CONCLUSION', 'RECOMMENDATIONS', 'PRESENTATION OF FINDINGS',
        'RESULTS AND DISCUSSION', 'ANALYSIS AND FINDINGS'
    }
    
    # Parent sections and their expected children (for semantic hierarchy detection)
    # Format: 'parent_keyword': ['child_keyword1', 'child_keyword2', ...]
    PARENT_CHILD_PATTERNS = {
        'research objectives': ['main research objective', 'specific research objective', 'general objective', 'specific objectives'],
        'research questions': ['main research question', 'specific research question', 'general question', 'specific questions'],
        'research hypothesis': ['main research hypothes', 'specific research hypothes', 'main hypothes', 'specific hypothes', 'null hypothesis', 'alternative hypothesis'],
        'research hypotheses': ['main research hypothes', 'specific research hypothes', 'main hypothes', 'specific hypothes', 'null hypothesis', 'alternative hypothesis'],
        'justification of the study': ['policy maker', 'accountability', 'knowledge gap', 'evaluation', 'stakeholder', 'researcher', 'student', 'government'],
        'significance of the study': ['policy maker', 'accountability', 'knowledge gap', 'evaluation', 'stakeholder', 'researcher', 'student', 'government'],
        'operational definition': ['cost sharing', 'tax payer', 'mechanism', 'quality', 'grant', 'private universit', 'government subsid', 'community fund', 'private fund', 'tuition', 'fee', 'scholarship', 'bursary'],
        'definition of terms': ['cost sharing', 'tax payer', 'mechanism', 'quality', 'grant', 'private universit', 'government subsid', 'community fund', 'private fund', 'tuition', 'fee', 'scholarship', 'bursary'],
        'theoretical review': ['equity theory', 'human capital theory', 'revenue theory', 'resource dependency', 'stakeholder theory', 'agency theory', 'institutional theory'],
        'theoretical framework': ['equity theory', 'human capital theory', 'revenue theory', 'resource dependency', 'stakeholder theory', 'agency theory', 'institutional theory'],
        'conceptual review': ['concept of', 'conceptual framework', 'nexus', 'relationship between'],
        'conceptual framework': ['concept of', 'nexus', 'relationship between'],
        'empirical review': ['studies on', 'research on', 'findings', 'previous studies'],
        'delimitation': ['geographical', 'scope', 'population', 'time frame', 'period'],
        'limitation': ['sample size', 'time constraint', 'access', 'generalizability'],
        # New hierarchical patterns
        'week': ['day', 'session', 'class', 'lecture'],
        'month': ['week', 'phase'],
        'year': ['quarter', 'semester', 'term'],
        'chapter': ['section', 'topic', 'subtopic'],
        'unit': ['lesson', 'module', 'exercise'],
        'part': ['chapter', 'section'],
        'module': ['topic', 'subtopic', 'activity'],
        'theory': ['principle', 'concept', 'model'],
        'method': ['step', 'procedure', 'technique'],
        'analysis': ['result', 'finding', 'interpretation'],
        'framework': ['component', 'element', 'dimension'],
        'system': ['subsystem', 'module', 'component'],
        'process': ['stage', 'phase', 'step'],
        'model': ['variable', 'component', 'element'],
        'strategy': ['tactic', 'approach', 'method'],
        'category': ['type', 'class', 'form'],
        'principle': ['rule', 'guideline', 'standard'],
    }
    
    # Patterns that indicate content should NOT be a main section (should be subsection)
    SUBSECTION_INDICATORS = [
        r'^main\s+',  # "Main Research Objective"
        r'^specific\s+',  # "Specific Research Questions"
        r'^general\s+',  # "General Objective"
        r'^primary\s+',  # "Primary Research Question"
        r'^secondary\s+',  # "Secondary Questions"
        r'^null\s+',  # "Null Hypothesis"
        r'^alternative\s+',  # "Alternative Hypothesis"
        r'^\w+\s+can\s+be\s+expressed\s+as',  # "The production function can be expressed as"
        r'^where\s*:?\s*$',  # "Where:" (mathematical notation)
        r'^note\s*:?\s*$',  # "Note:"
        r'^example\s*:?\s*$',  # "Example:"
    ]
    
    # Patterns that indicate this is a definition term (should be under definitions section)
    DEFINITION_TERMS = [
        'cost sharing', 'tax payer', 'taxpayer', 'tuition', 'fee', 'fees',
        'scholarship', 'bursary', 'grant', 'loan', 'subsidy', 'subsidies',
        'funding', 'revenue', 'expenditure', 'budget', 'allocation',
        'quality education', 'quality of education', 'accreditation',
        'enrollment', 'enrolment', 'retention', 'graduation rate',
        'private university', 'public university', 'state university',
        'community funding', 'private funding', 'government funding',
    ]
    
    def __init__(self):
        self.reset()
        self.hierarchy_corrector = HierarchyCorrector()
        
    def reset(self):
        """Reset all counters for a new document."""
        self.current_chapter = 0
        self.current_section = 0  # X.1, X.2, etc.
        self.current_subsection = 0  # X.Y.1, X.Y.2, etc.
        self.current_subsubsection = 0  # X.Y.Z.1, etc.
        self.in_appendix = False
        self.appendix_letter = 'A'
        self.last_level = 0  # Track the last heading level for hierarchy
        self.last_heading_text = ''  # Track last heading for parent-child detection
        self.last_heading_normalized = ''  # Normalized version for matching
        self.in_parent_section = None  # Current parent section (e.g., 'research objectives')
        self.parent_section_number = ''  # Number of current parent section (e.g., '1.4')
        
    def _normalize_text(self, text):
        """Normalize text for comparison (lowercase, remove punctuation, extra spaces)."""
        clean = re.sub(r'^#+\s*', '', text).strip().lower()
        clean = re.sub(r'\*\*', '', clean)  # Remove markdown bold
        clean = re.sub(r'[\d\.]+\s*', '', clean)  # Remove existing numbers
        clean = re.sub(r'[^\w\s]', ' ', clean)  # Remove punctuation
        clean = re.sub(r'\s+', ' ', clean).strip()  # Normalize whitespace
        return clean
    
    def _is_child_of_parent(self, heading_text, parent_text):
        """
        Check if heading_text should be a child of parent_text based on semantic patterns.
        
        Returns:
            bool: True if heading should be a subsection of parent
        """
        heading_norm = self._normalize_text(heading_text)
        parent_norm = self._normalize_text(parent_text)
        
        # Check predefined parent-child patterns
        for parent_key, children in self.PARENT_CHILD_PATTERNS.items():
            if parent_key in parent_norm:
                for child_pattern in children:
                    if child_pattern in heading_norm:
                        return True
        
        return False
    
    def _is_subsection_indicator(self, text):
        """Check if text has patterns indicating it should be a subsection."""
        text_lower = self._normalize_text(text)
        
        for pattern in self.SUBSECTION_INDICATORS:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        
        return False
    
    def _is_definition_term(self, text):
        """Check if text looks like a definition term that should be under Definitions section."""
        text_lower = self._normalize_text(text)
        
        # Check if this matches known definition terms
        for term in self.DEFINITION_TERMS:
            if term in text_lower or text_lower in term:
                return True
        
        # Short titles (1-3 words) after a definitions section are likely definition terms
        if self.in_parent_section and 'definition' in self.in_parent_section:
            word_count = len(text_lower.split())
            if word_count <= 4:
                return True
        
        return False
    
    def _detect_parent_section(self, text):
        """
        Detect if this heading starts a parent section that will have children.
        
        Returns:
            str or None: The parent section key if detected, None otherwise
        """
        text_norm = self._normalize_text(text)
        
        for parent_key in self.PARENT_CHILD_PATTERNS.keys():
            if parent_key in text_norm:
                return parent_key
        
        return None
    
    def _should_be_subsection(self, text):
        """
        Determine if this heading should be a subsection of the previous heading.
        Uses semantic analysis to detect parent-child relationships.
        
        Returns:
            bool: True if this should be a subsection
        """
        text_norm = self._normalize_text(text)
        
        # Check for keywords that indicate a NEW main section (should NOT be subsection)
        # These are full section titles, not partial matches
        main_section_keywords = [
            'chapter summary', 'summary of the chapter', 'conclusion of the chapter',
            'delimitation of the study', 'limitations of the study', 'delimitations',
            'significance of the study', 'scope of the study', 'organization of the study',
            'structure of the study', 'structure of the thesis', 'structure of the dissertation',
            'statement of the problem', 'problem statement',
            'conceptual review', 'conceptual framework',
            'theoretical review', 'theoretical framework',
            'empirical review', 'empirical studies', 'review of empirical',
            'research design', 'research methodology', 'methodology',
            'data presentation', 'data analysis and interpretation',
            'summary conclusion and recommendation', 'summary and conclusion',
            'recommendations for further', 'recommendations',
        ]
        for keyword in main_section_keywords:
            if keyword in text_norm:
                # This is a main section - exit any parent context
                self.in_parent_section = None
                self.parent_section_number = ''
                return False
        
        # Check if it's a known subsection indicator
        if self._is_subsection_indicator(text):
            return True
        
        # Check if current parent section expects this as a child
        if self.in_parent_section:
            if self._is_child_of_parent(text, self.in_parent_section):
                return True
            
            # Definition terms under definitions section
            if 'definition' in self.in_parent_section and self._is_definition_term(text):
                return True
        
        # Check if it's a child of the last heading
        if self.last_heading_normalized:
            if self._is_child_of_parent(text, self.last_heading_text):
                return True
        
        return False
        
    def parse_chapter_number(self, text):
        """
        Extract chapter number from chapter heading text.
        
        Args:
            text: The heading text (e.g., "CHAPTER ONE", "CHAPTER 1", "CHAPTER IV")
            
        Returns:
            int: Chapter number, or 0 if not a chapter heading
        """
        text_upper = text.upper().strip()
        
        # Remove markdown heading markers
        text_upper = re.sub(r'^#+\s*', '', text_upper).strip()
        
        # Pattern: CHAPTER + word (ONE, TWO, etc.)
        match = re.match(r'^CHAPTER\s+([A-Z]+)\b', text_upper)
        if match:
            word = match.group(1)
            if word in self.WORD_TO_INT:
                return self.WORD_TO_INT[word]
        
        # Pattern: CHAPTER + Roman numeral
        match = re.match(r'^CHAPTER\s+([IVXLCDM]+)\b', text_upper)
        if match:
            roman = match.group(1)
            if roman in self.ROMAN_TO_INT:
                return self.ROMAN_TO_INT[roman]
        
        # Pattern: CHAPTER + digit
        match = re.match(r'^CHAPTER\s+(\d+)', text_upper)
        if match:
            return int(match.group(1))
        
        return 0
    
    def is_chapter_heading(self, text):
        """Check if text is a chapter heading."""
        return self.parse_chapter_number(text) > 0
    
    def is_chapter_title(self, text):
        """
        Check if text is a chapter title section (ALL CAPS section after chapter).
        These typically don't get numbered.
        """
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        return clean in self.CHAPTER_TITLE_SECTIONS
    
    def is_unnumbered_section(self, text):
        """Check if text is a front matter or special section that shouldn't be numbered."""
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        return clean in self.UNNUMBERED_SECTIONS
    
    def is_appendix_heading(self, text):
        """Check if text starts an appendix section."""
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        return clean.startswith('APPENDIX') or clean.startswith('APPENDICES')
    
    def already_has_number(self, text):
        """
        Check if heading already has a hierarchical number.
        Matches patterns like: 1.1, 2.1.1, A.1, etc.
        """
        clean = re.sub(r'^#+\s*', '', text).strip()
        # Match: "1.2 Title" or "1.2.3 Title" or "A.1 Title"
        return bool(re.match(r'^(\d+\.)+\d+\s+', clean) or re.match(r'^[A-Z]\.(\d+\.)*\d+\s+', clean))
    
    def extract_existing_number(self, text):
        """
        Extract existing hierarchical number from heading.
        
        Returns:
            tuple: (number_string, title) or (None, text) if no number
        """
        clean = re.sub(r'^#+\s*', '', text).strip()
        
        # Match hierarchical number
        match = re.match(r'^((?:\d+\.)+\d+)\s+(.+)$', clean)
        if match:
            return match.group(1), match.group(2)
        
        # Match appendix number
        match = re.match(r'^([A-Z]\.(?:\d+\.)*\d+)\s+(.+)$', clean)
        if match:
            return match.group(1), match.group(2)
        
        return None, clean
    
    def get_heading_level_from_markdown(self, text):
        """
        Get heading level from markdown markers.
        
        Returns:
            int: 1 for #, 2 for ##, 3 for ###, 0 if no markers
        """
        match = re.match(r'^(#+)\s*', text)
        if match:
            return len(match.group(1))
        return 0
    
    def determine_heading_level(self, text, is_short=True, is_all_caps=False):
        """
        Determine the appropriate heading level based on text characteristics.
        
        Returns:
            int: 1, 2, or 3 indicating the heading level
        """
        clean = re.sub(r'^#+\s*', '', text).strip()
        clean_upper = clean.upper()
        
        # Chapter headings and chapter titles are level 1
        if self.is_chapter_heading(text) or self.is_chapter_title(text):
            return 1
        
        # ALL CAPS short headings are usually level 1 or 2
        if is_all_caps and len(clean.split()) <= 6:
            return 1
        
        # Already numbered - determine level from number
        existing_num, _ = self.extract_existing_number(text)
        if existing_num:
            dots = existing_num.count('.')
            if dots == 1:
                return 2  # X.Y
            elif dots >= 2:
                return 3  # X.Y.Z or deeper
        
        # Title case, moderate length - likely level 2
        if clean[0].isupper() and not is_all_caps:
            return 2
        
        return 2  # Default to level 2
    
    def number_heading(self, text, target_level=None):
        """
        Apply hierarchical numbering to a heading with semantic hierarchy detection.
        
        Args:
            text: The heading text (may include markdown markers)
            target_level: The desired heading level (1, 2, or 3). If None, auto-detect.
            
        Returns:
            dict: {
                'original': original text,
                'numbered': text with number applied (or unchanged if shouldn't be numbered),
                'number': the number string (e.g., "1.2.1"),
                'level': the heading level,
                'chapter': current chapter number,
                'was_renumbered': True if number was added/changed
            }
        """
        result = {
            'original': text,
            'numbered': text,
            'number': None,
            'level': 1,
            'chapter': self.current_chapter,
            'was_renumbered': False
        }
        

        # Preprocess: apply hierarchy correction to heading lines if needed
        # (Assume text is a single heading, but if batch, use correct_lines)
        # This is a placeholder for batch correction integration
        md_level = self.get_heading_level_from_markdown(text)
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Check for chapter heading
        chapter_num = self.parse_chapter_number(text)
        if chapter_num > 0:
            self.current_chapter = chapter_num
            self.current_section = 0
            self.current_subsection = 0
            self.current_subsubsection = 0
            self.in_appendix = False
            self.in_parent_section = None
            self.parent_section_number = ''
            self.last_heading_text = ''
            self.last_heading_normalized = ''
            result['level'] = 1
            result['chapter'] = chapter_num
            return result  # Don't number the chapter heading itself
        
        # Check for appendix
        if self.is_appendix_heading(text):
            self.in_appendix = True
            self.current_section = 0
            self.current_subsection = 0
            self.in_parent_section = None
            result['level'] = 1
            return result  # Don't number appendix heading itself
        
        # Check for unnumbered sections (front matter)
        if self.is_unnumbered_section(text):
            result['level'] = 1
            return result
        
        # Check for chapter title sections (ALL CAPS after chapter)
        if self.is_chapter_title(text):
            result['level'] = 1
            return result  # Keep as is without numbering
        
        # If no chapter context yet, don't number
        if self.current_chapter == 0 and not self.in_appendix:
            return result
        
        # Check if already has proper numbering - extract existing number and title
        existing_num, title = self.extract_existing_number(text)
        
        # Semantic hierarchy detection - check if this should be a subsection
        should_be_child = self._should_be_subsection(text)
        
        # Detect if this heading starts a new parent section
        detected_parent = self._detect_parent_section(text)
        
        # Determine target level using semantic analysis
        if target_level is None:
            if should_be_child:
                # This should be a subsection of the previous section
                target_level = 3  # X.Y.Z
            elif detected_parent:
                # This is a new parent section
                target_level = 2  # X.Y
            elif md_level > 0:
                target_level = min(md_level, 3)
            else:
                target_level = self.determine_heading_level(text)
        
        # Override level if semantic analysis says this should be a child
        if should_be_child and target_level < 3:
            target_level = 3
        
        result['level'] = target_level
        
        # Generate new number based on level
        if target_level == 1 or target_level == 2:
            # Level 1/2 treated as X.Y (main section within chapter)
            # Check if we're leaving a parent section
            if self.in_parent_section and not should_be_child:
                # Moving to a new section - reset parent context if heading is different type
                new_parent = self._detect_parent_section(text)
                if new_parent != self.in_parent_section:
                    self.in_parent_section = None
                    self.parent_section_number = ''
            
            self.current_section += 1
            self.current_subsection = 0
            self.current_subsubsection = 0
            
            if self.in_appendix:
                new_number = f"{self.appendix_letter}.{self.current_section}"
            else:
                new_number = f"{self.current_chapter}.{self.current_section}"
            
            # Update parent section tracking if this starts a parent section
            if detected_parent:
                self.in_parent_section = detected_parent
                self.parent_section_number = new_number
                
        else:  # target_level >= 3
            # Level 3: X.Y.Z (subsection)
            if self.current_section == 0:
                # If no section yet, start with section 1
                self.current_section = 1
            self.current_subsection += 1
            self.current_subsubsection = 0
            
            if self.in_appendix:
                new_number = f"{self.appendix_letter}.{self.current_section}.{self.current_subsection}"
            else:
                new_number = f"{self.current_chapter}.{self.current_section}.{self.current_subsection}"
        
        result['number'] = new_number
        self.last_level = target_level
        
        # Update heading tracking for next iteration
        self.last_heading_text = clean_text
        self.last_heading_normalized = self._normalize_text(clean_text)
        
        # Apply new number if different from existing
        if existing_num != new_number:
            # Reconstruct with markdown markers
            prefix = '#' * md_level + ' ' if md_level > 0 else ''
            result['numbered'] = f"{prefix}{new_number} {title}"
            result['was_renumbered'] = True
        
        return result
    
    def process_document_headings(self, lines):
        """
        Process all headings in a document and apply consistent numbering.
        
        Args:
            lines: List of text lines from the document
            
        Returns:
            list: List of dicts with numbered headings and metadata
        """
        self.reset()
        
        # Pre-process lines to correct hierarchical numbering issues
        # This handles cases where numbering already exists but is incorrect (e.g. 3.6 WEEK TWO, 3.7 DAY ONE)
        if hasattr(self, 'hierarchy_corrector') and self.hierarchy_corrector:
            lines = self.hierarchy_corrector.correct_lines(lines)
            
        results = []
        
        for i, line in enumerate(lines):
            line = line.strip() if isinstance(line, str) else str(line).strip()
            
            # Skip empty lines
            if not line:
                results.append({
                    'original': line,
                    'numbered': line,
                    'is_heading': False,
                    'line_num': i
                })
                continue
            
            # Check if this is a heading (starts with # or matches heading patterns)
            is_heading = False
            
            if line.startswith('#'):
                is_heading = True
            elif re.match(r'^[A-Z][A-Z\s]+$', line) and len(line) <= 60:
                # ALL CAPS line, likely a heading
                is_heading = True
            elif re.match(r'^(\d+\.)+\d*\s+[A-Z]', line):
                # Numbered heading like "1.2 Background"
                is_heading = True
            
            if is_heading:
                numbered_result = self.number_heading(line)
                numbered_result['is_heading'] = True
                numbered_result['line_num'] = i
                results.append(numbered_result)
            else:
                results.append({
                    'original': line,
                    'numbered': line,
                    'is_heading': False,
                    'line_num': i
                })
        
        return results


# ============================================================================
# FIGURE DETECTION AND FORMATTING SYSTEM
# ============================================================================

class FigureFormatter:
    """
    Detect, validate, and format figures in academic documents.
    
    Features:
    - Comprehensive pattern-based figure detection
    - Sequential numbering validation with gap/duplicate detection
    - Proper formatting (Times New Roman, 12pt, italic)
    - Caption placement validation (below figure)
    - List of Figures tracking for LOF generation
    
    Formatting rules:
    - Font: Times New Roman
    - Size: 12pt
    - Style: Italic for captions
    - Placement: Caption BELOW figure
    - Numbering: "Figure X:" or "Figure X.Y:" for sub-figures
    """
    
    # ============================================================
    # ENHANCED DETECTION PATTERNS (December 31, 2025)
    # ============================================================
    
    # Pattern A: Enhanced Primary Detection - All figure variations with decimal numbers
    # Matches: Figure 1.18:, Fig 4.18: Blasting operations, figure 2.1, Fig. 3.5:, FIG 2.0
    FIGURE_REFERENCE_PATTERN = re.compile(
        r'(?:Figure|Fig\.?|Fig\s+)\s*(\d+(?:\.\d+)?)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern B: Specific pattern for "Fig X.YY:" format (targeted for missed formats)
    # Matches: Fig 4.18: Blasting operations dashboard, Fig 2.01: Temperature variations
    FIG_DECIMAL_PATTERN = re.compile(
        r'Fig\s+(\d+\.\d{1,3})\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern C: Pattern with any whitespace variations (including no space)
    # Matches: Fig4.18:, Figure  2.1  :, Fig.  4.18  :, Figure2.1
    WHITESPACE_VARIATION_PATTERN = re.compile(
        r'(?:Figure|Fig\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern C2: No-space figure patterns
    # Matches: Fig4.18:title, Figure2.1:description
    NO_SPACE_PATTERN = re.compile(
        r'(?:Figure|Fig\.?)(\d+(?:\.\d+)*)\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern D: Comprehensive Figure Title Capture - captures entire reference with title
    # Matches: Fig 4.18: Blasting operations dashboard for mining operations
    FIGURE_TITLE_PATTERN = re.compile(
        r'^\s*(?:Figure|Fig(?:\.|\s+)?)\s*(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)(?:\n|$)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern E: Context-Sensitive Detection - avoids false positives
    # Avoids: "configuration", "figurative" - correctly identifies: Fig 4.18: Dashboard
    CONTEXT_SENSITIVE_PATTERN = re.compile(
        r'(?<![A-Za-z])(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)+)\s*[\.:]\s*(.+?)(?=\n{2}|\t|  |$)',
        re.IGNORECASE
    )
    
    # Pattern F: Multi-line figure titles
    MULTILINE_TITLE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*((?:[^\n]+(?:\n(?![A-Z][a-z]+:|\s*(?:Figure|Table|Appendix)\s+|\s*$)[^\n]*)*))',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern G: Figure References in Bulleted/Numbered Lists
    # Matches: • Fig 4.18: Dashboard overview, 1. Figure 2.1: Framework diagram
    LIST_FIGURE_PATTERN = re.compile(
        r'^\s*(?:[-•*\d]\.?\s+)?(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern H: Sub-figures with (a), (b), (c) notation
    # Matches: Fig 4.18: Blasting operations dashboard (a)., Figure 2.1: Conceptual framework (b):
    SUBFIGURE_LETTER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)\s*\([a-z]\)\s*(?:\.|:)?',
        re.IGNORECASE
    )
    
    # Pattern I: Figure with Multiple Sub-figures
    # Matches: Fig 4.18: Blasting operations (a, b, c), Figure 2.1: Analysis results (a-c)
    MULTI_SUBFIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)\s*\(?(?:[a-d](?:\s*,\s*[a-d])*|[a-d]-[a-d])\)?',
        re.IGNORECASE
    )
    
    # Pattern J: Technical Figures - dashboard, chart, graph, diagram, etc.
    TECHNICAL_FIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*([^:\n]*?(?:dashboard|chart|graph|diagram|map|model|framework|flow|process)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern K: Figure Titles with Units/Measurements
    # Matches: Fig 4.18: Temperature variations (0-100°C), Figure 2.1: Speed distribution (0-120 km/h)
    MEASUREMENT_FIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:\d+\s*(?:m|kg|s|Hz|%|°C|km/h)[^:\n]*))',
        re.IGNORECASE
    )
    
    # Pattern L: Sequential Decimal Numbering - for validation
    # Captures: Fig 4.18 → Group1=4, Group2=18
    DECIMAL_NUMBER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+)\.(\d+)\b',
        re.IGNORECASE
    )
    
    # Pattern for sub-figures (Figure 1.1, Figure 2.3)
    SUBFIGURE_PATTERN = re.compile(
        r'Figure\s+(\d+)\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern for figures in appendices (Figure A.1, Figure B.2)
    APPENDIX_FIGURE_PATTERN = re.compile(
        r'Figure\s+([A-Z])\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern M: Inline figure references (see Figure 1, refer to Fig. 2)
    INLINE_REFERENCE_PATTERN = re.compile(
        r'(?:see|refer to|as shown in|shown in|presented in|illustrated in|depicted in)\s+(?:Figure|Fig\.?)\s+(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern N: Figure in parentheses ((Figure 1), (Fig. 2))
    PARENTHESES_PATTERN = re.compile(
        r'\((?:Figure|Fig\.?)\s+(\d+(?:\.\d+)?)\)',
        re.IGNORECASE
    )
    
    # Pattern O: Extract figure numbers for validation
    FIGURE_NUMBER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s*(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # ============================================================
    # FIGURE TYPE DETECTION PATTERNS
    # ============================================================
    
    # Conceptual frameworks
    CONCEPTUAL_FRAMEWORK_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:conceptual\s+framework|theoretical\s+model|research\s+paradigm)',
        re.IGNORECASE
    )
    
    # Statistical charts
    STATISTICAL_CHART_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:bar\s+chart|pie\s+chart|line\s+graph|histogram|scatter\s+plot|trend\s+analysis)',
        re.IGNORECASE
    )
    
    # Process flows
    PROCESS_FLOW_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:flow\s*chart|process\s+diagram|workflow|sequence\s+diagram)',
        re.IGNORECASE
    )
    
    # Technical diagrams (dashboard, operations, monitoring)
    TECHNICAL_DIAGRAM_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:dashboard|operations|monitoring|interface|system|architecture)',
        re.IGNORECASE
    )
    
    # ============================================================
    # ENHANCED CAPTION DETECTION PATTERNS
    # ============================================================
    
    # Caption patterns for detection - comprehensive list
    CAPTION_PATTERNS = [
        # Standard Figure X: format
        re.compile(r'^Figure\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig. X: format
        re.compile(r'^Fig\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig X: format (no period after Fig)
        re.compile(r'^Fig\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # FIG X: format (all caps)
        re.compile(r'^FIG\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Diagram X: format
        re.compile(r'^Diagram\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Chart X: format
        re.compile(r'^Chart\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Graph X: format
        re.compile(r'^Graph\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Illustration X: format
        re.compile(r'^Illustration\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig X.YY: format (decimal without period after Fig)
        re.compile(r'^Fig\s+\d+\.\d+\s*:', re.IGNORECASE),
        # Bulleted/numbered list figures
        re.compile(r'^\s*[-•*\d]\.?\s+(?:Figure|Fig)\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # No-space format: Fig4.18:, Figure2.1:
        re.compile(r'^(?:Figure|Fig\.?)(\d+(?:\.\d+)*)\s*:', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.reset()
        
    def reset(self):
        """Reset all tracking for a new document."""
        self.figures = []  # All detected figures with metadata
        self.figure_numbers = []  # Just the numbers for validation
        self.figure_entries = []  # Entries for List of Figures
        self.current_chapter = 0  # For chapter-based numbering
        self.numbering_issues = []  # Track gaps, duplicates, etc.
        
    def detect_figures(self, text):
        """
        Detect all figure references in a block of text using multi-pass detection.
        Uses multiple patterns to catch all variations.
        
        Args:
            text: Document text content
            
        Returns:
            list: List of detected figures with metadata
        """
        figures = []
        seen_positions = set()  # Avoid duplicates
        
        # Multi-pass detection with different patterns
        detection_patterns = [
            # Pass 1: Standard figure titles/captions
            ('FIGURE_TITLE_PATTERN', self.FIGURE_TITLE_PATTERN),
            # Pass 2: Fig X.YY: format specifically
            ('FIG_DECIMAL_PATTERN', self.FIG_DECIMAL_PATTERN),
            # Pass 3: No-space format (Fig4.18:)
            ('NO_SPACE_PATTERN', self.NO_SPACE_PATTERN),
            # Pass 4: List figures
            ('LIST_FIGURE_PATTERN', self.LIST_FIGURE_PATTERN),
            # Pass 5: Technical figures
            ('TECHNICAL_FIGURE_PATTERN', self.TECHNICAL_FIGURE_PATTERN),
            # Pass 6: Whitespace variations
            ('WHITESPACE_VARIATION_PATTERN', self.WHITESPACE_VARIATION_PATTERN),
        ]
        
        for pattern_name, pattern in detection_patterns:
            for match in pattern.finditer(text):
                # Skip if we've already detected this position
                pos_key = (match.start(), match.end())
                if pos_key in seen_positions:
                    continue
                seen_positions.add(pos_key)
                
                figure_num = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ''
                
                figures.append({
                    'number': figure_num,
                    'title': title,
                    'full_match': match.group(0),
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'caption',
                    'figure_type': self._classify_figure_type(title),
                    'pattern_used': pattern_name
                })
        
        # Track all figure numbers for validation using comprehensive pattern
        for match in self.FIGURE_NUMBER_PATTERN.finditer(text):
            num = match.group(1)
            if num not in self.figure_numbers:
                self.figure_numbers.append(num)
        
        return figures
    
    def detect_figure_caption(self, paragraph_text):
        """
        Detect if a paragraph is a figure caption using enhanced multi-pattern detection.
        
        Args:
            paragraph_text: Text of the paragraph
            
        Returns:
            dict or None: Figure metadata if caption detected, None otherwise
        """
        if not paragraph_text:
            return None
        
        text = paragraph_text.strip()
        
        # First check with standard caption patterns
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(text):
                # Try multiple extraction patterns for the number and title
                extraction_patterns = [
                    self.FIGURE_TITLE_PATTERN,
                    self.FIG_DECIMAL_PATTERN,
                    self.NO_SPACE_PATTERN,
                    self.TECHNICAL_FIGURE_PATTERN,
                ]
                
                for extract_pattern in extraction_patterns:
                    match = extract_pattern.match(text)
                    if match:
                        return {
                            'number': match.group(1),
                            'title': match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else '',
                            'full_text': text,
                            'figure_type': self._classify_figure_type(text)
                        }
                
                # Fallback: extract using general number pattern
                num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                return {
                    'number': num_match.group(1) if num_match else '?',
                    'title': text,
                    'full_text': text,
                    'figure_type': 'unknown'
                }
        
        return None
    
    def _classify_figure_type(self, title):
        """Classify the type of figure based on title content."""
        title_lower = title.lower()
        
        if any(term in title_lower for term in ['framework', 'model', 'paradigm']):
            return 'conceptual_framework'
        elif any(term in title_lower for term in ['chart', 'graph', 'histogram', 'plot']):
            return 'statistical_chart'
        elif any(term in title_lower for term in ['flow', 'diagram', 'workflow', 'process']):
            return 'process_flow'
        elif any(term in title_lower for term in ['map', 'location', 'geographic']):
            return 'map'
        elif any(term in title_lower for term in ['photo', 'photograph', 'image', 'picture']):
            return 'photograph'
        elif any(term in title_lower for term in ['dashboard', 'interface', 'screen', 'ui']):
            return 'technical_interface'
        elif any(term in title_lower for term in ['blast', 'operation', 'equipment', 'machinery']):
            return 'technical_operation'
        elif any(term in title_lower for term in ['structure', 'architecture', 'layout']):
            return 'structure'
        else:
            return 'general'
    
    def validate_numbering(self):
        """
        Validate sequential figure numbering.
        Detects gaps, duplicates, and out-of-order numbers.
        
        Returns:
            dict: Validation results with issues found
        """
        self.numbering_issues = []
        
        if not self.figure_numbers:
            return {'valid': True, 'issues': []}
        
        # Parse numbers (handle both "1" and "1.1" formats)
        parsed_numbers = []
        for num in self.figure_numbers:
            if '.' in num:
                parts = num.split('.')
                parsed_numbers.append((int(parts[0]), int(parts[1]), num))
            else:
                parsed_numbers.append((int(num), 0, num))
        
        # Sort by chapter then sub-number
        parsed_numbers.sort(key=lambda x: (x[0], x[1]))
        
        # Check for gaps in main figure numbers
        main_numbers = sorted(set(n[0] for n in parsed_numbers))
        expected = list(range(1, max(main_numbers) + 1)) if main_numbers else []
        missing = [n for n in expected if n not in main_numbers]
        
        if missing:
            self.numbering_issues.append({
                'type': 'missing_numbers',
                'numbers': missing,
                'message': f"Missing figure number(s): {', '.join(map(str, missing))}"
            })
        
        # Check for duplicates
        seen = set()
        duplicates = []
        for num in self.figure_numbers:
            if num in seen:
                duplicates.append(num)
            seen.add(num)
        
        if duplicates:
            self.numbering_issues.append({
                'type': 'duplicates',
                'numbers': duplicates,
                'message': f"Duplicate figure number(s): {', '.join(duplicates)}"
            })
        
        return {
            'valid': len(self.numbering_issues) == 0,
            'issues': self.numbering_issues,
            'total_figures': len(self.figure_numbers),
            'figure_numbers': self.figure_numbers
        }
    
    def add_figure_entry(self, number, title, page_number=None):
        """
        Add a figure entry for List of Figures generation.
        
        Args:
            number: Figure number (e.g., "1" or "1.2")
            title: Figure title/caption text
            page_number: Optional page number (will be updated by Word)
        """
        self.figure_entries.append({
            'number': number,
            'title': title,
            'page': page_number,
            'full_entry': f"Figure {number}: {title}"
        })
    
    def get_lof_entries(self):
        """
        Get all figure entries for List of Figures.
        
        Returns:
            list: Sorted list of figure entries
        """
        # Sort by figure number
        def sort_key(entry):
            num = entry['number']
            if '.' in num:
                parts = num.split('.')
                return (int(parts[0]), int(parts[1]))
            return (int(num), 0)
        
        return sorted(self.figure_entries, key=sort_key)
    
    def format_caption_text(self, number, title):
        """
        Generate properly formatted caption text.
        
        Args:
            number: Figure number
            title: Figure title
            
        Returns:
            str: Formatted caption text "Figure X: Title"
        """
        # Ensure proper capitalization
        formatted_title = title.strip()
        if formatted_title and not formatted_title[0].isupper():
            formatted_title = formatted_title[0].upper() + formatted_title[1:]
        
        return f"Figure {number}: {formatted_title}"
    
    def is_figure_caption(self, text):
        """
        Quick check if text is a figure caption.
        
        Args:
            text: Paragraph text to check
            
        Returns:
            bool: True if text is a figure caption
        """
        if not text:
            return False
        
        text = text.strip()
        return any(pattern.match(text) for pattern in self.CAPTION_PATTERNS)
    
    def extract_inline_references(self, text):
        """
        Extract all inline figure references from text.
        Used to validate all referenced figures exist.
        
        Args:
            text: Document text
            
        Returns:
            list: List of referenced figure numbers
        """
        references = []
        
        # Pattern B: Inline references
        for match in self.INLINE_REFERENCE_PATTERN.finditer(text):
            references.append(match.group(1))
        
        # Pattern C: Parenthetical references
        for match in self.PARENTHESES_PATTERN.finditer(text):
            references.append(match.group(1))
        
        return list(set(references))  # Remove duplicates
    
    def renumber_figures(self, start_from=1):
        """
        Renumber all figures sequentially starting from a given number.
        Used to fix numbering gaps.
        
        Args:
            start_from: Starting figure number (default 1)
            
        Returns:
            dict: Mapping of old numbers to new numbers
        """
        if not self.figure_entries:
            return {}
        
        # Sort entries by current number
        sorted_entries = self.get_lof_entries()
        
        renumber_map = {}
        current_num = start_from
        
        for entry in sorted_entries:
            old_num = entry['number']
            if '.' not in old_num:  # Only renumber main figures
                renumber_map[old_num] = str(current_num)
                current_num += 1
        
        return renumber_map


# ============================================================
# TABLE FORMATTER CLASS - December 31, 2025
# ============================================================

class TableFormatter:
    """
    Detect, validate, and format tables in academic documents.
    
    Features:
    - Comprehensive pattern-based table detection
    - Sequential numbering validation with gap/duplicate detection
    - Proper formatting (Times New Roman, 12pt, bold)
    - Caption placement validation (below table)
    - List of Tables tracking for LOT generation
    
    Formatting rules:
    - Font: Times New Roman
    - Size: 12pt
    - Style: Bold for captions
    - Placement: Caption BELOW table
    - Numbering: "Table X:" or "Table X.Y:" for sub-tables
    """
    
    # ============================================================
    # ENHANCED DETECTION PATTERNS
    # ============================================================
    
    # Pattern A: Standard Table References - Primary detection
    # Matches: Table 1:, Tbl. 2, Tab. 3.1:, Table 4 -, table 6
    TABLE_REFERENCE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)\s*[\.:\-–—]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern B: Comprehensive Table Detection with Titles
    # Matches: Table 1: Summary of Research Variables, Tbl. 2.1: Demographics
    TABLE_TITLE_PATTERN = re.compile(
        r'^\s*(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)(?:\n|$)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern C: Inline Table References
    # Matches: see Table 1, refer to Tbl. 2, as shown in Table 3.1
    INLINE_REFERENCE_PATTERN = re.compile(
        r'(?:see|refer to|as shown in|presented in|shown in|displayed in)\s+(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern D: Table References in Parentheses
    # Matches: (Table 1), (Tbl. 2.1), (table 3)
    PARENTHESES_PATTERN = re.compile(
        r'\((?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)\)',
        re.IGNORECASE
    )
    
    # Pattern E: Multi-line Table Titles
    MULTILINE_TITLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*((?:[^\n]+(?:\n(?![A-Z][a-z]+:|\s*(?:Table|Figure|Appendix)\s+|\s*$)[^\n]*)*))',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern F: Tables in Lists/Numbering Systems
    # Matches: • Table 1: Variables, 1. Table 2.1: Data
    LIST_TABLE_PATTERN = re.compile(
        r'^\s*(?:[-•*\d]\.?\s+)?(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern G: Tables with Statistical Notation
    STATISTICAL_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:mean|SD|SE|p-value|F-value|t-test|ANOVA|regression|correlation)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern H: No-space variations
    # Matches: Table1.2:, Tbl.3:, Tab4.5:
    NO_SPACE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)(\d+(?:\.\d+)*)\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern I: Whitespace variations
    # Matches: Table  2.1  :, Tbl.  4.18  :
    WHITESPACE_VARIATION_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern J: Summary/Descriptive Tables
    SUMMARY_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:summary|descriptive|demographic|characteristics|overview)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern K: Comparison Tables
    COMPARISON_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:comparison|contrast|vs\.|versus|between|among)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern L: Correlation/Regression Tables
    CORRELATION_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:correlation|regression|coefficient|model|results)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern M: Context-Sensitive Table Detection
    # Avoids false positives like "tablespoon", "tablet"
    CONTEXT_SENSITIVE_PATTERN = re.compile(
        r'(?<![A-Za-z])(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)+)\s*[\.:]\s*(.+?)(?=\n{2}|\t|  |$)',
        re.IGNORECASE
    )
    
    # Pattern N: Extract Table Numbers for Validation
    TABLE_NUMBER_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern O: Tables in Appendices
    APPENDIX_TABLE_PATTERN = re.compile(
        r'Table\s+([A-Z])\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # ============================================================
    # TABLE TYPE DETECTION PATTERNS
    # ============================================================
    
    # Summary tables
    SUMMARY_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:summary|descriptive|demographic|characteristics)',
        re.IGNORECASE
    )
    
    # Statistical tables
    STATISTICAL_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:mean|SD|ANOVA|regression|correlation|t-test|p-value)',
        re.IGNORECASE
    )
    
    # Comparison tables
    COMPARISON_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:comparison|contrast|vs\.|versus|difference)',
        re.IGNORECASE
    )
    
    # ============================================================
    # ENHANCED CAPTION DETECTION PATTERNS
    # ============================================================
    
    CAPTION_PATTERNS = [
        # Standard Table X: format
        re.compile(r'^Table\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tbl. X: format
        re.compile(r'^Tbl\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tbl X: format (no period after Tbl)
        re.compile(r'^Tbl\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tab. X: format
        re.compile(r'^Tab\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tab X: format (no period after Tab)
        re.compile(r'^Tab\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # TBL X: format (all caps)
        re.compile(r'^TBL\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # TABLE X: format (all caps)
        re.compile(r'^TABLE\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Bulleted/numbered list tables
        re.compile(r'^\s*[-•*\d]\.?\s+(?:Table|Tbl)\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # No-space format: Table1.2:, Tbl.3:
        re.compile(r'^(?:Table|Tbl\.?|Tab\.?)(\d+(?:\.\d+)*)\s*:', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.reset()
        
    def reset(self):
        """Reset all tracking for a new document."""
        self.tables = []  # All detected tables with metadata
        self.table_numbers = []  # Just the numbers for validation
        self.table_entries = []  # Entries for List of Tables
        self.current_chapter = 0  # For chapter-based numbering
        self.numbering_issues = []  # Track gaps, duplicates, etc.
        
    def detect_tables(self, text):
        """
        Detect all table references in a block of text using multi-pass detection.
        Uses multiple patterns to catch all variations.
        
        Args:
            text: Document text content
            
        Returns:
            list: List of detected tables with metadata
        """
        tables = []
        seen_positions = set()  # Avoid duplicates
        
        # Multi-pass detection with different patterns
        detection_patterns = [
            # Pass 1: Standard table titles/captions
            ('TABLE_TITLE_PATTERN', self.TABLE_TITLE_PATTERN),
            # Pass 2: Statistical tables
            ('STATISTICAL_TABLE_PATTERN', self.STATISTICAL_TABLE_PATTERN),
            # Pass 3: No-space format
            ('NO_SPACE_PATTERN', self.NO_SPACE_PATTERN),
            # Pass 4: List tables
            ('LIST_TABLE_PATTERN', self.LIST_TABLE_PATTERN),
            # Pass 5: Summary tables
            ('SUMMARY_TABLE_PATTERN', self.SUMMARY_TABLE_PATTERN),
            # Pass 6: Comparison tables
            ('COMPARISON_TABLE_PATTERN', self.COMPARISON_TABLE_PATTERN),
            # Pass 7: Whitespace variations
            ('WHITESPACE_VARIATION_PATTERN', self.WHITESPACE_VARIATION_PATTERN),
        ]
        
        for pattern_name, pattern in detection_patterns:
            for match in pattern.finditer(text):
                # Skip if we've already detected this position
                pos_key = (match.start(), match.end())
                if pos_key in seen_positions:
                    continue
                seen_positions.add(pos_key)
                
                table_num = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ''
                
                tables.append({
                    'number': table_num,
                    'title': title,
                    'full_match': match.group(0),
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'caption',
                    'table_type': self._classify_table_type(title),
                    'pattern_used': pattern_name
                })
        
        # Track all table numbers for validation
        for match in self.TABLE_NUMBER_PATTERN.finditer(text):
            num = match.group(1)
            if num not in self.table_numbers:
                self.table_numbers.append(num)
        
        return tables
    
    def detect_table_caption(self, paragraph_text):
        """
        Detect if a paragraph is a table caption using enhanced multi-pattern detection.
        
        Args:
            paragraph_text: Text of the paragraph
            
        Returns:
            dict or None: Table metadata if caption detected, None otherwise
        """
        if not paragraph_text:
            return None
        
        text = paragraph_text.strip()
        
        # First check with standard caption patterns
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(text):
                # Try multiple extraction patterns for the number and title
                extraction_patterns = [
                    self.TABLE_TITLE_PATTERN,
                    self.STATISTICAL_TABLE_PATTERN,
                    self.NO_SPACE_PATTERN,
                    self.SUMMARY_TABLE_PATTERN,
                ]
                
                for extract_pattern in extraction_patterns:
                    match = extract_pattern.match(text)
                    if match:
                        return {
                            'number': match.group(1),
                            'title': match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else '',
                            'full_text': text,
                            'table_type': self._classify_table_type(text)
                        }
                
                # Fallback: extract using general number pattern
                num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                return {
                    'number': num_match.group(1) if num_match else '?',
                    'title': text,
                    'full_text': text,
                    'table_type': 'unknown'
                }
        
        return None
    
    def _classify_table_type(self, title):
        """Classify the type of table based on title content."""
        title_lower = title.lower()
        
        if any(term in title_lower for term in ['summary', 'descriptive', 'overview']):
            return 'summary'
        elif any(term in title_lower for term in ['demographic', 'characteristics', 'profile']):
            return 'demographic'
        elif any(term in title_lower for term in ['comparison', 'contrast', 'vs.', 'versus']):
            return 'comparison'
        elif any(term in title_lower for term in ['correlation', 'regression', 'coefficient']):
            return 'correlation'
        elif any(term in title_lower for term in ['mean', 'sd', 'anova', 't-test', 'p-value']):
            return 'statistical'
        elif any(term in title_lower for term in ['results', 'findings', 'analysis']):
            return 'results'
        elif any(term in title_lower for term in ['frequency', 'distribution', 'percentage']):
            return 'frequency'
        elif any(term in title_lower for term in ['cost', 'budget', 'financial', 'expenditure']):
            return 'financial'
        else:
            return 'general'
    
    def validate_numbering(self):
        """
        Validate sequential table numbering.
        Detects gaps, duplicates, and out-of-order numbers.
        
        Returns:
            dict: Validation results with issues found
        """
        self.numbering_issues = []
        
        if not self.table_numbers:
            return {'valid': True, 'issues': []}
        
        # Parse numbers (handle both "1" and "1.1" formats)
        parsed_numbers = []
        for num in self.table_numbers:
            if '.' in num:
                parts = num.split('.')
                parsed_numbers.append((int(parts[0]), int(parts[1]), num))
            else:
                parsed_numbers.append((int(num), 0, num))
        
        # Sort by chapter then sub-number
        parsed_numbers.sort(key=lambda x: (x[0], x[1]))
        
        # Check for gaps in main table numbers
        main_numbers = sorted(set(n[0] for n in parsed_numbers))
        expected = list(range(1, max(main_numbers) + 1)) if main_numbers else []
        missing = [n for n in expected if n not in main_numbers]
        
        if missing:
            self.numbering_issues.append({
                'type': 'gap',
                'message': f'Missing table numbers: {missing}',
                'missing': missing
            })
        
        # Check for duplicates
        seen = set()
        duplicates = []
        for num in self.table_numbers:
            if num in seen:
                duplicates.append(num)
            seen.add(num)
        
        if duplicates:
            self.numbering_issues.append({
                'type': 'duplicate',
                'message': f'Duplicate table numbers: {duplicates}',
                'duplicates': duplicates
            })
        
        return {
            'valid': len(self.numbering_issues) == 0,
            'issues': self.numbering_issues,
            'total_tables': len(self.table_numbers),
            'table_numbers': self.table_numbers
        }
    
    def add_table_entry(self, number, title, page_number=None):
        """
        Add a table entry for List of Tables generation.
        
        Args:
            number: Table number (e.g., "1" or "1.2")
            title: Table title/caption text
            page_number: Optional page number (will be updated by Word)
        """
        self.table_entries.append({
            'number': number,
            'title': title,
            'page': page_number,
            'full_entry': f"Table {number}: {title}"
        })
    
    def get_lot_entries(self):
        """
        Get all table entries for List of Tables.
        
        Returns:
            list: Sorted list of table entries
        """
        # Sort by table number
        def sort_key(entry):
            num = entry['number']
            if '.' in num:
                parts = num.split('.')
                return (int(parts[0]), int(parts[1]))
            return (int(num), 0)
        
        return sorted(self.table_entries, key=sort_key)
    
    def format_caption_text(self, number, title):
        """
        Generate properly formatted caption text.
        
        Args:
            number: Table number
            title: Table title
            
        Returns:
            str: Formatted caption text "Table X: Title"
        """
        # Ensure proper capitalization
        formatted_title = title.strip()
        if formatted_title and not formatted_title[0].isupper():
            formatted_title = formatted_title[0].upper() + formatted_title[1:]
        
        return f"Table {number}: {formatted_title}"
    
    def is_table_caption(self, text):
        """
        Quick check if text is a table caption.
        
        Args:
            text: Paragraph text to check
            
        Returns:
            bool: True if text is a table caption
        """
        if not text:
            return False
        
        text = text.strip()
        return any(pattern.match(text) for pattern in self.CAPTION_PATTERNS)
    
    def extract_inline_references(self, text):
        """
        Extract all inline table references from text.
        Used to validate all referenced tables exist.
        
        Args:
            text: Document text
            
        Returns:
            list: List of referenced table numbers
        """
        references = []
        
        # Inline references
        for match in self.INLINE_REFERENCE_PATTERN.finditer(text):
            references.append(match.group(1))
        
        # Parenthetical references
        for match in self.PARENTHESES_PATTERN.finditer(text):
            references.append(match.group(1))
        
        return list(set(references))  # Remove duplicates
    
    def renumber_tables(self, start_from=1):
        """
        Renumber all tables sequentially starting from a given number.
        Used to fix numbering gaps.
        
        Args:
            start_from: Starting table number (default 1)
            
        Returns:
            dict: Mapping of old numbers to new numbers
        """
        if not self.table_entries:
            return {}
        
        # Sort entries by current number
        sorted_entries = self.get_lot_entries()
        
        renumber_map = {}
        current_num = start_from
        
        for entry in sorted_entries:
            old_num = entry['number']
            if '.' not in old_num:  # Only renumber main tables
                renumber_map[old_num] = str(current_num)
                current_num += 1
        
        return renumber_map


# =================================================================================
# BULLET IMPLEMENTATION HELPER FUNCTIONS
# =================================================================================

def detect_bullet_type(line_text):
    """
    Detect bullet type and extract content with character mapping to Word equivalents
    """
    bullet_patterns = {
        # Standard bullets (most common)
        r'^\s*[•]\s+(.+)$': 'standard',
        r'^\s*[○]\s+(.+)$': 'white_circle',
        r'^\s*[●]\s+(.+)$': 'black_circle',
        r'^\s*[▪]\s+(.+)$': 'small_square',
        r'^\s*[■]\s+(.+)$': 'square',
        
        # Dash/arrow bullets
        r'^\s*[-–—]\s+(.+)$': 'dash',
        r'^\s*[→➔➜➤➢]\s+(.+)$': 'arrow',
        
        # Asterisk variants
        r'^\s*\*\s+(.+)$': 'asterisk',
        r'^\s*[⁎⁑※]\s+(.+)$': 'asterisk_variant',
        
        # Checkbox bullets
        r'^\s*[☐]\s+(.+)$': 'checkbox_empty',
        r'^\s*[☑]\s+(.+)$': 'checkbox_checked',
        r'^\s*[✓✔]\s+(.+)$': 'checkmark',
        
        # Number-like bullets
        r'^\s*[⓿①②③④⑤⑥⑦⑧⑨]\s+(.+)$': 'circled_number',
        r'^\s*[❶❷❸❹❺❻❼❽❾❿]\s+(.+)$': 'dingbat_number',
        
        # Creative/AI-generated bullets
        r'^\s*[✦✧]\s+(.+)$': 'sparkle',
        r'^\s*[★☆]\s+(.+)$': 'star',
        r'^\s*[♥♡]\s+(.+)$': 'heart',
        r'^\s*[◆◇]\s+(.+)$': 'diamond',
        r'^\s*[◉◎]\s+(.+)$': 'bullseye',
        r'^\s*[▸▹►]\s+(.+)$': 'right_triangle',
        r'^\s*[◂◃◄]\s+(.+)$': 'left_triangle',
        
        # Comprehensive catch-all (should be last)
        r'^\s*([•○●▪■□◆◇→➔➜➤➢–—*⁎⁑※✱✲✳✴☐☑✓✔✗✘⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿➀-➉✦✧★☆♥♡◉◎▸▹►◂◃◄⦿⁍-])\s+(.+)$': 'generic_bullet'
    }
    
    for pattern, bullet_type in bullet_patterns.items():
        match = re.match(pattern, line_text, re.UNICODE)
        if match:
            content = match.group(1) if len(match.groups()) == 1 else match.group(2)
            bullet_char = match.group(1) if len(match.groups()) > 1 else None
            
            return {
                'type': bullet_type,
                'bullet_char': bullet_char,
                'content': content.strip(),
                'original_line': line_text.strip(),
                'indentation': len(line_text) - len(line_text.lstrip())
            }
    
    return None

def map_to_word_bullet_style(bullet_info):
    """
    Map detected bullet to Microsoft Word bullet style
    Returns tuple: (bullet_char, font_name, bullet_type_code)
    """
    bullet_mapping = {
        # Standard → Change to Square ■
        'standard': ('■', 'Arial', 'square'),
        'asterisk': ('■', 'Arial', 'square'),
        
        # Circle bullets → Change to Square ■
        'white_circle': ('■', 'Arial', 'square'),
        'black_circle': ('■', 'Arial', 'square'),
        
        # Square bullets
        'small_square': ('■', 'Arial', 'square'),
        'square': ('■', 'Arial', 'square'),
        
        # Dash/arrow bullets
        'dash': ('–', 'Arial', 'dash'),
        'arrow': ('■', 'Arial', 'square'),
        
        # Checkbox bullets
        'checkbox_empty': ('☐', 'Segoe UI Symbol', 'checkbox'),
        'checkbox_checked': ('☑', 'Segoe UI Symbol', 'checkbox'),
        'checkmark': ('✓', 'Segoe UI Symbol', 'checkmark'),
        
        # Number-like bullets (preserve original)
        'circled_number': (bullet_info.get('bullet_char', '①'), 'Arial Unicode MS', 'number'),
        'dingbat_number': (bullet_info.get('bullet_char', '❶'), 'Wingdings', 'number'),
        
        # Creative bullets
        'sparkle': ('✦', 'Segoe UI Symbol', 'star'),
        'star': ('★', 'Segoe UI Symbol', 'star'),
        'heart': ('♥', 'Segoe UI Symbol', 'heart'),
        'diamond': ('◆', 'Symbol', 'diamond'),
        'bullseye': ('◉', 'Arial Unicode MS', 'circle'),
        'right_triangle': ('►', 'Symbol', 'triangle'),
        'left_triangle': ('◄', 'Symbol', 'triangle'),
        
        # Generic fallback
        'generic_bullet': (bullet_info.get('bullet_char', '•'), 'Symbol', 'bullet')
    }
    
    bullet_type = bullet_info['type']
    if bullet_type in bullet_mapping:
        return bullet_mapping[bullet_type]
    else:
        # Default to standard bullet
        return ('•', 'Symbol', 'bullet')

def is_nested_bullet(line_text, previous_indent=0):
    """
    Determine if bullet is nested based on indentation
    """
    current_indent = len(line_text) - len(line_text.lstrip())
    
    if current_indent > previous_indent + 2:  # At least 2 spaces more
        return True, current_indent
    elif current_indent < previous_indent - 2:  # At least 2 spaces less
        return False, current_indent  # Outdented (higher level)
    else:
        return False, current_indent  # Same level

def process_bullet_list(lines, start_index):
    """
    Process consecutive bullet lines and apply consistent formatting
    """
    bullets = []
    i = start_index
    
    while i < len(lines):
        bullet_info = detect_bullet_type(lines[i])
        if bullet_info:
            bullets.append(bullet_info)
            i += 1
        else:
            break
    
    # Apply Word bullet styling
    if bullets:
        formatted_bullets = []
        previous_indent = 0
        
        for idx, bullet in enumerate(bullets):
            # Determine nesting level
            is_nested, current_indent = is_nested_bullet(
                bullet['original_line'], 
                previous_indent
            )
            
            # Map to Word bullet style
            bullet_char, font_name, bullet_type = map_to_word_bullet_style(bullet)
            
            formatted_bullets.append({
                'level': 1 if not is_nested else 2,
                'bullet_char': bullet_char,
                'font_name': font_name,
                'bullet_type': bullet_type,
                'content': bullet['content'],
                'original_bullet': bullet['bullet_char'] if bullet['bullet_char'] else '•',
                'indentation': current_indent
            })
            
            previous_indent = current_indent
        
        return formatted_bullets, i  # Return processed bullets and next index
    else:
        return [], start_index


class ImpliedBulletDetector:
    """
    Intelligent system to detect implied bullet points in text blocks.
    Uses heuristics, context analysis, and scoring to identify lists that aren't explicitly formatted.
    """
    
    def __init__(self):
        self.patterns = self._initialize_patterns()
        
    def _initialize_patterns(self):
        return {
            # Group 1: Imperative Verbs (Start of line)
            'imperative_start': re.compile(r'^\s*(?:Ensure|Verify|Check|Analyze|Review|Create|Update|Delete|Insert|Select|Run|Execute|Test|Validate|Monitor|Configure|Install|Deploy|Build|Compile|Debug|Fix|Resolve|Close|Open|Save|Print|Export|Import|Send|Receive|Get|Set|Put|Post|Patch|Head|Options|Trace|Connect)\b', re.IGNORECASE),
            
            # Group 2: Sentence Fragments (No subject/verb structure)
            'fragment_start': re.compile(r'^\s*(?:For|To|With|By|From|In|On|At|About|Under|Over|Between|Among|Through|During|Before|After|While|When|Where|Why|How|If|Unless|Until|Since|Because|Although|Though|Whereas|While)\b', re.IGNORECASE),
            
            # Group 3: Visual Separation (Short lines)
            'short_line': re.compile(r'^.{5,60}$'),  # 5-60 chars is typical for bullet points
            
            # Group 4: Sequential Transition Words
            'sequential_transition': re.compile(r'^\s*(?:First|Second|Third|Fourth|Fifth|Next|Then|Finally|Lastly|Moreover|Furthermore|Additionally|Also|In addition|Consequently|Therefore|Thus|Hence|Accordingly|As a result|For example|For instance|Specifically|In particular|Notably|Significantly|Importantly|Crucially|Essentially|Fundamentally|Ultimately|Eventually|Meanwhile|Simultaneously|Concurrently|Subsequently|Previously|Initially|Primarily|Secondarily|Tertiary)\b', re.IGNORECASE),
            
            # Group 5: Common List Starters
            'list_starter': re.compile(r'^\s*(?:[-*•>+]|\d+\.|\w\.)\s+'),
            
            # Group 6: Negative Indicators (Things that suggest it's NOT a bullet)
            'continuation_marker': re.compile(r'^\s*(?:and|or|but|so|nor|yet)\b', re.IGNORECASE),
            'pronoun_start': re.compile(r'^\s*(?:I|We|You|He|She|It|They|This|That|These|Those)\b', re.IGNORECASE),
        }

    def calculate_line_list_score(self, line, prev_line=None, next_line=None):
        """
        Calculate a score (0-100) indicating likelihood of a line being a list item.
        """
        if not line or not line.strip():
            return 0
            
        text = line.strip()
        score = 0
        
        # 1. Base Pattern Matching (Max 40 points)
        if self.patterns['imperative_start'].match(text):
            score += 30
        elif self.patterns['fragment_start'].match(text):
            score += 20
        elif self.patterns['sequential_transition'].match(text):
            score += 25
            
        # 2. Visual/Structural Heuristics (Max 30 points)
        if self.patterns['short_line'].match(text):
            score += 20
        
        # Check for capitalization
        if text[0].isupper():
            score += 10
            
        # 3. Contextual Analysis (Max 30 points)
        # If previous line ended with colon, high probability
        if prev_line and prev_line.strip().endswith(':'):
            score += 25
            
        # If next line starts similarly (parallelism)
        if next_line:
            next_text = next_line.strip()
            # Check for same starting word type
            if (self.patterns['imperative_start'].match(text) and self.patterns['imperative_start'].match(next_text)):
                score += 15
            elif (self.patterns['sequential_transition'].match(text) and self.patterns['sequential_transition'].match(next_text)):
                score += 15
                
        # 4. Penalties
        if self.patterns['continuation_marker'].match(text):
            score -= 30
        if self.patterns['pronoun_start'].match(text):
            score -= 20
        if text.endswith('.'): # Full sentences are less likely to be bullets in some contexts, but not always
            score -= 5
            
        return max(0, min(100, score))

    def detect_implied_bullet_blocks(self, lines):
        """
        Scans lines to identify blocks that should be converted to bullets.
        Returns a list of (start_index, end_index, bullet_type) tuples.
        """
        blocks = []
        i = 0
        current_block_start = -1
        current_block_scores = []
        
        while i < len(lines):
            line = lines[i]
            if isinstance(line, dict):
                line = line.get('text', '')
            
            prev = lines[i-1] if i > 0 else None
            if isinstance(prev, dict): prev = prev.get('text', '')
            
            nxt = lines[i+1] if i < len(lines)-1 else None
            if isinstance(nxt, dict): nxt = nxt.get('text', '')
            
            score = self.calculate_line_list_score(line, prev, nxt)
            
            # Thresholds
            DEFINITE_THRESHOLD = 85
            LIKELY_THRESHOLD = 70
            
            if score >= LIKELY_THRESHOLD:
                if current_block_start == -1:
                    current_block_start = i
                current_block_scores.append(score)
            else:
                if current_block_start != -1:
                    # End of block
                    # Validate block: needs at least 2 items or 1 very high confidence item
                    if len(current_block_scores) >= 2 or (len(current_block_scores) == 1 and current_block_scores[0] >= DEFINITE_THRESHOLD):
                        blocks.append((current_block_start, i - 1, 'square')) # Default to square for implied
                    
                    current_block_start = -1
                    current_block_scores = []
            
            i += 1
            
        # Check if block continues to end
        if current_block_start != -1:
             if len(current_block_scores) >= 2 or (len(current_block_scores) == 1 and current_block_scores[0] >= DEFINITE_THRESHOLD):
                blocks.append((current_block_start, len(lines) - 1, 'square'))
                
        return blocks


class PatternEngine:
    """Ultra-fast pattern matching engine for document analysis"""
    
    def __init__(self):
        self.patterns = self._initialize_patterns()
        self.implied_detector = ImpliedBulletDetector()
        
    def _initialize_patterns(self):
        """Initialize all recognition patterns - 40+ regex patterns"""
        return {
            # Heading Level 1 Patterns (ALL CAPS, Major Sections)
            'heading_1': [
                re.compile(r'^([A-Z][A-Z\s]{2,49})$'),  # ALL CAPS (3-50 chars)
                re.compile(r'^(CHAPTER\s+\d+.*)$', re.IGNORECASE),  # CHAPTER 1: Title
                re.compile(r'^(PART\s+[IVX]+.*)$', re.IGNORECASE),  # PART I, PART II
                re.compile(r'^(PART\s+\d+.*)$', re.IGNORECASE),  # PART 1, PART 2
                re.compile(r'^(\d+\.\s+[A-Z][A-Z\s]+)$'),  # "1. INTRODUCTION"
                re.compile(r'^(ACKNOWLEDGEMENT|ABSTRACT|INTRODUCTION|CONCLUSION|REFERENCES|BIBLIOGRAPHY|APPENDIX|APPENDICES|GLOSSARY|INDEX|PREFACE|FOREWORD|DEDICATION|TABLE OF CONTENTS|LIST OF FIGURES|LIST OF TABLES)S?$', re.IGNORECASE),
                re.compile(r'^(EXECUTIVE\s+SUMMARY)$', re.IGNORECASE),
                re.compile(r'^(LITERATURE\s+REVIEW)$', re.IGNORECASE),
                re.compile(r'^(RESEARCH\s+METHODOLOGY)$', re.IGNORECASE),
                re.compile(r'^(DATA\s+ANALYSIS)$', re.IGNORECASE),
                re.compile(r'^(FINDINGS\s+AND\s+DISCUSSION)$', re.IGNORECASE),
                re.compile(r'^(RECOMMENDATIONS)$', re.IGNORECASE),
            ],
            
            # Heading Level 2 Patterns (Title Case, Numbered Sections)
            'heading_2': [
                re.compile(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,6})$'),  # Title Case
                re.compile(r'^\d+\.\d+\s+([A-Z].{3,80})$'),  # "1.1 Background"
                re.compile(r'^([A-Z][a-z]+\s+and\s+[A-Z][a-z]+)$'),  # "Methods and Results"
                re.compile(r'^([A-Z][a-z]+\s+of\s+[A-Z][a-z]+.*)$'),  # "Analysis of Data"
                re.compile(r'^(Section\s+\d+[:\s].*)$', re.IGNORECASE),  # "Section 1: Overview"
                re.compile(r'^(\d+\s+[A-Z][a-z]+(?:\s+[A-Z]?[a-z]+)*)$'),  # "1 Introduction"
            ],
            
            # Heading Level 3 Patterns (Sub-sections)
            'heading_3': [
                re.compile(r'^\d+\.\d+\.\d+\s+(.+)$'),  # "1.1.1 Details"
                re.compile(r'^([a-z]\)\s+.+)$'),  # "a) Subsection"
                re.compile(r'^(\([a-z]\)\s+.+)$'),  # "(a) Subsection"
                re.compile(r'^([A-Z][a-z]+:)\s*$'),  # "Definition:"
                re.compile(r'^([ivx]+\.\s+.+)$', re.IGNORECASE),  # "i. First point"
                re.compile(r'^\(\d+\)\s+(.+)$'),  # "(1) Subsection"
            ],
            
            # Reference Patterns (APA, MLA, Chicago, IEEE)
            'reference': [
                re.compile(r'^([A-Z][a-z]+,?\s+[A-Z]\..*\(\d{4}\))'),  # "Smith, J. (2024)"
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z]\.\s+[A-Z]?\.?\s*\(\d{4}\))'),  # "Smith, J. A. (2024)"
                re.compile(r'^([A-Z][a-z]+\s+et\s+al\..*\d{4})'),  # "Smith et al. 2024"
                re.compile(r'^\[\d+\]'),  # "[1] Reference" - IEEE style
                re.compile(r'^([A-Z][a-z]+.*Retrieved from)'),  # Web reference
                re.compile(r'^([A-Z][a-z]+.*https?://)'),  # URL reference
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z]\.\s+\(\d{4}\)\.)'),  # APA format
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z][a-z]+\.\s+".+"\s+.+\d{4})'),  # MLA format
                re.compile(r'^\d+\.\s+[A-Z][a-z]+,?\s+[A-Z]'),  # Numbered reference
                re.compile(r'^([A-Z][a-z]+,?\s+[A-Z]\.\s*(&|and)\s+[A-Z][a-z]+)'),  # Multiple authors
                # NEW PATTERNS FOR ORGANIZATIONS AND LEGAL DOCUMENTS
                re.compile(r'^[\w\s\.\-&]+\s*\(\d{4}(?:/\d{4})?\).*$'), # Organization (Year)
                re.compile(r'^(?:Decree|Law|Order|Decision|Arrete)\s+No\.?.*$', re.IGNORECASE), # Legal
            ],
            
            # List Patterns - Bullet (Enhanced)
            'bullet_list': [
                # Standard bullets
                re.compile(r'^\s*[•○●▪■]\s+(.+)$'),
                # Dash/arrow bullets
                re.compile(r'^\s*[-–—]\s+(.+)$'),
                re.compile(r'^\s*[→➔➜➤➢]\s+(.+)$'),
                # Asterisk variants
                re.compile(r'^\s*[\*⁎⁑※]\s+(.+)$'),
                # Checkbox bullets
                re.compile(r'^\s*[☐☑✓✔]\s+(.+)$'),
                # Number-like bullets
                re.compile(r'^\s*[⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿]\s+(.+)$'),
                # Creative bullets
                re.compile(r'^\s*[✦✧★☆♥♡◆◇◉◎▸▹►◂◃◄]\s+(.+)$'),
            ],
            
            # List Patterns - Numbered
            'numbered_list': [
                re.compile(r'^(\d+[\.)]\s+.+)$'),  # 1. or 1)
                re.compile(r'^([a-z][\.)]\s+.+)$'),  # a. or a)
                re.compile(r'^([ivxlcdm]+[\.)]\s+.+)$', re.IGNORECASE),  # Roman numerals
                re.compile(r'^\(\d+\)\s+.+$'),  # (1) format
                re.compile(r'^\([a-z]\)\s+.+$'),  # (a) format
                re.compile(r'^[A-Z][\.)]\s+.+$'),  # A. or A)
                re.compile(r'^\d+\)\s+.+$'),  # 1) format
            ],
            
            # Table Patterns
            'table_marker': [
                re.compile(r'^\[TABLE\s*START\]', re.IGNORECASE),
                re.compile(r'^\[TABLE\s*END\]', re.IGNORECASE),
                re.compile(r'^Table\s+\d+', re.IGNORECASE),
                re.compile(r'^TABLE\s+\d+', re.IGNORECASE),
                re.compile(r'^Tabel\s+\d+', re.IGNORECASE),  # Common typo
            ],
            
            'table_row': [
                re.compile(r'^\|(.+\|)+$'),  # Markdown table row |cell|cell|
                re.compile(r'^\|[\s\-:]+\|$'),  # Markdown table separator |---|---|
            ],
            
            # Table separator (markdown)
            'table_separator': [
                re.compile(r'^\|[\s\-:]+\|[\s\-:]+.*\|$'),  # |---|---|
                re.compile(r'^[\-\|:\s]+$'),  # Pure separator line
            ],
            
            # Definition/Key Term Patterns
            'definition': [
                re.compile(r'^(Definition|Objective|Task|Goal|Purpose|Aim|Method|Result|Conclusion|Note|Important|Key Point|Summary|Overview|Background|Context|Example|Theorem|Lemma|Corollary|Proposition|Proof|Remark|Observation|Hypothesis|Assumption|Constraint|Limitation|Scope|Significance|Implication|Application|Contribution|Finding|Evidence|Data|Analysis|Interpretation|Explanation|Description|Specification|Requirement|Criteria|Criterion|Parameter|Variable|Constant|Factor|Element|Component|Aspect|Feature|Property|Characteristic|Attribute|Quality|Measure|Metric|Indicator|Index|Ratio|Rate|Percentage|Value|Score|Level|Degree|Extent|Amount|Quantity|Size|Scale|Range|Interval|Duration|Period|Phase|Stage|Step|Process|Procedure|Protocol|Algorithm|Formula|Equation|Model|Framework|Theory|Concept|Principle|Law|Rule|Guideline|Standard|Norm|Benchmark|Baseline|Reference|Source|Origin|Cause|Effect|Impact|Outcome|Consequence|Benefit|Advantage|Disadvantage|Risk|Challenge|Problem|Issue|Question|Answer|Solution|Strategy|Approach|Technique|Tool|Instrument|Device|System|Structure|Organization|Classification|Category|Type|Kind|Class|Group|Set|Collection|Series|Sequence|Order|Pattern|Trend|Distribution|Correlation|Relationship|Connection|Link|Association|Comparison|Contrast|Difference|Similarity|Analogy|Metaphor|Symbol|Sign|Signal|Indicator|Warning|Caution|Attention|Focus|Priority|Emphasis|Highlight|Point|Argument|Claim|Assertion|Statement|Proposition|Premise|Inference|Deduction|Induction|Generalization|Specialization|Abstraction|Instantiation|Implementation|Execution|Operation|Function|Action|Activity|Task|Job|Work|Effort|Attempt|Trial|Experiment|Test|Evaluation|Assessment|Review|Examination|Inspection|Investigation|Inquiry|Study|Research|Survey|Poll|Interview|Questionnaire|Form|Document|Report|Paper|Article|Book|Chapter|Section|Paragraph|Sentence|Word|Term|Phrase|Expression|Language|Text|Content|Information|Data|Knowledge|Wisdom|Intelligence|Understanding|Comprehension|Interpretation|Meaning|Significance|Importance|Relevance|Value|Worth|Merit|Quality|Standard|Excellence|Performance|Efficiency|Effectiveness|Productivity|Output|Input|Resource|Asset|Capital|Investment|Cost|Expense|Price|Fee|Charge|Payment|Revenue|Income|Profit|Loss|Balance|Budget|Forecast|Projection|Estimate|Calculation|Computation|Measurement|Quantification|Qualification|Certification|Accreditation|Validation|Verification|Confirmation|Approval|Authorization|Permission|License|Right|Privilege|Responsibility|Duty|Obligation|Commitment|Promise|Guarantee|Warranty|Assurance|Insurance|Protection|Security|Safety|Risk|Hazard|Danger|Threat|Vulnerability|Weakness|Strength|Opportunity|Challenge):\s*(.+)?$', re.IGNORECASE),
            ],
            
            # Figure/Image Caption
            'figure': [
                re.compile(r'^Figure\s+\d+', re.IGNORECASE),
                re.compile(r'^Fig\.\s*\d+', re.IGNORECASE),
                re.compile(r'^Image\s+\d+', re.IGNORECASE),
                re.compile(r'^Diagram\s+\d+', re.IGNORECASE),
                re.compile(r'^Chart\s+\d+', re.IGNORECASE),
                re.compile(r'^Graph\s+\d+', re.IGNORECASE),
                re.compile(r'^Illustration\s+\d+', re.IGNORECASE),
            ],
            
            # Equation patterns
            'equation': [
                re.compile(r'^Equation\s+\d+', re.IGNORECASE),
                re.compile(r'^Eq\.\s*\d+', re.IGNORECASE),
                re.compile(r'^\(\d+\)$'),  # Just (1) on its own line - equation number
            ],
            
            # Quote patterns
            'quote': [
                re.compile(r'^["\"].*["\"]$'),  # Quoted text
                re.compile(r'^>\s+.+$'),  # Blockquote markdown style
            ],
            
            # Code block patterns
            'code': [
                re.compile(r'^```'),  # Markdown code fence
                re.compile(r'^~~~'),  # Alternative code fence
                re.compile(r'^\t{2,}.+$'),  # Heavily indented (code-like)
            ],
            
            # Section Keywords for special detection
            'section_abstract': re.compile(r'^(abstract|executive\s+summary)$', re.IGNORECASE),
            'section_intro': re.compile(r'^(introduction|background|overview|motivation|context)$', re.IGNORECASE),
            'section_methods': re.compile(r'^(method|methodology|approach|procedure|materials|design|implementation)s?$', re.IGNORECASE),
            'section_results': re.compile(r'^(results?|findings?|outcomes?|data|observations?)$', re.IGNORECASE),
            'section_discussion': re.compile(r'^(discussion|analysis|interpretation|evaluation)$', re.IGNORECASE),
            'section_conclusion': re.compile(r'^(conclusions?|summary|final\s+remarks|future\s+work|recommendations?)$', re.IGNORECASE),
            'section_references': re.compile(r'^(references|bibliography|works\s+cited|citations|sources)$', re.IGNORECASE),
            'section_appendix': re.compile(r'^(appendix|appendices|supplementary|annexure)$', re.IGNORECASE),
            
            # NEW PATTERNS (December 30, 2025)
            
            # Inline Formatting Patterns (Bold/Italic/Bold-Italic)
            'inline_formatting': [
                re.compile(r'\*\*\*(.+?)\*\*\*|___(.+?)___'),  # Bold+Italic (check first)
                re.compile(r'\*\*(.+?)\*\*|__(.+?)__'),       # Bold
                re.compile(r'(?<!\*)\*([^*\n]+?)\*(?!\*)|(?<!_)_([^_\n]+?)_(?!_)'),  # Italic (avoid matching bold)
            ],
            
            # Page Header/Footer Patterns
            'page_metadata': [
                re.compile(r'^\s*(?:page|p|pg)\.?\s*\d+\s*(?:of\s*\d+)?\s*$', re.IGNORECASE),
                re.compile(r'^\s*(?:header|footer|running head):?\s*.{1,50}$', re.IGNORECASE),
                re.compile(r'^\s*(?:confidential|draft|version\s*\d+|date:|©|copyright)\b.*$', re.IGNORECASE),
                re.compile(r'^\s*-\s*\d+\s*-\s*$'),  # Centered page numbers like - 1 -
            ],
            
            # Academic Metadata Patterns (Title, Author, Affiliation)
            'academic_metadata': [
                re.compile(r'^\s*(?:by|authors?:)\s+([A-Z][a-z]+(?:\s+[A-Z]\.)*(?:\s+[A-Z][a-z]+)*(?:\s*,\s*[A-Z][a-z]+(?:\s+[A-Z]\.)*(?:\s+[A-Z][a-z]+)*)*)$', re.IGNORECASE),
                re.compile(r'^\s*(?:department|school|college|university|institute|faculty|center|laboratory)\s+of\s+.+$', re.IGNORECASE),
                re.compile(r'^\s*(?:email|e-mail|correspondence\s+to|contact):?\s*\S+@\S+\.\S+\s*$', re.IGNORECASE),
                re.compile(r'^\s*(?:submitted\s+to|prepared\s+for|in\s+partial\s+fulfillment)\s+.+$', re.IGNORECASE),
            ],
            
            # Mathematical Expression Patterns
            'math_expression': [
                re.compile(r'^\$\$[^$]+\$\$$'),  # Display math $$...$$
                re.compile(r'\$[^$\n]+\$'),      # Inline math $...$
                re.compile(r'^\\\[.*\\\]$'),   # LaTeX display \[...\]
                re.compile(r'\\\(.*\\\)'),    # LaTeX inline \(...\)
                re.compile(r'^\s*[A-Za-z]\s*[=<>≤≥≠≈]\s*.+$'),  # Simple equations like x = 5
            ],
            
            # Footnote/Endnote Patterns
            'footnote_endnote': [
                re.compile(r'^\s*(?:endnotes?|footnotes?)\s*$', re.IGNORECASE),  # Section header
                re.compile(r'^\s*(\d{1,3}|[a-z]|\*{1,3})\s*[\.\)]\s+.{10,}$'),  # Footnote entry
                re.compile(r'^\s*\[\d+\]\s+.{10,}$'),  # Bracketed footnote [1] text...
                re.compile(r'\[\^\d+\]'),  # Markdown footnote reference [^1]
            ],
            
            # ============================================================
            # NEW PATTERNS - December 30, 2025 (20 Academic Formatting Patterns)
            # ============================================================
            
            # 1. HEADING_HIERARCHY - Markdown-style hierarchical headings
            'heading_hierarchy': [
                re.compile(r'^#\s+CHAPTER\s+\w+[:\s]', re.IGNORECASE),  # # CHAPTER X: Title
                re.compile(r'^##\s+\d+\.\d+\s+'),  # ## 1.1 Section
                re.compile(r'^###\s+\d+\.\d+\.\d+\s+'),  # ### 1.1.1 Subsection
                re.compile(r'^####\s+\d+\.\d+\.\d+\.\d+\s+'),  # #### 1.1.1.1 Sub-subsection
                re.compile(r'^#####\s+[a-z]\)\s+', re.IGNORECASE),  # ##### a) Point
                re.compile(r'^#{1,6}\s+'),  # Generic markdown heading
            ],
            
            # 2. ACADEMIC_TABLE - Enhanced table detection
            'academic_table': [
                re.compile(r'^\|[-\s:]+\|[-\s:]+\|'),  # Table separator row
                re.compile(r'^\|\s*\*\*[^|]+\*\*\s*\|'),  # Bold header cells
                re.compile(r'^\|[^|]+\|[^|]+\|[^|]+\|'),  # 3+ column table
                re.compile(r'^Table\s+\d+\.\d+[:\s]', re.IGNORECASE),  # Table X.Y: caption
            ],
            
            # 3. LIST_NORMALIZER - Enhanced list detection with indentation
            'list_nested': [
                re.compile(r'^(\s{2,})[•\-\*]\s+'),  # Indented bullet
                re.compile(r'^(\s{2,})\d+[\.\)]\s+'),  # Indented numbered
                re.compile(r'^(\s{2,})[a-z][\.\)]\s+'),  # Indented lettered
                re.compile(r'^\s*□\s+'),  # Checkbox empty
                re.compile(r'^\s*[☐☑✓✗]\s+'),  # Various checkbox symbols
            ],
            
            # 4. FIGURE_EQUATION - Enhanced figure/equation captions
            'figure_equation': [
                re.compile(r'^[Ff]igure\s+\d+\.\d+[:\s]'),  # Figure X.Y: caption
                re.compile(r'^\*\*Figure\s+\d+'),  # **Figure X**
                re.compile(r'^\$\$\s*$'),  # Start of display equation
                re.compile(r'\\begin\{equation\}'),  # LaTeX equation environment
                re.compile(r'\\end\{equation\}'),
            ],
            
            # 5. CITATION_STYLE - In-text citation patterns
            'citation_inline': [
                re.compile(r'\([A-Z][a-z]+,\s*\d{4}\)'),  # (Smith, 2020)
                re.compile(r'\([A-Z][a-z]+\s*&\s*[A-Z][a-z]+,\s*\d{4}\)'),  # (Smith & Jones, 2020)
                re.compile(r'\([A-Z][a-z]+\s+et\s+al\.,\s*\d{4}\)'),  # (Smith et al., 2020)
                re.compile(r'\([A-Z][a-z]+,\s*\d{4};\s*[A-Z][a-z]+,\s*\d{4}\)'),  # Multiple citations
                re.compile(r'\([A-Z][a-z]+\s+and\s+[A-Z][a-z]+,\s*\d{4}\)'),  # (Smith and Jones, 2020)
            ],
            
            # 6. HEADER_FOOTER - Running header/footer detection
            'running_header': [
                re.compile(r'^[A-Z][A-Z\s]+\s*\|\s*[A-Z]'),  # CHAPTER | TITLE format
                re.compile(r'^\s*Page\s+\d+\s+of\s+\d+\s*$', re.IGNORECASE),  # Page X of Y
                re.compile(r'^Chapter\s+\d+\s*\|\s*', re.IGNORECASE),  # Chapter X | format
            ],
            
            # 7. APPENDIX_FORMAT - Appendix-specific patterns
            'appendix_format': [
                re.compile(r'^APPENDIX\s+[A-Z]$', re.IGNORECASE),  # APPENDIX A
                re.compile(r'^[A-Z]\.\d+[:\s]'),  # A.1: Appendix section
                re.compile(r'^[A-Z]\.\d+\.\d+[:\s]'),  # A.1.1: Appendix subsection
                re.compile(r'^Appendix\s+[A-Z][:\s]', re.IGNORECASE),  # Appendix A: Title
            ],
            
            # 8. BLOCK_QUOTE - Extended block quote detection
            'block_quote': [
                re.compile(r'^>\s+.+$'),  # Markdown blockquote
                re.compile(r'^\s{4,}".+$'),  # Indented quoted text
                re.compile(r'^"[^"]{50,}'),  # Long quoted text
                re.compile(r"^'[^']{50,}"),  # Single-quote long text
            ],
            
            # 9. MATH_MODEL - Statistical/Mathematical model patterns
            'math_model': [
                re.compile(r'[Yy]\s*=\s*[βα]'),  # Y = β... regression
                re.compile(r'\\beta_\d'),  # LaTeX beta subscript
                re.compile(r'[Rr]²\s*=\s*[\d\.]'),  # R² = 0.xx
                re.compile(r'[Ff]\s*\(\s*\d+\s*,\s*\d+\s*\)'),  # F(df1, df2)
                re.compile(r'[Pp]\s*[<>=]\s*[\d\.]+'),  # p < 0.05
                re.compile(r'\\epsilon|\\sigma|\\mu'),  # Greek letters
            ],
            
            # 10. FONT_CONSISTENCY - Text emphasis detection
            'text_emphasis': [
                re.compile(r'\*\*[A-Z][^*]+\*\*:'),  # **Term**: definition
                re.compile(r'\*[A-Za-z][^*]+\*'),  # *italicized text*
                re.compile(r'`[^`]+`'),  # `monospace`
                re.compile(r'\*\*\*.+\*\*\*'),  # ***bold-italic***
            ],
            
            # 11. REFERENCE_FORMAT - Reference list patterns (APA enhanced)
            'reference_apa': [
                re.compile(r'^[A-Z][a-z]+,\s*[A-Z]\.\s*[A-Z]?\.\s*\(\d{4}\)'),  # Author, A. B. (Year)
                re.compile(r'\*[^*]+\*\.'),  # Italicized title
                re.compile(r'https?://doi\.org/'),  # DOI URL
                re.compile(r'doi:\s*[\d\.]+/'),  # DOI reference
                re.compile(r'Retrieved\s+\w+\s+\d+,\s*\d{4}'),  # Retrieved Month Day, Year
            ],
            
            # 12. TOC_GENERATOR - Table of contents patterns
            'toc_entry': [
                re.compile(r'^[\d\.]+\s+[A-Z].+\.{3,}\s*\d+$'),  # 1.1 Title....... 5
                re.compile(r'^[A-Z].+\.{5,}\s*\d+$'),  # Title........ 5
                re.compile(r'^(CHAPTER|APPENDIX)\s+\w+.+\d+$', re.IGNORECASE),  # TOC chapter entry
                re.compile(r'^\s{2,}\d+\.\d+.+\d+$'),  # Indented subsection entry
            ],
            
            # 13. FOOTNOTE_FORMAT - Enhanced footnote patterns
            'footnote_marker': [
                re.compile(r'\[\^\d+\]'),  # [^1] markdown footnote
                re.compile(r'\^\d+'),  # ^1 superscript style
                re.compile(r'\(\d+\)$'),  # (1) at end of line
                re.compile(r'^\[\^\d+\]:'),  # [^1]: footnote definition
            ],
            
            # 14. ABBREVIATION_MANAGER - Abbreviation detection
            'abbreviation': [
                re.compile(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\s+\([A-Z]{2,}\)'),  # Full Term (ABBR)
                re.compile(r'\([A-Z]{2,}\)'),  # (ABBR) alone
                re.compile(r'^[A-Z]{2,}:\s+'),  # ABBR: definition
                re.compile(r'\b[A-Z]{2,}s?\b'),  # Standalone abbreviation
            ],
            
            # 15. CAPTION_STYLE - Table/Figure caption formatting
            'caption_format': [
                re.compile(r'^\*\*Table\s+\d+'),  # **Table X**
                re.compile(r'^\*\*Figure\s+\d+'),  # **Figure X**
                re.compile(r'^Table\s+\d+\.\d+:'),  # Table X.Y: caption
                re.compile(r'^Figure\s+\d+\.\d+:'),  # Figure X.Y: caption
                re.compile(r'^Source:\s+'),  # Source: attribution
                re.compile(r'^Note:\s+', re.IGNORECASE),  # Note: table note
            ],
            
            # 16. PAGE_BREAK - Page break indicators
            'page_break': [
                re.compile(r'^-{3,}$'),  # --- horizontal rule
                re.compile(r'^\*{3,}$'),  # *** horizontal rule
                re.compile(r'^_{3,}$'),  # ___ horizontal rule
                re.compile(r'^\[PAGE\s*BREAK\]', re.IGNORECASE),  # [PAGE BREAK] marker
                re.compile(r'^\\newpage', re.IGNORECASE),  # LaTeX newpage
                re.compile(r'^\\pagebreak', re.IGNORECASE),  # LaTeX pagebreak
            ],
            
            # 17. STATS_FORMAT - Statistical results patterns
            'statistical_result': [
                re.compile(r'β\s*=\s*[-\d\.]+'),  # β = 0.45
                re.compile(r'[Pp]\s*[<>=]\s*\.?\d+'),  # p < .001
                re.compile(r'[Ff]\s*\(\d+,\s*\d+\)\s*='),  # F(3, 56) =
                re.compile(r'[Tt]\s*\(\d+\)\s*='),  # t(45) =
                re.compile(r'χ²\s*\(\d+\)\s*='),  # χ²(5) =
                re.compile(r'[Rr]²?\s*=\s*\.?\d+'),  # R = .67 or R² = .45
                re.compile(r'[Nn]\s*=\s*\d+'),  # N = 100
                re.compile(r'[Mm]\s*=\s*[\d\.]+,\s*[Ss][Dd]\s*='),  # M = 3.5, SD =
                re.compile(r'CI\s*=?\s*\[[\d\.\-,\s]+\]'),  # CI = [0.5, 1.2]
            ],
            
            # 18. QUESTIONNAIRE_STYLE - Survey/questionnaire patterns
            'questionnaire': [
                re.compile(r'^Section\s+[A-Z][:\s]', re.IGNORECASE),  # Section A:
                re.compile(r'^\d+\.\s+\*\*[^*]+\*\*'),  # 1. **Question**
                re.compile(r'[□☐☑✓✗○●]\s+[A-Za-z]'),  # Checkbox options
                re.compile(r'\|\s*SA\s*\|\s*A\s*\|\s*N\s*\|\s*D\s*\|\s*SD\s*\|'),  # Likert header
                re.compile(r'^\s*□\s+[A-Za-z]'),  # □ Option
            ],
            
            # 19. GLOSSARY_FORMAT - Glossary/definition list patterns
            'glossary_entry': [
                re.compile(r'^\*\*[A-Z][^*]+\*\*:\s+'),  # **Term**: Definition
                re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:\s+[A-Z]'),  # Term: Definition sentence
                re.compile(r'^•\s+\*\*[^*]+\*\*'),  # • **Term**
            ],
            
            # 20. CROSS_REFERENCE - Cross-reference patterns
            'cross_reference': [
                re.compile(r'[Ss]ee\s+[Tt]able\s+\d+'),  # See Table X
                re.compile(r'[Ss]ee\s+[Ff]igure\s+\d+'),  # See Figure X
                re.compile(r'[Ss]ee\s+[Ss]ection\s+\d+'),  # See Section X
                re.compile(r'[Aa]s\s+shown\s+in\s+[Tt]able'),  # As shown in Table
                re.compile(r'[Aa]s\s+discussed\s+in\s+[Ss]ection'),  # As discussed in Section
                re.compile(r'\([Ss]ee\s+[Pp]age\s+\d+\)'),  # (See page X)
                re.compile(r'\(p\.\s*\d+\)'),  # (p. 45)
            ],
            
            # 21. ACADEMIC_SECTION_PAGE_BREAKS - Sections that must start on new pages
            'academic_section_page_breaks': [
                # Front matter sections (case-insensitive)
                re.compile(r'^#+\s*(ACKNOWLEDGEMENTS?|ACKNOWLEDGMENTS?)\s*$', re.IGNORECASE),
                re.compile(r'^(ACKNOWLEDGEMENTS?|ACKNOWLEDGMENTS?)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*DEDICATION\s*$', re.IGNORECASE),
                re.compile(r'^DEDICATION\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*ABSTRACT\s*$', re.IGNORECASE),
                re.compile(r'^ABSTRACT\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(TABLE\s+OF\s+CONTENTS|CONTENTS)\s*$', re.IGNORECASE),
                re.compile(r'^(TABLE\s+OF\s+CONTENTS|CONTENTS)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*LIST\s+OF\s+TABLES\s*$', re.IGNORECASE),
                re.compile(r'^LIST\s+OF\s+TABLES\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*LIST\s+OF\s+FIGURES\s*$', re.IGNORECASE),
                re.compile(r'^LIST\s+OF\s+FIGURES\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*GLOSSARY\s*$', re.IGNORECASE),
                re.compile(r'^GLOSSARY\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(LIST\s+OF\s+ABBREVIATIONS|ABBREVIATIONS)\s*$', re.IGNORECASE),
                re.compile(r'^(LIST\s+OF\s+ABBREVIATIONS|ABBREVIATIONS)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(APPENDICES|APPENDIX)\s*$', re.IGNORECASE),
                re.compile(r'^(APPENDICES|APPENDIX)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*REFERENCES?\s*$', re.IGNORECASE),
                re.compile(r'^REFERENCES?\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*BIBLIOGRAPHY\s*$', re.IGNORECASE),
                re.compile(r'^BIBLIOGRAPHY\s*$', re.IGNORECASE),
                
                # Chapter headings (various formats)
                re.compile(r'^#+\s*CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^#+\s*CHAPTER\s+[1-9][0-9]?\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAPTER\s+[1-9][0-9]?\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^#+\s*CHAPTER\s+[IVXLC]+\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAPTER\s+[IVXLC]+\s*[:\.]?\s*.*$', re.IGNORECASE),
            ],
            
            # 23. DISSERTATION_CHAPTER_TITLES - Chapter heading + title patterns
            'dissertation_chapter': [
                # Chapter heading only (CHAPTER ONE, CHAPTER 1, CHAPTER I)
                re.compile(r'^#+?\s*CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAPTER\s+[1-9][0-9]?\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAPTER\s+[IVXLC]+\s*$', re.IGNORECASE),
                # Chapter heading with title on same line
                re.compile(r'^#+?\s*CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\-\.]\s*.+$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAPTER\s+[1-9][0-9]?\s*[:\-\.]\s*.+$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAPTER\s+[IVXLC]+\s*[:\-\.]\s*.+$', re.IGNORECASE),
            ],
            
            # 24. COPYRIGHT_PAGE - Copyright notice patterns
            'copyright_content': [
                re.compile(r'^\s*©\s*[Cc]opyright\s+', re.IGNORECASE),
                re.compile(r'^\s*[Cc]opyright\s+©', re.IGNORECASE),
                re.compile(r'^\s*[Cc]opyright\s+\d{4}', re.IGNORECASE),
                re.compile(r'^\s*All\s+Rights\s+Reserved\s*\.?\s*$', re.IGNORECASE),
                re.compile(r'^\s*No\s+part\s+of\s+this\s+(document|thesis|dissertation|work)', re.IGNORECASE),
            ],
            
            # 25. DECLARATION_PAGE - Declaration patterns
            'declaration_content': [
                re.compile(r'^#+?\s*DECLARATION\s*$', re.IGNORECASE),
                re.compile(r'I\s+(hereby\s+)?declare\s+that', re.IGNORECASE),
                re.compile(r'This\s+(thesis|dissertation|work)\s+is\s+(my\s+own|original)', re.IGNORECASE),
                re.compile(r'has\s+not\s+been\s+(previously\s+)?submitted', re.IGNORECASE),
            ],
            
            # 26. CERTIFICATION_PAGE - Certification/Approval patterns
            'certification_content': [
                re.compile(r'^#+?\s*CERTIFICATION\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*APPROVAL\s+PAGE\s*$', re.IGNORECASE),
                re.compile(r'This\s+is\s+to\s+certify\s+that', re.IGNORECASE),
                re.compile(r'has\s+been\s+(read\s+and\s+)?approved', re.IGNORECASE),
                re.compile(r'meets\s+the\s+requirements?', re.IGNORECASE),
                re.compile(r'^#+?\s*COMMITTEE\s+APPROVAL\s*$', re.IGNORECASE),
            ],
            
            # 27. CHAPTER_TITLE_FOLLOWING - Chapter title that follows chapter heading
            'chapter_title_following': [
                # All caps title (common for chapter titles)
                re.compile(r'^#+?\s*[A-Z][A-Z\s]+$'),
                # Common chapter title keywords
                re.compile(r'^#+?\s*(INTRODUCTION|LITERATURE\s+REVIEW|METHODOLOGY|RESULTS|DISCUSSION|CONCLUSION|RECOMMENDATIONS?|THEORETICAL\s+FRAMEWORK|SUMMARY\s+AND\s+CONCLUSIONS?|SUMMARY\s+OF\s+FINDINGS)', re.IGNORECASE),
            ],
            
            # 28. SIGNATURE_LINE - Signature block patterns
            'signature_line': [
                re.compile(r'^\s*_{10,}\s*$'),  # Long underscore line
                re.compile(r'^\s*\.{10,}\s*$'),  # Long dotted line
                re.compile(r'^\s*-{10,}\s*$'),  # Long dash line
                re.compile(r'^\s*(Signed?|Signature)\s*:\s*_{5,}', re.IGNORECASE),
                re.compile(r'^\s*Date\s*:\s*_{5,}', re.IGNORECASE),
                re.compile(r'^\s*Name\s*:\s*_{5,}', re.IGNORECASE),
            ],
            
            # 29. SUPERVISOR_INFO - Supervisor/Advisor patterns
            'supervisor_info': [
                re.compile(r'^\s*(Supervisor|Advisor|Chair)\s*:', re.IGNORECASE),
                re.compile(r'^\s*(First|Second|Third)\s+(Reader|Supervisor|Advisor)', re.IGNORECASE),
                re.compile(r'^\s*(Dr\.|Prof\.|Professor)\s+[A-Z]', re.IGNORECASE),
                re.compile(r'^\s*Committee\s+(Member|Chair)', re.IGNORECASE),
            ],
            
            # 30. TOC_ENTRY - Table of Contents entry patterns
            'toc_entry': [
                re.compile(r'^.+\.{3,}\s*\d+\s*$'),  # Title......... Page
                re.compile(r'^.+\s+\.{2,}\s*\d+\s*$'),  # Title .. Page
                re.compile(r'^\s*[IVXLC]+\s+.+\s+\d+\s*$'),  # Roman numeral entry
                re.compile(r'^\s*\d+\.\d*\s+.+\s+\d+\s*$'),  # Numbered entry
            ],
            
            # 31. HEADING_SPACE_ISSUES - Patterns for detecting spacing issues in headings
            'heading_space_issues': [
                # Trailing spaces on headings
                re.compile(r'^#+\s+.+\s{2,}$'),  # Heading with 2+ trailing spaces
                # Multiple spaces before punctuation
                re.compile(r'^#+\s+.+\s{2,}[:\.]\s*'),  # Heading with spaces before colon/period
                # Multiple internal spaces
                re.compile(r'^#+\s+.+\s{3,}.+$'),  # Heading with 3+ consecutive spaces inside
            ],
            
            # 32. SPACING_CLEANUP - General spacing issue patterns
            'spacing_cleanup': [
                re.compile(r'\s{2,}'),  # Multiple spaces (for replacement)
                re.compile(r'\s+,'),  # Space before comma
                re.compile(r'\s+\.'),  # Space before period
                re.compile(r'\s+:'),  # Space before colon
                re.compile(r'\s+$'),  # Trailing whitespace
            ],
            
            # 33. ACADEMIC_TABLE_CONTENT - Table content type detection patterns
            'table_content_numeric': [
                re.compile(r'^\s*[\-\+]?\d+\.?\d*\s*$'),  # Plain numbers: 123, 123.45, -45.6
                re.compile(r'^\s*[\-\+]?\d{1,3}(,\d{3})*(\.\d+)?\s*$'),  # Comma-separated: 1,234.56
            ],
            
            'table_content_percentage': [
                re.compile(r'^\s*[\-\+]?\d+\.?\d*\s*%\s*$'),  # 25.5%, -10%
            ],
            
            'table_content_statistical': [
                re.compile(r'[βαγδ]'),  # Greek letters
                re.compile(r'p\s*[<>=]\s*[\d\.]+'),  # p-values: p < 0.05
                re.compile(r'[Ff]\s*\(\s*\d+\s*,\s*\d+\s*\)'),  # F-statistic: F(3, 56)
                re.compile(r'[Rr]²?\s*=\s*[\d\.]+'),  # R or R²
                re.compile(r'[Tt]\s*\(\s*\d+\s*\)'),  # t-statistic: t(45)
                re.compile(r'χ²\s*\(\s*\d+\s*\)'),  # Chi-square
                re.compile(r'SE\s*=?\s*[\d\.]+'),  # Standard error
                re.compile(r'CI\s*=?\s*\['),  # Confidence interval
                re.compile(r'\*{1,3}$'),  # Significance asterisks
            ],
            
            'table_content_ci': [
                re.compile(r'\[[\d\.\-,\s]+\]'),  # Confidence interval: [0.5, 1.2]
                re.compile(r'\([\d\.\-,\s]+\)'),  # Alternative: (0.5, 1.2)
            ],
            
            # 34. TABLE_CAPTION_FORMAT - Table caption patterns
            'table_caption_format': [
                re.compile(r'^\s*[Tt]able\s+\d+\.\d+[\:\.]?\s*.+$'),  # Table 4.1: Title
                re.compile(r'^\s*[Tt]able\s+\d+[\:\.]?\s*.+$'),  # Table 4: Title
                re.compile(r'^\s*\*\*[Tt]able\s+\d+\.?\d*\*\*[\:\.]?\s*.+$'),  # **Table 4.1**: Title
                re.compile(r'^\s*TABLE\s+\d+\.?\d*[\:\.]?\s*.+$'),  # TABLE 4.1: Title
            ],
            
            # 35. SHORT_DOCUMENT_INDICATORS - Detect if document is short/assignment type
            'short_doc_indicators': [
                re.compile(r'^Assignment\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Homework\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Course\s+Material[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Exercise\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Lab\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Worksheet[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Quiz\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem\s+Set\s*\d*[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 36. LONG_DOCUMENT_INDICATORS - Indicators of formal academic documents
            'long_doc_indicators': [
                re.compile(r'Dissertation', re.IGNORECASE),
                re.compile(r'\bThesis\b', re.IGNORECASE),
                re.compile(r'\bPhD\b', re.IGNORECASE),
                re.compile(r"Master's\s+(?:Thesis|Dissertation)", re.IGNORECASE),
                re.compile(r'Literature\s+Review', re.IGNORECASE),
                re.compile(r'Abstract', re.IGNORECASE),
            ],
            
            # 37. TOC_PATTERNS - Detect Table of Contents for removal
            'toc_header': [
                re.compile(r'^#+\s*TABLE\s+OF\s+CONTENTS\s*$', re.IGNORECASE),
                re.compile(r'^TABLE\s+OF\s+CONTENTS\s*$', re.IGNORECASE),
                re.compile(r'^Contents\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*Contents\s*$', re.IGNORECASE),
            ],
            
            'toc_content_line': [
                re.compile(r'^\d+\.\s+.+\.{3,}\s*\d+$'),  # 1. Introduction ........ 5
                re.compile(r'^[A-Za-z].+\.{3,}\s*\d+$'),  # Introduction ........ 5
                re.compile(r'^.+\s+\.{5,}\s*\d+$'),  # Anything with dot leaders
                re.compile(r'^\s*\d+\.\d+\s+.+\s+\d+\s*$'),  # 1.1 Section  5
            ],
            
            # 38. KEY_POINT_LEARNING_OBJECTIVES - Learning goals/objectives
            'key_point_learning': [
                re.compile(r'^Learning\s+Objectives?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Objectives?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Goals?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^By\s+the\s+end.*you\s+will\s+be\s+able\s+to[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Learning\s+Outcomes?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Outcome\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Objective\s+\d+[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 39. KEY_POINT_DEFINITIONS - Important definitions
            'key_point_definitions': [
                re.compile(r'^Definition[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^[A-Z][a-z]+\s+is\s+(?:defined\s+as|refers\s+to)', re.IGNORECASE),
                re.compile(r'^The\s+term\s+["\'][^"\']+["\']\s+means', re.IGNORECASE),
                re.compile(r'^[A-Z][a-z]+[\:\s]+(?:A|An|The)\s+', re.IGNORECASE),  # Term: A definition
            ],
            
            # 40. KEY_POINT_KEY_CONCEPTS - Critical information markers
            'key_point_concepts': [
                re.compile(r'^Key\s+(?:Concept|Idea|Point)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Important[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Note[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Remember[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Critical[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Essential[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Key\s+Takeaway[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 41. KEY_POINT_PROCEDURES - Steps and procedures
            'key_point_procedures': [
                re.compile(r'^Steps?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Procedure[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Process[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Algorithm[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Method[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Step\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Instructions?[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 42. KEY_POINT_EXAMPLES - Example markers
            'key_point_examples': [
                re.compile(r'^Example[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^For\s+instance[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^For\s+example[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Consider[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Sample[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Example\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Case\s+Study[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 43. KEY_POINT_WARNINGS - Caution/warning markers
            'key_point_warnings': [
                re.compile(r'^Warning[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Caution[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Avoid[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Common\s+(?:mistake|error)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Do\s+not[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Never[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Pitfall[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 44. KEY_POINT_EXERCISES - Exercise/question markers
            'key_point_exercises': [
                re.compile(r'^Exercise[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Question[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Task[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Challenge[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Exercise\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Question\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Practice[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 45. KEY_POINT_SUMMARY - Summary/conclusion markers
            'key_point_summary': [
                re.compile(r'^Summary[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Conclusion[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Takeaways?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^In\s+summary[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^To\s+(?:summarize|recap)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Main\s+(?:points?|takeaways?)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Key\s+Points?[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 46. ASSIGNMENT_HEADER_ELEMENTS - Assignment metadata fields
            'assignment_header': [
                re.compile(r'^Student(?:\s+Name)?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Student\s+ID[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Name[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Course(?:\s+Code)?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Instructor[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Professor[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Due\s+Date[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Date[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Submitted\s+(?:by|to)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Class[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Section[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # ============================================================
            # 47. POINT_FORM_CONTENT_FORMATTING - Detect and format point-form content
            # ============================================================
            
            # 47a. List patterns - Various list markers
            'point_form_numbered': [
                re.compile(r'^\s*(\d+)[\.\)]\s+.+$'),  # 1. Item or 1) Item
                re.compile(r'^\s*([a-z])[\.\)]\s+.+$'),  # a. Item or a) Item
                re.compile(r'^\s*([ivxIVX]+)[\.\)]\s+.+$'),  # i. Item, ii. Item (Roman)
                re.compile(r'^\s*\((\d+)\)\s+.+$'),  # (1) Item
                re.compile(r'^\s*\(([a-z])\)\s+.+$'),  # (a) Item
                re.compile(r'^\s*([A-Z])[\.\)]\s+.+$'),  # A. Item or A) Item
            ],
            
            'point_form_bulleted': [
                re.compile(r'^\s*[\-\*\•◦→○▪▫■□►▸◆◇]\s+.+$'),  # Various bullet symbols
                re.compile(r'^\s*—\s+.+$'),  # Em dash
                re.compile(r'^\s*–\s+.+$'),  # En dash
            ],
            
            'point_form_checkbox': [
                re.compile(r'^\s*□\s+.+$'),  # □ Item
                re.compile(r'^\s*☐\s+.+$'),  # ☐ Item
                re.compile(r'^\s*☑\s+.+$'),  # ☑ Item (checked)
                re.compile(r'^\s*✓\s+.+$'),  # ✓ Item
                re.compile(r'^\s*✗\s+.+$'),  # ✗ Item
                re.compile(r'^\s*\[\s*\]\s+.+$'),  # [ ] Item
                re.compile(r'^\s*\[[xX]\]\s+.+$'),  # [x] Item (checked)
            ],
            
            # 47b. Context clues - Phrases that precede point lists
            'point_form_context_clues': [
                re.compile(r'(?:the\s+)?following(?:\s+\w+)?:', re.IGNORECASE),  # "The following:", "Following items:"
                re.compile(r'below\s+are:', re.IGNORECASE),  # "Below are:"
                re.compile(r'as\s+follows:', re.IGNORECASE),  # "As follows:"
                re.compile(r'namely:', re.IGNORECASE),  # "Namely:"
                re.compile(r'such\s+as:', re.IGNORECASE),  # "Such as:"
                re.compile(r'including:', re.IGNORECASE),  # "Including:"
                re.compile(r':\s*$'),  # Ends with colon
            ],
            
            # 47c. Structured content headings that should have point-form content
            'point_form_objectives': [
                re.compile(r'^(?:Learning\s+)?Objectives?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Goals?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Aims?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Purposes?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Learning\s+Outcomes?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_steps': [
                re.compile(r'^Steps?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Procedures?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Process[\s:]*$', re.IGNORECASE),
                re.compile(r'^Methods?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Algorithms?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Phases?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Instructions?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_key_points': [
                re.compile(r'^Key\s+Points?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Main\s+Ideas?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Important[\s:]*$', re.IGNORECASE),
                re.compile(r'^Critical[\s:]*$', re.IGNORECASE),
                re.compile(r'^Essential[\s:]*$', re.IGNORECASE),
                re.compile(r'^Takeaways?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_features': [
                re.compile(r'^Features?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Characteristics?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Properties?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Attributes?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Aspects?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_pros_cons': [
                re.compile(r'^Advantages?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Disadvantages?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Benefits?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Drawbacks?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Pros?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cons?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Strengths?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Weaknesses?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_examples': [
                re.compile(r'^Examples?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cases?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Instances?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Scenarios?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Illustrations?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_questions': [
                re.compile(r'^Questions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Problems?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Exercises?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Tasks?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Challenges?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_requirements': [
                re.compile(r'^Requirements?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Criteria[\s:]*$', re.IGNORECASE),
                re.compile(r'^Conditions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Prerequisites?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Necessities?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Materials?\s+(?:Needed|Required)[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_tips': [
                re.compile(r'^Tips?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Suggestions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Recommendations?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Advice[\s:]*$', re.IGNORECASE),
                re.compile(r'^Hints?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_warnings': [
                re.compile(r'^Warnings?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cautions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Avoid[\s:]*$', re.IGNORECASE),
                re.compile(r'^Do\s+Not[\s:]*$', re.IGNORECASE),
                re.compile(r'^Never[\s:]*$', re.IGNORECASE),
                re.compile(r'^Common\s+Mistakes?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_components': [
                re.compile(r'^Components?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Parts?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Elements?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Sections?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Modules?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_rules': [
                re.compile(r'^Rules?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Principles?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Laws?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Theorems?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Guidelines?[\s:]*$', re.IGNORECASE),
            ],
            
            # 47d. Serial comma detection - Content in sentence that should be points
            'point_form_serial_comma': [
                re.compile(r'^[^:]+:\s*[A-Z][^,]+,\s*[^,]+,\s*(?:and|or)\s*[^,.]+[\.:]?$'),  # X: item1, item2, and item3
                re.compile(r'^[^:]+:\s*[^;]+;\s*[^;]+;\s*[^;]+[;.]?$'),  # X: item1; item2; item3
            ],
            
            # 47e. Implicit list in sentence - First/Second/Third pattern
            'point_form_ordinal_sentence': [
                re.compile(r'(?:First|Firstly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Second|Secondly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Third|Thirdly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Fourth|Fourthly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Fifth|Finally|Lastly),?\s+[^.]+\.', re.IGNORECASE),
            ],
        }
    
    def get_cell_content_type(self, cell_text):
        """
        Determine the content type of a table cell for alignment purposes.
        Returns: 'numeric', 'percentage', 'statistical', 'text'
        """
        if not cell_text:
            return 'text'
        
        text = str(cell_text).strip()
        
        # Check for percentage first (most specific)
        for pattern in self.patterns.get('table_content_percentage', []):
            if pattern.match(text):
                return 'percentage'
        
        # Check for statistical notation
        for pattern in self.patterns.get('table_content_statistical', []):
            if pattern.search(text):
                return 'statistical'
        
        # Check for confidence intervals
        for pattern in self.patterns.get('table_content_ci', []):
            if pattern.search(text):
                return 'statistical'
        
        # Check for plain numeric
        for pattern in self.patterns.get('table_content_numeric', []):
            if pattern.match(text):
                return 'numeric'
        
        return 'text'
    
    def get_column_content_types(self, rows):
        """
        Analyze all rows to determine the predominant content type for each column.
        Returns list of content types, one per column.
        """
        if not rows or len(rows) == 0:
            return []
        
        num_cols = max(len(row) for row in rows)
        column_types = []
        
        for col_idx in range(num_cols):
            type_counts = {'numeric': 0, 'percentage': 0, 'statistical': 0, 'text': 0}
            
            # Skip header row (index 0), analyze data rows only
            for row_idx, row in enumerate(rows):
                if row_idx == 0:  # Skip header
                    continue
                if col_idx < len(row):
                    cell_type = self.get_cell_content_type(row[col_idx])
                    type_counts[cell_type] += 1
            
            # Determine predominant type (statistical and percentage take precedence)
            if type_counts['statistical'] > 0:
                column_types.append('statistical')
            elif type_counts['percentage'] > 0:
                column_types.append('percentage')
            elif type_counts['numeric'] > type_counts['text']:
                column_types.append('numeric')
            else:
                column_types.append('text')
        
        return column_types
    
    def get_alignment_for_content_type(self, content_type):
        """
        Get the appropriate alignment for a content type.
        Returns: WD_ALIGN_PARAGRAPH constant
        """
        alignment_map = {
            'text': WD_ALIGN_PARAGRAPH.LEFT,
            'numeric': WD_ALIGN_PARAGRAPH.RIGHT,
            'percentage': WD_ALIGN_PARAGRAPH.RIGHT,
            'statistical': WD_ALIGN_PARAGRAPH.CENTER,
            'header': WD_ALIGN_PARAGRAPH.CENTER,
        }
        return alignment_map.get(content_type, WD_ALIGN_PARAGRAPH.LEFT)

    def clean_heading_spaces(self, text):
        """
        Clean up spacing issues in heading text.
        Removes trailing spaces, normalizes internal spacing, fixes punctuation spacing.
        """
        if not text:
            return text
        
        original = text
        
        # Remove trailing whitespace
        text = text.rstrip()
        
        # Check if this is a heading line
        if not text.lstrip().startswith('#'):
            return text
        
        # Extract markdown prefix and heading content
        match = re.match(r'^(\s*#+\s*)', text)
        if not match:
            return text
        
        prefix = match.group(1)
        content = text[len(prefix):]
        
        # Clean the content
        # 1. Remove trailing spaces and punctuation that shouldn't be there
        content = content.rstrip(' .,:;')
        
        # 2. Normalize multiple spaces to single space
        content = re.sub(r'\s{2,}', ' ', content)
        
        # 3. Fix spacing around colons (for chapter headings like "CHAPTER ONE: INTRO")
        content = re.sub(r'\s+:', ':', content)
        content = re.sub(r':\s{2,}', ': ', content)
        
        # 4. Fix spacing around hyphens
        content = re.sub(r'\s+-\s+', ' - ', content)
        content = re.sub(r'\s{2,}-', ' -', content)
        content = re.sub(r'-\s{2,}', '- ', content)
        
        # Reconstruct the heading
        cleaned = prefix + content
        
        return cleaned
    
    def clean_document_spacing(self, text):
        """
        Clean up spacing issues throughout a document.
        Applies to all lines, with special handling for headings.
        """
        if not text:
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Check if it's a heading
            if line.lstrip().startswith('#'):
                cleaned_line = self.clean_heading_spaces(line)
            else:
                # For non-heading lines, just remove trailing spaces
                cleaned_line = line.rstrip()
                
                # Don't modify spacing in code blocks or special formatting
                if not cleaned_line.startswith('```') and not cleaned_line.startswith('    '):
                    # Fix space before punctuation (except in tables)
                    if not '|' in cleaned_line:
                        cleaned_line = re.sub(r'\s+,', ',', cleaned_line)
                        cleaned_line = re.sub(r'\s+\.(?!\d)', '.', cleaned_line)  # Preserve decimals
                        cleaned_line = re.sub(r'\s+;', ';', cleaned_line)
            
            cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    # ========== SHORT DOCUMENT FORMATTING METHODS ==========
    
    def is_short_document(self, text):
        """
        Determine if a document is 'short' (assignment, course material, exercise).
        Uses multiple criteria: word count, indicators, and structure.
        Returns: (is_short, reason)
        """
        if not text:
            return False, "Empty document"
        
        # Calculate document metrics
        word_count = len(text.split())
        line_count = len(text.split('\n'))
        char_count = len(text)
        estimated_pages = char_count / 2000  # ~2000 chars per page
        
        # Count headings/sections
        heading_count = len(re.findall(r'^#+\s+', text, re.MULTILINE))
        if heading_count == 0:
            # Count uppercase headings
            heading_count = len(re.findall(r'^[A-Z][A-Z\s]{5,50}$', text, re.MULTILINE))
        
        # Check for LONG document indicators (Dissertation, Thesis, etc.)
        for pattern in self.patterns.get('long_doc_indicators', []):
            if pattern.search(text):
                return False, "Contains long document indicators"
        
        # Check for SHORT document indicators (Assignment, Homework, etc.)
        has_short_indicator = False
        for pattern in self.patterns.get('short_doc_indicators', []):
            if pattern.search(text):
                has_short_indicator = True
                break
        
        # Multiple threshold checks
        is_short_by_words = word_count < 3000
        is_short_by_pages = estimated_pages < 10
        is_short_by_lines = line_count < 500
        is_short_by_sections = heading_count < 10
        
        # Decision logic
        if has_short_indicator:
            return True, "Contains short document indicator (Assignment/Homework/Exercise)"
        
        if is_short_by_words and is_short_by_pages:
            return True, f"Short by word count ({word_count} words) and pages ({estimated_pages:.1f} pages)"
        
        if is_short_by_lines and is_short_by_sections:
            return True, f"Short by line count ({line_count} lines) and sections ({heading_count} sections)"
        
        return False, "Document does not meet short document criteria"
    
    def is_toc_header_line(self, text):
        """Check if a line is a Table of Contents header."""
        if not text:
            return False
        
        clean_text = text.strip()
        for pattern in self.patterns.get('toc_header', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def is_toc_content_line(self, text):
        """Check if a line is a Table of Contents content line (entries with page numbers)."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        # Check for dot leaders pattern (common in TOC)
        if '.....' in clean_text or '…..' in clean_text or '…' in clean_text:
            return True
        
        for pattern in self.patterns.get('toc_content_line', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def remove_toc_from_lines(self, lines):
        """
        Remove Table of Contents section from document lines.
        Returns lines without TOC.
        """
        cleaned_lines = []
        in_toc = False
        toc_blank_count = 0
        
        for line in lines:
            text = line.strip() if isinstance(line, str) else line.get('text', '').strip()
            
            # Check if this line starts TOC
            if self.is_toc_header_line(text):
                in_toc = True
                toc_blank_count = 0
                continue  # Skip TOC header
            
            if in_toc:
                # Track blank lines to detect TOC end
                if not text:
                    toc_blank_count += 1
                    if toc_blank_count >= 2:
                        # Two consecutive blank lines - end of TOC
                        in_toc = False
                    continue
                
                # Check if still in TOC content
                if self.is_toc_content_line(text):
                    continue  # Skip TOC content
                
                # Check if this looks like a real heading (end of TOC)
                if re.match(r'^#+\s+', text) or (text.isupper() and len(text) > 5):
                    # Real content starting - end of TOC
                    in_toc = False
                    cleaned_lines.append(line)
                    continue
                
                # Check if this line has chapter/section-like content (end of TOC)
                if re.match(r'^\d+\.\s+[A-Z]', text) and not re.search(r'\.\.\.|…|\s+\d+\s*$', text):
                    # Looks like actual content, not TOC entry
                    in_toc = False
                    cleaned_lines.append(line)
                    continue
            else:
                cleaned_lines.append(line)
        
        return cleaned_lines
    
    def get_key_point_type(self, text):
        """
        Detect if a line is a key point that should be emphasized.
        Returns: (key_point_type, prefix_emoji) or (None, None)
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Check each key point category
        key_point_categories = [
            ('key_point_learning', 'learning', '📚 '),
            ('key_point_definitions', 'definition', ''),
            ('key_point_concepts', 'concept', '🔑 '),
            ('key_point_procedures', 'procedure', '📝 '),
            ('key_point_examples', 'example', '📋 '),
            ('key_point_warnings', 'warning', '⚠️ '),
            ('key_point_exercises', 'exercise', '💪 '),
            ('key_point_summary', 'summary', '📊 '),
        ]
        
        for pattern_key, point_type, emoji in key_point_categories:
            for pattern in self.patterns.get(pattern_key, []):
                if pattern.match(clean_text):
                    return point_type, emoji
        
        return None, None
    
    def is_assignment_header_field(self, text):
        """Check if a line is an assignment header field (Student Name, Course, etc.)."""
        if not text:
            return False
        
        clean_text = text.strip()
        for pattern in self.patterns.get('assignment_header', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def emphasize_key_point(self, text, point_type, emoji):
        """
        Apply emphasis formatting to a key point line.
        Returns the emphasized text.
        """
        if not text:
            return text
        
        clean_text = text.strip()
        
        # Already formatted (has ** markers)
        if clean_text.startswith('**') or clean_text.startswith('*'):
            return text
        
        # Apply formatting based on type
        if point_type in ['warning']:
            # Warnings get bold
            emphasized = f"**{clean_text}**"
        elif point_type in ['example']:
            # Examples get italic
            emphasized = f"*{clean_text}*"
        else:
            # Most key points get bold
            emphasized = f"**{clean_text}**"
        
        # Add emoji prefix if provided
        if emoji:
            emphasized = emoji + emphasized
        
        return emphasized
    
    # ========== POINT FORM CONTENT FORMATTING METHODS ==========
    
    def is_point_form_line(self, text):
        """
        Check if a line is already in point form (numbered, bulleted, checkbox).
        Returns: (is_point_form, point_type) - type is 'numbered', 'bulleted', 'checkbox', or None
        """
        if not text:
            return False, None
        
        clean_text = text.strip()
        
        # Check numbered patterns
        for pattern in self.patterns.get('point_form_numbered', []):
            if pattern.match(clean_text):
                return True, 'numbered'
        
        # Check bulleted patterns
        for pattern in self.patterns.get('point_form_bulleted', []):
            if pattern.match(clean_text):
                return True, 'bulleted'
        
        # Check checkbox patterns
        for pattern in self.patterns.get('point_form_checkbox', []):
            if pattern.match(clean_text):
                return True, 'checkbox'
        
        return False, None
    
    def get_point_form_heading_type(self, text):
        """
        Check if a line is a heading that should have point-form content following it.
        Returns: (heading_type, format_type) or (None, None)
        - heading_type: objectives, steps, key_points, features, pros_cons, examples, etc.
        - format_type: 'numbered' or 'bulleted'
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Mapping of pattern keys to (heading_type, format_type)
        heading_patterns = [
            ('point_form_objectives', 'objectives', 'bulleted'),
            ('point_form_steps', 'steps', 'numbered'),
            ('point_form_key_points', 'key_points', 'bulleted'),
            ('point_form_features', 'features', 'bulleted'),
            ('point_form_pros_cons', 'pros_cons', 'bulleted'),
            ('point_form_examples', 'examples', 'bulleted'),
            ('point_form_questions', 'questions', 'numbered'),
            ('point_form_requirements', 'requirements', 'bulleted'),
            ('point_form_tips', 'tips', 'bulleted'),
            ('point_form_warnings', 'warnings', 'bulleted'),
            ('point_form_components', 'components', 'bulleted'),
            ('point_form_rules', 'rules', 'numbered'),
        ]
        
        for pattern_key, heading_type, format_type in heading_patterns:
            for pattern in self.patterns.get(pattern_key, []):
                if pattern.match(clean_text):
                    return heading_type, format_type
        
        return None, None
    
    def has_list_context_clue(self, text):
        """
        Check if a line ends with context that suggests a list follows.
        Returns True if line has context clue (e.g., ends with colon after "following").
        """
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('point_form_context_clues', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def extract_serial_comma_items(self, text):
        """
        Extract items from a sentence with serial commas.
        Example: "Tools needed: hammer, nails, and screwdriver" -> ["hammer", "nails", "screwdriver"]
        Returns: (heading, items_list) or (None, None)
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Check for serial comma pattern with colon
        for pattern in self.patterns.get('point_form_serial_comma', []):
            if pattern.match(clean_text):
                # Split on colon
                if ':' in clean_text:
                    parts = clean_text.split(':', 1)
                    heading = parts[0].strip()
                    content = parts[1].strip()
                    
                    # Try to split by comma or semicolon
                    if ';' in content:
                        items = [item.strip() for item in content.split(';')]
                    else:
                        # Handle "and" or "or" at the end
                        content = re.sub(r',?\s+(?:and|or)\s+', ', ', content)
                        items = [item.strip().rstrip('.') for item in content.split(',')]
                    
                    # Clean up items
                    items = [item for item in items if item and len(item) > 1]
                    
                    if len(items) >= 2:
                        return heading, items
        
        return None, None
    
    def extract_ordinal_steps(self, text):
        """
        Extract steps from a sentence with ordinal words (First, Second, Third...).
        Returns: list of steps or None
        """
        if not text:
            return None
        
        # Pattern to find ordinal sentences
        ordinal_pattern = re.compile(
            r'(?:First(?:ly)?|Second(?:ly)?|Third(?:ly)?|Fourth(?:ly)?|Fifth(?:ly)?|Finally|Lastly),?\s+([^.]+)\.',
            re.IGNORECASE
        )
        
        matches = ordinal_pattern.findall(text)
        if len(matches) >= 2:
            return [match.strip() for match in matches]
        
        return None
    
    def clean_point_content(self, text):
        """
        Clean a point by removing existing markers and fixing formatting.
        """
        if not text:
            return text
        
        clean = text.strip()
        
        # Remove common list markers
        marker_patterns = [
            r'^\s*\d+[\.\)]\s*',      # 1. or 1)
            r'^\s*[a-z][\.\)]\s*',    # a. or a)
            r'^\s*[ivxIVX]+[\.\)]\s*',  # Roman numerals
            r'^\s*\(\d+\)\s*',        # (1)
            r'^\s*\([a-z]\)\s*',      # (a)
            r'^\s*[A-Z][\.\)]\s*',    # A. or A)
            r'^\s*[\-\*\•◦→○▪▫■□►▸◆◇]\s*',  # Bullet symbols
            r'^\s*[—–]\s*',           # Dashes
            r'^\s*□\s*',              # Checkbox empty
            r'^\s*[☐☑✓✗]\s*',        # Checkbox symbols
            r'^\s*\[\s*[xX]?\s*\]\s*',  # [ ] or [x]
        ]
        
        for pattern in marker_patterns:
            clean = re.sub(pattern, '', clean)
        
        # Capitalize first letter
        if clean and clean[0].islower():
            clean = clean[0].upper() + clean[1:]
        
        return clean.strip()
    
    def format_as_numbered_list(self, items):
        """
        Format items as a numbered list.
        """
        formatted = []
        for i, item in enumerate(items, 1):
            clean_item = self.clean_point_content(item)
            if clean_item:
                formatted.append(f"{i}. {clean_item}")
        return formatted
    
    def format_as_bulleted_list(self, items):
        """
        Format items as a bulleted list using dash.
        """
        formatted = []
        for item in items:
            clean_item = self.clean_point_content(item)
            if clean_item:
                formatted.append(f"- {clean_item}")
        return formatted
    
    def standardize_existing_list(self, lines, start_idx, format_type):
        """
        Standardize an existing list to consistent formatting.
        Returns: (formatted_lines, end_idx)
        """
        formatted_lines = []
        i = start_idx
        point_number = 1
        
        while i < len(lines):
            line = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
            line_text = line.strip()
            
            # Check if this line is a point-form line
            is_point, point_type = self.is_point_form_line(line_text)
            
            if is_point:
                clean_content = self.clean_point_content(line_text)
                if clean_content:
                    if format_type == 'numbered':
                        formatted_lines.append(f"{point_number}. {clean_content}")
                        point_number += 1
                    else:
                        formatted_lines.append(f"- {clean_content}")
                i += 1
            elif not line_text:
                # Empty line - might be end of list
                # Look ahead to see if list continues
                if i + 1 < len(lines):
                    next_line = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                    next_is_point, _ = self.is_point_form_line(next_line.strip())
                    if next_is_point:
                        # List continues after blank
                        formatted_lines.append('')
                        i += 1
                        continue
                # End of list
                break
            else:
                # Non-point line - end of list
                break
        
        return formatted_lines, i
    
    def process_point_form_content(self, text):
        """
        Process document text to detect and format point-form content.
        Returns processed text with properly formatted lists.
        """
        if not text:
            return text
        
        lines = text.split('\n')
        
        # NEW: Detect implied bullet blocks and convert them to explicit bullets
        try:
            # Use the implied detector to find blocks
            implied_blocks = self.implied_detector.detect_implied_bullet_blocks(lines)
            
            # Apply bullets to detected blocks
            for start, end, bullet_type in implied_blocks:
                for i in range(start, end + 1):
                    line = lines[i]
                    # Skip if already a bullet or empty
                    if not line.strip() or self.is_point_form_line(line)[0]:
                        continue
                        
                    # Add bullet marker
                    lines[i] = f"■ {line.strip()}"
            
            # Update text with implied bullets applied
            # We re-join and re-split to ensure consistency
            text = '\n'.join(lines)
            lines = text.split('\n')
            
        except Exception as e:
            logger.error(f"Error in implied bullet detection: {e}")
            # Continue with original text if detection fails
            lines = text.split('\n')

        processed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i]
            line_text = line.strip()
            
            # Skip empty lines
            if not line_text:
                processed_lines.append(line)
                i += 1
                continue
            
            # Check 1: Is this a structured content heading (Objectives:, Steps:, etc.)?
            heading_type, format_type = self.get_point_form_heading_type(line_text)
            if heading_type:
                # Add heading with bold formatting
                heading_clean = line_text.rstrip(':')
                if not heading_clean.startswith('**'):
                    processed_lines.append(f"**{heading_clean}:**")
                else:
                    processed_lines.append(line)
                i += 1
                
                # Process following lines as point-form content
                point_lines = []
                while i < len(lines):
                    next_line = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
                    next_text = next_line.strip()
                    
                    if not next_text:
                        # Blank line might end list
                        if i + 1 < len(lines):
                            following = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                            is_point, _ = self.is_point_form_line(following.strip())
                            if is_point or self.could_be_list_item(following.strip()):
                                point_lines.append('')
                                i += 1
                                continue
                        break
                    
                    # Check if it's already a point-form line or could be
                    is_point, _ = self.is_point_form_line(next_text)
                    could_be_item = self.could_be_list_item(next_text)
                    
                    if is_point or could_be_item:
                        point_lines.append(next_text)
                        i += 1
                    else:
                        # End of list
                        break
                
                # Format the collected points
                if point_lines:
                    formatted = self.format_point_block(point_lines, format_type)
                    processed_lines.extend(formatted)
                continue
            
            # Check 2: Does this line have context clue and content to convert?
            heading, items = self.extract_serial_comma_items(line_text)
            if heading and items:
                # Convert to list format
                processed_lines.append(f"**{heading}:**")
                formatted_items = self.format_as_bulleted_list(items)
                processed_lines.extend(formatted_items)
                i += 1
                continue
            
            # Check 3: Does this line have ordinal steps (First, Second, Third...)?
            steps = self.extract_ordinal_steps(line_text)
            if steps:
                # Check if there's a heading phrase before "First"
                first_match = re.search(r'^(.+?)(?:First(?:ly)?)', line_text, re.IGNORECASE)
                if first_match and len(first_match.group(1).strip()) > 3:
                    heading = first_match.group(1).strip().rstrip(':,.')
                    processed_lines.append(f"**{heading}:**")
                formatted_steps = self.format_as_numbered_list(steps)
                processed_lines.extend(formatted_steps)
                i += 1
                continue
            
            # Check 4: Is this an existing point-form line that needs standardization?
            is_point, point_type = self.is_point_form_line(line_text)
            if is_point:
                # Standardize the list starting from here
                format_type = 'numbered' if point_type == 'numbered' else 'bulleted'
                formatted_block, new_idx = self.standardize_existing_list(lines, i, format_type)
                processed_lines.extend(formatted_block)
                i = new_idx
                continue
            
            # Check 5: Does this line have a context clue (ends with colon)?
            if self.has_list_context_clue(line_text):
                # Check if next lines could be list items
                if i + 1 < len(lines):
                    next_line = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                    next_text = next_line.strip()
                    is_next_point, _ = self.is_point_form_line(next_text)
                    could_be_item = self.could_be_list_item(next_text)
                    
                    if is_next_point or could_be_item:
                        # Format heading
                        heading = line_text.rstrip(':')
                        if not heading.startswith('**'):
                            processed_lines.append(f"**{heading}:**")
                        else:
                            processed_lines.append(line)
                        i += 1
                        
                        # Collect and format following points
                        point_lines = []
                        while i < len(lines):
                            l = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
                            l_text = l.strip()
                            
                            if not l_text:
                                if i + 1 < len(lines):
                                    following = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                                    is_pt, _ = self.is_point_form_line(following.strip())
                                    if is_pt or self.could_be_list_item(following.strip()):
                                        point_lines.append('')
                                        i += 1
                                        continue
                                break
                            
                            is_pt, _ = self.is_point_form_line(l_text)
                            if is_pt or self.could_be_list_item(l_text):
                                point_lines.append(l_text)
                                i += 1
                            else:
                                break
                        
                        if point_lines:
                            formatted = self.format_point_block(point_lines, 'bulleted')
                            processed_lines.extend(formatted)
                        continue
            
            # Default: keep line as-is
            processed_lines.append(line)
            i += 1
        
        return '\n'.join(processed_lines)
    
    def could_be_list_item(self, text):
        """
        Check if text could be a list item (short, starts with capital, no complex structure).
        """
        if not text:
            return False
        
        clean = text.strip()
        
        # Too long for a list item
        if len(clean) > 150:
            return False
        
        # Too short
        if len(clean) < 3:
            return False
        
        # Contains multiple sentences (not a simple list item)
        sentence_count = len(re.findall(r'[.!?]\s+[A-Z]', clean))
        if sentence_count > 0:
            return False
        
        # Starts with capital letter or common list patterns
        if re.match(r'^[A-Z]', clean):
            return True
        
        return False
    
    def format_point_block(self, lines, format_type):
        """
        Format a block of lines as either numbered or bulleted list.
        """
        formatted = []
        point_number = 1
        
        for line in lines:
            if not line or not line.strip():
                formatted.append('')
                continue
            
            clean_content = self.clean_point_content(line)
            if clean_content:
                if format_type == 'numbered':
                    formatted.append(f"{point_number}. {clean_content}")
                    point_number += 1
                else:
                    formatted.append(f"- {clean_content}")
        
        return formatted

    def process_short_document(self, text):
        """
        Process a short document: remove TOC, emphasize key points, and format point-form content.
        Returns processed text.
        """
        if not text:
            return text
        
        # Check if document is short
        is_short, reason = self.is_short_document(text)
        if not is_short:
            return text  # No changes for long documents
        
        lines = text.split('\n')
        
        # Step 1: Remove Table of Contents
        lines = self.remove_toc_from_lines(lines)
        
        # Convert lines back to text for point-form processing
        text = '\n'.join([line if isinstance(line, str) else line.get('text', '') for line in lines])
        
        # Step 2: Process point-form content (convert serial lists to bullet points, standardize lists)
        text = self.process_point_form_content(text)
        
        # Step 3: Process each line for key point emphasis
        lines = text.split('\n')
        processed_lines = []
        for line in lines:
            line_text = line if isinstance(line, str) else line.get('text', '')
            
            # Skip already formatted point-form lines (starting with - or number.)
            if re.match(r'^\s*[-\*•]\s+', line_text) or re.match(r'^\s*\d+\.\s+', line_text):
                processed_lines.append(line_text)
                continue
            
            # Check if it's an assignment header field
            if self.is_assignment_header_field(line_text):
                # Bold the field
                if not line_text.strip().startswith('**'):
                    processed_lines.append(f"**{line_text.strip()}**")
                else:
                    processed_lines.append(line_text)
                continue
            
            # Check if it's a key point
            point_type, emoji = self.get_key_point_type(line_text)
            if point_type:
                emphasized = self.emphasize_key_point(line_text, point_type, emoji)
                processed_lines.append(emphasized)
            else:
                processed_lines.append(line_text)
        
        return '\n'.join(processed_lines)
    
    def is_main_heading(self, text):
        """
        Check if text is a main heading (level 1) that should have space cleaning applied.
        """
        if not text:
            return False
        
        # Must start with single # (level 1)
        if not re.match(r'^\s*#\s+', text) or re.match(r'^\s*##', text):
            return False
        
        # Extract heading content
        clean_text = re.sub(r'^#+\s*', '', text).strip().upper()
        clean_text = clean_text.rstrip(' .:,;')
        
        # Main heading keywords
        main_headings = [
            'DEDICATION', 'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'ABSTRACT', 'RESUME', 'RÉSUMÉ', 'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES', 'LIST OF FIGURES', 'LIST OF ABBREVIATIONS', 'LIST OF ACRONYMS',
            'GLOSSARY', 'APPENDICES', 'APPENDIX',
            'DECLARATION', 'CERTIFICATION', 'APPROVAL PAGE', 'COMMITTEE APPROVAL',
            'REFERENCES', 'REFERENCE', 'BIBLIOGRAPHY',
        ]
        
        # Check exact match
        for heading in main_headings:
            if clean_text == heading:
                return True
        
        # Check for chapter headings
        if re.match(r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)', clean_text, re.IGNORECASE):
            return True
        
        return False

    def should_start_on_new_page(self, text):
        """
        Check if a heading should start on a new page.
        Returns True for major academic sections like chapters, abstract, etc.
        """
        if not text:
            return False
            
        # Strip markdown heading markers and whitespace
        clean_text = re.sub(r'^#+\s*', '', text).strip().upper()
        
        # Front matter sections that require new pages
        front_matter_sections = [
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'DEDICATION',
            'ABSTRACT', 'RESUME', 'RÉSUMÉ',
            'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES',
            'LIST OF FIGURES',
            'GLOSSARY',
            'LIST OF ABBREVIATIONS', 'ABBREVIATIONS', 'LIST OF ACRONYMS',
            'APPENDICES', 'APPENDIX',
            'REFERENCES', 'REFERENCE',
            'BIBLIOGRAPHY',
        ]
        
        # Check for exact front matter match
        for section in front_matter_sections:
            if clean_text == section or clean_text.startswith(section + ':'):
                return True
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        chapter_pattern = r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)'
        if re.match(chapter_pattern, clean_text, re.IGNORECASE):
            return True
        
        return False
    
    def should_be_centered(self, text, heading_level=1):
        """
        Determine if a heading should be centered.
        Only level 1 headings (#) can be centered.
        Returns True for major academic sections that should be centered.
        """
        if not text:
            return False
        
        # Only top-level headings (level 1) can be centered
        if heading_level != 1:
            return False
        
        # Strip markdown heading markers and whitespace
        clean_text = re.sub(r'^#+\s*', '', text).strip().upper()
        
        # Remove any trailing colon or punctuation for matching
        match_text = clean_text.rstrip(':.')
        
        # Front matter sections to center
        front_matter_center = [
            'DEDICATION',
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'ABSTRACT',
            'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES',
            'LIST OF FIGURES',
            'LIST OF ABBREVIATIONS', 'ABBREVIATIONS',
            'GLOSSARY',
            'APPENDICES', 'APPENDIX',
        ]
        
        # Back matter sections to center
        back_matter_center = [
            'REFERENCES', 'REFERENCE',
            'BIBLIOGRAPHY',
        ]
        
        # Check for exact front/back matter match
        all_center_sections = front_matter_center + back_matter_center
        for section in all_center_sections:
            if match_text == section:
                return True
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        # These should be centered regardless of having a subtitle
        if clean_text.startswith('CHAPTER'):
            chapter_pattern = r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)'
            if re.match(chapter_pattern, clean_text, re.IGNORECASE):
                return True
        
        return False
    
    def is_chapter_heading(self, text):
        """
        Check if text is a chapter heading (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        Returns (is_chapter, chapter_num, chapter_title) tuple
        """
        if not text:
            return False, None, None
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Chapter with title on same line: CHAPTER ONE: INTRODUCTION
        pattern_with_title = r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)\s*[:\-\.]\s*(.+)$'
        match = re.match(pattern_with_title, clean_text, re.IGNORECASE)
        if match:
            return True, match.group(1), match.group(2).strip()
        
        # Chapter heading only: CHAPTER ONE
        pattern_only = r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)\s*$'
        match = re.match(pattern_only, clean_text, re.IGNORECASE)
        if match:
            return True, match.group(1), None
        
        return False, None, None
    
    def is_chapter_title(self, text, prev_was_chapter=False):
        """
        Check if text is a chapter title following a chapter heading.
        Chapter titles are typically: ALL CAPS, or common academic section names
        Returns True if it looks like a chapter title.
        """
        if not text:
            return False
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Skip if too long (likely a paragraph)
        if len(clean_text) > 100:
            return False
        
        # Skip if empty
        if not clean_text:
            return False
        
        # Common chapter title keywords
        chapter_title_keywords = [
            'INTRODUCTION', 'LITERATURE REVIEW', 'METHODOLOGY', 'METHODS',
            'RESULTS', 'DISCUSSION', 'CONCLUSION', 'CONCLUSIONS',
            'RECOMMENDATIONS', 'THEORETICAL FRAMEWORK', 'RESEARCH METHODOLOGY',
            'SUMMARY AND CONCLUSIONS', 'SUMMARY OF FINDINGS', 'DATA ANALYSIS',
            'BACKGROUND', 'PROBLEM STATEMENT', 'RESEARCH DESIGN',
            'FINDINGS AND DISCUSSION', 'ANALYSIS AND INTERPRETATION',
        ]
        
        # Check for exact match with common titles
        if clean_text.upper() in chapter_title_keywords:
            return True
        
        # Check if all caps (common for chapter titles)
        if clean_text == clean_text.upper() and len(clean_text) > 5 and any(c.isalpha() for c in clean_text):
            # All caps title, likely a chapter title
            return True
        
        # If previous line was a chapter heading, be more lenient
        if prev_was_chapter:
            # Title case or heading-like
            if re.match(r'^[A-Z][A-Za-z\s]+$', clean_text):
                return True
        
        return False
    
    def is_copyright_content(self, text):
        """Check if text is copyright-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('copyright_content', []):
            if pattern.match(clean_text):
                return True
        
        return False
    
    def is_declaration_content(self, text):
        """Check if text is declaration-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('declaration_content', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def is_certification_content(self, text):
        """Check if text is certification/approval-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('certification_content', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def is_signature_line(self, text):
        """Check if text is a signature line or date line."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('signature_line', []):
            if pattern.match(clean_text):
                return True
        
        return False
    
    def is_toc_entry(self, text):
        """Check if text is a Table of Contents entry line."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('toc_entry', []):
            if pattern.match(clean_text):
                return True
        
        # Also check for dot leaders pattern
        if '.....' in clean_text or '…..' in clean_text:
            return True
        
        return False

    def get_front_matter_section_type(self, text):
        """
        Get the type of front matter section if text is a front matter heading.
        Returns section type (declaration, certification, dedication, etc.) or None.
        """
        if not text:
            return None
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip().upper()
        
        front_matter_sections = {
            'DECLARATION': 'declaration',
            'CERTIFICATION': 'certification',
            'APPROVAL PAGE': 'certification',
            'COMMITTEE APPROVAL': 'certification',
            'DEDICATION': 'dedication',
            'ACKNOWLEDGEMENTS': 'acknowledgements',
            'ACKNOWLEDGMENTS': 'acknowledgements',
            'ACKNOWLEDGEMENT': 'acknowledgements',
            'ACKNOWLEDGMENT': 'acknowledgements', # Missing 'e' variation
            'ABSTRACT': 'abstract',
            'RESUME': 'resume',  # French equivalent of abstract
            'RÉSUMÉ': 'resume',  # With accent
            'RÉSUME': 'resume',  # With first accent only
            'RESUMÉ': 'resume',  # With last accent only
            'TABLE OF CONTENTS': 'toc',
            'CONTENTS': 'toc',
            'LIST OF TABLES': 'list_of_tables',
            'LIST OF FIGURES': 'list_of_figures',
            'LIST OF ABBREVIATIONS': 'abbreviations',
            'ABBREVIATIONS': 'abbreviations',
            'LIST OF ACRONYMS': 'abbreviations',
            'GLOSSARY': 'glossary',
            'APPENDIX': 'appendix',
            'APPENDICES': 'appendix',
            'REFERENCES': 'references',
            'BIBLIOGRAPHY': 'bibliography',
        }
        
        for section, section_type in front_matter_sections.items():
            if clean_text == section:
                return section_type
        
        return None

    def analyze_line(self, line, line_num, prev_line='', next_line='', context=None):
        """Analyze a single line with multiple pattern checks"""
        # FIRST: Clean heading spaces before analysis
        if line.lstrip().startswith('#'):
            line = self.clean_heading_spaces(line)
        
        trimmed = line.strip()
        
        if not trimmed:
            return {'type': 'empty', 'content': '', 'line_num': line_num}
        
        analysis = {
            'line_num': line_num,
            'type': 'paragraph',
            'content': trimmed,
            'original': line,
            'level': 0,
            'confidence': 0.0,
        }
        
        # Get line characteristics
        length = len(trimmed)
        is_short = length < 100
        is_very_short = length < 60
        is_all_caps = trimmed == trimmed.upper() and any(c.isalpha() for c in trimmed) and length > 2
        is_title_case = re.match(r'^[A-Z][a-z]+(?:\s+[A-Za-z][a-z]*)*$', trimmed) is not None
        has_period = trimmed.endswith('.')
        word_count = len(trimmed.split())
        
        # Skip very long lines for heading detection (likely paragraphs)
        # But check for paragraph with colon that might be a definition
        
        # Priority 1: Check for table patterns (highest priority to preserve structure)
        for pattern in self.patterns['table_marker']:
            if pattern.match(trimmed):
                if re.search(r'START', trimmed, re.IGNORECASE):
                    analysis['type'] = 'table_start'
                elif re.search(r'END', trimmed, re.IGNORECASE):
                    analysis['type'] = 'table_end'
                else:
                    analysis['type'] = 'table_caption'
                analysis['confidence'] = 1.0
                return analysis
        
        for pattern in self.patterns['table_row']:
            if pattern.match(trimmed):
                # Check if it's a separator row
                if re.match(r'^\|[\s\-:]+\|', trimmed):
                    analysis['type'] = 'table_separator'
                else:
                    analysis['type'] = 'table_row'
                cells = [c.strip() for c in trimmed.split('|') if c.strip()]
                analysis['cells'] = cells
                analysis['confidence'] = 1.0
                return analysis
        
        # Priority 1.5: DISSERTATION-SPECIFIC PATTERNS
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I)
        is_chapter, chapter_num, chapter_title = self.is_chapter_heading(trimmed)
        if is_chapter:
            analysis['type'] = 'chapter_heading'
            analysis['level'] = 1
            analysis['chapter_num'] = chapter_num
            analysis['chapter_title'] = chapter_title  # May be None if title is on separate line
            analysis['confidence'] = 1.0
            analysis['needs_page_break'] = True
            analysis['should_center'] = True
            return analysis
        
        # Check for chapter title following a chapter heading (using context)
        if context and context.get('prev_was_chapter'):
            if self.is_chapter_title(trimmed, prev_was_chapter=True):
                analysis['type'] = 'chapter_title'
                analysis['level'] = 2  # Slightly below chapter heading
                analysis['confidence'] = 0.95
                analysis['should_center'] = True  # Chapter titles are centered
                analysis['needs_page_break'] = False  # Already on new page with chapter heading
                return analysis
        
        # Check for front matter section headings (Declaration, Certification, etc.)
        front_matter_type = self.get_front_matter_section_type(trimmed)
        if front_matter_type:
            analysis['type'] = 'front_matter_heading'
            analysis['front_matter_type'] = front_matter_type
            analysis['level'] = 1
            analysis['confidence'] = 1.0
            analysis['needs_page_break'] = True
            analysis['should_center'] = True
            return analysis
        
        # Check for copyright content
        if self.is_copyright_content(trimmed):
            analysis['type'] = 'copyright_content'
            analysis['confidence'] = 0.95
            analysis['should_center'] = True  # Copyright content is often centered
            return analysis
        
        # Check for signature lines
        if self.is_signature_line(trimmed):
            analysis['type'] = 'signature_line'
            analysis['confidence'] = 0.95
            return analysis
        
        # Check for TOC entries
        if self.is_toc_entry(trimmed):
            analysis['type'] = 'toc_entry'
            analysis['confidence'] = 0.90
            return analysis

        # Priority 2: Check for heading patterns
        if is_short and not has_period:
            # H1 detection - major sections
            for pattern in self.patterns['heading_1']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 1
                    analysis['confidence'] = 0.95
                    # Check if this heading needs a page break and/or centering
                    analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                    analysis['should_center'] = self.should_be_centered(trimmed, 1)
                    return analysis
            
            # H2 detection - sub-sections
            for pattern in self.patterns['heading_2']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 2
                    analysis['confidence'] = 0.90
                    analysis['needs_page_break'] = False  # Sub-sections don't get page breaks
                    analysis['should_center'] = False  # Sub-sections don't get centered
                    return analysis
            
            # H3 detection - sub-sub-sections
            for pattern in self.patterns['heading_3']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 3
                    analysis['confidence'] = 0.85
                    analysis['needs_page_break'] = False  # Sub-sub-sections don't get page breaks
                    analysis['should_center'] = False  # Sub-sub-sections don't get centered
                    return analysis
            
            # Heuristic heading detection for ALL CAPS
            if is_all_caps and is_very_short and word_count <= 6:
                analysis['type'] = 'heading'
                analysis['level'] = 1
                analysis['confidence'] = 0.80
                # Check if this heading needs a page break and/or centering
                analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                analysis['should_center'] = self.should_be_centered(trimmed, 1)
                return analysis
            
            # Heuristic heading detection for Title Case (no ending punctuation, short)
            if is_title_case and is_very_short and word_count >= 2 and word_count <= 8:
                analysis['type'] = 'heading'
                analysis['level'] = 2
                analysis['confidence'] = 0.75
                analysis['needs_page_break'] = False  # Title case headings don't get page breaks
                analysis['should_center'] = False  # Title case headings don't get centered
                return analysis
        
        # Priority 3: Check for reference patterns
        for pattern in self.patterns['reference']:
            if pattern.match(trimmed):
                analysis['type'] = 'reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 4: Check for list patterns
        # Use the enhanced bullet detection logic
        bullet_info = detect_bullet_type(line)  # Use original line to preserve indentation
        if bullet_info:
            analysis['type'] = 'bullet_list'
            analysis['content'] = bullet_info['content']
            analysis['bullet_info'] = bullet_info  # Store full info for WordGenerator
            analysis['confidence'] = 0.95
            return analysis

        for pattern in self.patterns['bullet_list']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'bullet_list'
                analysis['content'] = match.group(1) if match.lastindex else trimmed.lstrip('•●○▪▫■□◆◇-–—* →➤➢').strip()
                analysis['confidence'] = 0.95
                return analysis
        
        for pattern in self.patterns['numbered_list']:
            if pattern.match(trimmed):
                analysis['type'] = 'numbered_list'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 5: Check for definition patterns
        for pattern in self.patterns['definition']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'definition'
                analysis['term'] = match.group(1)
                analysis['definition'] = match.group(2) if match.lastindex > 1 and match.group(2) else ''
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 6: Check for figure captions
        for pattern in self.patterns['figure']:
            if pattern.match(trimmed):
                analysis['type'] = 'figure'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 7: Check for equation labels
        for pattern in self.patterns['equation']:
            if pattern.match(trimmed):
                analysis['type'] = 'equation'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 8: Check for quotes
        for pattern in self.patterns['quote']:
            if pattern.match(trimmed):
                analysis['type'] = 'quote'
                analysis['confidence'] = 0.85
                return analysis
        
        # Priority 9: Check for code blocks
        for pattern in self.patterns['code']:
            if pattern.match(trimmed):
                analysis['type'] = 'code'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 10: Check for page metadata (headers/footers/page numbers)
        for pattern in self.patterns['page_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'page_metadata'
                analysis['confidence'] = 0.90
                # Determine subtype
                if re.search(r'page|p\.|pg\.', trimmed, re.IGNORECASE) or re.match(r'^\s*-?\s*\d+\s*-?\s*$', trimmed):
                    analysis['subtype'] = 'page_number'
                elif re.search(r'header|running head', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'header'
                elif re.search(r'footer', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'footer'
                else:
                    analysis['subtype'] = 'document_metadata'
                return analysis
        
        # Priority 11: Check for academic metadata
        for pattern in self.patterns['academic_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'academic_metadata'
                analysis['confidence'] = 0.80
                # Determine subtype
                if re.search(r'\bby\b|authors?:', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'author'
                elif re.search(r'@', trimmed):
                    analysis['subtype'] = 'contact'
                elif re.search(r'department|school|college|university|institute', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'affiliation'
                else:
                    analysis['subtype'] = 'metadata'
                return analysis
        
        # Priority 12: Check for mathematical expressions
        for pattern in self.patterns['math_expression']:
            if pattern.search(trimmed):
                # Avoid false positives with currency
                if re.match(r'^\$\d+', trimmed) and not re.search(r'\$[^$]+\$', trimmed):
                    continue  # This is likely currency, not math
                analysis['type'] = 'math_expression'
                # Determine subtype
                if trimmed.startswith('$$') and trimmed.endswith('$$'):
                    analysis['subtype'] = 'display_math'
                    analysis['confidence'] = 0.95
                elif '$' in trimmed or '\\(' in trimmed or '\\[' in trimmed:
                    analysis['subtype'] = 'inline_math'
                    analysis['confidence'] = 0.85
                else:
                    analysis['subtype'] = 'equation'
                    analysis['confidence'] = 0.75
                return analysis
        
        # Priority 13: Check for footnotes/endnotes
        for pattern in self.patterns['footnote_endnote']:
            if pattern.match(trimmed):
                analysis['type'] = 'footnote_endnote'
                # Determine subtype
                if re.match(r'^\s*(?:endnotes?|footnotes?)\s*$', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'section_header'
                    analysis['confidence'] = 0.95
                else:
                    analysis['subtype'] = 'footnote_entry'
                    analysis['confidence'] = 0.90
                return analysis
        
        # Priority 14: Check for inline formatting (bold/italic)
        # Check for markdown-style formatting but exclude lines that start with list markers
        if not re.match(r'^[\*\-•]\s', trimmed):  # Not a bullet list
            for pattern in self.patterns['inline_formatting']:
                matches = pattern.findall(trimmed)
                if matches and any(m for m in matches if any(g for g in (m if isinstance(m, tuple) else (m,)) if g)):
                    analysis['type'] = 'inline_formatting'
                    analysis['content'] = trimmed
                    # Determine formatting type
                    if '***' in trimmed or '___' in trimmed:
                        analysis['formatting'] = {'bold_italic': True, 'bold': False, 'italic': False}
                        analysis['confidence'] = 0.90
                    elif '**' in trimmed or '__' in trimmed:
                        analysis['formatting'] = {'bold': True, 'italic': False, 'bold_italic': False}
                        analysis['confidence'] = 0.85
                    else:
                        analysis['formatting'] = {'italic': True, 'bold': False, 'bold_italic': False}
                        analysis['confidence'] = 0.85
                    return analysis
        
        # ============================================================
        # NEW PATTERN DETECTION - December 30, 2025 (20 Academic Patterns)
        # ============================================================
        
        # Priority 15: Check for markdown heading hierarchy
        for pattern in self.patterns['heading_hierarchy']:
            if pattern.match(trimmed):
                analysis['type'] = 'heading_hierarchy'
                # Determine level by counting # symbols
                hash_count = len(trimmed) - len(trimmed.lstrip('#'))
                analysis['level'] = min(hash_count, 6)
                analysis['confidence'] = 0.95
                analysis['content'] = trimmed.lstrip('#').strip()
                # Check if this heading needs a page break and/or centering (only for level 1 headings)
                if hash_count == 1:
                    analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                    analysis['should_center'] = self.should_be_centered(trimmed, 1)
                else:
                    analysis['needs_page_break'] = False
                    analysis['should_center'] = False
                return analysis
        
        # Priority 16: Check for academic table patterns
        for pattern in self.patterns['academic_table']:
            if pattern.match(trimmed):
                analysis['type'] = 'academic_table'
                if re.match(r'^Table\s+\d+', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'caption'
                elif re.match(r'^\|[-\s:]+\|', trimmed):
                    analysis['subtype'] = 'separator'
                elif '**' in trimmed:
                    analysis['subtype'] = 'header_row'
                else:
                    analysis['subtype'] = 'data_row'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 17: Check for nested list patterns
        for pattern in self.patterns['list_nested']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'list_nested'
                # Calculate indent level (2 spaces per level)
                leading_spaces = len(trimmed) - len(trimmed.lstrip())
                analysis['indent_level'] = leading_spaces // 2
                if '□' in trimmed or '☐' in trimmed or '☑' in trimmed or '✓' in trimmed:
                    analysis['subtype'] = 'checkbox'
                else:
                    analysis['subtype'] = 'nested_item'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 18: Check for figure/equation patterns
        for pattern in self.patterns['figure_equation']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'figure_equation'
                if re.match(r'^[Ff]igure', trimmed):
                    analysis['subtype'] = 'figure_caption'
                elif '$$' in trimmed or 'equation' in trimmed.lower():
                    analysis['subtype'] = 'equation_block'
                else:
                    analysis['subtype'] = 'math_content'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 19: Check for inline citations
        for pattern in self.patterns['citation_inline']:
            if pattern.search(trimmed):
                analysis['type'] = 'citation_inline'
                # Count citations in line
                citations = pattern.findall(trimmed)
                analysis['citation_count'] = len(citations)
                analysis['confidence'] = 0.85
                # Don't return - this is inline, keep processing
                break
        
        # Priority 20: Check for appendix formatting
        for pattern in self.patterns['appendix_format']:
            if pattern.match(trimmed):
                analysis['type'] = 'appendix_format'
                if re.match(r'^APPENDIX\s+[A-Z]$', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'appendix_header'
                    analysis['level'] = 1
                elif re.match(r'^[A-Z]\.\d+\.\d+', trimmed):
                    analysis['subtype'] = 'appendix_subsection'
                    analysis['level'] = 3
                elif re.match(r'^[A-Z]\.\d+', trimmed):
                    analysis['subtype'] = 'appendix_section'
                    analysis['level'] = 2
                else:
                    analysis['subtype'] = 'appendix_content'
                    analysis['level'] = 1
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 21: Check for block quotes
        for pattern in self.patterns['block_quote']:
            if pattern.match(trimmed):
                analysis['type'] = 'block_quote'
                analysis['confidence'] = 0.85
                analysis['content'] = trimmed.lstrip('> ').strip('"\'')
                return analysis
        
        # Priority 22: Check for mathematical models
        for pattern in self.patterns['math_model']:
            if pattern.search(trimmed):
                analysis['type'] = 'math_model'
                analysis['confidence'] = 0.85
                if re.search(r'[Yy]\s*=\s*[βα]', trimmed):
                    analysis['subtype'] = 'regression_model'
                elif re.search(r'[Rr]²', trimmed):
                    analysis['subtype'] = 'r_squared'
                elif re.search(r'[Pp]\s*[<>=]', trimmed):
                    analysis['subtype'] = 'p_value'
                else:
                    analysis['subtype'] = 'statistical_notation'
                return analysis
        
        # Priority 23: Check for text emphasis patterns
        for pattern in self.patterns['text_emphasis']:
            if pattern.search(trimmed):
                analysis['type'] = 'text_emphasis'
                if '`' in trimmed:
                    analysis['subtype'] = 'monospace'
                elif '***' in trimmed:
                    analysis['subtype'] = 'bold_italic'
                elif '**' in trimmed:
                    analysis['subtype'] = 'bold'
                else:
                    analysis['subtype'] = 'italic'
                analysis['confidence'] = 0.80
                # Don't return - inline emphasis, continue processing
                break
        
        # Priority 24: Check for APA reference format
        for pattern in self.patterns['reference_apa']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'reference_apa'
                if 'doi' in trimmed.lower():
                    analysis['subtype'] = 'doi_reference'
                elif 'Retrieved' in trimmed:
                    analysis['subtype'] = 'web_reference'
                else:
                    analysis['subtype'] = 'standard_reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 25: Check for TOC entries
        for pattern in self.patterns['toc_entry']:
            if pattern.match(trimmed):
                analysis['type'] = 'toc_entry'
                analysis['confidence'] = 0.95
                # Extract page number if present
                page_match = re.search(r'(\d+)\s*$', trimmed)
                if page_match:
                    analysis['page_number'] = int(page_match.group(1))
                return analysis
        
        # Priority 26: Check for footnote markers
        for pattern in self.patterns['footnote_marker']:
            if pattern.search(trimmed):
                analysis['type'] = 'footnote_marker'
                if trimmed.startswith('[^'):
                    analysis['subtype'] = 'footnote_definition'
                else:
                    analysis['subtype'] = 'footnote_reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 27: Check for abbreviations
        for pattern in self.patterns['abbreviation']:
            match = pattern.search(trimmed)
            if match:
                # Only classify if it looks like a definition
                if re.search(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\s+\([A-Z]{2,}\)', trimmed):
                    analysis['type'] = 'abbreviation'
                    analysis['subtype'] = 'definition'
                    analysis['confidence'] = 0.85
                    return analysis
        
        # Priority 28: Check for caption formatting
        for pattern in self.patterns['caption_format']:
            if pattern.match(trimmed):
                analysis['type'] = 'caption_format'
                if 'Table' in trimmed:
                    analysis['subtype'] = 'table_caption'
                elif 'Figure' in trimmed:
                    analysis['subtype'] = 'figure_caption'
                elif trimmed.startswith('Source:'):
                    analysis['subtype'] = 'source_attribution'
                elif trimmed.lower().startswith('note:'):
                    analysis['subtype'] = 'table_note'
                else:
                    analysis['subtype'] = 'caption'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 29: Check for page breaks
        for pattern in self.patterns['page_break']:
            if pattern.match(trimmed):
                analysis['type'] = 'page_break'
                analysis['confidence'] = 1.0
                return analysis
        
        # Priority 30: Check for statistical results
        for pattern in self.patterns['statistical_result']:
            if pattern.search(trimmed):
                analysis['type'] = 'statistical_result'
                analysis['confidence'] = 0.85
                # Identify specific stat types
                stats_found = []
                if re.search(r'β\s*=', trimmed):
                    stats_found.append('beta')
                if re.search(r'[Pp]\s*[<>=]', trimmed):
                    stats_found.append('p_value')
                if re.search(r'[Ff]\s*\(', trimmed):
                    stats_found.append('f_statistic')
                if re.search(r'[Rr]²?\s*=', trimmed):
                    stats_found.append('r_value')
                if re.search(r'CI\s*=', trimmed):
                    stats_found.append('confidence_interval')
                analysis['stats_types'] = stats_found
                return analysis
        
        # Priority 31: Check for questionnaire patterns
        for pattern in self.patterns['questionnaire']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'questionnaire'
                if re.match(r'^Section\s+[A-Z]', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'section_header'
                elif '□' in trimmed or '☐' in trimmed:
                    analysis['subtype'] = 'checkbox_item'
                elif 'SA' in trimmed and 'SD' in trimmed:
                    analysis['subtype'] = 'likert_header'
                else:
                    analysis['subtype'] = 'question_item'
                analysis['confidence'] = 0.85
                return analysis
        
        # Priority 32: Check for glossary entries
        for pattern in self.patterns['glossary_entry']:
            if pattern.match(trimmed):
                analysis['type'] = 'glossary_entry'
                # Extract term and definition
                term_match = re.match(r'\*\*([^*]+)\*\*:\s*(.+)', trimmed)
                if term_match:
                    analysis['term'] = term_match.group(1)
                    analysis['definition'] = term_match.group(2)
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 33: Check for cross-references
        for pattern in self.patterns['cross_reference']:
            if pattern.search(trimmed):
                analysis['type'] = 'cross_reference'
                refs_found = []
                if re.search(r'[Tt]able\s+\d+', trimmed):
                    refs_found.append('table')
                if re.search(r'[Ff]igure\s+\d+', trimmed):
                    refs_found.append('figure')
                if re.search(r'[Ss]ection\s+\d+', trimmed):
                    refs_found.append('section')
                if re.search(r'[Pp]age\s+\d+', trimmed):
                    refs_found.append('page')
                analysis['reference_types'] = refs_found
                analysis['confidence'] = 0.80
                # Don't return - cross-references are inline
                break
        
        # Priority 34: Check for running headers
        for pattern in self.patterns['running_header']:
            if pattern.match(trimmed):
                analysis['type'] = 'running_header'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 35: SHORT DOCUMENT KEY POINT DETECTION
        # Check for key point markers (learning objectives, definitions, warnings, etc.)
        key_point_type, emoji = self.get_key_point_type(trimmed)
        if key_point_type:
            analysis['type'] = 'key_point'
            analysis['key_point_type'] = key_point_type
            analysis['emoji_prefix'] = emoji
            analysis['confidence'] = 0.90
            return analysis
        
        # Priority 36: Assignment header fields (Student Name, Course, etc.)
        if self.is_assignment_header_field(trimmed):
            analysis['type'] = 'assignment_header_field'
            analysis['confidence'] = 0.90
            return analysis
        
        # Default: paragraph
        analysis['confidence'] = 0.70
        return analysis


# ============================================================================
# COVER PAGE HANDLER - December 30, 2025
# Detects cover pages and extracts info from FIRST PAGE ONLY
# Creates a standardized cover page using template layout
# ============================================================================

class CoverPageHandler:
    """
    Detect and extract cover page information from academic documents.
    Only extracts from the first page (before DECLARATION/CERTIFICATION).
    Extracts text from shapes/textboxes for topics.
    """
    
    # Path to logo image
    LOGO_PATH = os.path.join(os.path.dirname(__file__), 'coverpage_template', 'cover_logo.png')
    
    # Default academic year if date not found
    DEFAULT_DATE = '2025/2026'
    
    def __init__(self):
        self.extracted_data = {
            'university': 'THE UNIVERSITY OF BAMENDA',
            'faculty': None,
            'department': None,
            'topic': None,
            'degree': None,
            'name': None,
            'registration_number': None,
            'supervisor': None,
            'co_supervisor': None,
            'field_supervisor': None,
            'month_year': None
        }
        self.has_cover_page = False
        self.cover_page_end_index = 0
        self.shape_texts = []  # Text extracted from shapes/textboxes
    
    def _extract_text_from_shapes(self, doc):
        """
        Extract text from shapes and textboxes in the document.
        Topics are often placed inside shapes/textboxes on cover pages.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            list: List of text strings found in shapes
        """
        shape_texts = []
        
        try:
            # Define namespaces for parsing
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                'v': 'urn:schemas-microsoft-com:vml',
                'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
                'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
            }
            
            # Search in document body for shapes
            body = doc.element.body
            
            # Find all text in wps:txbx (textbox content in shapes)
            for txbx in body.iter('{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'):
                for txbxContent in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
                    texts = []
                    for t in txbxContent.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            texts.append(t.text)
                    if texts:
                        full_text = ''.join(texts).strip()
                        if full_text and len(full_text) > 5:
                            shape_texts.append(full_text)
            
            # Also check for VML textboxes (older format)
            for textbox in body.iter('{urn:schemas-microsoft-com:vml}textbox'):
                for txbxContent in textbox.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
                    texts = []
                    for t in txbxContent.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            texts.append(t.text)
                    if texts:
                        full_text = ''.join(texts).strip()
                        if full_text and len(full_text) > 5:
                            shape_texts.append(full_text)
            
            # Also check for drawing textboxes
            for drawing in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                for t in drawing.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}t'):
                    if t.text:
                        text = t.text.strip()
                        if text and len(text) > 5:
                            shape_texts.append(text)
            
            # Alternative: iterate through inline shapes
            for p in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                for pict in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                    for shape in pict.iter('{urn:schemas-microsoft-com:vml}shape'):
                        for textbox in shape.iter('{urn:schemas-microsoft-com:vml}textbox'):
                            texts = []
                            for t in textbox.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                if t.text:
                                    texts.append(t.text)
                            if texts:
                                full_text = ''.join(texts).strip()
                                if full_text and len(full_text) > 5:
                                    shape_texts.append(full_text)
            
            logger.info(f"Extracted {len(shape_texts)} text items from shapes/textboxes")
            
        except Exception as e:
            logger.warning(f"Error extracting text from shapes: {e}")
        
        return shape_texts
    
    def detect_and_extract(self, doc):
        """
        Detect cover page and extract info from FIRST PAGE ONLY.
        Stops at DECLARATION, CERTIFICATION, or similar markers.
        Also extracts text from shapes/textboxes for topics.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            tuple: (has_cover_page, extracted_data, cover_page_end_index)
        """
        paragraphs = doc.paragraphs
        cover_indicators = 0
        
        # First, extract text from shapes (topics are often in shapes)
        self.shape_texts = self._extract_text_from_shapes(doc)
        
        # Stop markers - these indicate content AFTER cover page
        stop_markers = ['DECLARATION', 'CERTIFICATION', 'DEDICATION', 'ACKNOWLEDGEMENT', 
                        'ABSTRACT', 'TABLE OF CONTENTS', 'ALL RIGHTS RESERVED', 'RIGHTS RESERVED']
        
        # Find where cover page ends (before declaration/certification)
        # Also stop at ©Copyright line which indicates end of cover page
        cover_end = 0
        for i, para in enumerate(paragraphs[:80]):
            text = para.text.strip().upper()
            
            # Check for copyright symbol - indicates end of cover page
            if '©' in para.text or 'COPYRIGHT' in text:
                cover_end = i
                break
            
            # Check for stop markers - these indicate end of cover page
            for marker in stop_markers:
                if text == marker or text.startswith(marker + ' '):
                    cover_end = i
                    break
            if cover_end > 0:
                break
        
        # If no stop marker found, assume first 45 paragraphs max
        if cover_end == 0:
            cover_end = 45
        
        # Collect all paragraph texts for better name detection
        all_cover_texts = [para.text.strip() for para in paragraphs[:cover_end]]
        
        # Extract from first page paragraphs only
        for i, para in enumerate(paragraphs[:cover_end]):
            text = para.text.strip()
            text_upper = text.upper()
            is_bold = any(run.bold for run in para.runs if run.bold)
            
            # University name
            if 'UNIVERSITY OF BAMENDA' in text_upper:
                cover_indicators += 3
                self.extracted_data['university'] = 'THE UNIVERSITY OF BAMENDA'
            
            # Faculty/Institute detection
            if not self.extracted_data['faculty']:
                if 'HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT' in text_upper:
                    self.extracted_data['faculty'] = 'HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT'
                    cover_indicators += 2
                elif 'NATIONAL HIGHER POLYTECHNIC INSTITUTE' in text_upper:
                    self.extracted_data['faculty'] = 'NATIONAL HIGHER POLYTECHNIC INSTITUTE'
                    cover_indicators += 2
                elif 'HIGHER TEACHERS TRAINING COLLEGE' in text_upper or 'HIGHER TEACHER TRAINING COLLEGE' in text_upper:
                    self.extracted_data['faculty'] = 'HIGHER TEACHERS TRAINING COLLEGE'
                    cover_indicators += 2
                elif 'ECOLE NORMALE' in text_upper:
                    self.extracted_data['faculty'] = text.upper()
                    cover_indicators += 2
            
            # Department
            dept_match = re.search(r'DEPARTMENT\s+OF\s+([A-Z][A-Z\s&]+)', text_upper)
            if dept_match and not self.extracted_data['department']:
                # Clean up - stop at common words
                dept = dept_match.group(1).strip()
                # Truncate at "IN THE" or similar
                for stop_word in [' IN THE ', ' OF THE ', ' FOR THE ']:
                    if stop_word in dept:
                        dept = dept.split(stop_word)[0].strip()
                self.extracted_data['department'] = dept
                cover_indicators += 2
            
            # Submission statement - extract faculty and degree
            if 'DISSERTATION SUBMITTED' in text_upper or 'THESIS SUBMITTED' in text_upper or 'PROJECT SUBMITTED' in text_upper:
                cover_indicators += 3
                self._extract_from_submission(text)
            
            # Name after BY
            if text_upper == 'BY' or text_upper.startswith('PRESENTED BY'):
                cover_indicators += 2
                self._extract_name_after_by(paragraphs, i, cover_end)
            
            # Registration number
            reg_match = re.search(r'REGISTRATION\s+NUMBER[:\s]+([A-Z0-9]+)', text_upper)
            if reg_match and not self.extracted_data['registration_number']:
                self.extracted_data['registration_number'] = reg_match.group(1).strip()
                cover_indicators += 2
            
            # Supervisor
            if 'SUPERVISOR' in text_upper and 'CO' not in text_upper:
                cover_indicators += 1
                self._extract_supervisor(paragraphs, i, cover_end, is_co=False)
            
            # Co-supervisor
            if 'CO-SUPERVISOR' in text_upper or 'CO - SUPERVISOR' in text_upper or 'CO SUPERVISOR' in text_upper:
                self._extract_supervisor(paragraphs, i, cover_end, is_co=True)
            
            # Field supervisor
            if 'FIELD SUPERVISOR' in text_upper or 'FIELD-SUPERVISOR' in text_upper:
                self._extract_field_supervisor(paragraphs, cover_end)
            
            # Date (month year) - also check for academic year format
            date_match = re.search(
                r'\b(JANUARY|FEBRUARY|MARCH|APRIL|MAY|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER)\s*,?\s*(\d{4})\b',
                text_upper
            )
            if date_match and len(text) < 30 and not self.extracted_data['month_year']:
                self.extracted_data['month_year'] = f"{date_match.group(1).title()} {date_match.group(2)}"
                cover_indicators += 1
            
            # Also check for year format like "2024/2025" or "2025"
            if not self.extracted_data['month_year']:
                year_match = re.search(r'\b(20\d{2})\s*/\s*(20\d{2})\b', text)
                if year_match:
                    self.extracted_data['month_year'] = f"{year_match.group(1)}/{year_match.group(2)}"
                    cover_indicators += 1
                elif re.match(r'^\s*20\d{2}\s*$', text):
                    self.extracted_data['month_year'] = text.strip()
                    cover_indicators += 1
        
        # Check if we have enough indicators for a cover page
        if cover_indicators >= 5:
            self.has_cover_page = True
            self.cover_page_end_index = cover_end
            
            # Second pass: find topic (longest bold centered text or from shapes)
            self._extract_topic(paragraphs[:cover_end])
            
            # Third pass: comprehensive name extraction if not found
            if not self.extracted_data['name']:
                self._extract_name_comprehensive(paragraphs, cover_end)
                
            # Fourth pass: check shapes for name if still not found
            if not self.extracted_data['name']:
                self._extract_name_from_shapes()
            
            # Fifth pass: field supervisor if not found
            if not self.extracted_data['field_supervisor']:
                self._extract_field_supervisor(paragraphs, cover_end)
            
            # Set default date if not found
            if not self.extracted_data['month_year']:
                self.extracted_data['month_year'] = self.DEFAULT_DATE
            
            logger.info(f"Cover page detected with {cover_indicators} indicators, ends at index {self.cover_page_end_index}")
            logger.info(f"Extracted data: {self.extracted_data}")
        
        return self.has_cover_page, self.extracted_data, self.cover_page_end_index
    
    def _extract_from_submission(self, text):
        """Extract faculty, department, degree from submission statement"""
        text_upper = text.upper()
        
        # Extract faculty from submission
        if not self.extracted_data['faculty']:
            if 'HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT' in text_upper:
                self.extracted_data['faculty'] = 'HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT'
            elif 'NATIONAL HIGHER POLYTECHNIC INSTITUTE' in text_upper:
                self.extracted_data['faculty'] = 'NATIONAL HIGHER POLYTECHNIC INSTITUTE'
            else:
                # Try generic pattern
                match = re.search(r'(HIGHER\s+INSTITUTE\s+OF\s+[A-Z]+(?:\s+AND\s+[A-Z]+)?)', text_upper)
                if match:
                    self.extracted_data['faculty'] = match.group(1)
        
        # Extract department if not found
        if not self.extracted_data['department']:
            match = re.search(r'DEPARTMENT\s+OF\s+([^,]+),', text, re.IGNORECASE)
            if match:
                dept = match.group(1).strip().upper()
                self.extracted_data['department'] = dept
        
        # Extract degree (MBA, M.Eng, M.Sc, etc.)
        if not self.extracted_data['degree']:
            # Look for degree in parentheses like (MBA) or (M.Eng)
            match = re.search(r'\(([A-Z]\.?[A-Za-z\.]+)\)', text)
            if match:
                self.extracted_data['degree'] = match.group(1)
            else:
                # Look for Master's, Bachelor's, etc.
                match = re.search(r"(MASTER'?S?|BACHELOR'?S?|DOCTORATE|PHD|PH\.D)", text_upper)
                if match:
                    self.extracted_data['degree'] = match.group(1).title()
    
    def _extract_name_after_by(self, paragraphs, by_index, max_index):
        """Extract author name from paragraphs after BY - improved patterns"""
        if self.extracted_data['name']:
            return
        
        for j in range(1, 6):  # Check more paragraphs
            if by_index + j < min(len(paragraphs), max_index):
                para = paragraphs[by_index + j]
                text = para.text.strip()
                is_bold = any(run.bold for run in para.runs if run.bold)
                
                # Skip empty or short text
                if not text or len(text) < 4:
                    continue
                
                # Skip if it's registration number line
                if 'REGISTRATION' in text.upper():
                    continue
                
                # Skip if contains supervisor keywords
                if 'SUPERVISOR' in text.upper():
                    continue
                
                # Skip date patterns
                if re.match(r'^(JANUARY|FEBRUARY|MARCH|APRIL|MAY|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER)', text.upper()):
                    continue
                
                # Name is usually in CAPS or bold
                if text.isupper() or is_bold:
                    # Clean up the name
                    name = text.replace('**', '').strip()
                    # Remove any trailing codes containing digits (like UBA23CP043)
                    name = re.sub(r'\s+[A-Z]*\d+[A-Z0-9]*$', '', name)
                    # Remove any matricule or ID patterns
                    name = re.sub(r'\s*\([A-Z0-9]+\)\s*$', '', name)
                    
                    # Validate: should have at least 2 words (first and last name)
                    if len(name.split()) >= 2 and len(name) >= 5:
                        self.extracted_data['name'] = name
                        return
    
    def _extract_name_comprehensive(self, paragraphs, max_index):
        """Comprehensive name extraction - looks for name patterns across cover page"""
        if self.extracted_data['name']:
            return
        
        # Name patterns to look for
        name_indicators = []
        
        for i, para in enumerate(paragraphs[:max_index]):
            text = para.text.strip()
            text_upper = text.upper()
            is_bold = any(run.bold for run in para.runs if run.bold)
            
            # Skip empty text
            if not text or len(text) < 4:
                continue
            
            # Look for "BY" followed by name
            if text_upper == 'BY' or 'PRESENTED BY' in text_upper or 'SUBMITTED BY' in text_upper:
                # Check next few paragraphs
                for j in range(1, 5):
                    if i + j < max_index:
                        next_para = paragraphs[i + j]
                        next_text = next_para.text.strip()
                        next_upper = next_text.upper()
                        
                        if not next_text or len(next_text) < 5:
                            continue
                        
                        # Skip non-name indicators
                        if any(skip in next_upper for skip in ['REGISTRATION', 'SUPERVISOR', 'DEPARTMENT', 'FACULTY']):
                            continue
                        
                        # Check if looks like a name (all caps, 2+ words)
                        if next_text.isupper() and len(next_text.split()) >= 2:
                            # Clean it
                            name = re.sub(r'\s+[A-Z]*\d+[A-Z0-9]*$', '', next_text)
                            name = re.sub(r'\s*\([A-Z0-9]+\)\s*$', '', name)
                            if len(name.split()) >= 2:
                                name_indicators.append((len(name), name))
                                break
            
            # Look for name pattern: ALL CAPS text that's not a header
            if text.isupper() and is_bold:
                # Check it's not an institutional header
                skip_patterns = ['UNIVERSITY', 'DEPARTMENT', 'INSTITUTE', 'FACULTY', 'SCHOOL',
                               'SUPERVISOR', 'REGISTRATION', 'SUBMITTED', 'DISSERTATION', 
                               'THESIS', 'PROJECT', 'BAMENDA', 'HIGHER', 'NATIONAL']
                
                if not any(p in text_upper for p in skip_patterns):
                    # Check if looks like a name (2-4 words, no numbers)
                    words = text.split()
                    if 2 <= len(words) <= 5 and not re.search(r'\d', text):
                        name = text.replace('**', '').strip()
                        name_indicators.append((len(name), name))
        
        # Use the best candidate (longest reasonable name)
        if name_indicators:
            # Sort by length, pick the best
            name_indicators.sort(reverse=True)
            for length, name in name_indicators:
                if 8 <= length <= 50:  # Reasonable name length
                    self.extracted_data['name'] = name
                    return
    
    def _extract_supervisor(self, paragraphs, sup_index, max_index, is_co=False):
        """Extract supervisor name from paragraphs after SUPERVISOR header - improved"""
        key = 'co_supervisor' if is_co else 'supervisor'
        if self.extracted_data[key]:
            return
        
        # Check the current paragraph for inline supervisor name
        current_text = paragraphs[sup_index].text.strip()
        
        # Check for "SUPERVISOR: Dr. Name" format
        sup_match = re.search(r'(?:CO-?)?SUPERVISOR[:\s]+(.+)', current_text, re.IGNORECASE)
        if sup_match:
            name = sup_match.group(1).strip()
            if name and len(name) > 3:
                self.extracted_data[key] = name.replace('**', '').strip()
                return
        
        for j in range(1, 6):  # Check more paragraphs
            if sup_index + j < min(len(paragraphs), max_index):
                para = paragraphs[sup_index + j]
                text = para.text.strip()
                
                if not text:
                    continue
                
                # Stop if we hit another section header
                if any(h in text.upper() for h in ['ACKNOWLEDGEMENT', 'DEDICATION', 'DECLARATION', 'ABSTRACT']):
                    break
                
                # Skip header lines
                if 'SUPERVISOR' in text.upper() and not re.search(r'(Dr\.|Prof\.|Engr\.|Mr\.|Mrs\.)', text, re.IGNORECASE):
                    continue
                
                # Look for Dr., Prof., Engr., Mr., Mrs. patterns
                if re.search(r'(Dr\.|Prof\.|Engr\.|Mr\.|Mrs\.|Ms\.)', text, re.IGNORECASE):
                    self.extracted_data[key] = text.replace('**', '').strip()
                    return
                
                # Also accept CAPS names after supervisor header
                if text.isupper() and len(text.split()) >= 2:
                    self.extracted_data[key] = text.replace('**', '').strip()
                    return
    
    def _extract_field_supervisor(self, paragraphs, max_index):
        """Extract field supervisor if present"""
        for i, para in enumerate(paragraphs[:max_index]):
            text = para.text.strip()
            text_upper = text.upper()
            
            if 'FIELD SUPERVISOR' in text_upper or 'FIELD-SUPERVISOR' in text_upper:
                # Check inline
                match = re.search(r'FIELD[\s-]*SUPERVISOR[:\s]+(.+)', text, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    if name and len(name) > 3:
                        self.extracted_data['field_supervisor'] = name.replace('**', '').strip()
                        return
                
                # Check next paragraphs
                for j in range(1, 4):
                    if i + j < max_index:
                        next_para = paragraphs[i + j]
                        next_text = next_para.text.strip()
                        if next_text and re.search(r'(Dr\.|Prof\.|Engr\.|Mr\.|Mrs\.)', next_text, re.IGNORECASE):
                            self.extracted_data['field_supervisor'] = next_text.replace('**', '').strip()
                            return
    
    def _extract_topic(self, paragraphs):
        """Extract topic - first try bold paragraphs, then check shapes for actual topic text"""
        if self.extracted_data['topic']:
            return
        
        # PRIORITY 1: Look for longest bold/centered text in paragraphs (most reliable)
        best_topic = None
        best_length = 0
        
        for para in paragraphs:
            text = para.text.strip()
            is_bold = any(run.bold for run in para.runs if run.bold)
            
            # Skip short text
            if len(text) < 25:
                continue
            
            # Skip institutional headers
            skip_patterns = ['UNIVERSITY', 'DEPARTMENT OF', 'INSTITUTE', 'FACULTY', 
                           'SUPERVISOR', 'REGISTRATION', 'SUBMITTED', 'BY', 'BAMENDA',
                           'HIGHER INSTITUTE', 'NATIONAL', 'PRESENTED', 'PARTIAL FULFILLMENT',
                           'REQUIREMENTS', 'AWARD OF', 'DEGREE']
            if any(p in text.upper() for p in skip_patterns):
                continue
            
            # Topic is usually bold and longer than 25 chars
            if is_bold and len(text) > best_length:
                best_topic = text.replace('**', '').strip()
                best_length = len(text)
        
        if best_topic:
            self.extracted_data['topic'] = best_topic
            logger.info(f"Topic extracted from paragraph: {best_topic[:50]}...")
            return
        
        # PRIORITY 2: Check shape texts for topic (only if no paragraph topic found)
        # Note: Shape extraction may get text from throughout the document, so filter carefully
        best_shape_topic = None
        best_shape_length = 0
        
        for shape_text in self.shape_texts:
            text = shape_text.strip()
            text_upper = text.upper()
            
            # Skip short texts (topics are usually substantial)
            if len(text) < 15:
                continue
            
            # Skip texts that are mostly underscores or dashes (signature lines)
            underscore_ratio = (text.count('_') + text.count('-')) / len(text)
            if underscore_ratio > 0.15:
                continue
            
            # Skip if it has many underscores (signature/form fields)
            if text.count('_') > 2:
                continue
            
            # Skip diagram/figure labels (short labels with specific keywords at START)
            diagram_labels = ['INDEPENDENT', 'DEPENDENT', 'CONTROL', 'VARIABLE', 'FIGURE', 
                             'TABLE', 'SOURCE', 'CHART', 'HYPOTHESIS', 'FORMAL', 'INFORMAL']
            if any(text_upper.startswith(lbl) for lbl in diagram_labels):
                continue
            
            # Skip exact institutional headers (but allow longer text containing these words)
            exact_skip = ['HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT',
                         'DEPARTMENT OF MANAGEMENT AND ENTREPRENEURSHIP',
                         'NATIONAL HIGHER POLYTECHNIC INSTITUTE',
                         'THE UNIVERSITY OF BAMENDA']
            if text_upper in exact_skip or any(text_upper == s for s in exact_skip):
                continue
            
            # Skip short administrative labels
            admin_labels = ['COORDINATOR', 'POSTGRADUATE', 'DIRECTOR', 'HEAD OF DEPARTMENT',
                          'SUPERVISOR', 'DATE:', 'NAME:', 'APPROVED', 'REJECTED', 'GENERAL']
            if any(lbl in text_upper for lbl in admin_labels):
                continue
            
            # Topic should have topic-like keywords or patterns
            topic_keywords = ['ANALYSIS', 'STUDY', 'IMPACT', 'EFFECT OF', 'EFFECTS OF',
                            'ROLE OF', 'ASSESSMENT', 'EVALUATION', 'INVESTIGATION', 
                            'INFLUENCE', 'RELATIONSHIP', 'FACTORS AFFECTING',
                            'STRATEGIES', 'IMPLEMENTATION', 'EFFECTIVENESS', 
                            'CASE OF', 'CASE STUDY', 'CHALLENGES', 'PERFORMANCE IN',
                            'PERFORMANCE OF', 'DETERMINANTS', 'CONTRIBUTION']
            
            has_topic_keyword = any(kw in text_upper for kw in topic_keywords)
            
            # Accept if it has keyword OR is sufficiently long and looks like a title (all caps or title case)
            is_likely_title = len(text) > 40 and (text.isupper() or text.istitle())
            
            if (has_topic_keyword or is_likely_title) and len(text) > best_shape_length:
                best_shape_topic = text.replace('**', '').strip()
                best_shape_length = len(text)
                logger.info(f"Found potential topic in shape: {text[:60]}...")
        
        if best_shape_topic:
            self.extracted_data['topic'] = best_shape_topic
            logger.info(f"Topic extracted from shape: {best_shape_topic[:50]}...")

    def _extract_name_from_shapes(self):
        """Extract author name from shapes if not found in paragraphs"""
        if self.extracted_data['name']:
            return
            
        # Look for short, all-caps text in shapes that isn't the topic
        for shape_text in self.shape_texts:
            text = shape_text.strip()
            
            # Check for "BY:" prefix (common in shapes)
            by_match = re.search(r'BY[:\s]+([A-Z\s]+)', text, re.IGNORECASE)
            if by_match:
                name_candidate = by_match.group(1).strip()
                # Clean up (remove degrees, numbers, etc.)
                name_candidate = re.split(r'[\(\[\{0-9]', name_candidate)[0].strip()
                if len(name_candidate) > 3:
                    self.extracted_data['name'] = name_candidate
                    logger.info(f"Name extracted from shape (BY prefix): {name_candidate}")
                    return
            
            # Skip if it's the topic
            if text == self.extracted_data.get('topic'):
                continue
                
            # Skip if it contains "Department", "University", etc.
            if any(w in text.upper() for w in ['DEPARTMENT', 'UNIVERSITY', 'INSTITUTE', 'FACULTY', 'SCHOOL', 'COLLEGE', 'HUMAN RESOURCE']):
                continue
                
            # Skip if it contains "Supervisor"
            if 'SUPERVISOR' in text.upper():
                continue
                
            # Name criteria: 2-4 words, all caps (usually), reasonable length
            words = text.split()
            if 2 <= len(words) <= 4 and 10 <= len(text) <= 40:
                # Check if it looks like a name (no numbers, no special chars except maybe dot)
                if not re.search(r'[0-9!@#$%^&*()_+={}\[\]|\\:;"<>,?/~`]', text.replace('.', '')):
                    # Prefer all caps
                    if text.isupper():
                        self.extracted_data['name'] = text
                        logger.info(f"Name extracted from shape: {text}")
                        return
                    # Or Title Case
                    elif text.istitle():
                        self.extracted_data['name'] = text
                        logger.info(f"Name extracted from shape: {text}")
                        return


class CertificationPageHandler:
    """
    Detect and extract certification page information from academic documents.
    Certification pages typically appear after the cover page.
    Generates a standardized certification page with extracted data.
    """
    
    # Regex patterns for certification page extraction
    PATTERNS = {
        # Extract Topic/Title (quoted text or text between keywords)
        'topic': re.compile(r'["“”]([^"“”]+)["“”]|titled\s+["“”]([^"“”]+)["“”]|research\s+titled\s+["“”]([^"“”]+)["“”]|dissertation\s*["“”]([^"“”]+)["“”]|titled\s+(?!["“”])(.+?)\s+(?:submitted\s+by|is\s+the\s+original\s+work\s+of)|report\s+on\s+(.+?)\s+(?:carried\s+out|was\s+done)|internship\s+carried\s+out.*?\s+at\s+([A-Z0-9\s&]+?)(?:\s+was|\s+by)', re.IGNORECASE | re.DOTALL),
        
        # Extract Author Name (after "work of" or "done by")
        'author': re.compile(r'(?:original\s+work\s+of|was\s+done\s+by)\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)+)', re.IGNORECASE),
        
        # Extract Degree Program  
        'degree': re.compile(r'award\s+of\s+(?:a\s+)?([^\.]+?)\s+(?:in\s+[A-Z]|degree)', re.IGNORECASE),
        
        # Extract Supervisor Name
        'supervisor': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Supervisor\)', re.IGNORECASE),
        
        # Extract Head of Department Name
        'hod': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Head\s+[Oo]f\s+Department\)', re.IGNORECASE),
        
        # Extract Director Name
        'director': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Director\)', re.IGNORECASE),
        
        # Certification Header Detection
        'header': re.compile(r'^\s*CERTIFICATION\s*$', re.IGNORECASE | re.MULTILINE),
        
        # Institution extraction
        'institution': re.compile(r'(?:Higher\s+Institute\s+of\s+[A-Za-z\s]+|University\s+of\s+[A-Za-z]+)', re.IGNORECASE),
    }
    
    def __init__(self):
        self.extracted_data = {
            'topic': None,
            'author': None,
            'degree': None,
            'program': None,
            'supervisor': None,
            'head_of_department': None,
            'director': None,
            'institution': 'The Higher Institute of Commerce and Management of The University of Bamenda'
        }
        self.has_certification_page = False
        self.certification_start_index = 0
        self.certification_end_index = 0
    
    def detect_certification_page(self, paragraphs, start_index=0):
        """
        Detect if there's a certification page in the document.
        
        Args:
            paragraphs: List of paragraph objects from python-docx
            start_index: Index to start searching from (after cover page)
            
        Returns:
            tuple: (has_certification, start_index, end_index)
        """
        # Search in paragraphs after cover page, within first 100 paragraphs
        search_end = min(start_index + 100, len(paragraphs))
        
        for i in range(start_index, search_end):
            para = paragraphs[i]
            text = para.text.strip().upper()
            
            # Look for CERTIFICATION header
            if text == 'CERTIFICATION':
                self.has_certification_page = True
                self.certification_start_index = i
                
                # Find where certification page ends
                # Look for next major section (DEDICATION, ACKNOWLEDGEMENT, ABSTRACT, etc.)
                end_markers = ['DEDICATION', 'ACKNOWLEDGEMENT', 'ACKNOWLEDGMENT', 'ABSTRACT', 
                              'TABLE OF CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES',
                              'CHAPTER', 'DECLARATION']
                
                for j in range(i + 1, min(i + 50, len(paragraphs))):
                    end_text = paragraphs[j].text.strip().upper()
                    if any(end_text.startswith(marker) for marker in end_markers):
                        self.certification_end_index = j
                        break
                else:
                    # If no end marker found, estimate based on content
                    self.certification_end_index = min(i + 40, len(paragraphs))
                
                logger.info(f"Certification page detected at index {i}, ends at {self.certification_end_index}")
                return True, self.certification_start_index, self.certification_end_index
        
        return False, 0, 0
    
    def extract_from_paragraphs(self, paragraphs):
        """
        Extract certification data from paragraph objects.
        
        Args:
            paragraphs: List of paragraph objects (certification page paragraphs only)
        """
        # Combine all text for pattern matching
        full_text = '\n'.join([p.text for p in paragraphs])
        
        # 1. Extract Topic (quoted text)
        topic_match = self.PATTERNS['topic'].search(full_text)
        if topic_match:
            for group in topic_match.groups():
                if group:
                    topic_text = group.strip()
                    # If it looks like a place (internship), prefix it
                    if 'BACCCUL' in topic_text or 'BANK' in topic_text or 'COUNCIL' in topic_text or len(topic_text) < 20:
                         # Check if it's likely a place name rather than a full title
                         if not topic_text.lower().startswith('the effect') and not topic_text.lower().startswith('an analysis'):
                             topic_text = f"Internship Report at {topic_text}"
                    
                    self.extracted_data['topic'] = topic_text
                    break
        
        # 2. Extract Author Name
        author_match = self.PATTERNS['author'].search(full_text)
        if author_match:
            author_text = author_match.group(1).strip()
            # Clean up author name (remove trailing "with", "registration", etc.)
            for stop_word in [' with', ' registration', ' student', ' matriculation', ' level']:
                if stop_word in author_text.lower():
                    author_text = re.split(stop_word, author_text, flags=re.IGNORECASE)[0].strip()
            self.extracted_data['author'] = author_text
        
        # 3. Extract Degree and Program
        degree_match = self.PATTERNS['degree'].search(full_text)
        if degree_match:
            degree_text = degree_match.group(1).strip()
            # Split degree and program (e.g., "Master's in Business Administration (MBA) in Management")
            if ' in ' in degree_text.lower():
                parts = re.split(r'\s+in\s+', degree_text, maxsplit=1, flags=re.IGNORECASE)
                if len(parts) >= 2:
                    self.extracted_data['degree'] = parts[0].strip()
                    self.extracted_data['program'] = parts[1].strip()
                else:
                    self.extracted_data['degree'] = degree_text
            else:
                self.extracted_data['degree'] = degree_text
        
        # 4. Extract Supervisor Name
        supervisor_match = self.PATTERNS['supervisor'].search(full_text)
        if supervisor_match:
            self.extracted_data['supervisor'] = self._format_name(supervisor_match.group(1).strip())
        
        # 5. Extract Head of Department Name
        hod_match = self.PATTERNS['hod'].search(full_text)
        if hod_match:
            self.extracted_data['head_of_department'] = self._format_name(hod_match.group(1).strip())
        
        # 6. Extract Director Name
        director_match = self.PATTERNS['director'].search(full_text)
        if director_match:
            self.extracted_data['director'] = self._format_name(director_match.group(1).strip())
        
        # 7. Extract Institution
        inst_match = self.PATTERNS['institution'].search(full_text)
        if inst_match:
            self.extracted_data['institution'] = inst_match.group(0).strip()
        
        logger.info(f"Certification data extracted: {self.extracted_data}")
        return self.extracted_data
    
    def _format_name(self, name):
        """Format name with proper title abbreviation."""
        if not name:
            return name
        
        # Standardize title abbreviations
        name = re.sub(r'^Prof\.?\s+', 'Prof. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Dr\.?\s+', 'Dr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Engr\.?\s+', 'Engr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Mr\.?\s+', 'Mr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Mrs\.?\s+', 'Mrs. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Ms\.?\s+', 'Ms. ', name, flags=re.IGNORECASE)
        
        return name
    
    def detect_and_extract(self, doc, start_index=0):
        """
        Detect certification page and extract all data.
        
        Args:
            doc: python-docx Document object
            start_index: Index to start searching from (after cover page)
            
        Returns:
            tuple: (has_certification, extracted_data, start_index, end_index)
        """
        paragraphs = doc.paragraphs
        
        has_cert, start_idx, end_idx = self.detect_certification_page(paragraphs, start_index)
        
        if has_cert:
            cert_paragraphs = paragraphs[start_idx:end_idx]
            self.extract_from_paragraphs(cert_paragraphs)
        
        return self.has_certification_page, self.extracted_data, self.certification_start_index, self.certification_end_index


def format_questionnaire_in_word(doc, questionnaire_data):
    """
    Format a questionnaire document in Word based on extracted structure.
    """
    # Clear existing content if any (except styles)
    # Note: We don't clear everything because we might want to keep headers/footers
    # But for the body content, we start fresh or append.
    # If doc is empty, fine. If not, we append.
    
    # Set up styles
    styles = doc.styles
    
    # Title Style
    if 'Questionnaire Title' not in styles:
        style = styles.add_style('Questionnaire Title', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(16)
        font.bold = True
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(12)
    
    # Section Header Style
    if 'Questionnaire Section' not in styles:
        style = styles.add_style('Questionnaire Section', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True
        pf = style.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        
    # Question Style
    if 'Question Text' not in styles:
        style = styles.add_style('Question Text', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        pf = style.paragraph_format
        pf.space_after = Pt(3)
        
    # Option Style
    if 'Question Option' not in styles:
        style = styles.add_style('Question Option', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        pf = style.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_after = Pt(2)

    # Add Title
    if questionnaire_data.get('title'):
        doc.add_paragraph(questionnaire_data['title'], style='Questionnaire Title')
        
    # Add Instructions
    if questionnaire_data.get('instructions'):
        for instruction in questionnaire_data['instructions']:
            p = doc.add_paragraph(instruction)
            p.paragraph_format.space_after = Pt(12)
            p.italic = True
            
    # Process Sections
    for section in questionnaire_data.get('sections', []):
        # Add Section Header
        if section.get('title') and section.get('title') != 'Questions':
            doc.add_paragraph(section['title'], style='Questionnaire Section')
            
        # Process Questions
        for question in section.get('questions', []):
            # Question Text
            q_num = question.get('number', '')
            q_text_content = question.get('text', '')
            full_text = f"{q_num} {q_text_content}".strip() if q_num else q_text_content
            
            doc.add_paragraph(full_text, style='Question Text')
            
            # Question Options/Input
            q_type = question.get('type', 'open_ended')
            
            if q_type == 'likert_table':
                # Create a full Likert Table (Matrix)
                scale = question.get('scale', {})
                items = scale.get('items', [])
                sub_questions = question.get('sub_questions', [])
                
                if items and sub_questions:
                    # Create table: 1 col for statement + N cols for scale items
                    table = doc.add_table(rows=len(sub_questions) + 1, cols=len(items) + 1)
                    table.style = 'Table Grid'
                    table.autofit = True
                    
                    # Header Row
                    # First cell is "Statement"
                    header_cell = table.cell(0, 0)
                    header_cell.text = "Statement"
                    header_cell.paragraphs[0].runs[0].bold = True
                    
                    # Scale headers
                    for i, item in enumerate(items):
                        cell = table.cell(0, i + 1)
                        cell.text = item
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p.runs:
                            p.runs[0].bold = True
                    
                    # Data Rows
                    for r, statement in enumerate(sub_questions):
                        # Statement cell
                        row_idx = r + 1
                        cell = table.cell(row_idx, 0)
                        cell.text = statement
                        
                        # Radio buttons for each scale item
                        for c in range(len(items)):
                            cell = table.cell(row_idx, c + 1)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run('○')
                            run.font.size = Pt(14)
                    
                    # Add spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(12)

            elif q_type == 'scale' and question.get('scale'):
                # Create Likert Scale Table (Single Question)
                scale = question['scale']
                items = scale.get('items', [])
                if items:
                    table = doc.add_table(rows=2, cols=len(items))
                    table.style = 'Table Grid'
                    table.autofit = True
                    
                    # Header Row
                    for i, item in enumerate(items):
                        cell = table.cell(0, i)
                        cell.text = item
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    # Radio Button Row
                    for i in range(len(items)):
                        cell = table.cell(1, i)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run('○')
                        run.font.size = Pt(14)
                    
                    # Add spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
            elif q_type in ['multiple_choice', 'single_select', 'radio']:
                options = question.get('options', [])
                # Check if we should use horizontal layout (short options)
                # Criteria: max length < 20 chars, total length < 80 chars
                is_short = all(len(opt.get('text', '')) < 20 for opt in options)
                total_len = sum(len(opt.get('text', '')) for opt in options)
                
                if is_short and total_len < 80 and len(options) > 1:
                    # Horizontal Layout
                    p = doc.add_paragraph(style='Question Option')
                    p.paragraph_format.left_indent = Inches(0.25)
                    
                    for i, option in enumerate(options):
                        run = p.add_run('○ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                        
                        # Add spacing between options (except last)
                        if i < len(options) - 1:
                            p.add_run('\t\t') # Double tab for spacing
                else:
                    # Vertical Layout
                    for option in options:
                        p = doc.add_paragraph(style='Question Option')
                        
                        # Handle text input option (Other: ___)
                        if option.get('type') == 'text_input':
                            run = p.add_run('○ ')
                            run.font.name = 'Segoe UI Symbol'
                            p.add_run('Other: ' + '_' * 30)
                        else:
                            run = p.add_run('○ ')
                            run.font.name = 'Segoe UI Symbol'
                            p.add_run(option.get('text', ''))
                    
            elif q_type in ['multiple_select', 'checkbox', 'check_all']:
                options = question.get('options', [])
                # Check if we should use horizontal layout
                is_short = all(len(opt.get('text', '')) < 20 for opt in options)
                total_len = sum(len(opt.get('text', '')) for opt in options)
                
                if is_short and total_len < 80 and len(options) > 1:
                    # Horizontal Layout
                    p = doc.add_paragraph(style='Question Option')
                    p.paragraph_format.left_indent = Inches(0.25)
                    
                    for i, option in enumerate(options):
                        run = p.add_run('☐ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                        
                        # Add spacing between options
                        if i < len(options) - 1:
                            p.add_run('\t\t')
                else:
                    # Vertical Layout
                    for option in options:
                        p = doc.add_paragraph(style='Question Option')
                        run = p.add_run('☐ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                    
            elif q_type == 'open_ended':
                # Add lines for writing
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run('_' * 60)
                p.paragraph_format.space_after = Pt(6)
                
    return doc


class QuestionnaireProcessor:
    """
    Detect, parse, and format questionnaires, surveys, and assessment forms.
    """
    def __init__(self):
        self.questionnaire_data = {
            'is_questionnaire': False,
            'title': '',
            'sections': [],
            'questions': [],
            'demographics': [],
            'instructions': [],
            'scale_type': None
        }
    
    def detect_questionnaire(self, text):
        """
        Determine if document is a questionnaire and extract structure
        """
        if not text:
            return self.questionnaire_data
            
        lines = text.split('\n')
        questionnaire_indicators = 0
        total_indicators = 0
        
        # Quick check for questionnaire keywords in first few lines
        header_check = '\n'.join(lines[:20]).upper()
        if not any(k in header_check for k in ['QUESTIONNAIRE', 'SURVEY', 'ASSESSMENT', 'EVALUATION', 'FEEDBACK FORM']):
            # If no explicit title, check for question density
            question_count = sum(1 for line in lines[:50] if self.is_question_line(line))
            if question_count < 3:
                return self.questionnaire_data
        
        for i, line in enumerate(lines):
            # Check for title patterns
            if self.is_questionnaire_title(line):
                questionnaire_indicators += 3
                self.questionnaire_data['title'] = line.strip()
            
            # Check for question patterns
            if self.is_question_line(line):
                questionnaire_indicators += 2
            
            # Check for Likert scales
            if self.is_likert_scale(line):
                questionnaire_indicators += 2
            
            # Check for demographic sections
            if self.is_demographic_section(line):
                questionnaire_indicators += 2
            
            total_indicators += 1
        
        # Calculate confidence score
        if total_indicators > 0:
            # Normalize confidence
            confidence = min((questionnaire_indicators / 15) * 100, 100)
            self.questionnaire_data['is_questionnaire'] = confidence > 40 # Lower threshold as indicators are specific
            self.questionnaire_data['confidence'] = confidence
        
        return self.questionnaire_data
    
    def parse_questionnaire_structure(self, text):
        """
        Parse questionnaire into structured format
        """
        if not text:
            return self.questionnaire_data
            
        lines = text.split('\n')
        current_section = None
        current_question = None
        
        # State tracking for Likert tables
        in_likert_table = False
        current_likert_scale = []
        
        # Extract instructions first
        for i, line in enumerate(lines[:20]):
            if self.is_instruction_line(line):
                self.questionnaire_data['instructions'].append(line.strip())
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            section_match = self.detect_section_header(line)
            if section_match:
                if current_question:
                    self._finalize_question(current_question)
                    current_question = None
                
                # Reset table state
                in_likert_table = False
                current_likert_scale = []
                
                current_section = {
                    'type': section_match['type'],
                    'title': section_match['title'],
                    'questions': []
                }
                self.questionnaire_data['sections'].append(current_section)
                continue
            
            # Detect Likert Table Header (e.g., "Table 1: Income Adequacy")
            table_match = re.match(r'^(?:Table|Tbl)\s*\d+[:.]\s*(.+)$', line, re.IGNORECASE)
            if table_match:
                if current_question:
                    self._finalize_question(current_question)
                    current_question = None
                
                # Start a new "question" that is actually a table
                current_question = {
                    'number': '',
                    'text': table_match.group(1).strip(),
                    'type': 'likert_table',
                    'options': [],
                    'scale': None,
                    'sub_questions': [] # For table rows
                }
                if current_section:
                    current_section['questions'].append(current_question)
                else:
                    current_section = {'type': 'main', 'title': 'Questions', 'questions': [current_question]}
                    self.questionnaire_data['sections'].append(current_section)
                
                in_likert_table = True
                continue
            
            # Handle content inside Likert Table
            if in_likert_table and current_question:
                # Check if this line is actually a new question
                # If so, we should exit table mode
                if self.detect_question(line):
                    in_likert_table = False
                    # Fall through to main question detection
                else:
                    # Check for instructions inside table
                    if self.is_instruction_line(line) or line.strip().lower().startswith('instructions:'):
                        current_question['instructions'] = line.strip()
                        continue

                    # Check if this line defines the scale (header row)
                    # e.g. "Statement Strongly Disagree Disagree Neutral Agree Strongly Agree"
                    likert_indicators = [
                        'strongly disagree', 'disagree', 'neutral', 'agree', 'strongly agree',
                        'never', 'rarely', 'sometimes', 'often', 'always',
                        'very dissatisfied', 'dissatisfied', 'satisfied', 'very satisfied',
                        'not at all', 'slightly', 'moderately', 'very', 'extremely'
                    ]
                    line_lower = line.lower()
                    
                    # If line contains multiple scale indicators, it's likely the header row
                    found_indicators = [ind for ind in likert_indicators if ind in line_lower]
                    # Also check for abbreviated headers (SD D N A SA)
                    has_abbrev = re.search(r'\b(?:SD|D|N|A|SA)\b', line) and len(line.split()) <= 10
                    
                    if len(found_indicators) >= 3 or has_abbrev:
                        # Extract the scale items properly (splitting by tab, multiple spaces, or pipes)
                        # Remove leading/trailing pipes first
                        clean_line = line.strip('|').strip()
                        parts = re.split(r'\||\t|\s{2,}', clean_line)
                        # Filter out "Statement", empty strings, and whitespace
                        scale_items = [p.strip() for p in parts if p.strip() and p.strip().lower() != 'statement']
                        current_question['scale'] = {'items': scale_items, 'type': 'likert'}
                        continue
                    
                    # Otherwise, it's a row in the table (a sub-question)
                    # e.g. "Our household income is sufficient... [ ] [ ] [ ]"
                    # We want to extract the statement text.
                    # Remove the [ ] parts and pipes
                    clean_line = re.sub(r'\[\s*[xX]?\s*\]', '', line)
                    clean_line = clean_line.replace('|', '').strip()
                    
                    if clean_line:
                        current_question['sub_questions'].append(clean_line)
                    continue

            # Detect questions
            question_match = self.detect_question(line)
            if question_match:
                if current_question:
                    self._finalize_question(current_question)
                
                # Reset table state
                in_likert_table = False
                current_likert_scale = []

                current_question = {
                    'number': question_match['number'],
                    'text': question_match['text'],
                    'type': question_match['type'],
                    'options': [],
                    'scale': None,
                    'instructions': ''
                }
                
                if current_section:
                    current_section['questions'].append(current_question)
                else:
                    # Create default section if none exists
                    current_section = {
                        'type': 'main',
                        'title': 'Questions',
                        'questions': [current_question]
                    }
                    self.questionnaire_data['sections'].append(current_section)
                continue
            
            # Detect options for current question
            if current_question:
                option_match = self.detect_option(line)
                if option_match:
                    current_question['options'].append({
                        'label': option_match['label'],
                        'text': option_match['text'],
                        'type': option_match['type']
                    })
                    continue
                
                # Detect "Other: ____" as an option
                other_match = re.match(r'^(?:Other|Specify)\s*[:.]\s*(_+)?$', line, re.IGNORECASE)
                if other_match:
                     current_question['options'].append({
                        'label': '',
                        'text': 'Other: ________________',
                        'type': 'text_input'
                    })
                     continue

                # Detect scale for current question
                scale_match = self.detect_scale(line)
                if scale_match and not current_question.get('scale'):
                    current_question['scale'] = scale_match
                    continue
                
                # Capture question continuation or instructions
                if line.strip() and not self.is_section_header(line) and not self.is_instruction_line(line):
                    # Check if it's an implicit option
                    # If question ends with ?, treat subsequent lines as options if they are not new questions
                    question_ended = current_question['text'].strip().endswith('?') or current_question['text'].strip().endswith(':')
                    
                    # Also check if the line looks like an option (short, capitalized)
                    is_likely_option = len(line) < 100 and line[0].isupper() or line[0].isdigit()
                    
                    if (question_ended or len(current_question['options']) > 0) and not self.is_question_line(line):
                         # If we already have options, or the question seems finished, treat as option
                         current_question['options'].append({
                                'label': '',
                                'text': line.strip(),
                                'type': 'implicit_option'
                            })
                         continue

                    if not current_question['options'] and not current_question.get('scale'):
                        # Likely continuation of question text
                        # Only append if it doesn't look like a new question start
                        if not self.is_question_line(line):
                            current_question['text'] += ' ' + line.strip()
                    elif line.strip().endswith('?'):
                        # Additional question part
                        current_question['text'] += ' ' + line.strip()


        
        # Finalize last question
        if current_question:
            self._finalize_question(current_question)
        
        return self.questionnaire_data
    
    def _finalize_question(self, question):
        """Finalize question processing"""
        # Auto-detect question type if not specified
        if not question.get('type'):
            question['type'] = self._infer_question_type(question)
        
        # Ensure options have proper types
        for option in question['options']:
            if not option.get('type'):
                option['type'] = self._infer_option_type(question['type'])
            
            # Fix implicit options
            if option.get('type') == 'implicit_option':
                option['type'] = self._infer_option_type(question['type'])

    def _infer_option_type(self, question_type):
        """Infer option type based on question type"""
        if question_type in ['multiple_select', 'check_all']:
            return 'checkbox'
        elif question_type in ['single_select', 'multiple_choice', 'scale']:
            return 'radio'
        else:
            return 'text'

    # DETECTION HELPER METHODS
    
    def is_questionnaire_title(self, line):
        patterns = [
            r'^(?:#+\s+)?(?:QUESTIONNAIRE|SURVEY|ASSESSMENT|EVALUATION)\b',
            r'^(?:#+\s+)?(?:Research|Study|Data)\s+(?:Collection|Gathering)\s+(?:Tool|Instrument)\b',
            r'^(?:#+\s+)?(?:Student|Teacher|Parent|Employee|Customer)\s+(?:Feedback|Satisfaction)\s+(?:Form|Survey)\b',
            r'^(?:#+\s+)?(?:Self-[Aa]ssessment|Self-[Ee]valuation)\b',
            r'^(?:#+\s+)?([A-Z][A-Za-z\s]+(?:Assessment|Evaluation|Survey|Study)\s+(?:Questionnaire|Survey|Form))\s*$'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_question_line(self, line):
        patterns = [
            r'^\s*(?:Q|Question)\s*\d+[:.)]',
            r'^\s*\d+[:.)]\s+.*[?]',
            r'^\s*Item\s+\d+[:.)]',
            r'^\s*[A-Z][:.)]\s+.*[?]'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_likert_scale(self, line):
        likert_indicators = [
            'strongly disagree', 'disagree', 'neutral', 'agree', 'strongly agree',
            'never', 'rarely', 'sometimes', 'often', 'always',
            'very dissatisfied', 'dissatisfied', 'satisfied', 'very satisfied',
            'not at all', 'slightly', 'moderately', 'very', 'extremely'
        ]
        line_lower = line.lower()
        count = sum(1 for indicator in likert_indicators if indicator in line_lower)
        
        # Also check for abbreviated headers (SD D N A SA)
        if re.search(r'\b(?:SD|D|N|A|SA)\b', line) and len(line.split()) <= 10:
             return True
             
        return count >= 3  # At least 3 Likert items present
    
    def is_demographic_section(self, line):
        patterns = [
            r'^(?:#+\s+)?(?:DEMOGRAPHIC|BACKGROUND|PERSONAL)\s+(?:INFORMATION|DATA)\b',
            r'^\s*(?:Name|Age|Gender|School|University|College|Class|Grade)\s*[:.]',
            r'^PART\s+[A-Z]\s*[:.]?\s*(?:Demographic|Personal)'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_instruction_line(self, line):
        patterns = [
            r'^\s*(?:Instructions?|Directions?|Guidelines?)\s*[:.]?\s*$',
            r'^\s*(?:Please|Kindly)\s+(?:read|answer|complete|fill|respond)',
            r'^\s*(?:All\s+questions|Each\s+item)\s+(?:must|should)\s+be',
            r'^\s*(?:Mark|Tick|Check|Circle)\s+(?:your|the)\s+(?:answer|response|choice)'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
        
    def is_section_header(self, line):
        return self.detect_section_header(line) is not None

    def detect_section_header(self, line):
        patterns = {
            r'^(?:#+\s+)?PART\s+([A-Z])\s*[:.]?\s*(.+)$': {'type': 'part', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?SECTION\s+(\d+)\s*[:.]?\s*(.+)$': {'type': 'section', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(?:Module|Segment)\s+([A-Z\d])\s*[:.]?\s*(.+)$': {'type': 'module', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(I+|[IVXLCDM]+)\s*\.\s*(.+)$': {'type': 'roman_section', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(?:DEMOGRAPHIC|BACKGROUND|PERSONAL)\s+(?:INFORMATION|DATA)\b': {'type': 'demographic', 'group': 0, 'title': 0},
            r'^(?:#+\s+)?(?:Introduction|Purpose|Background)\s*$': {'type': 'intro', 'group': 0, 'title': 0},
            r'^(?:#+\s+)?(?:Additional\s+Comments|Optional\s+Contact\s+Information)\s*$': {'type': 'meta', 'group': 0, 'title': 0}
        }
        
        for pattern, config in patterns.items():
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                if config['group'] == 0:
                    return {
                        'type': config['type'],
                        'identifier': '',
                        'title': match.group(0).strip()
                    }
                return {
                    'type': config['type'],
                    'identifier': match.group(config['group']),
                    'title': match.group(config['title']).strip()
                }
        return None

    def detect_question(self, line):
        patterns = [
            (r'^\s*(?:Q|Question)\s*(\d+)[:.)]\s+(.+)$', 'standard'),
            (r'^\s*(\d+)[:.)]\s+(.+[?])(?:\s*\(.+\))?$', 'numbered'),
            (r'^\s*Item\s+(\d+)[:.)]\s+(.+)$', 'item'),
            (r'^\s*([A-Z])[:.)]\s+(.+[?])(?:\s*\(.+\))?$', 'lettered'),
            (r'^\s*(\d+[a-z]?)[:.)]\s+(.+)$', 'subnumbered'),
            (r'^\s*(\d+\.\d+)[:.)]\s+(.+)$', 'decimal'),
            # Enhanced patterns for unnumbered questions
            (r'^\s*([A-Z][A-Za-z\s?!,\'()\-]+[?])\s*$', 'unnumbered'),
            (r'^\s*([A-Z][a-z]+(?:\s+[A-Z]?[a-z]+){1,15}[?])\s*$', 'unnumbered'),
            (r'^\s*(.+\?)\s*(?:\([^)]+\))?\s*$', 'unnumbered'),
            # Pattern for questions ending in colon (common in surveys)
            (r'^\s*([A-Z].+[:])\s*$', 'unnumbered')
        ]
        
        for pattern, q_type in patterns:
            match = re.match(pattern, line) # Removed IGNORECASE to respect capitalization for unnumbered
            if match:
                # For unnumbered, the group is 1 (text), for others it's 2 (text)
                if q_type == 'unnumbered':
                    question_text = match.group(1).strip()
                    number = ''
                else:
                    question_text = match.group(2).strip()
                    number = match.group(1)
                
                # Determine question type from content
                inferred_type = self._infer_question_type_from_text(question_text)
                
                return {
                    'number': number,
                    'text': question_text,
                    'format': q_type,
                    'type': inferred_type
                }
        return None
    
    def _infer_question_type_from_text(self, text):
        """Infer question type from text content"""
        text_lower = text.lower()
        
        if any(phrase in text_lower for phrase in ['rate', 'on a scale', '1 to 5', '1-5']):
            return 'scale'
        elif any(phrase in text_lower for phrase in ['select all', 'all that apply', 'choose all']):
            return 'multiple_select'
        elif any(phrase in text_lower for phrase in ['select one', 'choose one', 'circle one']):
            return 'single_select'
        elif text_lower.endswith('?'):
            # Ends with question mark but not scale-related
            return 'multiple_choice'
        else:
            return 'open_ended'
    
    def detect_option(self, line):
        patterns = [
            (r'^\s*([a-zA-Z])\s*[.)]\s*(.+)$', 'lettered', None),
            (r'^\s*\(([a-zA-Z])\)\s+(.+)$', 'parenthesized', None),
            (r'^\s*(\d+)\s*[.)]\s*(.+)$', 'numbered', None),
            (r'^\s*\[[\s]\](?:\s*([a-zA-Z0-9])(?:[.)]|\s+))?\s*(.+)$', 'checkbox', 'checkbox'),
            (r'^\s*\[[xX]\]\s*(.+)$', 'checked', 'checked'),
            (r'^\s*(?:□|☐)\s*(.+)$', 'checkbox_symbol', 'checkbox'),
            (r'^\s*(?:○|◯)\s*(.+)$', 'radio_empty', 'radio'),
            (r'^\s*(?:●|◉)\s*(.+)$', 'radio_selected', 'selected'),
            # Enhanced Option Patterns
            (r'^\s*\$\s*([\d,]+)\s*[-–]\s*\$\s*([\d,]+)\s*$', 'range_currency', 'radio'),
            (r'^\s*([\d.%]+)\s*[-–]\s*([\d.%]+)\s*$', 'range_numeric', 'radio'),
            (r'^\s*(?:Less than|More than|Above|Below)\s+.+$', 'range_descriptive', 'radio'),
            (r'^\s*(?:Yes|No),\s+[A-Za-z].+$', 'yes_no_extended', 'radio'),
            (r'^\s*(?:We|I|Our)\s+[A-Za-z].+$', 'first_person', 'radio'),
            (r'^\s*(?:Much|Somewhat|Slightly|Significantly)\s+(?:worse|better)\s*$', 'comparative', 'radio'),
            (r'^\s*About the same\s*$', 'comparative', 'radio')
        ]
        
        for pattern, opt_type, opt_symbol in patterns:
            match = re.match(pattern, line)
            if match:
                groups = match.groups()
                if len(groups) == 2:
                    label, text = groups[0], groups[1]
                    # For ranges, combine groups into text
                    if opt_type in ['range_currency', 'range_numeric']:
                        text = line.strip()
                        label = ''
                elif len(groups) == 1:
                    label, text = '', groups[0]
                else:
                    # No capturing groups, use the whole line/match as text
                    label, text = '', match.group(0).strip()
                
                # For descriptive/comparative, ensure we use the full text if needed
                if opt_type in ['range_descriptive', 'yes_no_extended', 'first_person', 'comparative']:
                    text = line.strip()
                
                return {
                    'label': label.strip() if label else '',
                    'text': text.strip(),
                    'format': opt_type,
                    'type': opt_symbol or 'text'
                }
        return None
    
    def detect_scale(self, line):
        # Likert scale detection
        likert_pattern = r'(?:Strongly\s+Disagree|Disagree|Neutral|Agree|Strongly\s+Agree)'
        if re.search(likert_pattern, line, re.IGNORECASE):
            matches = re.findall(likert_pattern, line, re.IGNORECASE)
            if len(matches) >= 3:
                return {
                    'type': 'likert',
                    'items': matches,
                    'min': 1,
                    'max': 5,
                    'description': 'Likert scale'
                }
        
        # Numeric scale detection
        numeric_pattern = r'(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)'
        match = re.search(numeric_pattern, line)
        if match:
            return {
                'type': 'numeric',
                'items': [match.group(i) for i in range(1, 6)],
                'min': int(match.group(1)),
                'max': int(match.group(5)),
                'description': 'Numeric scale'
            }
        
        # Rating scale detection
        rating_pattern = r'(?:Poor|Fair|Average|Good|Excellent)'
        if re.search(rating_pattern, line, re.IGNORECASE):
            matches = re.findall(rating_pattern, line, re.IGNORECASE)
            if len(matches) >= 3:
                return {
                    'type': 'rating',
                    'items': matches,
                    'min': 1,
                    'max': len(matches),
                    'description': 'Rating scale'
                }
        
        return None


class DocumentProcessor:
    """Process documents using pattern engine"""
    
    def __init__(self):
        self.engine = PatternEngine()
        self.image_extractor = ImageExtractor()
        self.extracted_images = []  # Store extracted images
        self.cover_page_handler = CoverPageHandler()  # Cover page detection
        self.cover_page_data = None  # Extracted cover page data
        self.cover_page_end_index = 0  # Where cover page ends
        self.certification_handler = CertificationPageHandler()  # Certification page detection
        self.certification_data = None  # Extracted certification page data
        self.certification_start_index = 0  # Where certification page starts
        self.certification_end_index = 0  # Where certification page ends
        self.questionnaire_processor = QuestionnaireProcessor() # Questionnaire detection
        self.questionnaire_data = None # Extracted questionnaire data
        self.heading_numberer = HeadingNumberer()  # Auto-number headings based on chapter context
        
    def process_docx(self, file_path):
        """Process Word document line by line, preserving table and image positions"""
        doc = Document(file_path)
        
        # Step 0: Extract all images from document FIRST
        self.extracted_images = self.image_extractor.extract_all_images(file_path)
        logger.info(f"Extracted {len(self.extracted_images)} images from document")
        
        # Step 0.5: COVER PAGE DETECTION - Extract cover page info from first page only
        has_cover_page, cover_data, cover_end_idx = self.cover_page_handler.detect_and_extract(doc)
        if has_cover_page:
            self.cover_page_data = cover_data
            self.cover_page_end_index = cover_end_idx
            logger.info(f"Cover page detected! Skipping first {cover_end_idx} paragraphs")
        
        # Step 0.6: CERTIFICATION PAGE DETECTION - After cover page
        search_start = self.cover_page_end_index if has_cover_page else 0
        has_cert, cert_data, cert_start, cert_end = self.certification_handler.detect_and_extract(doc, search_start)
        if has_cert:
            self.certification_data = cert_data
            self.certification_start_index = cert_start
            self.certification_end_index = cert_end
            logger.info(f"Certification page detected! From {cert_start} to {cert_end}")
        
        # Step 0.7: QUESTIONNAIRE DETECTION
        # Check if the document is a questionnaire
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        q_data = self.questionnaire_processor.detect_questionnaire(full_text)
        if q_data['is_questionnaire']:
            self.questionnaire_data = self.questionnaire_processor.parse_questionnaire_structure(full_text)
            logger.info(f"Questionnaire detected! Title: {self.questionnaire_data['title']}")
            # We will process this differently in WordGenerator
        
        # Create image position lookup for tracking
        image_positions = {}
        for img in self.extracted_images:
            if img['position_type'] == 'paragraph':
                key = img['paragraph_index']
                if key not in image_positions:
                    image_positions[key] = []
                image_positions[key].append(img)
        
        # Extract all content in document order (paragraphs and tables)
        lines = []
        paragraph_index = 0
        
        # Iterate through document body elements in order
        for element in doc.element.body:
            # Check if element is a paragraph
            if element.tag.endswith('p'):
                # Find the corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        # SKIP COVER PAGE PARAGRAPHS - they will be replaced with standardized cover
                        if has_cover_page and paragraph_index < cover_end_idx:
                            paragraph_index += 1
                            break
                        
                        # SKIP CERTIFICATION PAGE PARAGRAPHS - they will be replaced with standardized certification
                        if has_cert and self.certification_start_index <= paragraph_index < self.certification_end_index:
                            paragraph_index += 1
                            break
                        
                        text = para.text.strip()
                        
                        # Check for automatic numbering/bullets (Word automatic lists)
                        # If present, prepend a bullet so PatternEngine detects it as a list
                        try:
                            if para._element.pPr is not None and para._element.pPr.numPr is not None:
                                # Check if text already has a bullet-like start (manual numbering)
                                if text and not re.match(r'^[\s•○●▪■□◆◇→➔➜➤➢–—*⁎⁑※✱✲✳✴☐☑✓✔✗✘⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿➀-➉✦✧★☆♥♡◉◎▸▹►◂◃◄⦿⁍-]', text):
                                    text = f"• {text}"
                        except Exception:
                            pass
                        
                        # Check if this paragraph has images (skip cover page images)
                        if paragraph_index in image_positions:
                            for img in image_positions[paragraph_index]:
                                # Insert image placeholder
                                lines.append({
                                    'text': f'[IMAGE:{img["image_id"]}]',
                                    'style': 'Image',
                                    'bold': False,
                                    'font_size': 12,
                                    'type': 'image_placeholder',
                                    'image_id': img['image_id'],
                                })
                                logger.info(f"Added image placeholder for {img['image_id']} at paragraph {paragraph_index}")
                        
                        if text:
                            font_size = 12  # Default
                            is_bold = False
                            if para.runs:
                                is_bold = any(run.bold for run in para.runs if run.bold)
                                if para.runs[0].font.size:
                                    font_size = para.runs[0].font.size.pt
                            
                            lines.append({
                                'text': text,
                                'style': para.style.name if para.style else 'Normal',
                                'bold': is_bold,
                                'font_size': font_size,
                            })
                        
                        paragraph_index += 1
                        break
            
            # Check if element is a table
            elif element.tag.endswith('tbl'):
                # Find the corresponding table object
                table_index = 0
                for table in doc.tables:
                    if table._element is element:
                        lines.append({'text': '[TABLE START]', 'style': 'Table', 'bold': False, 'font_size': 12})
                        
                        # Check for images in table cells
                        for row_idx, row in enumerate(table.rows):
                            row_cells = []
                            for cell_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                
                                # Check for images in this cell
                                for img in self.extracted_images:
                                    if img['position_type'] == 'table':
                                        loc = img['table_location']
                                        if (loc['table_index'] == table_index and 
                                            loc['row_index'] == row_idx and 
                                            loc['cell_index'] == cell_idx):
                                            cell_text = f'[IMAGE:{img["image_id"]}] {cell_text}'
                                
                                row_cells.append(cell_text)
                            
                            row_text = ' | '.join(row_cells)
                            lines.append({'text': f'| {row_text} |', 'style': 'Table', 'bold': False, 'font_size': 12})
                        
                        lines.append({'text': '[TABLE END]', 'style': 'Table', 'bold': False, 'font_size': 12})
                        break
                    table_index += 1
        
        return self.process_lines(lines), self.extracted_images
    
    def process_text(self, text):
        """Process plain text (no images in plain text)"""
        # FIRST: Apply document-wide spacing cleanup
        text = self.engine.clean_document_spacing(text)
        
        # SECOND: Apply short document processing (TOC removal, key point emphasis)
        text = self.engine.process_short_document(text)
        
        lines = [{'text': line, 'style': 'Normal', 'bold': False, 'font_size': 12} 
                 for line in text.split('\n')]
        return self.process_lines(lines), []  # No images in plain text
    
    def process_lines(self, lines):
        """Core line-by-line processing"""
        analyzed = []
        stats = {
            'total_lines': len(lines),
            'headings': 0,
            'paragraphs': 0,
            'references': 0,
            'tables': 0,
            'lists': 0,
            'definitions': 0,
            'figures': 0,
            'images': 0,  # Add image count
            'h1_count': 0,
            'h2_count': 0,
            'h3_count': 0,
            # New pattern stats
            'inline_formatting': 0,
            'page_metadata': 0,
            'academic_metadata': 0,
            'math_expressions': 0,
            'footnotes': 0,
        }
        
        # Reset heading numberer for new document
        self.heading_numberer.reset()
        
        # Analyze each line
        for i, line_data in enumerate(lines):
            text = line_data['text'] if isinstance(line_data, dict) else line_data
            
            prev_line = lines[i-1]['text'] if i > 0 and isinstance(lines[i-1], dict) else (lines[i-1] if i > 0 else '')
            next_line = lines[i+1]['text'] if i < len(lines) - 1 and isinstance(lines[i+1], dict) else (lines[i+1] if i < len(lines) - 1 else '')
            
            if isinstance(prev_line, dict):
                prev_line = prev_line.get('text', '')
            if isinstance(next_line, dict):
                next_line = next_line.get('text', '')
            
            # Check for image placeholder FIRST (before pattern analysis)
            if isinstance(line_data, dict) and line_data.get('type') == 'image_placeholder':
                analysis = {
                    'type': 'image_placeholder',
                    'text': text,
                    'image_id': line_data.get('image_id'),
                    'confidence': 1.0,
                }
                analyzed.append(analysis)
                stats['images'] = stats.get('images', 0) + 1
                continue
            
            # Build context for dissertation-specific detection
            context = {
                'prev_was_chapter': False,
                'prev_front_matter': None,
                'prev_was_author_header': False,  # For cover page author name detection
            }
            
            # Check if previous line was a chapter heading or author header
            if i > 0 and analyzed:
                prev_analysis = analyzed[-1]
                if prev_analysis.get('type') == 'chapter_heading':
                    context['prev_was_chapter'] = True
                elif prev_analysis.get('type') == 'front_matter_heading':
                    context['prev_front_matter'] = prev_analysis.get('front_matter_type')
                elif prev_analysis.get('type') == 'cover_page_author_header':
                    context['prev_was_author_header'] = True
            
            analysis = self.engine.analyze_line(
                text, 
                i, 
                prev_line, 
                next_line,
                context=context
            )
            
            # SAFETY CHECK: Ensure analysis is a dictionary
            if not isinstance(analysis, dict):
                logger.error(f"analyze_line returned non-dict for line {i}: {type(analysis)}")
                # Create a fallback analysis
                analysis = {
                    'type': 'paragraph',
                    'content': text,
                    'line_num': i,
                    'confidence': 0.0
                }
            
            # Enhance with original formatting if available
            if isinstance(line_data, dict):
                analysis['original_style'] = line_data.get('style', 'Normal')
                analysis['original_bold'] = line_data.get('bold', False)
                analysis['original_font_size'] = line_data.get('font_size', 12)
            
            # Apply automatic heading numbering for chapter-based content
            if analysis['type'] in ['heading', 'heading_hierarchy', 'chapter_heading', 'chapter_title']:
                # Use the heading numberer to apply hierarchical numbering
                heading_level = analysis.get('level', 2)
                number_result = self.heading_numberer.number_heading(text, target_level=heading_level)
                
                # If heading was renumbered, update the analysis
                if number_result['was_renumbered']:
                    analysis['original_text'] = text  # Preserve original
                    analysis['text'] = number_result['numbered']
                    analysis['heading_number'] = number_result['number']
                    # Also update 'content' which is used by _structure_document
                    clean_numbered = re.sub(r'^#+\s*', '', number_result['numbered']).strip()
                    analysis['content'] = clean_numbered
                    logger.debug(f"Auto-numbered heading: '{text}' -> '{number_result['numbered']}'")
                
                # Store chapter context info
                analysis['chapter'] = number_result['chapter']
                analysis['in_appendix'] = self.heading_numberer.in_appendix
            elif analysis['type'] == 'chapter_heading':
                # Detect chapter and update numberer state (even if not renumbered)
                self.heading_numberer.number_heading(text)
            
            analyzed.append(analysis)
            
            # Update stats
            if analysis['type'] == 'heading':
                stats['headings'] += 1
                if analysis.get('level') == 1:
                    stats['h1_count'] += 1
                elif analysis.get('level') == 2:
                    stats['h2_count'] += 1
                elif analysis.get('level') == 3:
                    stats['h3_count'] += 1
            elif analysis['type'] == 'paragraph':
                stats['paragraphs'] += 1
            elif analysis['type'] == 'reference':
                stats['references'] += 1
            elif analysis['type'] in ['table_start', 'table_caption']:
                stats['tables'] += 1
            elif 'list' in analysis['type']:
                stats['lists'] += 1
            elif analysis['type'] == 'definition':
                stats['definitions'] += 1
            elif analysis['type'] == 'figure':
                stats['figures'] += 1
            # New pattern type stats
            elif analysis['type'] == 'inline_formatting':
                stats['inline_formatting'] += 1
            elif analysis['type'] == 'page_metadata':
                stats['page_metadata'] += 1
            elif analysis['type'] == 'academic_metadata':
                stats['academic_metadata'] += 1
            elif analysis['type'] == 'math_expression':
                stats['math_expressions'] += 1
            elif analysis['type'] == 'footnote_endnote':
                stats['footnotes'] += 1
            # New pattern type stats (December 30, 2025 - 20 Academic Patterns)
            elif analysis['type'] == 'heading_hierarchy':
                stats['headings'] += 1
                level = analysis.get('level', 1)
                if level == 1:
                    stats['h1_count'] += 1
                elif level == 2:
                    stats['h2_count'] += 1
                elif level >= 3:
                    stats['h3_count'] += 1
            elif analysis['type'] == 'academic_table':
                stats['tables'] += 1
            elif analysis['type'] == 'list_nested':
                stats['lists'] += 1
            elif analysis['type'] == 'figure_equation':
                stats['figures'] += 1
            elif analysis['type'] in ['citation_inline', 'reference_apa']:
                stats['references'] += 1
            elif analysis['type'] == 'appendix_format':
                stats['headings'] += 1
            elif analysis['type'] in ['block_quote', 'math_model', 'statistical_result']:
                stats['paragraphs'] += 1
            elif analysis['type'] == 'toc_entry':
                pass  # Don't count TOC entries as paragraphs
            elif analysis['type'] in ['footnote_marker', 'questionnaire', 'glossary_entry']:
                stats['definitions'] += 1
            elif analysis['type'] == 'caption_format':
                if 'table' in analysis.get('subtype', ''):
                    stats['tables'] += 1
                elif 'figure' in analysis.get('subtype', ''):
                    stats['figures'] += 1
            # Dissertation-specific pattern stats (December 30, 2025)
            elif analysis['type'] == 'chapter_heading':
                stats['headings'] += 1
                stats['h1_count'] += 1
            elif analysis['type'] == 'chapter_title':
                stats['headings'] += 1  # Count as heading
            elif analysis['type'] == 'front_matter_heading':
                stats['headings'] += 1
                stats['h1_count'] += 1
            elif analysis['type'] == 'copyright_content':
                stats['paragraphs'] += 1  # Count as paragraph
            elif analysis['type'] == 'signature_line':
                pass  # Don't count signature lines in stats
            # Short document pattern stats (December 30, 2025)
            elif analysis['type'] == 'key_point':
                stats['paragraphs'] += 1  # Key points are paragraphs with emphasis
            elif analysis['type'] == 'assignment_header_field':
                stats['paragraphs'] += 1  # Header fields are metadata-like paragraphs
        
        # Structure the document
        structured = self._structure_document(analyzed)
        
        print("DEBUG: Returning dict from process_docx")
        return {
            'analyzed': analyzed,
            'structured': structured,
            'stats': stats,
        }
    
    def _structure_document(self, analyzed):
        """Group lines into logical sections"""
        sections = []
        current_section = None
        current_list = None
        current_table = None
        pending_table_caption = ''  # Track caption before table_start
        in_references = False
        in_code_block = False
        
        for line in analyzed:
            # SAFETY CHECK: Ensure line is a dictionary
            if not isinstance(line, dict):
                logger.warning(f"Skipping non-dict line in structure_document: {type(line)}")
                continue

            if line['type'] == 'empty':
                continue
            
            # Detect reference section
            if line['type'] == 'heading' and line['level'] == 1:
                content_lower = line['content'].lower()
                if 'reference' in content_lower or 'bibliography' in content_lower or 'works cited' in content_lower:
                    in_references = True
                else:
                    in_references = False
            
            # Handle code blocks
            if line['type'] == 'code':
                in_code_block = not in_code_block
                continue
            
            # Handle headings - start new section
            if line['type'] == 'heading':
                # Save current list/table if any
                if current_list and current_section:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table and current_section:
                    current_section['content'].append(current_table)
                    current_table = None
                
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'type': 'section',
                    'heading': line['content'],
                    'level': line['level'],
                    'content': [],
                    'needs_page_break': line.get('needs_page_break', False),  # Propagate page break flag
                    'should_center': line.get('should_center', False),  # Propagate centering flag
                    'is_references_section': in_references,  # Track if this is a references section
                }
                continue
            
            # Initialize first section if no heading found
            if current_section is None:
                current_section = {
                    'type': 'section',
                    'heading': 'Document',
                    'level': 1,
                    'content': [],
                }
            
            # Handle references
            if in_references and line['type'] == 'reference':
                # Close current table first if exists (preserve position)
                if current_table and current_table.get('rows'):
                    current_section['content'].append(current_table)
                    current_table = None
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                })
                continue
            
            # Handle lists
            if 'list' in line['type']:
                list_type = 'bullet_list' if 'bullet' in line['type'] else 'numbered_list'
                if not current_list or current_list['type'] != list_type:
                    # Save previous list if exists
                    if current_list:
                        current_section['content'].append(current_list)
                    current_list = {
                        'type': list_type,
                        'items': [],
                    }
                
                # Store full line info if available, otherwise just content
                if 'bullet_info' in line:
                    current_list['items'].append(line)
                else:
                    current_list['items'].append(line['content'])
                continue
            else:
                # End current list
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
            
            # Handle tables
            if line['type'] == 'table_start':
                # If we already have a current_table with a caption (from table_caption),
                # preserve that caption. Otherwise create new table.
                if current_table and current_table.get('caption'):
                    # Keep the caption, just reset rows
                    current_table['rows'] = []
                else:
                    current_table = {
                        'type': 'table',
                        'caption': pending_table_caption if pending_table_caption else '',
                        'rows': [],
                    }
                    pending_table_caption = ''  # Clear after use
                continue
            
            if line['type'] == 'table_caption':
                # Store caption for the upcoming table
                pending_table_caption = line['content']
                if current_table:
                    current_table['caption'] = line['content']
                else:
                    current_table = {
                        'type': 'table',
                        'caption': line['content'],
                        'rows': [],
                    }
                continue
            
            if line['type'] == 'table_row' and current_table is not None:
                current_table['rows'].append(line.get('cells', []))
                continue
            
            # Handle table row when no table exists yet - initialize new table
            if line['type'] == 'table_row' and current_table is None:
                current_table = {
                    'type': 'table',
                    'caption': '',
                    'rows': [],
                }
                current_table['rows'].append(line.get('cells', []))
                continue
            
            if line['type'] == 'table_separator':
                # Skip separator rows, they're just formatting
                continue
            
            if line['type'] == 'table_end':
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                continue
            
            # Close current table if we encounter non-table content
            # (Tables should be positioned where they appear in the document)
            if line['type'] not in ['table_start', 'table_caption', 'table_row', 'table_separator', 'table_end', 'academic_table']:
                if current_table and current_table.get('rows'):
                    current_section['content'].append(current_table)
                    current_table = None
            
            # Handle definitions
            if line['type'] == 'definition':
                current_section['content'].append({
                    'type': 'definition',
                    'term': line.get('term', ''),
                    'definition': line.get('definition', ''),
                })
                continue
            
            # Handle figures
            if line['type'] == 'figure':
                current_section['content'].append({
                    'type': 'figure',
                    'caption': line['content'],
                })
                continue
            
            # Handle image placeholders (extracted images from DOCX)
            if line['type'] == 'image_placeholder':
                current_section['content'].append({
                    'type': 'image_placeholder',
                    'image_id': line.get('image_id'),
                    'text': line.get('text', ''),
                })
                continue
            
            # Handle quotes
            if line['type'] == 'quote':
                current_section['content'].append({
                    'type': 'quote',
                    'text': line['content'],
                })
                continue
            
            # Handle equations
            if line['type'] == 'equation':
                current_section['content'].append({
                    'type': 'equation',
                    'label': line['content'],
                })
                continue
            
            # Handle page metadata (headers/footers/page numbers)
            if line['type'] == 'page_metadata':
                # Page metadata is typically excluded from main content
                # but we can add it to a special metadata section if needed
                current_section['content'].append({
                    'type': 'page_metadata',
                    'subtype': line.get('subtype', 'metadata'),
                    'text': line['content'],
                })
                continue
            
            # Handle academic metadata (author, affiliation, etc.)
            if line['type'] == 'academic_metadata':
                current_section['content'].append({
                    'type': 'academic_metadata',
                    'subtype': line.get('subtype', 'metadata'),
                    'text': line['content'],
                })
                continue
            
            # Handle mathematical expressions
            if line['type'] == 'math_expression':
                current_section['content'].append({
                    'type': 'math_expression',
                    'subtype': line.get('subtype', 'inline_math'),
                    'text': line['content'],
                })
                continue
            
            # Handle footnotes/endnotes
            if line['type'] == 'footnote_endnote':
                if line.get('subtype') == 'section_header':
                    # Start new footnote section
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'type': 'section',
                        'heading': line['content'],
                        'level': 2,
                        'content': [],
                    }
                else:
                    current_section['content'].append({
                        'type': 'footnote_endnote',
                        'subtype': line.get('subtype', 'footnote_entry'),
                        'text': line['content'],
                    })
                continue
            
            # Handle inline formatting
            if line['type'] == 'inline_formatting':
                current_section['content'].append({
                    'type': 'inline_formatting',
                    'formatting': line.get('formatting', {}),
                    'text': line['content'],
                })
                continue
            
            # ============================================================
            # NEW PATTERN STRUCTURING - December 30, 2025 (20 Academic Patterns)
            # ============================================================
            
            # Handle heading hierarchy (markdown-style)
            if line['type'] == 'heading_hierarchy':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'type': 'section',
                    'heading': line.get('content', line['content']),
                    'level': line.get('level', 1),
                    'content': [],
                    'needs_page_break': line.get('needs_page_break', False),  # Propagate page break flag
                    'should_center': line.get('should_center', False),  # Propagate centering flag
                }
                continue
            
            # Handle academic table
            if line['type'] == 'academic_table':
                subtype = line.get('subtype', 'data_row')
                if subtype == 'caption':
                    if current_table:
                        current_section['content'].append(current_table)
                    current_table = {
                        'type': 'table',
                        'caption': line['content'],
                        'rows': [],
                        'has_header': False,
                    }
                elif subtype == 'header_row':
                    if not current_table:
                        current_table = {'type': 'table', 'caption': '', 'rows': [], 'has_header': True}
                    cells = [c.strip().strip('*') for c in line['content'].split('|') if c.strip()]
                    current_table['rows'].append(cells)
                    current_table['has_header'] = True
                elif subtype == 'data_row':
                    if not current_table:
                        current_table = {'type': 'table', 'caption': '', 'rows': [], 'has_header': False}
                    cells = [c.strip() for c in line['content'].split('|') if c.strip()]
                    current_table['rows'].append(cells)
                # Skip separator rows
                continue
            
            # Handle nested lists
            if line['type'] == 'list_nested':
                list_type = 'bullet_list' if line.get('subtype') == 'checkbox' else 'nested_list'
                if not current_list or current_list['type'] != list_type:
                    if current_list:
                        current_section['content'].append(current_list)
                    current_list = {
                        'type': list_type,
                        'items': [],
                    }
                current_list['items'].append({
                    'text': line['content'].lstrip(' \t•-*□☐☑✓✗'),
                    'indent_level': line.get('indent_level', 0),
                    'is_checkbox': line.get('subtype') == 'checkbox',
                })
                continue
            
            # Handle figure/equation
            if line['type'] == 'figure_equation':
                current_section['content'].append({
                    'type': 'figure_equation',
                    'subtype': line.get('subtype', 'figure_caption'),
                    'text': line['content'],
                })
                continue
            
            # Handle inline citations
            if line['type'] == 'citation_inline':
                current_section['content'].append({
                    'type': 'citation_inline',
                    'text': line['content'],
                    'citation_count': line.get('citation_count', 1),
                })
                continue
            
            # Handle appendix format
            if line['type'] == 'appendix_format':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'type': 'appendix',
                    'heading': line['content'],
                    'level': line.get('level', 1),
                    'content': [],
                }
                continue
            
            # Handle block quotes
            if line['type'] == 'block_quote':
                current_section['content'].append({
                    'type': 'block_quote',
                    'text': line.get('content', line['content']),
                })
                continue
            
            # Handle math models
            if line['type'] == 'math_model':
                current_section['content'].append({
                    'type': 'math_model',
                    'subtype': line.get('subtype', 'statistical_notation'),
                    'text': line['content'],
                })
                continue
            
            # Handle text emphasis (inline - add to current paragraph or as standalone)
            if line['type'] == 'text_emphasis':
                current_section['content'].append({
                    'type': 'text_emphasis',
                    'subtype': line.get('subtype', 'bold'),
                    'text': line['content'],
                })
                continue
            
            # Handle APA references
            if line['type'] == 'reference_apa':
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                    'format': 'apa',
                    'subtype': line.get('subtype', 'standard_reference'),
                })
                continue
            
            # Handle TOC entries
            if line['type'] == 'toc_entry':
                current_section['content'].append({
                    'type': 'toc_entry',
                    'text': line['content'],
                    'page_number': line.get('page_number', None),
                })
                continue
            
            # Handle footnote markers
            if line['type'] == 'footnote_marker':
                current_section['content'].append({
                    'type': 'footnote_marker',
                    'subtype': line.get('subtype', 'footnote_reference'),
                    'text': line['content'],
                })
                continue
            
            # Handle abbreviations
            if line['type'] == 'abbreviation':
                current_section['content'].append({
                    'type': 'abbreviation',
                    'text': line['content'],
                })
                continue
            
            # Handle caption formatting
            if line['type'] == 'caption_format':
                current_section['content'].append({
                    'type': 'caption_format',
                    'subtype': line.get('subtype', 'caption'),
                    'text': line['content'],
                })
                continue
            
            # Handle page breaks
            if line['type'] == 'page_break':
                current_section['content'].append({
                    'type': 'page_break',
                })
                continue
            
            # Handle statistical results
            if line['type'] == 'statistical_result':
                current_section['content'].append({
                    'type': 'statistical_result',
                    'text': line['content'],
                    'stats_types': line.get('stats_types', []),
                })
                continue
            
            # Handle questionnaire items
            if line['type'] == 'questionnaire':
                current_section['content'].append({
                    'type': 'questionnaire',
                    'subtype': line.get('subtype', 'question_item'),
                    'text': line['content'],
                })
                continue
            
            # Handle glossary entries
            if line['type'] == 'glossary_entry':
                current_section['content'].append({
                    'type': 'glossary_entry',
                    'term': line.get('term', ''),
                    'definition': line.get('definition', line['content']),
                })
                continue
            
            # Handle cross-references
            if line['type'] == 'cross_reference':
                current_section['content'].append({
                    'type': 'cross_reference',
                    'text': line['content'],
                    'reference_types': line.get('reference_types', []),
                })
                continue
            
            # Handle running headers
            if line['type'] == 'running_header':
                current_section['content'].append({
                    'type': 'running_header',
                    'text': line['content'],
                })
                continue
            
            # ============================================================
            # DISSERTATION-SPECIFIC STRUCTURING - December 30, 2025
            # ============================================================
            
            # Handle chapter headings
            if line['type'] == 'chapter_heading':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                
                # Create chapter section with heading and optional title
                chapter_heading = line['content']
                chapter_title = line.get('chapter_title')
                
                current_section = {
                    'type': 'chapter',
                    'heading': chapter_heading,
                    'chapter_num': line.get('chapter_num'),
                    'chapter_title': chapter_title,  # May be None
                    'level': 1,
                    'content': [],
                    'needs_page_break': True,
                    'should_center': True,
                }
                continue
            
            # Handle chapter title (follows chapter heading)
            if line['type'] == 'chapter_title':
                # Add title to current chapter section if it exists
                if current_section and current_section.get('type') == 'chapter':
                    if not current_section.get('chapter_title'):
                        current_section['chapter_title'] = line['content']
                else:
                    # Standalone chapter title (treat as heading)
                    current_section['content'].append({
                        'type': 'chapter_title',
                        'text': line['content'],
                        'should_center': True,
                    })
                continue
            
            # Handle front matter headings (Declaration, Certification, etc.)
            if line['type'] == 'front_matter_heading':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                
                front_matter_type = line.get('front_matter_type', 'unknown')
                current_section = {
                    'type': 'front_matter',
                    'heading': line['content'],
                    'front_matter_type': front_matter_type,
                    'level': 1,
                    'content': [],
                    'needs_page_break': True,
                    'should_center': True,
                }
                continue
            
            # Handle copyright content
            if line['type'] == 'copyright_content':
                current_section['content'].append({
                    'type': 'copyright_content',
                    'text': line['content'],
                    'should_center': True,
                })
                continue
            
            # Handle signature lines
            if line['type'] == 'signature_line':
                current_section['content'].append({
                    'type': 'signature_line',
                    'text': line['content'],
                })
                continue
            
            # Handle TOC entries
            if line['type'] == 'toc_entry':
                current_section['content'].append({
                    'type': 'toc_entry',
                    'text': line['content'],
                })
                continue

            # Handle key points (treated as emphasized paragraphs)
            if line['type'] == 'key_point':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                    'is_key_point': True,
                    'key_point_type': line.get('key_point_type'),
                })
                continue

            # Handle assignment header fields
            if line['type'] == 'assignment_header_field':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                    'is_header_field': True,
                })
                continue

            # Handle paragraphs
            if line['type'] == 'paragraph':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                })
                continue
            
            # Handle regular references outside reference section
            if line['type'] == 'reference':
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                })
                continue
        
        # Add remaining list/table if any
        if current_list and current_section:
            current_section['content'].append(current_list)
        if current_table and current_section:
            current_section['content'].append(current_table)
        
        # Add final section
        if current_section:
            sections.append(current_section)
        
        return sections


class WordGenerator:
    """Generate formatted Word documents with image support"""
    
    # Path to cover page logo
    COVER_LOGO_PATH = os.path.join(os.path.dirname(__file__), 'coverpage_template', 'cover_logo.png')
    
    def __init__(self):
        self.doc = None
        self.images = []  # Extracted images
        self.image_lookup = {}  # image_id -> image_data
        self.image_inserter = None
        self.toc_entries = []  # Track headings for TOC generation
        self.toc_placeholder_index = None  # Index where TOC should be inserted
        self.heading_numberer = HeadingNumberer()  # For numbering headings
        self.figure_formatter = FigureFormatter()  # For figure detection and formatting
        self.table_formatter = TableFormatter()  # For table detection and formatting
        self.figure_entries = []  # Track figures for List of Figures
        self.table_entries = []  # Track tables for List of Tables
        self.has_figures = False  # Whether document contains figures
        self.has_tables = False  # Whether document contains tables
        self.use_continuous_arabic = False # Whether to use continuous Arabic numbering
        
    def _set_page_numbering(self, section, fmt='decimal', start=None):
        """Set page numbering format and start value for a section."""
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        
        pgNumType.set(qn('w:fmt'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))
            
    def _add_page_number_to_footer(self, section):
        """Add page number to the footer of a section."""
        # Ensure footer exists
        footer = section.footer
        # Use existing paragraph or create new one
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p.clear() # Clear existing content
        else:
            p = footer.add_paragraph()
            
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        
        # Add PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
    def generate(self, structured_data, output_path, images=None, cover_page_data=None, certification_data=None, questionnaire_data=None):
        """Generate Word document from structured data with images
        
        Args:
            structured_data: List of structured sections
            output_path: Path to save the output document
            images: List of extracted images
            cover_page_data: Dict with extracted cover page information (or None)
            certification_data: Dict with extracted certification page information (or None)
            questionnaire_data: Dict with extracted questionnaire structure (or None)
        """
        self.doc = Document()
        
        # Set default margins to 1 inch (2.54 cm)
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Store images for insertion
        if images:
            self.images = images
            self.image_lookup = {img['image_id']: img for img in images}
            
        # Initialize image inserter
        self.image_inserter = ImageInserter(self.doc, self.images)
        
        # Initialize figure/table formatters
        self.figure_formatter = FigureFormatter()
        self.table_formatter = TableFormatter()
        
        # Check if this is a questionnaire document
        if questionnaire_data and questionnaire_data.get('is_questionnaire'):
            logger.info("Generating questionnaire document...")
            self.doc = format_questionnaire_in_word(self.doc, questionnaire_data)
            self.doc.save(output_path)
            return output_path
            
        # --- STANDARD DOCUMENT GENERATION ---
        
        # 1. Generate Cover Page (if data available)
        if cover_page_data:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        self._setup_styles()
        
        # Define custom style for body content to ensure consistency
        try:
            styles = self.doc.styles
            if 'AcademicBody' not in styles:
                style = styles.add_style('AcademicBody', WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                pf = style.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.line_spacing = 1.5
                pf.space_after = Pt(0)
                pf.left_indent = Pt(0)
                pf.first_line_indent = Pt(0)
        except Exception as e:
            logger.warning(f"Could not create AcademicBody style: {e}")
        
        # Store images for reinsertion
        if images:
            self.images = images
            self.image_lookup = {img['image_id']: img for img in images}
            self.image_inserter = ImageInserter(self.doc, images)
            logger.info(f"WordGenerator initialized with {len(images)} images")
        
        # INSERT COVER PAGE FIRST if detected
        if cover_page_data:
            self._create_cover_page(cover_page_data)
            logger.info("Cover page created from extracted data")
        
        # INSERT CERTIFICATION PAGE AFTER COVER PAGE if detected
        if certification_data:
            self._create_certification_page(certification_data, cover_page_data)
            logger.info("Certification page created from extracted data")
        
        # Handle case where structured_data might be nested lists
        flat_structured_data = []
        for item in structured_data:
            if isinstance(item, dict):
                flat_structured_data.append(item)
            elif isinstance(item, list):
                flat_structured_data.extend([x for x in item if isinstance(x, dict)])
        structured_data = flat_structured_data if flat_structured_data else structured_data
        
        # Calculate if we need TOC (based on total content size)
        total_chars = 0
        for s in structured_data:
            if isinstance(s, dict):
                total_chars += len(str(s.get('heading', '')))
                content = s.get('content', [])
                if isinstance(content, list):
                    for c in content:
                        if isinstance(c, dict):
                            total_chars += len(str(c.get('text', '')))
        estimated_pages = total_chars / 2000  # ~2000 chars per page
        needs_toc = estimated_pages >= 10
        
        # Count preliminary pages/sections to determine numbering style
        prelim_count = 0
        if cover_page_data: prelim_count += 1
        if certification_data: prelim_count += 1
        if needs_toc: prelim_count += 1  # TOC itself counts as a page
        prelim_count += sum(1 for s in structured_data if s.get('type') == 'front_matter')
        
        # Determine numbering style
        has_preliminary = bool(cover_page_data or certification_data or needs_toc)
        
        # Override: If TOC is present but prelims are few (< 3), use Arabic throughout
        if needs_toc and prelim_count < 3:
            self.use_continuous_arabic = True
        elif not has_preliminary:
            self.use_continuous_arabic = True
        else:
            self.use_continuous_arabic = False
            
        # Set initial page numbering
        first_section = self.doc.sections[0]
        if self.use_continuous_arabic:
            self._set_page_numbering(first_section, fmt='decimal', start=1)
        else:
            self._set_page_numbering(first_section, fmt='lowerRoman', start=1)
        self._add_page_number_to_footer(first_section)
        
        # Initialize TOC entries and heading numberer BEFORE processing any sections
        self.toc_entries = []
        self.heading_numberer = HeadingNumberer()
        self.figure_formatter = FigureFormatter()  # Reset figure formatter
        self.table_formatter = TableFormatter()  # Reset table formatter
        self.figure_entries = []  # Reset figure entries
        self.table_entries = []  # Reset table entries
        self.has_figures = False  # Reset figure flag
        self.has_tables = False  # Reset table flag
        
        # Scan content for figures and tables to determine if LOF/LOT is needed
        for section in structured_data:
            for item in section.get('content', []):
                item_type = item.get('type', '')
                
                # Check for various figure-related types
                if item_type == 'image_placeholder':
                    # Only count if it has a caption
                    if item.get('caption'):
                        self.has_figures = True
                elif item_type == 'figure':
                    # Only count if it has a caption
                    if item.get('caption'):
                        self.has_figures = True
                elif item_type in ('figure_caption', 'figure_equation'):
                    self.has_figures = True
                
                # Check for table-related types
                if item_type == 'table':
                    # Only count if it has a caption (tables usually do, but check to be safe)
                    if item.get('caption') or item.get('title'):
                        self.has_tables = True
                elif item_type == 'table_caption':
                    self.has_tables = True
                
                # Also check paragraph text for figure and table captions
                if item_type == 'paragraph':
                    text = item.get('text', '')
                    if self.figure_formatter.is_figure_caption(text):
                        self.has_figures = True
                    if self.table_formatter.is_table_caption(text):
                        self.has_tables = True
            # Break early if both found
            if self.has_figures and self.has_tables:
                break
        
        # Track first section for TOC if it's a chapter
        first_section_for_toc = None
        
        # Add TOC placeholder if needed (uses Word's built-in TOC field)
        added_toc_break = False
        if needs_toc:
            # Add TOC field - will be updated by Word after saving
            self._add_toc_placeholder()
            added_toc_break = True
            
            # Add List of Figures after TOC if document has figures
            if self.has_figures:
                self._add_lof_placeholder()
                added_toc_break = True
            
            # Add List of Tables after LOF if document has tables
            if self.has_tables:
                self._add_lot_placeholder()
                added_toc_break = True
        
        # Add all sections
        rendered_section_count = 0
        for i, section in enumerate(structured_data):
            # Skip "Document" title section if it appears
            if section.get('heading', '').strip().lower() == 'document':
                continue

            # If we just added a TOC/LOF/LOT break, and this is the first RENDERED section,
            # we should suppress the section's page break to avoid double breaks.
            if rendered_section_count == 0 and added_toc_break and section.get('needs_page_break', False):
                # Create a copy to avoid modifying original data
                section = section.copy()
                section['needs_page_break'] = False
            
            self._add_section(section)
            rendered_section_count += 1
        
        # Save document first
        self.doc.save(output_path)
        logger.info(f"Document saved to {output_path}")
        
        # Update TOC using Microsoft Word COM automation
        if needs_toc:
            toc_updated = update_toc_with_word(output_path)
            if toc_updated:
                logger.info("Table of Contents updated automatically")
            else:
                logger.warning("TOC could not be auto-updated - user will need to update manually in Word")
        
        return output_path
    
    def _create_cover_page(self, data):
        """
        Create a standardized cover page with ALL MANDATORY textboxes.
        
        Layout (all elements are MANDATORY):
        1. University name at TOP (THE UNIVERSITY OF BAMENDA)
        2. Faculty (LEFT) | Logo (CENTER) | Department (RIGHT) - 3-column header
        3. TOPIC BOX - visible thick black outline, square, centered, ALL CAPS
        4. Submission statement
        5. BY label
        6. NAME BOX - no visible borders, centered
        7. Registration number
        8. SUPERVISOR BOXES - two columns (Supervisor | Co-Supervisor/Field Supervisor)
        9. DATE BOX - mandatory, centered at bottom, default "2025/2026" if not found
        """
        from docx.shared import Inches, Pt, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Default date if not found
        DEFAULT_DATE = '2025/2026'
        
        def remove_cell_borders(cell):
            """Remove all borders from a table cell."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'nil')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)
        
        def add_thick_black_borders(cell, width_pt=2):
            """Add thick black borders to a table cell (for topic box)."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            # Width in eighths of a point (24 = 3pt, 16 = 2pt, 12 = 1.5pt)
            border_width = str(int(width_pt * 8))
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'single')
                border_elem.set(qn('w:sz'), border_width)
                border_elem.set(qn('w:color'), '000000')
                border_elem.set(qn('w:space'), '0')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)
        
        def add_centered_text(doc, text, font_size=12, bold=False, italic=False, 
                             space_before=0, space_after=6, all_caps=False):
            """Add centered paragraph text."""
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(space_before)
            para.paragraph_format.space_after = Pt(space_after)
            
            display_text = text.upper() if all_caps else text
            run = para.add_run(display_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            return para
        
        def create_textbox_table(doc, text, has_border=True, font_size=14, bold=True, 
                                 all_caps=True, width_inches=5.5, padding_pt=12):
            """Create a centered table that acts as a textbox."""
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Inches(width_inches)
            
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            if has_border:
                add_thick_black_borders(cell, width_pt=2)
            else:
                remove_cell_borders(cell)
            
            # Add padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), str(int(padding_pt * 20)))  # Twips
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            display_text = text.upper() if all_caps else text
            run = para.add_run(display_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.bold = bold
            
            return table
        
        # ========== 1. UNIVERSITY NAME AT TOP (Above everything) ==========
        add_centered_text(self.doc, data.get('university', 'THE UNIVERSITY OF BAMENDA'), 
                         font_size=16, bold=True, space_after=6, all_caps=True)
        
        # ========== 2. FACULTY (LEFT) + LOGO (CENTER) + DEPARTMENT (RIGHT) ==========
        faculty_text = data.get('faculty') or 'HIGHER INSTITUTE OF COMMERCE AND MANAGEMENT'
        dept_text = data.get('department') or 'MANAGEMENT'
        
        # Create a 3-column table: Faculty | Logo | Department
        header_table = self.doc.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        
        # Set column widths
        header_table.columns[0].width = Inches(2.5)  # Faculty
        header_table.columns[1].width = Inches(1.5)  # Logo
        header_table.columns[2].width = Inches(2.5)  # Department
        
        # Remove borders from all cells
        for cell in header_table.rows[0].cells:
            remove_cell_borders(cell)
        
        # LEFT CELL: Faculty/School text
        faculty_cell = header_table.cell(0, 0)
        faculty_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        faculty_para = faculty_cell.paragraphs[0]
        faculty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        faculty_run = faculty_para.add_run(faculty_text.upper())
        faculty_run.bold = True
        faculty_run.font.name = 'Times New Roman'
        faculty_run.font.size = Pt(11)
        
        # CENTER CELL: Logo
        logo_cell = header_table.cell(0, 1)
        logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(self.COVER_LOGO_PATH):
            try:
                logo_run = logo_para.add_run()
                logo_run.add_picture(self.COVER_LOGO_PATH, width=Inches(1.2))
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")
        
        # RIGHT CELL: Department
        dept_cell = header_table.cell(0, 2)
        dept_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        dept_para = dept_cell.paragraphs[0]
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_run = dept_para.add_run(f"DEPARTMENT OF {dept_text.upper()}")
        dept_run.bold = True
        dept_run.font.name = 'Times New Roman'
        dept_run.font.size = Pt(11)
        
        # Spacing after header
        self.doc.add_paragraph().paragraph_format.space_after = Pt(18)
        
        # ========== 3. TOPIC BOX (MANDATORY - thick black outline, square, centered) ==========
        topic_text = data.get('topic') or 'RESEARCH TOPIC'
        create_textbox_table(self.doc, topic_text, has_border=True, font_size=13, 
                            bold=True, all_caps=True, width_inches=5.5, padding_pt=15)
        
        # Spacing after topic
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # ========== 4. SUBMISSION STATEMENT ==========
        degree_text = data.get('degree') or 'Master'
        dept_for_statement = (data.get('department') or 'Management').title()
        faculty_for_statement = (data.get('faculty') or 'Higher Institute of Commerce and Management').title()
        
        submission_text = (
            f"A Dissertation submitted to the Department of {dept_for_statement}, "
            f"{faculty_for_statement} of The University of Bamenda in partial fulfillment "
            f"of the requirements for the award of a {degree_text} Degree."
        )
        add_centered_text(self.doc, submission_text, font_size=11, italic=True, space_after=12)
        
        # ========== 5. "BY" LABEL ==========
        add_centered_text(self.doc, 'BY', font_size=12, bold=True, space_before=6, space_after=6)
        
        # ========== 6. NAME BOX (MANDATORY - no visible borders) ==========
        name_text = data.get('name') or 'STUDENT NAME'
        create_textbox_table(self.doc, name_text, has_border=False, font_size=14, 
                            bold=True, all_caps=True, width_inches=4.0, padding_pt=8)
        
        # ========== 7. REGISTRATION NUMBER ==========
        reg_num = data.get('registration_number') or ''
        if reg_num:
            add_centered_text(self.doc, f"({reg_num})", font_size=11, space_before=3, space_after=12)
        else:
            self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # ========== 8. SUPERVISOR BOXES (MANDATORY - two columns) ==========
        supervisor = data.get('supervisor') or ''
        co_supervisor = data.get('co_supervisor') or data.get('field_supervisor') or ''
        
        # Create a 2-column table for supervisors (always show both columns)
        sup_table = self.doc.add_table(rows=2, cols=2)
        sup_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sup_table.autofit = False
        
        sup_table.columns[0].width = Inches(3.0)
        sup_table.columns[1].width = Inches(3.0)
        
        # Remove borders from all cells
        for row in sup_table.rows:
            for cell in row.cells:
                remove_cell_borders(cell)
        
        # Row 0: Headers (always show)
        sup_header_cell = sup_table.cell(0, 0)
        sup_header_para = sup_header_cell.paragraphs[0]
        sup_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sup_h_run = sup_header_para.add_run('SUPERVISOR')
        sup_h_run.bold = True
        sup_h_run.underline = True
        sup_h_run.font.name = 'Times New Roman'
        sup_h_run.font.size = Pt(11)
        
        cosup_header_cell = sup_table.cell(0, 1)
        cosup_header_para = cosup_header_cell.paragraphs[0]
        cosup_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Determine co-supervisor label
        if data.get('field_supervisor'):
            cosup_label = 'FIELD SUPERVISOR'
        else:
            cosup_label = 'CO-SUPERVISOR'
        cosup_h_run = cosup_header_para.add_run(cosup_label)
        cosup_h_run.bold = True
        cosup_h_run.underline = True
        cosup_h_run.font.name = 'Times New Roman'
        cosup_h_run.font.size = Pt(11)
        
        # Row 1: Names (show placeholder if empty)
        sup_name_cell = sup_table.cell(1, 0)
        sup_name_para = sup_name_cell.paragraphs[0]
        sup_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sup_display = supervisor if supervisor else '________________'
        sup_n_run = sup_name_para.add_run(sup_display)
        sup_n_run.font.name = 'Times New Roman'
        sup_n_run.font.size = Pt(11)
        
        cosup_name_cell = sup_table.cell(1, 1)
        cosup_name_para = cosup_name_cell.paragraphs[0]
        cosup_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cosup_display = co_supervisor if co_supervisor else '________________'
        cosup_n_run = cosup_name_para.add_run(cosup_display)
        cosup_n_run.font.name = 'Times New Roman'
        cosup_n_run.font.size = Pt(11)
        
        # Spacing before date
        self.doc.add_paragraph().paragraph_format.space_after = Pt(24)
        
        # ========== 9. DATE BOX (MANDATORY - centered at bottom) ==========
        month_year = data.get('month_year') or DEFAULT_DATE
        create_textbox_table(self.doc, month_year.upper(), has_border=False, font_size=12, 
                            bold=True, all_caps=True, width_inches=2.5, padding_pt=6)
        
        # Add page break after cover page
        self.doc.add_page_break()
        logger.info("Cover page created with all mandatory elements")
    
    def _create_certification_page(self, cert_data, cover_data=None):
        """
        Create a standardized certification page.
        
        Layout:
        1. CERTIFICATION header (same format as dissertation headings)
        2. Certification text with topic (untouched, not bolded)
        3. Signature textboxes: Supervisor (left) | HOD (right)
        4. Director signature (left-aligned)
        5. Acceptance statement
        6. Date line
        7. General Coordinator section
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.table import _Cell
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        def remove_cell_borders(cell):
            """Remove all borders from a table cell."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'nil')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)
        
        def create_signature_textbox(doc, name, title, width_inches=2.8):
            """Create a textbox with signature line, name, and title."""
            table = doc.add_table(rows=3, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.columns[0].width = Inches(width_inches)
            
            # Remove borders from all cells
            for row in table.rows:
                for cell in row.cells:
                    remove_cell_borders(cell)
            
            # Row 0: Signature line
            line_cell = table.cell(0, 0)
            line_para = line_cell.paragraphs[0]
            line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            line_para.paragraph_format.space_before = Pt(18)
            line_para.paragraph_format.space_after = Pt(3)
            line_run = line_para.add_run('_' * 25)
            line_run.font.name = 'Times New Roman'
            line_run.font.size = Pt(12)
            
            # Row 1: Name
            name_cell = table.cell(1, 0)
            name_para = name_cell.paragraphs[0]
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_para.paragraph_format.space_before = Pt(3)
            name_para.paragraph_format.space_after = Pt(0)
            name_display = name if name else '________________'
            name_run = name_para.add_run(name_display)
            name_run.font.name = 'Times New Roman'
            name_run.font.size = Pt(12)
            name_run.bold = True
            
            # Row 2: Title
            title_cell = table.cell(2, 0)
            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(6)
            title_run = title_para.add_run(f'({title})')
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(11)
            title_run.italic = True
            
            return table
        
        # Get data from certification extraction or fall back to cover page data
        topic = cert_data.get('topic') or (cover_data.get('topic') if cover_data else None) or '[RESEARCH TOPIC]'
        author = cert_data.get('author') or (cover_data.get('name') if cover_data else None) or '[AUTHOR NAME]'
        degree = cert_data.get('degree') or (cover_data.get('degree') if cover_data else None) or "Master's in Business Administration (MBA)"
        program = cert_data.get('program') or (cover_data.get('department') if cover_data else None) or 'Management and Entrepreneurship'
        supervisor = cert_data.get('supervisor') or (cover_data.get('supervisor') if cover_data else None)
        hod = cert_data.get('head_of_department')
        director = cert_data.get('director')
        institution = cert_data.get('institution') or 'The Higher Institute of Commerce and Management of The University of Bamenda'
        
        # ========== 1. CERTIFICATION HEADER (same format as dissertation headings) ==========
        # Use Heading 1 style format: Times New Roman, 12pt, bold, centered, black
        cert_heading = self.doc.add_heading('CERTIFICATION', level=1)
        cert_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cert_heading.paragraph_format.space_before = Pt(12)
        cert_heading.paragraph_format.space_after = Pt(6)
        cert_heading.paragraph_format.line_spacing = 1.5
        for run in cert_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # ========== 2. CERTIFICATION TEXT (topic in bold) ==========
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        
        # First part of the paragraph
        run1 = para.add_run('This is to certify that this research titled ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        
        # Topic in bold with quotes
        topic_run = para.add_run(f'"{topic}"')
        topic_run.font.name = 'Times New Roman'
        topic_run.font.size = Pt(12)
        topic_run.font.bold = True
        
        # Rest of the paragraph
        run2 = para.add_run(f' is the original work of {author}. This work is submitted in partial fulfilment of the requirement for the award of a {degree} in {program} in {institution} Cameroon.')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        
        # ========== 3. SIGNATURE TEXTBOXES: Supervisor (left) | HOD (right) ==========
        # Create a 2-column table to hold the two textboxes side by side
        sig_container = self.doc.add_table(rows=1, cols=2)
        sig_container.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_container.autofit = False
        sig_container.columns[0].width = Inches(3.25)
        sig_container.columns[1].width = Inches(3.25)
        
        # Remove borders from container
        for cell in sig_container.rows[0].cells:
            remove_cell_borders(cell)
        
        # LEFT CELL: Supervisor textbox
        sup_cell = sig_container.cell(0, 0)
        
        # Signature line paragraph
        sup_line_para = sup_cell.paragraphs[0]
        sup_line_para.paragraph_format.space_before = Pt(18)
        sup_line_para.paragraph_format.space_after = Pt(3)
        sup_line = sup_line_para.add_run('_' * 15)
        sup_line.font.name = 'Times New Roman'
        sup_line.font.size = Pt(12)
        
        # Name paragraph (directly under line)
        sup_name_para = sup_cell.add_paragraph()
        sup_name_para.paragraph_format.space_before = Pt(0)
        sup_name_para.paragraph_format.space_after = Pt(0)
        sup_name_display = supervisor if supervisor else ''
        sup_name_run = sup_name_para.add_run(sup_name_display)
        sup_name_run.font.name = 'Times New Roman'
        sup_name_run.font.size = Pt(12)
        sup_name_run.bold = True
        
        # Title paragraph
        sup_title_para = sup_cell.add_paragraph()
        sup_title_para.paragraph_format.space_before = Pt(0)
        sup_title_para.paragraph_format.space_after = Pt(6)
        sup_title_run = sup_title_para.add_run('(Supervisor)')
        sup_title_run.font.name = 'Times New Roman'
        sup_title_run.font.size = Pt(11)
        sup_title_run.bold = True
        
        # RIGHT CELL: HOD textbox
        hod_cell = sig_container.cell(0, 1)
        
        # Signature line paragraph
        hod_line_para = hod_cell.paragraphs[0]
        hod_line_para.paragraph_format.space_before = Pt(18)
        hod_line_para.paragraph_format.space_after = Pt(3)
        hod_line = hod_line_para.add_run('_' * 15)
        hod_line.font.name = 'Times New Roman'
        hod_line.font.size = Pt(12)
        
        # Name paragraph (directly under line)
        hod_name_para = hod_cell.add_paragraph()
        hod_name_para.paragraph_format.space_before = Pt(0)
        hod_name_para.paragraph_format.space_after = Pt(0)
        hod_name_display = hod if hod else ''
        hod_name_run = hod_name_para.add_run(hod_name_display)
        hod_name_run.font.name = 'Times New Roman'
        hod_name_run.font.size = Pt(12)
        hod_name_run.bold = True
        
        # Title paragraph
        hod_title_para = hod_cell.add_paragraph()
        hod_title_para.paragraph_format.space_before = Pt(0)
        hod_title_para.paragraph_format.space_after = Pt(6)
        hod_title_run = hod_title_para.add_run('(Head Of Department)')
        hod_title_run.font.name = 'Times New Roman'
        hod_title_run.font.size = Pt(11)
        hod_title_run.bold = True
        
        # ========== 4. DIRECTOR SIGNATURE (left-aligned) ==========
        # Signature line
        dir_line_para = self.doc.add_paragraph()
        dir_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_line_para.paragraph_format.space_before = Pt(18)
        dir_line_para.paragraph_format.space_after = Pt(3)
        dir_line_run = dir_line_para.add_run('_' * 15)
        dir_line_run.font.name = 'Times New Roman'
        dir_line_run.font.size = Pt(12)
        
        # Name (directly under line)
        dir_name_para = self.doc.add_paragraph()
        dir_name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_name_para.paragraph_format.space_before = Pt(0)
        dir_name_para.paragraph_format.space_after = Pt(0)
        dir_name_display = director if director else ''
        dir_name_run = dir_name_para.add_run(dir_name_display)
        dir_name_run.font.name = 'Times New Roman'
        dir_name_run.font.size = Pt(12)
        dir_name_run.bold = True
        
        dir_title_para = self.doc.add_paragraph()
        dir_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_title_para.paragraph_format.space_before = Pt(0)
        dir_title_run = dir_title_para.add_run('(Director)')
        dir_title_run.font.name = 'Times New Roman'
        dir_title_run.font.size = Pt(11)
        dir_title_run.bold = True
        
        # ========== 4. ACCEPTANCE STATEMENT ==========
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        accept_para = self.doc.add_paragraph()
        accept_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        accept_para.paragraph_format.space_before = Pt(12)
        accept_para.paragraph_format.space_after = Pt(12)
        accept_para.paragraph_format.line_spacing = 1.5
        
        accept_run = accept_para.add_run('Having met the stipulated requirements, the dissertation has been accepted by the Postgraduate School')
        accept_run.font.name = 'Times New Roman'
        accept_run.font.size = Pt(12)
        
        # ========== 5. DATE LINE ==========
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.paragraph_format.space_before = Pt(18)
        date_para.paragraph_format.space_after = Pt(6)
        
        date_run = date_para.add_run('Date' + '_' * 25)
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        
        # ========== 6. GENERAL COORDINATOR SECTION ==========
        # Right-aligned signature section
        gc_table = self.doc.add_table(rows=3, cols=2)
        gc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gc_table.autofit = False
        gc_table.columns[0].width = Inches(3.0)
        gc_table.columns[1].width = Inches(3.0)
        
        # Remove borders
        for row in gc_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border_name}')
                    border_elem.set(qn('w:val'), 'nil')
                    tcBorders.append(border_elem)
                tcPr.append(tcBorders)
        
        # Right column only - signature line
        gc_line_cell = gc_table.cell(0, 1).paragraphs[0]
        gc_line_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_line_cell.paragraph_format.space_before = Pt(18)
        gc_line_cell.add_run('_' * 30).font.name = 'Times New Roman'
        
        # Title
        gc_title_cell = gc_table.cell(1, 1).paragraphs[0]
        gc_title_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_title_cell.paragraph_format.space_before = Pt(6)
        gc_title_run = gc_title_cell.add_run('The General Coordinator')
        gc_title_run.font.name = 'Times New Roman'
        gc_title_run.font.size = Pt(12)
        gc_title_run.bold = True
        
        # School
        gc_school_cell = gc_table.cell(2, 1).paragraphs[0]
        gc_school_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_school_run = gc_school_cell.add_run('Postgraduate School')
        gc_school_run.font.name = 'Times New Roman'
        gc_school_run.font.size = Pt(12)
        
        # Add page break after certification page
        self.doc.add_page_break()
        logger.info("Certification page created with all sections")
    
    def _setup_styles(self):
        """Configure document styles"""
        styles = self.doc.styles
        
        # Normal style - NO INDENTATION
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.left_indent = Pt(0)  # No left indent
        normal.paragraph_format.first_line_indent = Pt(0)  # No first line indent
        
        # Heading styles - All use Times New Roman, size 12, line spacing 1.5
        heading_configs = {
            1: {'size': 12, 'bold': True, 'space_before': 12, 'space_after': 6},
            2: {'size': 12, 'bold': True, 'space_before': 12, 'space_after': 6},
            3: {'size': 12, 'bold': True, 'space_before': 12, 'space_after': 6},
        }
        
        for level, config in heading_configs.items():
            try:
                heading = styles[f'Heading {level}']
                heading.font.name = 'Times New Roman'
                heading.font.bold = config['bold']
                heading.font.size = Pt(config['size'])
                heading.font.color.rgb = RGBColor(0, 0, 0)  # Black
                heading.paragraph_format.line_spacing = 1.5
                heading.paragraph_format.space_before = Pt(config['space_before'])
                heading.paragraph_format.space_after = Pt(config['space_after'])
                heading.paragraph_format.left_indent = Pt(0)  # No left indent
                heading.paragraph_format.first_line_indent = Pt(0)  # No first line indent
            except KeyError:
                pass  # Style doesn't exist, skip
    
    def _add_title(self, title_text):
        """Add document title"""
        title = self.doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = 1.5
        
        # Style the title - Times New Roman, size 12, bold, black
        for run in title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_toc_placeholder(self):
        """Add Table of Contents using Microsoft Word's built-in TOC field"""
        # Add TOC heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        toc_heading = self.doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_heading.add_run('TABLE OF CONTENTS')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        toc_heading.paragraph_format.space_after = Pt(0)
        toc_heading.paragraph_format.space_before = Pt(0)
        toc_heading.paragraph_format.line_spacing = 1.5
        
        # Add TOC field code directly (no blank paragraph in between)
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field - this generates the actual Word TOC
        # TOC \o "1-3" - include heading levels 1-3
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view
        # \u - use applied paragraph outline level
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        # This text will be replaced when the user updates the field
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get table of contents")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(12)
        run_instr.font.bold = True
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Single page break after TOC (not two)
        self.doc.add_page_break()
    
    def _add_lof_placeholder(self):
        """Add List of Figures using Microsoft Word's built-in TOC field for figures
        
        This creates a Word field that automatically lists all figure captions
        with their page numbers. Similar to TOC but specifically for figures.
        The field uses SEQ Figure field references for proper figure numbering.
        """
        # Add LOF heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        lof_heading = self.doc.add_paragraph()
        lof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lof_heading.add_run('LIST OF FIGURES')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        lof_heading.paragraph_format.space_after = Pt(0)
        lof_heading.paragraph_format.space_before = Pt(0)
        lof_heading.paragraph_format.line_spacing = 1.5
        
        # Add LOF field code directly
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field for figures
        # TOC \h \z \c "Figure" - Table of contents for Figure captions
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view  
        # \c "Figure" - Build TOC from SEQ Figure field captions
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Figure" '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get list of figures")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(12)
        run_instr.font.bold = True
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Page break after List of Figures
        self.doc.add_page_break()
    
    def _add_lot_placeholder(self):
        """Add List of Tables using Microsoft Word's built-in TOC field for tables
        
        This creates a Word field that automatically lists all table captions
        with their page numbers. Similar to TOC but specifically for tables.
        The field uses SEQ Table field references for proper table numbering.
        """
        # Add LOT heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        lot_heading = self.doc.add_paragraph()
        lot_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lot_heading.add_run('LIST OF TABLES')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        lot_heading.paragraph_format.space_after = Pt(0)
        lot_heading.paragraph_format.space_before = Pt(0)
        lot_heading.paragraph_format.line_spacing = 1.5
        
        # Add LOT field code directly
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field for tables
        # TOC \h \z \c "Table" - Table of contents for Table captions
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view  
        # \c "Table" - Build TOC from SEQ Table field captions
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Table" '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get list of tables")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(12)
        run_instr.font.bold = True
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Page break after List of Tables
        self.doc.add_page_break()
    
    def _add_table_caption(self, number, title, center=False):
        """
        Add a properly formatted table caption with SEQ field for LOT tracking.
        
        Args:
            number: Table number (e.g., "1" or "1.2")
            title: Table caption text
            center: Whether to center the caption (default False - left aligned)
        
        The caption is formatted as:
        - Times New Roman, 12pt
        - Bold
        - Left aligned (by default)
        - Includes SEQ field for Word's List of Tables to pick up
        
        Note: LOT entries are formatted as plain text via update_toc_with_word()
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        
        # Add "Table " text
        run1 = para.add_run('Table ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering and LOT tracking
        run_seq = para.add_run()
        
        # Create SEQ field - this allows Word to track tables for LOT
        # SEQ Table \* ARABIC - Sequential number for Table category
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Table \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # The actual number (will be updated by Word)
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Style the SEQ field run
        run_seq.font.name = 'Times New Roman'
        run_seq.font.size = Pt(12)
        run_seq.font.bold = True
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add colon and title
        run2 = para.add_run(f': {title}')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        return para

    def _add_figure_caption(self, number, title, center=True):
        """
        Add a properly formatted figure caption with SEQ field for LOF tracking.
        
        Args:
            number: Figure number (e.g., "1" or "1.2")
            title: Figure caption text
            center: Whether to center the caption (default True)
        
        The caption is formatted as:
        - Times New Roman, 12pt
        - Italic
        - Centered (by default)
        - Includes SEQ field for Word's List of Figures to pick up
        
        Note: LOF entries are formatted as plain text via update_toc_with_word()
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        
        # Add "Figure " text
        run1 = para.add_run('Figure ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.italic = True
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering and LOF tracking
        run_seq = para.add_run()
        
        # Create SEQ field - this allows Word to track figures for LOF
        # SEQ Figure \* ARABIC - Sequential number for Figure category
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Figure \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # The actual number (will be updated by Word)
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Style the SEQ field run (same approach as table captions)
        run_seq.font.name = 'Times New Roman'
        run_seq.font.size = Pt(12)
        run_seq.font.italic = True
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add ": " separator
        run2 = para.add_run(': ')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add the caption title
        run3 = para.add_run(title)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(12)
        run3.font.italic = True
        run3.font.color.rgb = RGBColor(0, 0, 0)
        
        # Track figure for validation
        self.figure_formatter.add_figure_entry(number, title)
        self.figure_entries.append({
            'number': number,
            'title': title
        })
        
        return para
    
    def _format_existing_figure_caption(self, para, text):
        """
        Reformat an existing figure caption paragraph with proper styling.
        
        Args:
            para: The paragraph object to format
            text: The original caption text
            
        Returns:
            The formatted paragraph
        """
        # Detect figure number and title from text
        figure_info = self.figure_formatter.detect_figure_caption(text)
        
        if not figure_info:
            return para  # Not a figure caption, return unchanged
        
        number = figure_info['number']
        title = figure_info['title']
        
        # Clear existing runs
        for run in para.runs:
            run.text = ''
        
        # Add formatted caption with SEQ field
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        
        # Add "Figure " text
        run1 = para.add_run('Figure ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.italic = True
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering
        run_seq = para.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Figure \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Add ": " and title
        run2 = para.add_run(': ' + title)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Track figure
        self.figure_formatter.add_figure_entry(number, title)
        self.figure_entries.append({
            'number': number,
            'title': title
        })
        
        return para
    
    def _track_heading_for_toc(self, heading_text, level, numbered_text=None):
        """Track a heading for inclusion in the Table of Contents
        
        Args:
            heading_text: The original heading text
            level: The heading level (1, 2, or 3)
            numbered_text: The numbered heading text (e.g., "1.2 Background")
        """
        if level > 3:
            return  # Only track levels 1-3 for TOC
        
        # Clean heading text (remove markdown markers)
        clean_text = re.sub(r'^#+\s*', '', heading_text).strip()
        
        # Skip TOC itself and certain front matter
        skip_headings = ['table of contents', 'contents', 'toc']
        if clean_text.lower() in skip_headings:
            return
        
        self.toc_entries.append({
            'text': clean_text,
            'numbered_text': numbered_text or clean_text,
            'level': level
        })
    
    def _populate_toc(self):
        """Populate the Table of Contents with actual entries by replacing the marker"""
        if not self.toc_entries:
            return
        
        body = self.doc.element.body
        marker_text = '<<<TOC_ENTRIES_PLACEHOLDER>>>'
        marker_element = None
        
        # Find the marker paragraph
        for element in body:
            if element.tag.endswith('p'):
                text = ''.join(node.text or '' for node in element.iter() if node.text)
                if marker_text in text:
                    marker_element = element
                    break
        
        if marker_element is None:
            logger.warning("Could not find TOC marker placeholder")
            return
        
        # Create TOC entry paragraphs and insert them BEFORE the marker
        for entry in self.toc_entries:
            # Create new paragraph element
            new_para = OxmlElement('w:p')
            
            # Create paragraph properties for indentation and spacing
            pPr = OxmlElement('w:pPr')
            
            # Set indentation based on level
            level = entry['level']
            indent = OxmlElement('w:ind')
            if level == 1:
                indent.set(qn('w:left'), '0')
            elif level == 2:
                indent.set(qn('w:left'), '360')  # 0.25 inch in twips
            else:  # level 3
                indent.set(qn('w:left'), '720')  # 0.5 inch in twips
            pPr.append(indent)
            
            # Set line spacing
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '60')  # 3pt in twips
            spacing.set(qn('w:line'), '360')  # 1.5 line spacing
            pPr.append(spacing)
            
            new_para.append(pPr)
            
            # Create run for text
            run = OxmlElement('w:r')
            
            # Run properties (font)
            rPr = OxmlElement('w:rPr')
            
            # Font name
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            
            # Font size (12pt = 24 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '24')
            rPr.append(szCs)
            
            # Bold for level 1
            if level == 1:
                bold = OxmlElement('w:b')
                rPr.append(bold)
            
            # Black color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '000000')
            rPr.append(color)
            
            run.append(rPr)
            
            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = entry['numbered_text']
            run.append(text_elem)
            
            new_para.append(run)
            
            # Insert before marker
            marker_element.addprevious(new_para)
        
        # Remove the marker paragraph
        body.remove(marker_element)
        
        logger.info(f"TOC populated with {len(self.toc_entries)} entries")
    
    def _add_toc(self):
        """Legacy method - now uses placeholder approach"""
        self._add_toc_placeholder()
    
    def _insert_image(self, image_id):
        """
        Insert an image into the document at current position.
        
        Args:
            image_id: The ID of the image to insert
        """
        if not image_id:
            logger.warning("No image_id provided to _insert_image")
            return
        
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found in image lookup")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Not found]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Create paragraph for image with minimal spacing to fit on page
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)  # Reduced spacing
            para.paragraph_format.space_after = Pt(3)   # Reduced spacing
            para.paragraph_format.keep_with_next = True  # Keep with caption if present
            
            # Add image from bytes
            run = para.add_run()
            
            # Create BytesIO stream from image data
            image_stream = BytesIO(img_data['data'])
            
            # Determine width and height - preserve original dimensions
            width = img_data.get('width', 4.0)
            height = img_data.get('height', 3.0)
            
            # Store original for logging
            original_width, original_height = width, height
            
            # Only scale down if exceeds page margins (6.5" width, 9" height for letter)
            max_width = 6.5  # Page width with margins
            max_height = 9.0  # Page height with margins
            
            # Scale proportionally only if too large
            if width > max_width:
                ratio = max_width / width
                width = max_width
                height = height * ratio
            
            if height > max_height:
                ratio = max_height / height
                height = max_height
                width = width * ratio
            
            # Preserve small images at original size (no minimum scaling up)
            # Only set minimum if dimensions are invalid (0 or negative)
            if width <= 0:
                width = original_width if original_width > 0 else 2.0
            if height <= 0:
                height = original_height if original_height > 0 else 1.5
            
            # Add picture to document
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            logger.info(f"Inserted image {image_id} ({width:.2f}x{height:.2f} inches)")
            
            # Add caption if exists
            if img_data.get('caption'):
                caption_para = self.doc.add_paragraph()
                caption_run = caption_para.add_run(img_data['caption'])
                caption_run.italic = True
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = Pt(10)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_before = Pt(0)
                caption_para.paragraph_format.space_after = Pt(6)  # Reduced spacing
            
        except Exception as e:
            logger.error(f"Error inserting image {image_id}: {str(e)}")
            # Add error placeholder
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Error: {str(e)}]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_section(self, section):
        """Add a document section"""
        section_type = section.get('type', 'section')
        
        # Check for Chapter 1 to switch numbering
        heading_text = section.get('heading', '').strip().upper()
        
        # Force page break for specific sections (Resume, Acknowledgements)
        force_break_headings = [
            'RESUME', 'RÉSUMÉ', 'RÉSUME', 'RESUMÉ', 
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT', 'ACKNOWLEDGMENT'
        ]
        # Check if heading matches any of these words
        if any(h in heading_text for h in force_break_headings):
             # Use page_break_before property instead of manual break for better reliability
             section['use_page_break_before'] = True
             section['needs_page_break'] = False # Disable manual break
        
        # Regex for "CHAPTER 1" or "CHAPTER ONE"
        is_chapter_one = bool(re.search(r'^CHAPTER\s+(1|ONE)\b', heading_text))
        
        if is_chapter_one and not self.use_continuous_arabic:
            # Add Section Break (Next Page) to switch to Arabic numbering
            new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
            self._set_page_numbering(new_section, fmt='decimal', start=1)
            new_section.footer.is_linked_to_previous = False
            self._add_page_number_to_footer(new_section)
        elif section.get('needs_page_break', False):
            self.doc.add_page_break()
        
        # Handle chapter sections (dissertation-specific)
        if section_type == 'chapter':
            self._add_chapter_section(section)
            return
        
        # Handle front matter sections (dissertation-specific)
        if section_type == 'front_matter':
            self._add_front_matter_section(section)
            return
        
        # Get heading text and number it
        heading_text = section['heading']
        level = min(section['level'], 3)
        
        # Number the heading using HeadingNumberer
        numbering_result = self.heading_numberer.number_heading(heading_text, target_level=level)
        numbered_heading = numbering_result.get('formatted', heading_text)
        
        # Track heading for TOC
        self._track_heading_for_toc(heading_text, level, numbered_heading)
        
        # Add heading for regular sections (use numbered heading)
        heading = self.doc.add_heading(numbered_heading, level=level)
        
        # Apply forced page break if requested
        if section.get('use_page_break_before'):
            heading.paragraph_format.page_break_before = True
        
        # Check if this heading should be centered
        if section.get('should_center', False):
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ensure heading is bold, black, Times New Roman, 12pt
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add content (pass section context for references handling)
        self._add_section_content(section)

    def _add_chapter_section(self, section):
        """Add a chapter section (dissertation-specific)"""
        # Add chapter number heading (CHAPTER ONE, CHAPTER 1, etc.)
        chapter_heading = section.get('heading', '')
        # Strip markdown heading markers
        clean_heading = re.sub(r'^#+\s*', '', chapter_heading).strip().upper()
        
        # Track chapter for TOC (combine chapter number and title)
        chapter_title = section.get('chapter_title', '')
        if chapter_title:
            toc_entry = f"{clean_heading}: {chapter_title.upper()}"
        else:
            toc_entry = clean_heading
        self._track_heading_for_toc(toc_entry, 1, toc_entry)
        
        heading = self.doc.add_heading(clean_heading, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.5
        
        # Add chapter title if present (centered, bold)
        if chapter_title:
            title_para = self.doc.add_heading(chapter_title.upper(), level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_para.paragraph_format.line_spacing = 1.5
        else:
            heading.paragraph_format.space_after = Pt(0)
        
        # Add content
        self._add_section_content(section)
    
    def _add_front_matter_section(self, section):
        """Add a front matter section (dissertation-specific)"""
        front_matter_type = section.get('front_matter_type', 'unknown')
        heading_text = section.get('heading', '')
        # Strip markdown heading markers
        clean_heading = re.sub(r'^#+\s*', '', heading_text).strip().upper()
        
        # Track front matter for TOC (skip certain types)
        skip_toc_types = ['toc', 'table_of_contents']
        if front_matter_type not in skip_toc_types:
            self._track_heading_for_toc(clean_heading, 1, clean_heading)
        
        # Add centered heading (level 1)
        heading = self.doc.add_heading(clean_heading, level=1)
        
        # Apply forced page break if requested
        if section.get('use_page_break_before'):
            heading.paragraph_format.page_break_before = True
            
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.5
        
        # Add content based on front matter type
        if front_matter_type == 'dedication':
            self._add_dedication_content(section)
        elif front_matter_type == 'declaration':
            self._add_declaration_content(section)
        elif front_matter_type == 'certification':
            self._add_certification_content(section)
        elif front_matter_type == 'acknowledgements':
            self._add_acknowledgements_content(section)
        elif front_matter_type == 'abstract':
            self._add_abstract_content(section)
        elif front_matter_type == 'resume':
            self._add_abstract_content(section)  # Same format as abstract
        elif front_matter_type == 'toc':
            self._add_toc_content(section)
        elif front_matter_type in ('list_of_tables', 'list_of_figures', 'abbreviations', 'glossary'):
            self._add_list_section_content(section)
        else:
            # Default: add content normally
            self._add_section_content(section)
    
    def _add_dedication_content(self, section):
        """Add dedication content (centered, not italic)"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            text = item.get('text', '')
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.italic = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)  # Reduced spacing
            para.paragraph_format.line_spacing = 1.5
    
    def _add_declaration_content(self, section):
        """Add declaration content with signature lines"""
        has_signature = False
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'signature_line':
                has_signature = True
                # Add signature line
                para = self.doc.add_paragraph()
                run = para.add_run(text if text else '________________________')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(36)
                para.paragraph_format.line_spacing = 1.5
            else:
                # Regular paragraph (block paragraph - no indent)
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        
        # Add signature line if not already present
        if not has_signature:
            # Add blank space
            blank = self.doc.add_paragraph()
            blank.paragraph_format.space_before = Pt(36)
            
            # Signature line
            sig_para = self.doc.add_paragraph()
            sig_para.add_run('Signed: ________________________')
            sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig_para.paragraph_format.line_spacing = 1.5
            for run in sig_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            # Date line
            date_para = self.doc.add_paragraph()
            date_para.add_run('Date: ________________________')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(12)
            date_para.paragraph_format.line_spacing = 1.5
            for run in date_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _add_certification_content(self, section):
        """Add certification content with multiple signature lines"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'signature_line':
                para = self.doc.add_paragraph()
                run = para.add_run(text if text else '________________________')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.line_spacing = 1.5
            else:
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        
        # Add multiple signature blocks for committee members
        self.doc.add_paragraph()  # Blank line
        
        signatures = [
            'Supervisor: ________________________',
            'Date: ________________________',
            '',
            'Head of Department: ________________________',
            'Date: ________________________',
        ]
        
        for sig_text in signatures:
            if sig_text:
                para = self.doc.add_paragraph()
                para.add_run(sig_text)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                self.doc.add_paragraph()  # Empty line between signature blocks
    
    def _add_acknowledgements_content(self, section):
        """Add acknowledgements content (block paragraphs - no indent)"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            text = item.get('text', '')
            para = self.doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _add_abstract_content(self, section):
        """Add abstract content with Keywords section"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            # Check for Keywords line
            if 'keywords' in text.lower() or item_type == 'keywords':
                para = self.doc.add_paragraph()
                run = para.add_run('Keywords: ')
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                # Extract keywords (after "Keywords:")
                keywords = re.sub(r'^[Kk]eywords?\s*[:\-]\s*', '', text)
                para.add_run(keywords)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.line_spacing = 1.5
            else:
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    def _add_toc_content(self, section):
        """Add Table of Contents entries with dot leaders"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'toc_entry':
                # Parse TOC entry for title and page number
                # Look for patterns like "Title........12" or "Title   12"
                match = re.match(r'^(.+?)\s*[\.…]+\s*(\d+)\s*$', text)
                if match:
                    title, page = match.groups()
                    para = self.doc.add_paragraph()
                    para.add_run(title.strip())
                    # Add tab with dot leader
                    para.add_run('\t')
                    para.add_run(page)
                    para.paragraph_format.line_spacing = 1.5
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.5
                    
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                para = self.doc.add_paragraph(text)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    def _add_list_section_content(self, section):
        """Add content for List of Tables, List of Figures, Abbreviations, Glossary sections."""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'toc_entry' or '.....' in text or '\t' in text:
                # Parse entry for title and page number (like TOC)
                match = re.match(r'^(.+?)\s*[\.…]+\s*(\d+)\s*$', text)
                if match:
                    title, page = match.groups()
                    para = self.doc.add_paragraph()
                    para.add_run(title.strip())
                    para.add_run('\t')
                    para.add_run(page)
                    para.paragraph_format.line_spacing = 1.5
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.5
            elif item_type == 'abbreviation' or ':' in text:
                # Abbreviation or glossary entry (Term: Definition format)
                para = self.doc.add_paragraph()
                if ':' in text:
                    parts = text.split(':', 1)
                    run = para.add_run(parts[0].strip() + ':')
                    run.bold = True
                    if len(parts) > 1:
                        para.add_run(' ' + parts[1].strip())
                else:
                    para.add_run(text)
                para.paragraph_format.line_spacing = 1.5
            else:
                para = self.doc.add_paragraph(text)
                para.paragraph_format.line_spacing = 1.5
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    def _format_single_apa_reference(self, text):
        """Apply APA formatting rules to a single reference string."""
        # Fix missing space before year parentheses: "UNESCO(2024)" -> "UNESCO (2024)"
        text = re.sub(r'([^\s\(])\((\d{4})', r'\1 (\2', text)
        
        # Fix capitalization in titles (ensure space after colon)
        text = re.sub(r'([A-Z][a-z]+)\s+([A-Z][a-z]+):', r'\1 \2:', text)
        
        # Ensure period after year
        text = re.sub(r'(\(\d{4}(?:/\d{4})?\))\s*([A-Z])', r'\1. \2', text)
        
        # Fix double periods
        text = re.sub(r'\.\.$', '.', text)
        
        # Fix journal formatting (italicize common journal patterns)
        # Matches "Journal of X Y", "International Journal of X", etc.
        text = re.sub(r'(\b(?:International\s+)?Journal\s+of\s+[A-Z][a-z]+(?:\s+[A-Z][a-zA-Z]+)*)', r'*\1*', text)
        
        # Fix "Available at:" formatting (optional, but good for consistency)
        text = re.sub(r'Available\s+at\s*:', 'Available at:', text, flags=re.IGNORECASE)
        
        # Fix lowercase titles (heuristic: if starts with lowercase after year)
        # UNESCO. (2020) global -> UNESCO. (2020) Global
        text = re.sub(r'(\(\d{4}(?:/\d{4})?\)\.?\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text)
        
        return text

    def _add_section_content(self, section):
        """Add content of a section"""
        # Check if this is a references section
        is_references_section = section.get('is_references_section', False)
        heading_lower = section.get('heading', '').lower()
        if 'reference' in heading_lower or 'bibliography' in heading_lower or 'works cited' in heading_lower:
            is_references_section = True
        
        # Collect and sort references if this is a references section
        content_items = list(section.get('content', []))
        if is_references_section:
            # Separate references from other content
            references = [item for item in content_items if item.get('type') == 'reference']
            non_references = [item for item in content_items if item.get('type') != 'reference']
            
            # Sort references alphabetically by text (A-Z)
            references.sort(key=lambda x: x.get('text', '').strip().lower())
            
            # Apply APA formatting to references
            for ref in references:
                ref['text'] = self._format_single_apa_reference(ref['text'])
            
            # Recombine: non-references first, then sorted references
            content_items = non_references + references
        
        for item in content_items:
            # SAFETY CHECK: Ensure item is a dictionary
            if not isinstance(item, dict):
                logger.warning(f"Skipping non-dict item in _add_section_content: {type(item)}")
                continue

            # Handle image placeholders FIRST
            if item.get('type') == 'image_placeholder':
                self._insert_image(item.get('image_id'))
                continue
            
            if item.get('type') == 'paragraph':
                text = item.get('text', '')
                
                # Check if this is a figure caption - format with SEQ field for LOF
                if self.figure_formatter.is_figure_caption(text):
                    figure_info = self.figure_formatter.detect_figure_caption(text)
                    if figure_info:
                        self._add_figure_caption(figure_info['number'], figure_info['title'])
                        self.has_figures = True
                        continue
                
                # Check if this is a table caption - format with SEQ field for LOT
                if self.table_formatter.is_table_caption(text):
                    table_info = self.table_formatter.detect_table_caption(text)
                    if table_info:
                        self._add_table_caption(table_info['number'], table_info['title'])
                        self.has_tables = True
                        continue
                
                para = self.doc.add_paragraph(text, style='AcademicBody')
                # Explicitly set properties again to be absolutely sure
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = 1.5
                # Explicitly set font for all runs to ensure consistency after merge
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'table_caption':
                # Explicit table caption type
                text = item.get('text', '')
                table_info = self.table_formatter.detect_table_caption(text)
                if table_info:
                    self._add_table_caption(table_info['number'], table_info['title'])
                else:
                    # Fallback - try to extract number from text
                    num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                    if num_match:
                        num = num_match.group(1)
                        title = re.sub(r'^(?:Table|Tbl\.?|Tab\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE)
                        self._add_table_caption(num, title)
                    else:
                        # Last resort - just add as bold centered
                        para = self.doc.add_paragraph(text)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                self.has_tables = True
                continue
            
            elif item.get('type') == 'figure_caption':
                # Explicit figure caption type
                text = item.get('text', '')
                figure_info = self.figure_formatter.detect_figure_caption(text)
                if figure_info:
                    self._add_figure_caption(figure_info['number'], figure_info['title'])
                else:
                    # Fallback - try to extract number from text
                    num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                    if num_match:
                        num = num_match.group(1)
                        title = re.sub(r'^(?:Figure|Fig\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE)
                        self._add_figure_caption(num, title)
                    else:
                        # Last resort - just add as italic centered
                        para = self.doc.add_paragraph(text)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.italic = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                self.has_figures = True
                continue
            
            elif item.get('type') == 'definition':
                para = self.doc.add_paragraph()
                run = para.add_run(f"{item.get('term', '')}: ")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if item.get('definition'):
                    run_def = para.add_run(item.get('definition', ''))
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(12)
            
            elif item.get('type') == 'bullet_list':
                # Process items to determine nesting
                list_items = item.get('items', [])
                
                # Check if we have enhanced bullet info in the first item
                has_enhanced_info = len(list_items) > 0 and isinstance(list_items[0], dict) and 'bullet_info' in list_items[0]
                
                if has_enhanced_info:
                    previous_indent = 0
                    
                    for list_item_data in list_items:
                        bullet_info = list_item_data.get('bullet_info', {})
                        content = bullet_info.get('content', list_item_data.get('content', ''))
                        
                        # Determine nesting
                        is_nested, current_indent = is_nested_bullet(
                            bullet_info.get('original_line', ''), 
                            previous_indent
                        )
                        
                        # Map to Word bullet style
                        bullet_char, font_name, bullet_type = map_to_word_bullet_style(bullet_info)
                        
                        # Create paragraph with appropriate style
                        # Use 'List Bullet' for level 1, 'List Bullet 2' for level 2, etc.
                        style_name = 'List Bullet 2' if is_nested else 'List Bullet'
                        
                        try:
                            para = self.doc.add_paragraph(content, style=style_name)
                        except:
                            # Fallback if style doesn't exist
                            para = self.doc.add_paragraph(content, style='List Bullet')
                        
                        # Apply custom indentation (slight indent for all bullets)
                        # Standard indent is usually 0.25 inch hanging. We want to push it in slightly.
                        # Level 1: 0.25 inch left indent
                        # Level 2: 0.5 inch left indent
                        
                        if is_nested:
                            para.paragraph_format.left_indent = Inches(0.5)
                            para.paragraph_format.first_line_indent = Inches(-0.25)
                        else:
                            para.paragraph_format.left_indent = Inches(0.25)
                            para.paragraph_format.first_line_indent = Inches(-0.25)
                        
                        # Apply custom bullet formatting if possible
                        # Note: Word's bullet formatting is complex via python-docx. 
                        # We rely on the style for the bullet character usually.
                        # But we can try to customize the font of the text.
                        
                        # If we want to force the bullet character to be ■, we might need to fake it 
                        # if the style doesn't support it, OR we rely on the user having a template.
                        # However, since we can't easily change the bullet char of a style via python-docx without low-level XML,
                        # we will prepend the character to the text and remove the list style if it's a special bullet.
                        
                        if bullet_char == '■':
                            # Remove list style and manually format
                            para.style = 'Normal'
                            para.clear() # Clear content added by add_paragraph
                            
                            # Add bullet run
                            run_bullet = para.add_run('■\t')
                            run_bullet.font.name = 'Arial'
                            run_bullet.font.size = Pt(12)
                            
                            # Add content run
                            run_text = para.add_run(content)
                            run_text.font.name = 'Times New Roman'
                            run_text.font.size = Pt(12)
                            
                            # Set indentation for manual bullet
                            if is_nested:
                                para.paragraph_format.tab_stops.add_tab_stop(Inches(0.75))
                                para.paragraph_format.left_indent = Inches(0.75)
                                para.paragraph_format.first_line_indent = Inches(-0.25)
                            else:
                                para.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
                                para.paragraph_format.left_indent = Inches(0.5)
                                para.paragraph_format.first_line_indent = Inches(-0.25)
                        else:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                            
                        previous_indent = current_indent
                else:
                    # Legacy handling
                    for list_item in list_items:
                        content = list_item if isinstance(list_item, str) else list_item.get('content', '')
                        para = self.doc.add_paragraph(content, style='List Bullet')
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
            
            elif item.get('type') == 'numbered_list':
                for list_item in item.get('items', []):
                    # Clean up the list item (remove the number prefix if present)
                    clean_item = re.sub(r'^[\d]+[\.\)]\s*', '', list_item)
                    clean_item = re.sub(r'^[a-z][\.\)]\s*', '', clean_item)
                    clean_item = re.sub(r'^\([a-z\d]+\)\s*', '', clean_item)
                    # Handle case where list_item is a dict
                    content_to_add = clean_item if clean_item else (list_item if isinstance(list_item, str) else list_item.get('content', ''))
                    para = self.doc.add_paragraph(content_to_add, style='List Number')
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            elif item.get('type') == 'table':
                self._add_table(item)
            
            elif item.get('type') == 'figure':
                # Figure with caption - use proper SEQ field formatting
                caption = item.get('caption', '')
                if caption:
                    figure_info = self.figure_formatter.detect_figure_caption(caption)
                    if figure_info:
                        self._add_figure_caption(figure_info['number'], figure_info['title'])
                    else:
                        # Try to extract number from caption
                        num_match = re.search(r'(\d+(?:\.\d+)?)', caption)
                        if num_match:
                            num = num_match.group(1)
                            title = re.sub(r'^(?:Figure|Fig\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', caption, flags=re.IGNORECASE)
                            self._add_figure_caption(num, title.strip() if title.strip() else caption)
                        else:
                            # Fallback to simple italic centered
                            para = self.doc.add_paragraph(caption)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.italic = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                self.has_figures = True
            
            elif item.get('type') == 'quote':
                para = self.doc.add_paragraph(item.get('text', ''))
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.right_indent = Inches(0.5)
                for run in para.runs:
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'equation':
                para = self.doc.add_paragraph(item.get('label', ''))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            elif item.get('type') == 'reference':
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                
                # Handle italics markers (*)
                parts = text.split('*')
                for i, part in enumerate(parts):
                    if not part: continue
                    run = para.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if i % 2 == 1:  # Odd parts are between * markers -> italic
                        run.italic = True
                
                # Only apply hanging indent for references in the references section
                if is_references_section:
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.5)
                else:
                    # In-text reference citations - no indent
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = 1.5
            
            # NEW PATTERN RENDERING (December 30, 2025)
            
            elif item.get('type') == 'page_metadata':
                # Page metadata - centered, italic
                para = self.doc.add_paragraph(item.get('text', ''))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.italic = True
            
            elif item.get('type') == 'academic_metadata':
                subtype = item.get('subtype', 'metadata')
                text = item.get('text', '')
                
                if subtype == 'author':
                    # Author names - centered, bold
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = True
                elif subtype == 'affiliation':
                    # Affiliation - normal text, centered
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_after = Pt(6)
                elif subtype == 'contact':
                    # Contact info - centered
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                else:
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
            
            elif item.get('type') == 'math_expression':
                subtype = item.get('subtype', 'inline_math')
                text = item.get('text', '')
                
                # Clean up LaTeX markers for display
                clean_text = re.sub(r'^\$\$|\$\$$', '', text)  # Remove $$
                clean_text = re.sub(r'^\$|\$$', '', clean_text)  # Remove $
                clean_text = re.sub(r'^\\\[|\\\]$', '', clean_text)  # Remove \[ \]
                clean_text = re.sub(r'^\\\(|\\\)$', '', clean_text)  # Remove \( \)
                
                if subtype == 'display_math':
                    # Display math - centered, with spacing
                    para = self.doc.add_paragraph(clean_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_before = Pt(12)
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                else:
                    # Inline math - just add as text
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            elif item.get('type') == 'footnote_endnote':
                subtype = item.get('subtype', 'footnote_entry')
                text = item.get('text', '')
                
                if subtype == 'footnote_entry':
                    # Footnote entry - hanging indent
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.25)
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                else:
                    para = self.doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            elif item.get('type') == 'inline_formatting':
                text = item.get('text', '')
                formatting = item.get('formatting', {})
                
                # Parse the text and apply formatting
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                
                if formatting.get('bold_italic'):
                    # Remove *** or ___ markers and apply both bold and italic
                    clean_text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
                    clean_text = re.sub(r'___(.+?)___', r'\1', clean_text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.italic = True
                elif formatting.get('bold'):
                    # Remove ** or __ markers and apply bold
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                    clean_text = re.sub(r'__(.+?)__', r'\1', clean_text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                elif formatting.get('italic'):
                    # Remove * or _ markers and apply italic
                    clean_text = re.sub(r'\*(.+?)\*', r'\1', text)
                    clean_text = re.sub(r'_(.+?)_', r'\1', clean_text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.italic = True
                else:
                    # Remove all formatting markers as fallback
                    clean_text = re.sub(r'[\*_]{1,3}(.+?)[\*_]{1,3}', r'\1', text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            # ================================================================
            # NEW 20 ACADEMIC PATTERN RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'figure_equation':
                # Figure or equation caption
                subtype = item.get('subtype', 'figure_caption')
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if subtype == 'figure_caption':
                        run.bold = True
            
            elif item.get('type') == 'citation_inline':
                # Inline citation - render as normal paragraph
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'appendix':
                # Appendix section - similar to heading but with distinct style
                text = item.get('heading', item.get('text', ''))
                heading = self.doc.add_heading(text, level=1)
                heading.paragraph_format.line_spacing = 1.5
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            
            elif item.get('type') == 'block_quote':
                # Block quote - indented, italic
                text = item.get('text', '')
                # Remove leading > markers
                clean_text = re.sub(r'^[>\s]+', '', text)
                para = self.doc.add_paragraph(clean_text)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.right_indent = Inches(0.5)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.italic = True
            
            elif item.get('type') == 'math_model':
                # Math model / statistical notation
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'text_emphasis':
                # Text emphasis - bold/italic/underline
                text = item.get('text', '')
                subtype = item.get('subtype', 'bold')
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                # Clean formatting markers
                clean_text = re.sub(r'[\*_~]{1,3}(.+?)[\*_~]{1,3}', r'\1', text)
                run = para.add_run(clean_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if 'bold' in subtype:
                    run.bold = True
                if 'italic' in subtype:
                    run.italic = True
                if 'strike' in subtype:
                    run.font.strike = True
            
            elif item.get('type') == 'toc_entry':
                # Table of contents entry
                text = item.get('text', '')
                page_num = item.get('page_number', '')
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                run = para.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if page_num:
                    para.add_run('\t')
                    run_num = para.add_run(str(page_num))
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = Pt(12)
            
            elif item.get('type') == 'footnote_marker':
                # Footnote marker / reference
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.25)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'abbreviation':
                # Abbreviation definition
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                # Try to extract abbreviation and definition
                match = re.match(r'([A-Z]{2,})\s*[-–:=]\s*(.+)', text)
                if match:
                    abbr, defn = match.groups()
                    run = para.add_run(abbr)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run_def = para.add_run(f' – {defn}')
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(12)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'caption_format':
                # Figure/table caption formatting
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'page_break':
                # Insert page break
                self.doc.add_page_break()
            
            elif item.get('type') == 'statistical_result':
                # Statistical result formatting
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'questionnaire':
                # Questionnaire item
                text = item.get('text', '')
                subtype = item.get('subtype', 'question_item')
                if subtype == 'likert_scale':
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 1.5
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'glossary_entry':
                # Glossary entry - term in bold, definition follows
                term = item.get('term', '')
                definition = item.get('definition', '')
                para = self.doc.add_paragraph()
                if term:
                    run = para.add_run(term)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run_def = para.add_run(': ' + definition if definition else '')
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(12)
                else:
                    run = para.add_run(definition)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.first_line_indent = Inches(-0.25)
            
            elif item.get('type') == 'cross_reference':
                # Cross-reference - render as normal text
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'running_header':
                # Running header - typically in document header
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'nested_list':
                # Nested list with indent levels
                for list_item in item.get('items', []):
                    indent = list_item.get('indent_level', 0)
                    para = self.doc.add_paragraph(list_item.get('text', ''), style='List Bullet')
                    para.paragraph_format.left_indent = Inches(0.25 + (0.25 * indent))
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            # ================================================================
            # DISSERTATION-SPECIFIC CONTENT RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'copyright_content':
                # Copyright content - centered, italic
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = 1.5
            
            elif item.get('type') == 'signature_line':
                # Signature line - right aligned
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                if text:
                    para.add_run(text)
                else:
                    para.add_run('________________________')
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            elif item.get('type') == 'chapter_title':
                # Chapter title - centered, bold (heading level 1)
                text = item.get('text', '')
                para = self.doc.add_heading(text.upper(), level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.5
            
            # ================================================================
            # SHORT DOCUMENT KEY POINT RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'key_point':
                # Key point with emphasis and optional emoji
                text = item.get('text', '')
                key_point_type = item.get('key_point_type', '')
                emoji = item.get('emoji_prefix', '')
                
                para = self.doc.add_paragraph()
                
                # Add emoji prefix if present
                if emoji:
                    para.add_run(emoji)
                
                # Parse and apply formatting (handle markdown-style bold/italic)
                clean_text = text.strip()
                
                # Apply formatting based on key point type
                if key_point_type in ['warning', 'concept', 'exercise', 'learning', 'summary', 'procedure']:
                    # Bold for emphasis
                    run = para.add_run(clean_text)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                elif key_point_type == 'example':
                    # Italic for examples
                    run = para.add_run(clean_text)
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                elif key_point_type == 'definition':
                    # Bold with definition format
                    # Try to extract term and definition
                    if ':' in clean_text:
                        parts = clean_text.split(':', 1)
                        term_run = para.add_run(parts[0] + ':')
                        term_run.bold = True
                        term_run.font.name = 'Times New Roman'
                        term_run.font.size = Pt(12)
                        if len(parts) > 1:
                            defn_run = para.add_run(' ' + parts[1].strip())
                            defn_run.font.name = 'Times New Roman'
                            defn_run.font.size = Pt(12)
                    else:
                        run = para.add_run(clean_text)
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                else:
                    # Default: bold
                    run = para.add_run(clean_text)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
            
            elif item.get('type') == 'assignment_header_field':
                # Assignment header field (Student Name, Course, etc.) - bold
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_after = Pt(3)
    
    def _add_table(self, table_data):
        """Add a table with academic formatting (proper alignment, column sizing)"""
        # Add caption if exists (above table, centered, bold) - WITH SEQ field for LOT
        if table_data.get('caption'):
            caption_text = table_data['caption']
            # Check if this is a table caption and extract number/title
            table_info = self.table_formatter.detect_table_caption(caption_text)
            if table_info:
                # Use proper SEQ field formatting for LOT
                self._add_table_caption(table_info['number'], table_info['title'])
                self.has_tables = True
            else:
                # Fallback - try to extract number from caption
                num_match = re.search(r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)?)', caption_text, re.IGNORECASE)
                if num_match:
                    num = num_match.group(1)
                    title = re.sub(r'^(?:Table|Tbl\.?|Tab\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', caption_text, flags=re.IGNORECASE)
                    self._add_table_caption(num, title.strip() if title.strip() else caption_text)
                    self.has_tables = True
                else:
                    # Last resort - just add as bold centered (no SEQ field)
                    caption = self.doc.add_paragraph()
                    caption_run = caption.add_run(caption_text)
                    caption_run.bold = True
                    caption_run.font.name = 'Times New Roman'
                    caption_run.font.size = Pt(12)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(6)
        
        # Add table
        if table_data.get('rows') and len(table_data['rows']) > 0:
            # Determine number of columns
            num_cols = max(len(row) for row in table_data['rows']) if table_data['rows'] else 1
            num_rows = len(table_data['rows'])
            
            if num_cols > 0 and num_rows > 0:
                # Analyze column content types for alignment
                engine = PatternEngine()
                column_types = engine.get_column_content_types(table_data['rows'])
                
                # Ensure we have types for all columns
                while len(column_types) < num_cols:
                    column_types.append('text')
                
                table = self.doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                # Set table to auto-fit contents
                table.autofit = True
                
                # Calculate column widths based on content length
                # Find max character count for each column
                col_max_lengths = [0] * num_cols
                for row_data in table_data['rows']:
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell_content = str(cell_text).strip() if cell_text else ''
                            # Count all characters (letters, numbers, spaces, etc.)
                            content_length = len(cell_content)
                            if content_length > col_max_lengths[col_idx]:
                                col_max_lengths[col_idx] = content_length
                
                # Calculate proportional widths based on content length
                total_content_length = sum(col_max_lengths) if sum(col_max_lengths) > 0 else 1
                # Total table width (in inches) - standard page width minus margins
                total_table_width = Inches(6.0)
                # Minimum column width
                min_col_width = Inches(0.5)
                
                # Set column widths proportionally
                for col_idx in range(num_cols):
                    # Calculate proportional width
                    proportion = col_max_lengths[col_idx] / total_content_length if total_content_length > 0 else 1 / num_cols
                    col_width = total_table_width * proportion
                    # Ensure minimum width
                    if col_width < min_col_width:
                        col_width = min_col_width
                    # Set width for all cells in this column
                    for row in table.rows:
                        row.cells[col_idx].width = col_width
                
                # Fill table with proper alignment
                for row_idx, row_data in enumerate(table_data['rows']):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:  # Safety check
                            cell = table.rows[row_idx].cells[col_idx]
                            cell_content = str(cell_text).strip() if cell_text else ''
                            cell.text = cell_content
                            
                            # Apply formatting to all paragraphs in cell
                            for paragraph in cell.paragraphs:
                                # Set font for all runs
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                                
                                # Header row (row 0): centered and bold
                                if row_idx == 0:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.bold = True
                                else:
                                    # Data rows: alignment based on column content type
                                    col_type = column_types[col_idx] if col_idx < len(column_types) else 'text'
                                    paragraph.alignment = engine.get_alignment_for_content_type(col_type)
                                
                                # Set line spacing
                                paragraph.paragraph_format.line_spacing = 1.0
                                paragraph.paragraph_format.space_before = Pt(2)
                                paragraph.paragraph_format.space_after = Pt(2)
        
        # Add spacing after table
        spacing = self.doc.add_paragraph()
        spacing.paragraph_format.space_before = Pt(6)


# Flask Routes
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'engine': 'pattern-based',
        'patterns_loaded': 40
    })


@app.route('/upload', methods=['POST'])
def upload_document():
    """Upload and process document"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400
    
    # Generate unique ID
    job_id = str(uuid.uuid4())
    
    # Save uploaded file
    file_ext = os.path.splitext(file.filename)[1].lower()
    input_path = os.path.join(UPLOAD_FOLDER, f"{job_id}{file_ext}")
    file.save(input_path)
    
    # Save metadata
    try:
        metadata = {
            'original_filename': file.filename,
            'upload_time': datetime.now().isoformat()
        }
        with open(os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json"), 'w') as f:
            json.dump(metadata, f)
    except Exception as e:
        logger.error(f"Failed to save metadata: {e}")
    
    try:
        # Process document
        processor = DocumentProcessor()
        images = []  # Extracted images
        
        if file_ext == '.docx':
            # Robust unpacking to handle both dict and tuple returns
            proc_result = processor.process_docx(input_path)
            if isinstance(proc_result, tuple) and len(proc_result) == 2:
                result, images = proc_result
            else:
                result = proc_result
                images = []
            logger.info(f"Processed document with {len(images)} images")
        else:
            # Read as text (txt, md, etc.) - no images in plain text
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            # Robust unpacking for text processing too
            proc_result = processor.process_text(text)
            if isinstance(proc_result, tuple) and len(proc_result) == 2:
                result, images = proc_result
            else:
                result = proc_result
                images = []
        
        # Validate result format
        if not isinstance(result, dict) or 'structured' not in result:
            logger.error(f"Invalid processing result format: {type(result)}")
            return jsonify({'error': 'Document processing failed to produce structured data'}), 500

        # Generate formatted Word document with images
        output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
        generator = WordGenerator()
        
        # Pass cover page data and certification data if detected
        cover_page_data = getattr(processor, 'cover_page_data', None)
        certification_data = getattr(processor, 'certification_data', None)
        generator.generate(
            result['structured'], 
            output_path, 
            images=images,
            cover_page_data=cover_page_data,
            certification_data=certification_data
        )
        
        # Generate preview markdown
        preview = generate_preview_markdown(result['structured'])
        
        # Add image count to stats
        result['stats']['images'] = len(images)
        
        return jsonify({
            'job_id': job_id,
            'stats': result['stats'],
            'structured': result['structured'],
            'preview': preview,
            'download_url': f'/download/{job_id}',
            'status': 'complete',
            'images_preserved': len(images),
        })
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"Error processing document: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up input file
        if os.path.exists(input_path):
            try:
                os.remove(input_path)
            except:
                pass


@app.route('/download/<job_id>', methods=['GET'])
def download_document(job_id):
    """Download formatted document"""
    output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    
    if not os.path.exists(output_path):
        return jsonify({'error': 'Document not found'}), 404
    
    # Try to get original filename from metadata
    download_name = 'formatted_document.docx'
    meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json")
    if os.path.exists(meta_path):
        try:
            with open(meta_path, 'r') as f:
                metadata = json.load(f)
                original_name = metadata.get('original_filename', '')
                if original_name:
                    name, ext = os.path.splitext(original_name)
                    download_name = f"{name}_formatted.docx"
        except Exception as e:
            logger.error(f"Error reading metadata for job {job_id}: {e}")
    
    inline = request.args.get('inline', 'false').lower() == 'true'
    as_attachment = not inline
    
    return send_file(
        output_path,
        as_attachment=as_attachment,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def generate_preview_markdown(structured):
    """Generate markdown preview from structured data"""
    markdown = ''
    
    for section in structured:
        # Add heading
        level = min(section.get('level', 1), 6)
        heading_prefix = '#' * level
        markdown += f"{heading_prefix} {section.get('heading', 'Untitled')}\n\n"
        
        # Add content
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
                
            if item.get('type') == 'paragraph':
                markdown += f"{item.get('text', '')}\n\n"
            
            elif item.get('type') == 'definition':
                markdown += f"**{item.get('term', '')}:** {item.get('definition', '')}\n\n"
            
            elif item.get('type') == 'bullet_list':
                for list_item in item.get('items', []):
                    markdown += f"- {list_item}\n"
                markdown += '\n'
            
            elif item.get('type') == 'numbered_list':
                for idx, list_item in enumerate(item.get('items', []), 1):
                    # Clean up numbering
                    clean_item = re.sub(r'^[\d]+[\.\)]\s*', '', list_item)
                    clean_item = re.sub(r'^[a-z][\.\)]\s*', '', clean_item)
                    markdown += f"{idx}. {clean_item if clean_item else list_item}\n"
                markdown += '\n'
            
            elif item.get('type') == 'table':
                rows = item.get('rows', [])
                if rows:
                    # Header
                    markdown += '| ' + ' | '.join(str(cell) for cell in rows[0]) + ' |\n'
                    markdown += '| ' + ' | '.join(['---'] * len(rows[0])) + ' |\n'
                    # Data rows
                    for row in rows[1:]:
                        markdown += '| ' + ' | '.join(str(cell) for cell in row) + ' |\n'
                    markdown += '\n'
            
            elif item.get('type') == 'reference':
                markdown += f"{item.get('text', '')}\n\n"
            
            elif item.get('type') == 'figure':
                markdown += f"*{item.get('caption', '')}*\n\n"
            
            elif item.get('type') == 'quote':
                markdown += f"> {item.get('text', '')}\n\n"
    
    return markdown


# --- Cover Page Generator Integration ---
from coverpage_generator import generate_cover_page, load_json

@app.route('/api/institutions', methods=['GET'])
def get_institutions():
    """Return list of institutions"""
    try:
        data = load_json('institutions.json')
        if not data:
            logger.error("institutions.json loaded empty or not found")
            return jsonify({"institutions": []}) # Return empty structure at least
        return jsonify(data)
    except Exception as e:
        logger.error(f"Error loading institutions: {e}")
        return jsonify({"institutions": []})

@app.route('/api/courses/search', methods=['GET'])
def search_courses():
    """Search courses by code or title"""
    query = request.args.get('q', '').lower()
    courses = load_json('courses_database.json')
    
    if not query:
        return jsonify([])
        
    matches = [
        c for c in courses 
        if query in c['code'].lower() or query in c['title'].lower()
    ]
    return jsonify(matches[:10]) # Limit to 10 results

@app.route('/api/coverpage/generate', methods=['POST'])
def api_generate_coverpage():
    """Generate cover page from form data"""
    data = request.json
    output_path, error = generate_cover_page(data)
    
    if error:
        return jsonify({'error': error}), 400
        
    # Check for merge request
    merge_job_id = data.get('mergeJobId')
    if merge_job_id:
        try:
            processed_path = os.path.join(OUTPUT_FOLDER, f"{merge_job_id}_formatted.docx")
            if os.path.exists(processed_path):
                logger.info(f"Merging cover page with processed document: {processed_path}")
                
                # Load documents
                cover_doc = Document(output_path)
                processed_doc = Document(processed_path)
                
                # Detect numbering style of processed doc (Section 0)
                # This helps us restore the correct numbering after merge
                processed_start_fmt = 'lowerRoman' # Default
                try:
                    if processed_doc.sections:
                        sectPr = processed_doc.sections[0]._sectPr
                        pgNumType = sectPr.find(qn('w:pgNumType'))
                        if pgNumType is not None:
                            fmt = pgNumType.get(qn('w:fmt'))
                            if fmt:
                                processed_start_fmt = fmt
                except Exception as e:
                    logger.warning(f"Could not detect processed doc numbering: {e}")
                
                # CRITICAL FIX: Ensure all content in processed_doc uses 'AcademicBody' style
                # instead of 'Normal'. This prevents the content from inheriting the Cover Page's
                # 'Normal' style (which might be Calibri/different spacing) during the merge.
                
                # 1. Ensure AcademicBody exists and has CORRECT properties
                academic_style = None
                academic_list_number = None
                academic_list_bullet = None
                
                try:
                    # --- AcademicBody ---
                    if 'AcademicBody' not in processed_doc.styles:
                        academic_style = processed_doc.styles.add_style('AcademicBody', WD_STYLE_TYPE.PARAGRAPH)
                        academic_style.base_style = processed_doc.styles['Normal']
                    else:
                        academic_style = processed_doc.styles['AcademicBody']
                    
                    font = academic_style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    pf = academic_style.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    pf.line_spacing = 1.5
                    pf.space_after = Pt(0)

                    # --- AcademicListNumber ---
                    # We create a custom list style to prevent merging with Cover Page's list styles
                    if 'AcademicListNumber' not in processed_doc.styles:
                        academic_list_number = processed_doc.styles.add_style('AcademicListNumber', WD_STYLE_TYPE.PARAGRAPH)
                        # Try to inherit from List Number if it exists to keep numbering definition
                        if 'List Number' in processed_doc.styles:
                            academic_list_number.base_style = processed_doc.styles['List Number']
                        else:
                            academic_list_number.base_style = processed_doc.styles['Normal']
                    else:
                        academic_list_number = processed_doc.styles['AcademicListNumber']
                    
                    font = academic_list_number.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    pf = academic_list_number.paragraph_format
                    pf.line_spacing = 1.5
                    
                    # --- AcademicListBullet ---
                    if 'AcademicListBullet' not in processed_doc.styles:
                        academic_list_bullet = processed_doc.styles.add_style('AcademicListBullet', WD_STYLE_TYPE.PARAGRAPH)
                        if 'List Bullet' in processed_doc.styles:
                            academic_list_bullet.base_style = processed_doc.styles['List Bullet']
                        else:
                            academic_list_bullet.base_style = processed_doc.styles['Normal']
                    else:
                        academic_list_bullet = processed_doc.styles['AcademicListBullet']
                        
                    font = academic_list_bullet.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    pf = academic_list_bullet.paragraph_format
                    pf.line_spacing = 1.5

                except Exception as e:
                    logger.warning(f"Error setting up Academic styles: {e}")

                if academic_style:
                    def process_para(para):
                        # Fix Normal -> AcademicBody
                        if para.style.name == 'Normal':
                            para.style = academic_style
                        
                        # Fix List Number -> AcademicListNumber
                        elif para.style.name == 'List Number' and academic_list_number:
                            para.style = academic_list_number
                            
                        # Fix List Bullet -> AcademicListBullet
                        elif para.style.name == 'List Bullet' and academic_list_bullet:
                            para.style = academic_list_bullet
                    
                    # Update Body Paragraphs
                    for para in processed_doc.paragraphs:
                        process_para(para)
                    
                    # Update Tables
                    for table in processed_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    process_para(para)
                                    
                    # Update Headers and Footers (All Sections)
                    for section in processed_doc.sections:
                        # Headers
                        for header in [section.header, section.first_page_header, section.even_page_header]:
                            if header:
                                for para in header.paragraphs:
                                    process_para(para)
                                for table in header.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                process_para(para)
                        # Footers
                        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                            if footer:
                                for para in footer.paragraphs:
                                    process_para(para)
                                for table in footer.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                process_para(para)

                # Add a section break to the cover doc to ensure separation of styles/margins
                # This forces the appended content to start on a new page with its own section properties
                new_section = cover_doc.add_section(WD_SECTION.NEW_PAGE)
                
                # Copy section properties from the processed document to the new section
                # This ensures the content retains its original formatting (margins, orientation, etc.)
                if processed_doc.sections:
                    src_section = processed_doc.sections[0]
                    new_section.left_margin = src_section.left_margin
                    new_section.right_margin = src_section.right_margin
                    new_section.top_margin = src_section.top_margin
                    new_section.bottom_margin = src_section.bottom_margin
                    new_section.page_width = src_section.page_width
                    new_section.page_height = src_section.page_height
                    new_section.orientation = src_section.orientation
                    new_section.gutter = src_section.gutter
                    new_section.header_distance = src_section.header_distance
                    new_section.footer_distance = src_section.footer_distance
                    new_section.different_first_page_header_footer = src_section.different_first_page_header_footer
                    # We don't copy start_type as we want NEW_PAGE

                
                # Merge using docxcompose
                composer = Composer(cover_doc)
                composer.append(processed_doc)
                
                # Save merged document
                composer.save(output_path)
                
                # --- POST-MERGE FIX: Restore Page Numbering ---
                # The merge process often breaks page numbering linkage.
                # We need to explicitly unlink the body section from the cover page
                # and restore the page numbering.
                try:
                    merged_doc = Document(output_path)
                    if len(merged_doc.sections) > 1:
                        # Section 0 is Cover Page (no number)
                        # Section 1 is the start of the Body (needs number)
                        body_section = merged_doc.sections[1]
                        
                        # 1. Unlink footer from Cover Page
                        body_section.footer.is_linked_to_previous = False
                        
                        # 2. Restore Page Numbering
                        # Use the format detected from the original processed document
                        sectPr = body_section._sectPr
                        pgNumType = sectPr.find(qn('w:pgNumType'))
                        if pgNumType is None:
                            pgNumType = OxmlElement('w:pgNumType')
                            sectPr.append(pgNumType)
                        
                        # Only set if not already set (preserve existing formatting if possible)
                        if not pgNumType.get(qn('w:fmt')):
                            pgNumType.set(qn('w:fmt'), processed_start_fmt)
                        if not pgNumType.get(qn('w:start')):
                            pgNumType.set(qn('w:start'), '1')
                            
                        # 3. Ensure PAGE field exists in footer
                        footer = body_section.footer
                        # Check if footer is empty
                        if not footer.paragraphs or not footer.paragraphs[0].text.strip():
                            if footer.paragraphs:
                                p = footer.paragraphs[0]
                                p.clear()
                            else:
                                p = footer.add_paragraph()
                            
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            fldChar1 = OxmlElement('w:fldChar')
                            fldChar1.set(qn('w:fldCharType'), 'begin')
                            run._r.append(fldChar1)
                            instrText = OxmlElement('w:instrText')
                            instrText.text = "PAGE"
                            run._r.append(instrText)
                            fldChar2 = OxmlElement('w:fldChar')
                            fldChar2.set(qn('w:fldCharType'), 'end')
                            run._r.append(fldChar2)
                        
                        merged_doc.save(output_path)
                        logger.info("Restored page numbering in merged document")
                except Exception as e:
                    logger.error(f"Error restoring page numbering: {e}")

                logger.info("Merge successful")
            else:
                logger.warning(f"Processed document not found for merge: {processed_path}")
        except Exception as e:
            logger.error(f"Error merging documents: {str(e)}")
            # We continue without merging if it fails, or we could return an error
            # For now, let's just log it and return the cover page
            
    # Return the file
    filename = os.path.basename(output_path)
    return jsonify({
        'success': True,
        'filename': filename,
        'downloadUrl': f'/api/download/{filename}'
    })

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download generated file"""
    # Check in cover pages output first
    # We need to reconstruct the path used in coverpage_generator.py
    # It uses: os.path.join(os.path.dirname(BASE_DIR), 'outputs', 'Cover Pages')
    # BASE_DIR in coverpage_generator is backend/
    # So it is backend/../outputs/Cover Pages -> outputs/Cover Pages relative to root
    
    # Here we are in backend/pattern_formatter_backend.py
    # So we can use a similar logic or hardcode relative to CWD if we know it.
    # Let's use absolute path logic similar to coverpage_generator
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    cover_page_dir = os.path.join(os.path.dirname(base_dir), 'outputs', 'Cover Pages')
    filepath = os.path.join(cover_page_dir, filename)
    
    inline = request.args.get('inline', 'false').lower() == 'true'
    as_attachment = not inline
    
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=as_attachment)
        
    # Fallback to standard output folder
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=as_attachment)
        
    return jsonify({'error': 'File not found'}), 404

# --- End Cover Page Generator Integration ---


import sys
if __name__ == '__main__':
    # Only start the server if not running tests
    # Check if running directly
    print(f"Starting backend server... (__name__={__name__})")
    if not any('unittest' in arg or 'test' in arg for arg in sys.argv):
        print("=" * 60)
        print("🚀 Pattern-Based Academic Document Formatter")
        print("=" * 60)
        print("✓ No AI dependencies - 100% pattern matching")
        print("✓ Ultra-fast processing - 1000+ lines/second")
        print("✓ 40+ regex patterns loaded")
        print("✓ Zero API costs")
        print("✓ 100% offline capability")
        print("=" * 60)
        print(f"Server starting at http://localhost:5000")
        print("=" * 60)
        app.run(debug=True, host='0.0.0.0', port=5000)
