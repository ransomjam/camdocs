import sys
import os
import json
import io
from docx import Document

# Add backend to path
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'backend'))

from pattern_formatter_backend import app, UPLOAD_FOLDER, OUTPUT_FOLDER

def test_merge_flow():
    client = app.test_client()
    
    # 1. Upload Document
    print("Step 1: Uploading document...")
    sample_file_path = os.path.join(os.path.dirname(__file__), 'sample_academic_paper.txt')
    
    with open(sample_file_path, 'rb') as f:
        data = {
            'file': (f, 'sample_academic_paper.txt')
        }
        response = client.post('/upload', data=data, content_type='multipart/form-data')
    
    if response.status_code != 200:
        print(f"Upload failed: {response.json}")
        return
        
    upload_result = response.json
    job_id = upload_result.get('job_id')
    print(f"Upload successful. Job ID: {job_id}")
    
    # Verify processed file exists
    processed_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    if os.path.exists(processed_path):
        print(f"Processed file exists at: {processed_path}")
    else:
        print("Processed file NOT found!")
        return

    # 2. Generate Cover Page with Merge
    print("\nStep 2: Generating Cover Page with Merge...")
    cover_data = {
        "documentType": "Assignment",
        "institution": "inst_001", # Assuming this ID exists or logic handles it, actually backend loads from json. 
        # If ID doesn't exist, it might fail or default. Let's provide raw text fields which usually work.
        "faculty": "Faculty of Science",
        "department": "Computer Science",
        "courseCode": "CSC 301",
        "courseTitle": "Cloud Computing",
        "title": "Cloud Computing Analysis",
        "studentName": "Test Student",
        "studentId": "TEST1234",
        "instructor": "Dr. Test Instructor",
        "date": "2025-10-10",
        "level": "300 Level",
        "mergeJobId": job_id
    }
    
    response = client.post('/api/coverpage/generate', 
                          data=json.dumps(cover_data),
                          content_type='application/json')
                          
    if response.status_code != 200:
        print(f"Cover page generation failed: {response.json}")
        return
        
    result = response.json
    print(f"Generation successful: {result}")
    
    # 3. Verify Merged Document
    print("\nStep 3: Verifying Merged Document...")
    # The download URL is /api/download/<filename>
    filename = result['filename']
    
    # We need to find where the file was saved. 
    # The backend saves it to OUTPUT_DIR in coverpage_generator, which is 'outputs/Cover Pages' relative to root usually.
    # But let's check the download logic.
    # It checks 'outputs/Cover Pages' then 'outputs'.
    
    # Let's try to find the file
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # Root
    cover_page_dir = os.path.join(base_dir, 'outputs', 'Cover Pages')
    filepath = os.path.join(cover_page_dir, filename)
    
    if not os.path.exists(filepath):
        print(f"File not found at expected path: {filepath}")
        # Try standard output
        filepath = os.path.join(base_dir, 'outputs', filename)
        if not os.path.exists(filepath):
             print(f"File not found at fallback path: {filepath}")
             return

    print(f"File found at: {filepath}")
    
    # Inspect content
    doc = Document(filepath)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    print(f"Document has {len(doc.paragraphs)} paragraphs.")
    
    # Check for Cover Page content
    has_student = "Test Student" in full_text
    has_instructor = "Dr. Test Instructor" in full_text
    
    # Check for Document content
    has_doc_content = "CLOUD COMPUTING AND ECONOMIES OF SCALE" in full_text or "Cloud computing refers to on-demand delivery" in full_text
    
    print(f"Contains Student Name: {has_student}")
    print(f"Contains Instructor Name: {has_instructor}")
    print(f"Contains Original Doc Content: {has_doc_content}")
    
    if has_student and has_doc_content:
        print("\nSUCCESS: Document merged successfully!")
    else:
        print("\nFAILURE: Merge incomplete.")

if __name__ == "__main__":
    test_merge_flow()
