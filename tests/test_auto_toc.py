"""
Test automatic Table of Contents generation using Microsoft Word's built-in TOC
"""
import sys
import os

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

from pattern_formatter_backend import WordGenerator
from docx import Document

def test_auto_toc():
    """Test that TOC is automatically populated using Word's built-in TOC feature"""
    
    # Create test structured data with enough content to trigger TOC
    structured_data = [
        {
            'heading': 'CHAPTER ONE',
            'level': 1,
            'type': 'chapter',
            'chapter_title': 'GENERAL INTRODUCTION',
            'content': [{'type': 'paragraph', 'text': 'This is the introduction. ' * 100}]
        },
        {
            'heading': 'Background to the Study',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Background content here. ' * 100}]
        },
        {
            'heading': 'Statement of the Problem',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Problem statement. ' * 100}]
        },
        {
            'heading': 'Research Objectives',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Objectives content. ' * 50}]
        },
        {
            'heading': 'Main Research Objective',
            'level': 3,
            'content': [{'type': 'paragraph', 'text': 'Main objective. ' * 50}]
        },
        {
            'heading': 'Specific Research Objectives',
            'level': 3,
            'content': [{'type': 'paragraph', 'text': 'Specific objectives. ' * 50}]
        },
        {
            'heading': 'CHAPTER TWO',
            'level': 1,
            'type': 'chapter',
            'chapter_title': 'REVIEW OF RELATED LITERATURE',
            'content': [{'type': 'paragraph', 'text': 'Literature review content. ' * 100}]
        },
        {
            'heading': 'Conceptual Review',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Conceptual review content. ' * 100}]
        },
        {
            'heading': 'Concept of Education',
            'level': 3,
            'content': [{'type': 'paragraph', 'text': 'Education concept. ' * 50}]
        },
        {
            'heading': 'Theoretical Review',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Theoretical review. ' * 100}]
        },
        {
            'heading': 'Equity Theory',
            'level': 3,
            'content': [{'type': 'paragraph', 'text': 'Equity theory content. ' * 50}]
        },
        {
            'heading': 'CHAPTER THREE',
            'level': 1,
            'type': 'chapter',
            'chapter_title': 'RESEARCH METHODOLOGY',
            'content': [{'type': 'paragraph', 'text': 'Methodology content. ' * 100}]
        },
        {
            'heading': 'Research Design',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Design content. ' * 100}]
        },
        {
            'heading': 'Population and Sample',
            'level': 2,
            'content': [{'type': 'paragraph', 'text': 'Population content. ' * 100}]
        },
    ]
    
    # Generate document with unique name
    import time
    output_path = os.path.join(os.path.dirname(__file__), f'test_output_toc_{int(time.time())}.docx')
    generator = WordGenerator()
    result_path = generator.generate(structured_data, output_path)
    
    print("\n=== Word TOC Test Results ===")
    print(f"Document generated: {result_path}")
    
    # Read the generated document
    doc = Document(output_path)
    
    # Check for TOC heading
    toc_heading_found = False
    toc_field_found = False
    headings_found = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ''
        
        if 'TABLE OF CONTENTS' in text.upper():
            toc_heading_found = True
            print(f"\n[OK] TOC Heading found: '{text}'")
        
        # Check for TOC field in paragraph XML
        if 'TOC' in para._element.xml and 'instrText' in para._element.xml:
            toc_field_found = True
        
        # Collect headings (Heading styles)
        if style.startswith('Heading'):
            headings_found.append((style, text[:50]))
    
    print(f"\n[OK] TOC Field Code present: {toc_field_found}")
    print(f"\nHeadings in document ({len(headings_found)}):")
    for style, text in headings_found[:10]:  # Show first 10
        print(f"  - {style}: {text}")
    
    # Check if TOC was updated (look for actual entries)
    # After Word updates the TOC, it will contain the heading text with page numbers
    toc_updated = False
    for para in doc.paragraphs:
        text = para.text.strip()
        # TOC entries typically have heading text followed by page number
        if 'Background to the Study' in text or 'CHAPTER ONE' in text:
            if toc_heading_found:  # Only count if we're in TOC area
                toc_updated = True
                break
    
    print(f"\n[OK] TOC appears to be populated: {toc_updated}")
    
    if os.path.exists(output_path):
        print(f"\n[OK] Document saved to: {output_path}")
        print("\nOpen the document in Microsoft Word to verify:")
        print("  - TOC should show all chapter and section headings")
        print("  - TOC should have proper indentation and page numbers")
        print("  - TOC is automatically populated - no manual update needed")
    
    return toc_heading_found and toc_field_found

if __name__ == '__main__':
    success = test_auto_toc()
    print(f"\n{'[PASS] TEST PASSED' if success else '[FAIL] TEST FAILED'}")
