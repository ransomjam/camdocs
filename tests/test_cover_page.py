"""
Test Cover Page Detection and Generation
Tests the CoverPageHandler class with sample_dissertation.docx
"""
import os
import sys

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

from docx import Document
from pattern_formatter_backend import CoverPageHandler, DocumentProcessor, WordGenerator


def test_cover_page_detection():
    """Test cover page detection with sample dissertation"""
    print("\n" + "="*60)
    print("TEST: Cover Page Detection")
    print("="*60)
    
    sample_path = os.path.join(os.path.dirname(__file__), '..', '..', 'Samples', 'sample_dissertation.docx')
    
    if not os.path.exists(sample_path):
        print(f"[SKIP] Sample file not found: {sample_path}")
        return 0, 1
    
    doc = Document(sample_path)
    handler = CoverPageHandler()
    
    has_cover, data, end_idx = handler.detect_and_extract(doc)
    
    passed = 0
    failed = 0
    
    # Test 1: Cover page detected
    if has_cover:
        print("[OK] Cover page detected")
        passed += 1
    else:
        print("[FAIL] Cover page not detected")
        failed += 1
    
    # Test 2: University extracted
    if data.get('university'):
        print(f"[OK] University: {data['university']}")
        passed += 1
    else:
        print("[FAIL] University not extracted")
        failed += 1
    
    # Test 3: Department extracted
    if data.get('department'):
        print(f"[OK] Department: {data['department']}")
        passed += 1
    else:
        print("[FAIL] Department not extracted")
        failed += 1
    
    # Test 4: Faculty extracted
    if data.get('faculty'):
        print(f"[OK] Faculty: {data['faculty']}")
        passed += 1
    else:
        print("[FAIL] Faculty not extracted")
        failed += 1
    
    # Test 5: Name extracted
    if data.get('name'):
        print(f"[OK] Name: {data['name']}")
        passed += 1
    else:
        print("[FAIL] Name not extracted")
        failed += 1
    
    # Test 6: Registration number
    if data.get('registration_number'):
        print(f"[OK] Registration: {data['registration_number']}")
        passed += 1
    else:
        print("[FAIL] Registration number not extracted")
        failed += 1
    
    # Test 7: Supervisor
    if data.get('supervisor'):
        print(f"[OK] Supervisor: {data['supervisor']}")
        passed += 1
    else:
        print("[FAIL] Supervisor not extracted")
        failed += 1
    
    # Test 8: Date
    if data.get('month_year'):
        print(f"[OK] Date: {data['month_year']}")
        passed += 1
    else:
        print("[FAIL] Date not extracted")
        failed += 1
    
    # Test 9: Cover page ends before DECLARATION (should be reasonable range)
    if end_idx > 0 and end_idx < 80:
        print(f"[OK] Cover page ends at index {end_idx}")
        passed += 1
    else:
        print(f"[FAIL] Invalid cover page end index: {end_idx}")
        failed += 1
    
    print(f"\nPassed: {passed}, Failed: {failed}")
    return passed, failed


def test_full_processing():
    """Test full document processing with cover page generation"""
    print("\n" + "="*60)
    print("TEST: Full Document Processing with Cover Page")
    print("="*60)
    
    sample_path = os.path.join(os.path.dirname(__file__), '..', '..', 'Samples', 'sample_dissertation.docx')
    output_path = os.path.join(os.path.dirname(__file__), 'outputs', 'test_cover_page_output.docx')
    
    # Ensure output folder exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    if not os.path.exists(sample_path):
        print(f"[SKIP] Sample file not found: {sample_path}")
        return 0, 1
    
    passed = 0
    failed = 0
    
    # Process document
    processor = DocumentProcessor()
    result, images = processor.process_docx(sample_path)
    
    # Check cover page was detected
    if processor.cover_page_data:
        print("[OK] Cover page data extracted during processing")
        passed += 1
        
        # Print extracted data
        print("\nExtracted cover page data:")
        for key, value in processor.cover_page_data.items():
            if value:
                print(f"  {key}: {value}")
    else:
        print("[FAIL] Cover page not detected during processing")
        failed += 1
    
    # Generate output document
    generator = WordGenerator()
    generator.generate(
        result['structured'],
        output_path,
        images=images,
        cover_page_data=processor.cover_page_data
    )
    
    # Verify output was created
    if os.path.exists(output_path):
        size = os.path.getsize(output_path)
        print(f"\n[OK] Output document created: {output_path}")
        print(f"     File size: {size} bytes")
        passed += 1
        
        # Verify cover page is in output
        output_doc = Document(output_path)
        first_para = output_doc.paragraphs[0].text if output_doc.paragraphs else ""
        
        # Check if logo paragraph or university name is first
        has_cover = False
        for para in output_doc.paragraphs[:5]:
            if 'UNIVERSITY OF BAMENDA' in para.text.upper():
                has_cover = True
                break
        
        if has_cover:
            print("[OK] Cover page found in output document")
            passed += 1
        else:
            print("[FAIL] Cover page not found in output document")
            failed += 1
    else:
        print(f"[FAIL] Output document not created")
        failed += 1
    
    print(f"\nPassed: {passed}, Failed: {failed}")
    return passed, failed


def main():
    """Run all tests"""
    print("="*60)
    print("COVER PAGE TEST SUITE")
    print("="*60)
    
    total_passed = 0
    total_failed = 0
    
    # Run tests
    p, f = test_cover_page_detection()
    total_passed += p
    total_failed += f
    
    p, f = test_full_processing()
    total_passed += p
    total_failed += f
    
    # Summary
    print("\n" + "="*60)
    print("FINAL SUMMARY")
    print("="*60)
    print(f"Total Passed: {total_passed}")
    print(f"Total Failed: {total_failed}")
    print(f"Success Rate: {100*total_passed/(total_passed+total_failed):.1f}%")
    
    return 0 if total_failed == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
