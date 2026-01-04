"""
Test script for image extraction and reinsertion.
Tests with sample_report_with images.docx
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

from docx import Document
from docx.oxml.ns import qn
from io import BytesIO

# Test document path
SAMPLE_DOC = r"C:\Users\user\Desktop\PATTERN\Samples\sample_report_with images.docx"
OUTPUT_DOC = r"C:\Users\user\Desktop\PATTERN\Samples\sample_report_output_with_images.docx"

def test_basic_image_detection():
    """Test basic image detection in the document"""
    print("=" * 60)
    print("TESTING IMAGE DETECTION")
    print("=" * 60)
    
    if not os.path.exists(SAMPLE_DOC):
        print(f"ERROR: Sample document not found at {SAMPLE_DOC}")
        return
    
    doc = Document(SAMPLE_DOC)
    
    print(f"\nDocument has {len(doc.paragraphs)} paragraphs")
    print(f"Document has {len(doc.tables)} tables")
    
    # Check document relationships for images
    print("\n--- Checking Document Relationships ---")
    image_rels = []
    for rel_id, rel in doc.part.rels.items():
        if 'image' in rel.target_ref.lower() or hasattr(rel, 'target_part') and rel.target_part and 'image' in str(rel.target_part.content_type):
            image_rels.append((rel_id, rel))
            print(f"  Found image relationship: {rel_id} -> {rel.target_ref}")
    
    print(f"\nTotal image relationships found: {len(image_rels)}")
    
    # Check for inline shapes in paragraphs
    print("\n--- Checking Paragraphs for Images ---")
    image_count = 0
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()[:50] if para.text.strip() else "(empty)"
        
        # Check for drawing elements
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            print(f"  Para {i}: Found {len(drawings)} drawing(s) - Text: {para_text}")
            image_count += len(drawings)
            
            # Check for blips inside drawings
            for drawing in drawings:
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        print(f"    -> Blip with embed rId: {embed}")
        
        # Check for pictures
        pics = para._element.findall('.//' + qn('pic:pic'))
        if pics:
            print(f"  Para {i}: Found {len(pics)} picture(s)")
    
    print(f"\nTotal images found in paragraphs: {image_count}")
    
    # Check inline shapes
    print("\n--- Checking for Inline Shapes ---")
    for i, para in enumerate(doc.paragraphs):
        inlines = para._element.findall('.//' + qn('wp:inline'))
        anchors = para._element.findall('.//' + qn('wp:anchor'))
        
        if inlines:
            print(f"  Para {i}: Found {len(inlines)} inline shape(s)")
        if anchors:
            print(f"  Para {i}: Found {len(anchors)} anchor shape(s)")


def test_image_extraction_improved():
    """Test improved image extraction method"""
    print("\n" + "=" * 60)
    print("TESTING IMPROVED EXTRACTION")
    print("=" * 60)
    
    if not os.path.exists(SAMPLE_DOC):
        print(f"ERROR: Sample document not found at {SAMPLE_DOC}")
        return
    
    doc = Document(SAMPLE_DOC)
    images = []
    
    # Method 1: Check all relationships
    print("\n--- Method 1: Extract from relationships ---")
    for rel_id, rel in doc.part.rels.items():
        if hasattr(rel, 'target_part') and rel.target_part:
            content_type = rel.target_part.content_type
            if 'image' in content_type:
                image_bytes = rel.target_part.blob
                print(f"  Image: {rel_id}, Type: {content_type}, Size: {len(image_bytes)} bytes")
                images.append({
                    'rel_id': rel_id,
                    'content_type': content_type,
                    'size': len(image_bytes),
                    'data': image_bytes
                })
    
    print(f"\nTotal images from relationships: {len(images)}")
    
    # Method 2: Track position in document
    print("\n--- Method 2: Track image positions ---")
    for para_idx, para in enumerate(doc.paragraphs):
        # Find all drawing elements
        for run in para.runs:
            drawings = run._element.findall('.//' + qn('w:drawing'))
            for drawing in drawings:
                # Find blip (binary large image pointer)
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        print(f"  Para {para_idx}: Image rId={embed}")
                        
                        # Get dimensions
                        extent = drawing.find('.//' + qn('wp:extent'))
                        if extent is not None:
                            cx = int(extent.get('cx', 0))
                            cy = int(extent.get('cy', 0))
                            width = cx / 914400  # EMU to inches
                            height = cy / 914400
                            print(f"    Dimensions: {width:.2f}\" x {height:.2f}\"")


def test_with_pattern_formatter():
    """Test using the actual pattern formatter"""
    print("\n" + "=" * 60)
    print("TESTING WITH PATTERN FORMATTER")
    print("=" * 60)
    
    try:
        from pattern_formatter_backend import ImageExtractor, DocumentProcessor
        
        extractor = ImageExtractor()
        images = extractor.extract_all_images(SAMPLE_DOC)
        
        print(f"\nImageExtractor found {len(images)} images:")
        for img in images:
            print(f"  - {img['image_id']}: {img['format']}, {img.get('width', 'N/A'):.2f}\"x{img.get('height', 'N/A'):.2f}\"")
            if img.get('caption'):
                print(f"    Caption: {img['caption'][:50]}...")
        
        print("\n--- Testing Full Processing ---")
        processor = DocumentProcessor()
        result, extracted_images = processor.process_docx(SAMPLE_DOC)
        
        print(f"\nProcessing results:")
        print(f"  Total lines: {result['stats']['total_lines']}")
        print(f"  Images: {result['stats'].get('images', 0)}")
        print(f"  Extracted images: {len(extracted_images)}")
        
        # Check for image placeholders in structured data
        image_placeholders = 0
        for section in result.get('structured', []):
            for item in section.get('content', []):
                if item.get('type') == 'image_placeholder':
                    image_placeholders += 1
                    print(f"  Found placeholder: {item.get('image_id')}")
        
        print(f"\nImage placeholders in structure: {image_placeholders}")
        
    except Exception as e:
        import traceback
        print(f"Error: {e}")
        traceback.print_exc()


def test_full_document_generation():
    """Test full document generation with images"""
    print("\n" + "=" * 60)
    print("TESTING FULL DOCUMENT GENERATION WITH IMAGES")
    print("=" * 60)
    
    if not os.path.exists(SAMPLE_DOC):
        print(f"ERROR: Sample document not found at {SAMPLE_DOC}")
        return
    
    from pattern_formatter_backend import DocumentProcessor, WordGenerator
    
    try:
        processor = DocumentProcessor()
        result, extracted_images = processor.process_docx(SAMPLE_DOC)
        
        print(f"\nProcessing complete:")
        print(f"  Sections: {len(result.get('structured', []))}")
        print(f"  Extracted images: {len(extracted_images)}")
        
        # Generate Word document
        generator = WordGenerator()
        generator.generate(result['structured'], OUTPUT_DOC, images=extracted_images)
        
        print(f"\nOutput document created: {OUTPUT_DOC}")
        
        # Verify output document has images
        output_doc = Document(OUTPUT_DOC)
        image_rels = []
        for rel_id, rel in output_doc.part.rels.items():
            if hasattr(rel, 'target_part') and rel.target_part:
                if 'image' in rel.target_part.content_type:
                    image_rels.append(rel_id)
        
        print(f"Images in output document: {len(image_rels)}")
        
        if len(image_rels) > 0:
            print("\n*** SUCCESS: Images preserved in output document! ***")
        else:
            print("\n*** WARNING: No images found in output document ***")
        
    except Exception as e:
        import traceback
        print(f"Error: {e}")
        traceback.print_exc()


if __name__ == '__main__':
    test_basic_image_detection()
    test_image_extraction_improved()
    test_with_pattern_formatter()
    test_full_document_generation()
