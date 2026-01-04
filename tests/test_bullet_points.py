import sys
import os
import json

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

try:
    from pattern_formatter_backend import PatternEngine, DocumentProcessor, WordGenerator
except ImportError:
    print("Error: Could not import backend modules.")
    sys.exit(1)

def test_bullet_points():
    input_path = r"c:\Users\user\Desktop\PATTERN\Samples\sample report with bullet points.docx"
    output_path = r"c:\Users\user\Desktop\PATTERN\Samples\test_bullet_output.docx"
    
    if not os.path.exists(input_path):
        print(f"Error: Input file not found at {input_path}")
        return

    print(f"Processing {input_path}...")
    processor = DocumentProcessor()
    
    # Process the document
    try:
        # First, let's inspect the raw lines to see what we are dealing with
        from docx import Document
        doc = Document(input_path)
        print("\n--- SCANNING FOR LIST ITEMS ---")
        list_items_found = 0
        for i, para in enumerate(doc.paragraphs):
            if para._element.pPr is not None and para._element.pPr.numPr is not None:
                numId = para._element.pPr.numPr.numId.val
                ilvl = para._element.pPr.numPr.ilvl.val
                print(f"Para {i}: '{para.text}' -> Has numbering! (numId: {numId}, ilvl: {ilvl})")
                list_items_found += 1
                if list_items_found > 20:
                    print("... (stopping scan after 20 items)")
                    break
        
        if list_items_found == 0:
            print("No automatic list items found. Checking for manual bullets...")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip().startswith(('•', '○', '-', '1.', 'a.')):
                     print(f"Para {i}: '{para.text}' -> Looks like manual bullet")
                     list_items_found += 1
                     if list_items_found > 20: break

        result_tuple = processor.process_docx(input_path)
        if isinstance(result_tuple, tuple):
            result, images = result_tuple
        else:
            result = result_tuple
            images = []
            
        print("Document processed successfully.")
        
        # Inspect the structure for bullet lists
        structured_data = result['structured']
        bullet_count = 0
        
        print("\n--- DETECTED BULLET LISTS ---")
        for section in structured_data:
            for item in section.get('content', []):
                if item['type'] == 'bullet_list':
                    bullet_count += 1
                    print(f"\nList found in section '{section.get('heading', 'Unknown')}'")
                    items = item.get('items', [])
                    print(f"Total items: {len(items)}")
                    
                    for i, list_item in enumerate(items):
                        if isinstance(list_item, dict) and 'bullet_info' in list_item:
                            info = list_item['bullet_info']
                            # Test mapping
                            from pattern_formatter_backend import map_to_word_bullet_style
                            mapped = map_to_word_bullet_style(info)
                            print(f"  Item {i+1}: Type='{info.get('type')}', Char='{info.get('bullet_char')}', Mapped={mapped}")
                            print(f"          Content: {info.get('content')[:50]}...")
                        else:
                            # Legacy or simple string
                            content = list_item if isinstance(list_item, str) else list_item.get('content', '')
                            print(f"  Item {i+1}: [Simple String] {content[:50]}...")
                            
        print(f"\nTotal bullet lists found: {bullet_count}")
        
        # Generate output
        print(f"\nGenerating output to {output_path}...")
        generator = WordGenerator()
        generator.generate(structured_data, output_path, images=images)
        print("Generation complete.")
        
    except Exception as e:
        print(f"Error during processing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_bullet_points()
